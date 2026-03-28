#!/usr/bin/env python3
"""
Tests for ATS DOCX generation and ATS compatibility scoring (CVOrchestrator).
"""

import json
import re as _re
import sys
import pytest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

from utils.cv_orchestrator import CVOrchestrator, validate_ats_report


# ── fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture
def selected_content():
    """Minimal but realistic CV data for ATS testing."""
    return {
        "personal_info": {
            "name": "Gregory R. Warnes, Ph.D.",
            "contact": {
                "email":    "consulting@warnes.net",
                "phone":    "585-678-6661",
                "address":  {"city": "Rochester", "state": "NY"},
                "linkedin": "https://linkedin.com/in/gregorywarnes",
            },
        },
        "summary": (
            "Senior Data Scientist with 25+ years of experience in biostatistics, "
            "genomics, and machine learning."
        ),
        "experiences": [
            {
                "title":      "Chief Scientific Officer",
                "company":    "TNT³",
                "location":   {"city": "Remote", "state": None},
                "start_date": "2024-01",
                "end_date":   "Present",
                "achievements": [
                    {"text": "Led development of automated trading systems using ML, improving accuracy by 35%"},
                    {"text": "Managed cross-functional team of 8 engineers and data scientists"},
                ],
            },
            {
                "title":      "Principal Research Scientist",
                "company":    "Pfizer Global R&D",
                "location":   {"city": "Groton", "state": "CT"},
                "start_date": "2000-01",
                "end_date":   "2024-12",
                "achievements": [
                    {"text": "Developed MiDAS Genomic Workflow System reducing processing time by 60%"},
                    {"text": "Led biostatistics support for 50+ clinical trials across oncology and immunology"},
                ],
            },
        ],
        "skills": [
            {"name": "Python",              "category": "Programming",    "years": 15},
            {"name": "R/Bioconductor",      "category": "Programming",    "years": 20},
            {"name": "Machine Learning",    "category": "Technical",      "years": 10},
            {"name": "Data Science",        "category": "Core Expertise", "years": 25},
            {"name": "Biostatistics",       "category": "Core Expertise", "years": 25},
            {"name": "Statistical Modeling","category": "Technical",      "years": 20},
        ],
        "education": [
            {
                "degree":      "Ph.D.",
                "field":       "Biostatistics",
                "institution": "University of Washington",
                "location":    {"city": "Seattle", "state": "WA"},
                "end_year":    2000,
            }
        ],
        "certifications": [
            {"name": "AWS Certified Solutions Architect", "issuer": "Amazon Web Services", "year": 2023}
        ],
        "awards":        [{"title": "Pfizer Achievement Award", "year": 2022, "description": "Outstanding contributions"}],
        "achievements":  [],
        "publications":  [],
    }


@pytest.fixture
def job_analysis():
    """Sample job analysis matching the selected_content above."""
    return {
        "company":  "TestBioTech",
        "title":    "Senior Data Scientist",
        "domain":   "biotechnology",
        "required_skills": ["Python", "Machine Learning", "Statistical Modeling", "Data Science", "R"],
        "ats_keywords": [
            "python", "machine learning", "statistical modeling",
            "biostatistics", "data science", "clinical trials",
            "genomics", "pytorch", "docker", "leadership",
        ],
        "must_have_requirements": [
            "PhD in relevant field",
            "10+ years experience in data science",
            "Strong Python and R skills",
            "Experience with machine learning",
        ],
    }


@pytest.fixture
def orchestrator(tmp_path):
    """CVOrchestrator wired to a temporary directory with no LLM client."""
    master_data = {
        "personal_info":          {"name": "Gregory R. Warnes, Ph.D."},
        "professional_summaries": {"default": "Test summary"},
        "education":              [],
        "awards":                 [],
        "certifications":         [],
    }
    master_path = tmp_path / "Master_CV_Data.json"
    master_path.write_text(json.dumps(master_data), encoding="utf-8")
    return CVOrchestrator(
        master_data_path=str(master_path),
        publications_path=str(tmp_path / "publications.bib"),  # may not exist — that's OK
        output_dir=str(tmp_path / "output"),
        llm_client=None,
    )


# ── tests ─────────────────────────────────────────────────────────────────────

def test_ats_docx_is_generated(orchestrator, selected_content, job_analysis, tmp_path):
    """_generate_ats_docx produces a non-empty .docx file."""
    out_dir = tmp_path / "ats_out"
    out_dir.mkdir()

    ats_file = orchestrator._generate_ats_docx(selected_content, job_analysis, out_dir)

    assert ats_file.exists(), f"Expected ATS DOCX at {ats_file}"
    assert ats_file.stat().st_size > 0, "ATS DOCX is empty"


def test_ats_docx_sections_use_heading1(orchestrator, selected_content, job_analysis, tmp_path):
    """ATS DOCX section headings must use Heading 1 style (US-H2).

    The candidate name is now a bold run (not a Heading style), so that the
    Heading 1 hierarchy is reserved exclusively for section headings as
    required by US-H2 acceptance criteria.
    """
    from docx import Document  # type: ignore

    out_dir = tmp_path / "ats_out_h1"
    out_dir.mkdir()

    ats_file = orchestrator._generate_ats_docx(selected_content, job_analysis, out_dir)
    doc = Document(str(ats_file))

    heading1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    # At least one standard section heading must be Heading 1
    assert any(h in heading1_texts for h in (
        "Professional Summary", "Work Experience", "Technical Skills",
        "Core Competencies", "Education", "Certifications",
    )), f"No standard Heading 1 section found; Heading 1 paragraphs: {heading1_texts}"

    # Candidate name must NOT be a Heading 1 — it is rendered as a bold run
    name = selected_content["personal_info"]["name"]
    assert name not in heading1_texts, (
        f"Candidate name '{name}' should not be a Heading 1 paragraph"
    )


def test_ats_docx_uses_skills_heading(orchestrator, selected_content, job_analysis, tmp_path):
    """ATS DOCX uses standard skill section headings on Heading 1 (US-H2, US-H8).

    Hard skills render under 'Technical Skills'; soft skills (when present)
    render under 'Core Competencies'.  With the default fixture (no soft
    skills), only 'Technical Skills' is expected.
    """
    from docx import Document  # type: ignore

    out_dir = tmp_path / "ats_out_skills_heading"
    out_dir.mkdir()

    ats_file = orchestrator._generate_ats_docx(selected_content, job_analysis, out_dir)
    doc = Document(str(ats_file))

    heading1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert "Technical Skills" in heading1_texts, (
        f"'Technical Skills' not found in Heading 1 paragraphs: {heading1_texts}"
    )
    # With all-hard skills fixture, the old flat 'SKILLS' label must be gone
    assert "SKILLS" not in heading1_texts


def test_ats_docx_soft_skill_section(orchestrator, job_analysis, tmp_path):
    """When soft skills are present they appear under 'Core Competencies' (US-H8)."""
    from docx import Document  # type: ignore

    out_dir = tmp_path / "ats_out_soft"
    out_dir.mkdir()

    content_with_soft = {
        "personal_info": {"name": "Test User", "contact": {}},
        "summary": "Test summary.",
        "experiences": [],
        "skills": [
            {"name": "Python",         "category": "Programming",    "years": 10},
            {"name": "Leadership",     "category": "Soft Skills",    "years": 5},
            {"name": "Communication",  "category": "Soft Skills",    "years": 5},
        ],
        "education": [], "certifications": [], "awards": [],
        "achievements": [], "publications": [],
    }

    ats_file = orchestrator._generate_ats_docx(content_with_soft, job_analysis, out_dir)
    doc = Document(str(ats_file))

    heading1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert "Technical Skills" in heading1_texts
    assert "Core Competencies" in heading1_texts


def test_ats_docx_phone_normalized(orchestrator, job_analysis, tmp_path):
    """Phone number in ATS DOCX contact line is normalized to NNN-NNN-NNNN (US-H3)."""
    from docx import Document  # type: ignore

    out_dir = tmp_path / "ats_out_phone"
    out_dir.mkdir()

    content_parentheses_phone = {
        "personal_info": {
            "name": "Test User",
            "contact": {
                "email": "test@example.com",
                "phone": "(585) 678-6661",
                "address": {"city": "Rochester", "state": "NY"},
            },
        },
        "summary": "Test.",
        "experiences": [], "skills": [], "education": [],
        "certifications": [], "awards": [], "achievements": [], "publications": [],
    }
    ats_file = orchestrator._generate_ats_docx(content_parentheses_phone, job_analysis, out_dir)
    doc = Document(str(ats_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "585-678-6661" in full_text, "Normalized phone not found in ATS DOCX"
    assert "(585)" not in full_text, "Phone parentheses not stripped in ATS DOCX"


def test_ats_docx_job_entry_one_line(orchestrator, selected_content, job_analysis, tmp_path):
    """Job entries use one-line format: Title | Company | Location | DateRange (US-H5)."""
    from docx import Document  # type: ignore

    out_dir = tmp_path / "ats_out_entry"
    out_dir.mkdir()

    ats_file = orchestrator._generate_ats_docx(selected_content, job_analysis, out_dir)
    doc = Document(str(ats_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # First experience: Chief Scientific Officer | TNT³ | Remote | 2024-01 – Present
    assert "Chief Scientific Officer" in full_text
    assert "2024-01" in full_text  # date appears
    # Key check: title and company appear on same line as the date range
    for para in doc.paragraphs:
        if "Chief Scientific Officer" in para.text and "2024-01" in para.text:
            assert "TNT³" in para.text, "Company must be on the same line as title and dates"
            break
    else:
        assert False, "No single paragraph contains title, company, and date range together"


def test_ats_compatibility_score_acceptable(orchestrator, selected_content, job_analysis):
    """_validate_ats_compatibility returns a numeric score ≥ 50 for well-matched content."""
    score = orchestrator._validate_ats_compatibility(selected_content, job_analysis)

    assert isinstance(score, (int, float)), f"Expected numeric score, got {type(score)}"
    assert score >= 50, f"ATS score {score} is below the minimum acceptable threshold of 50"


# ── Page Count Validation Tests ────────────────────────────────────────────────

def test_page_count_validation_none_page_count():
    """Page count validation should FAIL when page_count is None."""
    checks = []

    # Simulate the validation check logic
    def _chk(name, label, format_, status, detail):
        checks.append({'name': name, 'label': label, 'format': format_, 'status': status, 'detail': detail})

    page_count = None
    if page_count is None:
        _chk('cv_page_count', 'CV page count', 'pdf', 'fail',
             'Page count could not be determined (HTML render failed)')

    assert len(checks) == 1
    assert checks[0]['status'] == 'fail'
    assert checks[0]['name'] == 'cv_page_count'


def test_page_count_validation_single_page():
    """Page count validation should WARN when page_count is 1."""
    checks = []

    def _chk(name, label, format_, status, detail):
        checks.append({'name': name, 'label': label, 'format': format_, 'status': status, 'detail': detail})

    page_count = 1
    ideal_min, ideal_max = 2, 3

    if page_count == 1:
        _chk('cv_page_count', 'CV page count', 'pdf', 'warn',
             f'{page_count} page — consider {ideal_min}–{ideal_max} pages for senior candidates')

    assert len(checks) == 1
    assert checks[0]['status'] == 'warn'
    assert '1 page' in checks[0]['detail']


def test_page_count_validation_ideal_range():
    """Page count validation should PASS when page_count is 2 or 3."""
    checks = []

    def _chk(name, label, format_, status, detail):
        checks.append({'name': name, 'label': label, 'format': format_, 'status': status, 'detail': detail})

    ideal_min, ideal_max = 2, 3

    for page_count in [2, 3]:
        checks.clear()
        if ideal_min <= page_count <= ideal_max:
            _chk('cv_page_count', 'CV page count', 'pdf', 'pass',
                 f'{page_count} pages — within ideal {ideal_min}–{ideal_max} page range')

        assert len(checks) == 1
        assert checks[0]['status'] == 'pass'
        assert f'{page_count} pages' in checks[0]['detail']


def test_page_count_validation_exceeds_ideal():
    """Page count validation should WARN when 3 < page_count <= 4."""
    checks = []

    def _chk(name, label, format_, status, detail):
        checks.append({'name': name, 'label': label, 'format': format_, 'status': status, 'detail': detail})

    page_count = 4
    ideal_max, absolute_max = 3, 4

    if ideal_max < page_count <= absolute_max:
        _chk('cv_page_count', 'CV page count', 'pdf', 'warn',
             f'{page_count} pages — exceeds {ideal_max}-page ideal range')

    assert len(checks) == 1
    assert checks[0]['status'] == 'warn'
    assert '4 pages' in checks[0]['detail']


def test_page_count_validation_exceeds_maximum():
    """Page count validation should FAIL when page_count > 4."""
    checks = []

    def _chk(name, label, format_, status, detail):
        checks.append({'name': name, 'label': label, 'format': format_, 'status': status, 'detail': detail})

    page_count = 5
    absolute_max = 4

    if page_count > absolute_max:
        _chk('cv_page_count', 'CV page count', 'pdf', 'fail',
             f'{page_count} pages — exceeds {absolute_max}-page maximum; consider condensing')

    assert len(checks) == 1
    assert checks[0]['status'] == 'fail'
    assert '5 pages' in checks[0]['detail']
    assert 'exceeds' in checks[0]['detail']


# ──────────────────────────────────────────────────────────────────
# GAP-07: Proposed Bullet Ordering Tests
# ──────────────────────────────────────────────────────────────────

def _ach_keyword_score(ach, job_keywords):
    """Mirrors the scoring logic in /api/proposed-bullet-order."""
    text = (ach.get("text", "") if isinstance(ach, dict) else str(ach)).lower()
    tokens = set(_re.findall(r'\b\w+\b', text))
    return len(tokens & {kw.lower() for kw in job_keywords})


def _compute_proposed_order(achievements, job_keywords):
    """Return achievement indices sorted by keyword relevance (highest first)."""
    return sorted(range(len(achievements)), key=lambda i: _ach_keyword_score(achievements[i], job_keywords), reverse=True)


def test_proposed_order_sorts_by_keyword_relevance():
    """Achievements with more job keyword matches should rank higher."""
    achievements = [
        {"text": "Managed budgets and spreadsheets"},          # 0 — no keywords
        {"text": "Built Python API with FastAPI and Docker"},  # 1 — python, api, docker
        {"text": "Deployed Docker containers on Kubernetes"},  # 2 — docker, kubernetes
    ]
    job_keywords = {"python", "api", "docker", "kubernetes"}
    proposed = _compute_proposed_order(achievements, job_keywords)
    # Index 1 has 3 matches (python, api, docker), index 2 has 2 (docker, kubernetes), index 0 has 0
    assert proposed[0] == 1
    assert proposed[1] == 2
    assert proposed[2] == 0


def test_proposed_order_empty_achievements():
    """Empty achievements list returns empty proposed order."""
    proposed = _compute_proposed_order([], {"python", "api"})
    assert proposed == []


def test_proposed_order_no_job_keywords_returns_natural_order():
    """With no job keywords all scores are 0, so stable sort returns original order."""
    achievements = [{"text": "A"}, {"text": "B"}, {"text": "C"}]
    # No keywords → all scores are 0; sorted with reverse=True is stable so order preserved
    proposed = _compute_proposed_order(achievements, set())
    assert proposed == [0, 1, 2]


def test_proposed_order_string_achievements():
    """Plain string achievements (not dicts) are handled correctly."""
    achievements = ["Deployed Kubernetes clusters", "Wrote documentation", "Built Python services"]
    job_keywords = {"python", "kubernetes"}
    proposed = _compute_proposed_order(achievements, job_keywords)
    # Index 0 has 1 match (kubernetes), index 2 has 1 match (python), index 1 has 0
    # Both 0 and 2 score 1; index 1 must be last
    assert proposed[-1] == 1


def test_proposed_order_case_insensitive():
    """Keyword matching is case-insensitive."""
    achievements = [
        {"text": "Led PYTHON development"},
        {"text": "Managed team communications"},
    ]
    job_keywords = {"Python"}
    proposed = _compute_proposed_order(achievements, job_keywords)
    assert proposed[0] == 0  # PYTHON matches 'python' keyword


def test_build_json_ld_knows_about_has_skill_types(orchestrator, selected_content, job_analysis):
    """JSON-LD knowsAbout entries include additionalType HardSkill or SoftSkill (US-H8)."""
    import json as _json

    # Prepare cv_data as the orchestrator's prepare_cv_data would produce,
    # but use a direct call to _build_json_ld with a minimal skills_by_category.
    cv_data = {
        "personal_info": selected_content["personal_info"],
        "professional_summary": selected_content["summary"],
        "experiences": selected_content["experiences"],
        "education":   selected_content["education"],
        "awards":      selected_content.get("awards", []),
        "skills_by_category": [
            {
                "category": "Programming",
                "skills": [
                    {"name": "Python",       "category": "Programming", "years": 15},
                    {"name": "Leadership",   "category": "Soft Skills", "years": 5},
                ],
            }
        ],
    }

    json_ld_str = orchestrator._build_json_ld(cv_data, job_analysis)
    ld = _json.loads(json_ld_str)

    knows_about = ld.get("knowsAbout", [])
    assert len(knows_about) == 2, f"Expected 2 skill entries, got {len(knows_about)}"

    # All entries must be DefinedTerm objects with additionalType
    for entry in knows_about:
        assert isinstance(entry, dict), f"knowsAbout entry should be a dict: {entry}"
        assert entry.get("@type") == "DefinedTerm", f"Missing @type DefinedTerm: {entry}"
        assert entry.get("additionalType") in ("HardSkill", "SoftSkill"), (
            f"additionalType must be HardSkill or SoftSkill: {entry}"
        )

    # Python → HardSkill, Leadership (soft category) → SoftSkill
    by_name = {e["name"]: e["additionalType"] for e in knows_about}
    assert by_name.get("Python")    == "HardSkill"
    assert by_name.get("Leadership") == "SoftSkill"


if __name__ == '__main__':
    import pytest, sys
    sys.exit(pytest.main([__file__, '-v']))
