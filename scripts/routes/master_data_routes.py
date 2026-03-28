"""Master data management, summary, cover letter, and screening routes."""
import logging
import json
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from flask import Blueprint, jsonify, request
from werkzeug.utils import safe_join

# Live blueprint module registered by `scripts.web_app.create_app()`.

from utils.bibtex_parser import (
    bibtex_text_to_publications,
    format_publication,
    parse_bibtex_file,
    serialize_publications_to_bibtex,
)
from utils.llm_client import LLMError


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Module-level IO helpers (for testability)
# ---------------------------------------------------------------------------

def _load_master(master_data_path: str) -> "tuple[dict, Path]":
    """Read master CV JSON from disk and return (data, path)."""
    p = Path(master_data_path)
    with open(p, 'r', encoding='utf-8') as f:
        return json.load(f), p


def _save_master(master: Dict[str, Any], master_path: Path) -> None:
    """Write master CV data to disk, create a timestamped backup, and stage in git."""
    backup_dir = master_path.parent / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    backup_path = backup_dir / f"Master_CV_{ts}.json"
    if master_path.exists():
        shutil.copy2(master_path, backup_path)
    with open(master_path, 'w', encoding='utf-8') as f:
        json.dump(master, f, indent=2)
    subprocess.run(
        ['git', '-C', str(master_path.parent), 'add', master_path.name],
        capture_output=True, check=False,
    )


def _resolve_backup_path(backup_dir: Path, filename: str) -> Path | None:
    """Return a validated backup path constrained to ``backup_dir``."""
    safe_path = safe_join(str(backup_dir), filename)
    if safe_path is None:
        return None
    return Path(safe_path)


# ---------------------------------------------------------------------------
# Screening format guidance
# ---------------------------------------------------------------------------

_SCREENING_FORMAT_GUIDANCE: dict = {
    'direct':    ('Direct/Concise',    '150–200 words',
                  'Be clear and direct. State the answer, give one concrete example, close concisely.'),
    'star':      ('STAR',              '250–350 words',
                  'Use the STAR framework: Situation, Task, Action, Result. 1–2 sentences each.'),
    'technical': ('Technical Detail', '400–500 words',
                  'Provide full technical depth: context, methodology, tools/technologies, outcomes with metrics.'),
}

# Cover letter tone guidance
_TONE_GUIDANCE: Dict[str, str] = {
    'startup/tech':   'Energetic, direct, outcome-focused.  Emphasise velocity, impact, and technical depth.',
    'pharma/biotech': 'Precise, methodical, compliance-aware.  Reference domain credentials and regulatory rigour.',
    'academia':       'Scholarly, collaborative.  Highlight publications, teaching experience, and departmental service.',
    'financial':      'Professional, quantitative, risk-aware.  Emphasise fiduciary responsibility and data-driven decisions.',
    'leadership':     'Strategic, vision-focused, people-first.  Highlight team-building and organisational impact.',
}

# Text similarity helper (used in screening search)
def _text_similarity(query: str, target: str) -> float:
    """Simple word-overlap similarity score (0–1) for response library search."""
    _STOP = {
        'a', 'an', 'the', 'and', 'or', 'for', 'in', 'of', 'to', 'is',
        'are', 'was', 'were', 'i', 'my', 'your', 'we', 'our', 'this',
        'that', 'it', 'with', 'as', 'by', 'at', 'on', 'be', 'have',
        'has', 'had', 'do', 'does', 'did', 'will', 'would', 'can',
        'could', 'should', 'may', 'might', 'from', 'into', 'about',
    }
    def _tok(s: str) -> set:
        return {w.lower() for w in re.findall(r'\w+', s) if w.lower() not in _STOP and len(w) > 2}
    q_tok = _tok(query)
    t_tok = _tok(target)
    if not q_tok or not t_tok:
        return 0.0
    return len(q_tok & t_tok) / max(len(q_tok), len(t_tok))


def create_blueprint(deps):
    bp = Blueprint('master_data_routes', __name__)

    get_session = deps['get_session']
    validate_owner = deps['validate_owner']
    session_registry = deps['session_registry']
    llm_client_ref = deps['llm_client_ref']
    load_master = deps.get('load_master', _load_master)
    save_master = deps.get('save_master', _save_master)
    validate_master_data_file_fn = deps.get('validate_master_data_file')

    def _extract_year(value: Any) -> Optional[int]:
        if value is None:
            return None
        text = str(value).strip()
        if not text:
            return None
        m = re.search(r'(19|20)\d{2}', text)
        return int(m.group(0)) if m else None

    def _require_master_data_write_phase(entry):
        """Allow direct master-data writes only in the pre-job and post-job windows."""
        raw_phase = (entry.manager.state or {}).get('phase')
        current_phase = str(getattr(raw_phase, 'value', raw_phase) or '').strip()
        if current_phase in ('init', 'refinement'):
            return None
        return jsonify({
            "error": (
                "Master data can only be modified before job analysis begins or from the "
                "post-job finalise workflow."
            ),
            "phase": current_phase or None,
        }), 409

    # ------------------------------------------------------------------
    # master-fields (fast endpoint, no session required)
    # ------------------------------------------------------------------

    @bp.get("/api/master-fields")
    def master_fields():
        """Return selected_achievements and professional_summaries directly from the master CV file."""
        entry = get_session()
        orchestrator = entry.orchestrator
        try:
            data, _ = load_master(orchestrator.master_data_path)
            return jsonify({
                "ok": True,
                "selected_achievements":   data.get('selected_achievements', []),
                "professional_summaries":  data.get('professional_summaries', {}),
                "experiences":             data.get('experience', []),
            })
        except Exception as e:
            return jsonify({
                "ok": False,
                "error": str(e),
                "selected_achievements":  [],
                "professional_summaries": {},
                "experiences":            [],
            }), 500

    # ------------------------------------------------------------------
    # Master data CRUD
    # ------------------------------------------------------------------

    @bp.get("/api/master-data/overview")
    def master_data_overview():
        """Return a profile summary (counts + personal info) from the master CV file."""
        entry = get_session()
        orchestrator = entry.orchestrator
        try:
            data, _ = load_master(orchestrator.master_data_path)
            personal  = data.get('personal_info', {})
            skills    = data.get('skills', [])
            skill_count = (
                sum(len(v) if isinstance(v, list) else 0 for v in skills.values())
                if isinstance(skills, dict) else len(skills)
            )
            summaries = data.get('professional_summaries', {})
            return jsonify({
                "ok":                True,
                "name":              personal.get('name', ''),
                "headline":          personal.get('headline', personal.get('title', '')),
                "email":             (personal.get('contact') or {}).get('email',
                                     personal.get('email', '')),
                "experience_count":  len(data.get('experience', [])),
                "skill_count":       skill_count,
                "achievement_count": len(data.get('selected_achievements', [])),
                "summary_count":     len(summaries) if isinstance(summaries, dict)
                                     else len(summaries),
                "education_count":   len(data.get('education', [])),
                "publication_count": len(data.get('publications', [])),
            })
        except Exception:
            logger.exception("Failed to load master data overview")
            return jsonify({
                "ok": False,
                "error": "Failed to load master data overview.",
            }), 500

    @bp.post("/api/master-data/update-achievement")
    def master_data_update_achievement():
        """Update or delete a selected achievement, or add a new one to the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        ach_id = (req.get('id') or '').strip()
        if not ach_id:
            return jsonify({"error": "id is required"}), 400
        try:
            master, master_path = load_master(orchestrator.master_data_path)
            achievements = master.setdefault('selected_achievements', [])
            if req.get('action') == 'delete':
                original_len = len(achievements)
                master['selected_achievements'] = [a for a in achievements if a.get('id') != ach_id]
                if len(master['selected_achievements']) == original_len:
                    return jsonify({"ok": False, "error": "Achievement not found"}), 404
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted", "id": ach_id})
            existing = next((a for a in achievements if a.get('id') == ach_id), None)
            if existing:
                for field in ('title', 'description', 'relevant_for', 'importance'):
                    if field in req:
                        existing[field] = req[field]
                action = 'updated'
            else:
                new_ach = {'id': ach_id}
                for field in ('title', 'description', 'relevant_for', 'importance'):
                    if field in req:
                        new_ach[field] = req[field]
                achievements.append(new_ach)
                action = 'added'
            save_master(master, master_path)
            return jsonify({"ok": True, "action": action, "id": ach_id})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/update-summary")
    def master_data_update_summary():
        """Update, add, or delete a named professional summary variant in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req  = request.get_json() or {}
        key  = (req.get('key') or '').strip()
        text = (req.get('text') or '').strip()
        if not key:
            return jsonify({"error": "key is required"}), 400
        action = req.get('action')
        if action != 'delete' and not text:
            return jsonify({"error": "text is required for add/update"}), 400
        try:
            master, master_path = load_master(orchestrator.master_data_path)
            summaries = master.get('professional_summaries', {})
            if isinstance(summaries, list):
                summaries = {str(i): v for i, v in enumerate(summaries)}
                master['professional_summaries'] = summaries
            if action == 'delete':
                if key not in summaries:
                    return jsonify({"ok": False, "error": "Summary not found"}), 404
                del summaries[key]
                master['professional_summaries'] = summaries
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted", "key": key})
            is_new = key not in summaries
            summaries[key] = text
            master['professional_summaries'] = summaries
            save_master(master, master_path)
            return jsonify({"ok": True, "action": "added" if is_new else "updated", "key": key})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.get("/api/master-data/full")
    def master_data_full():
        """Return all editable sections of the master CV for the structured editor."""
        entry = get_session()
        orchestrator = entry.orchestrator
        try:
            master, _ = load_master(orchestrator.master_data_path)
            return jsonify({
                "ok":                     True,
                "personal_info":          master.get('personal_info', {}),
                "experience":             master.get('experience', master.get('experiences', [])),
                "skills":                 master.get('skills', []),
                "education":              master.get('education', []),
                "awards":                 master.get('awards', []),
                "selected_achievements":  master.get('selected_achievements', []),
                "professional_summaries": master.get('professional_summaries', {}),
            })
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.get("/api/master-data/validate")
    def master_data_validate():
        """Validate master data structure and return errors/warnings."""
        entry = get_session()
        validate_owner(entry)
        orchestrator = entry.orchestrator
        use_schema_param = (request.args.get('use_schema') or 'true').strip().lower()
        use_schema = use_schema_param in ('1', 'true', 'yes', 'on')
        schema_path = request.args.get('schema_path')

        validator = validate_master_data_file_fn
        if validator is None:
            from utils.master_data_validator import validate_master_data_file as validator

        result = validator(
            str(orchestrator.master_data_path),
            use_schema=use_schema,
            schema_path=schema_path,
        )
        return jsonify({'ok': result.valid, **result.to_dict()})

    @bp.post("/api/master-data/preview-diff")
    def master_data_preview_diff():
        """Return a read-only before/after diff preview for selected master-data edits."""
        entry = get_session()
        validate_owner(entry)
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        section = (req.get('section') or '').strip()
        if section not in ('personal_info', 'skill'):
            return jsonify({"error": "section must be one of: personal_info, skill"}), 400

        try:
            master, _ = load_master(orchestrator.master_data_path)

            if section == 'personal_info':
                pi = master.get('personal_info', {})
                contact = pi.get('contact', {}) if isinstance(pi, dict) else {}
                address = contact.get('address', {}) if isinstance(contact, dict) else {}

                existing = {
                    'name': pi.get('name', ''),
                    'title': pi.get('title', ''),
                    'email': contact.get('email', pi.get('email', '')),
                    'phone': contact.get('phone', ''),
                    'linkedin': contact.get('linkedin', ''),
                    'website': contact.get('website', ''),
                    'city': address.get('city', ''),
                    'state': address.get('state', ''),
                }

                changes = []
                for key in ('name', 'title', 'email', 'phone', 'linkedin', 'website', 'city', 'state'):
                    if key not in req:
                        continue
                    new_val = req.get(key)
                    old_val = existing.get(key, '')
                    if str(old_val or '').strip() == str(new_val or '').strip():
                        continue
                    changes.append({'field': key, 'old': old_val, 'new': new_val})

                return jsonify({'ok': True, 'section': section, 'changed': bool(changes), 'changes': changes})

            action = (req.get('action') or '').strip()
            if action not in ('add', 'delete', 'add_category', 'delete_category'):
                return jsonify({"error": "action must be add, delete, add_category, or delete_category"}), 400

            skills = master.get('skills', [])
            changes = []

            if action in ('add_category', 'delete_category'):
                cat_key = (req.get('category_key') or '').strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                exists = isinstance(skills, dict) and cat_key in skills
                if action == 'add_category' and not exists:
                    changes.append({'field': f'skills.category.{cat_key}', 'old': None, 'new': 'created'})
                if action == 'delete_category' and exists:
                    changes.append({'field': f'skills.category.{cat_key}', 'old': 'exists', 'new': None})
            else:
                skill_name = (req.get('skill') or '').strip()
                if not skill_name:
                    return jsonify({"error": "skill is required"}), 400

                if isinstance(skills, list):
                    existing_lower = {
                        (s.get('name', '') if isinstance(s, dict) else str(s)).strip().lower()
                        for s in skills
                    }
                    exists = skill_name.lower() in existing_lower
                    if action == 'add' and not exists:
                        changes.append({'field': 'skills', 'old': None, 'new': skill_name})
                    if action == 'delete' and exists:
                        changes.append({'field': 'skills', 'old': skill_name, 'new': None})
                elif isinstance(skills, dict):
                    cat_key = (req.get('category') or '').strip()
                    if not cat_key:
                        return jsonify({"error": "category is required for categorized skills"}), 400
                    cat_val = skills.get(cat_key)
                    if isinstance(cat_val, list):
                        cat_list = cat_val
                    elif isinstance(cat_val, dict):
                        raw = cat_val.get('skills', [])
                        if not isinstance(raw, list):
                            return jsonify({"error": "Category skills must be a list"}), 400
                        cat_list = raw
                    else:
                        cat_list = []
                    existing_lower = {
                        (s.get('name', '') if isinstance(s, dict) else str(s)).strip().lower()
                        for s in cat_list
                    }
                    exists = skill_name.lower() in existing_lower
                    if action == 'add' and not exists:
                        changes.append({'field': f'skills.{cat_key}', 'old': None, 'new': skill_name})
                    if action == 'delete' and exists:
                        changes.append({'field': f'skills.{cat_key}', 'old': skill_name, 'new': None})

            return jsonify({'ok': True, 'section': section, 'changed': bool(changes), 'changes': changes})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/personal-info")
    def master_data_update_personal_info():
        """Update personal_info fields in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}

        email = str(req.get('email') or '').strip()
        if email and '@' not in email:
            return jsonify({"error": "email must be a valid email address"}), 400

        for url_field in ('linkedin', 'website'):
            url_val = str(req.get(url_field) or '').strip()
            if url_val and not re.match(r'^https?://', url_val):
                return jsonify({"error": f"{url_field} must start with http:// or https://"}), 400

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            pi = master.setdefault('personal_info', {})
            for field in ('name', 'title'):
                if field in req:
                    pi[field] = req[field]
            contact = pi.setdefault('contact', {})
            for field in ('email', 'phone', 'linkedin', 'website'):
                if field in req:
                    contact[field] = req[field]
            if 'city' in req or 'state' in req:
                address = contact.setdefault('address', {})
                if 'city' in req:
                    address['city'] = req['city']
                if 'state' in req:
                    address['state'] = req['state']
            save_master(master, master_path)
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/experience")
    def master_data_update_experience():
        """Add, update, or delete an experience entry in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        if action in ('add', 'update'):
            exp_data = req.get('experience') or {}
            if not exp_data.get('title') or not exp_data.get('company'):
                return jsonify({"error": "title and company are required"}), 400

            try:
                importance_val = int(exp_data.get('importance') or 5)
            except (TypeError, ValueError):
                return jsonify({"error": "importance must be an integer"}), 400
            if importance_val < 1 or importance_val > 10:
                return jsonify({"error": "importance must be between 1 and 10"}), 400

            employment_type = str(exp_data.get('employment_type') or 'full_time').strip()
            allowed_types = {
                'full_time', 'part_time', 'contract', 'consulting',
                'internship', 'self_employed',
            }
            if employment_type not in allowed_types:
                return jsonify({"error": "employment_type is invalid"}), 400

            start_year = _extract_year(exp_data.get('start_date'))
            end_year = _extract_year(exp_data.get('end_date'))
            if start_year is not None and end_year is not None and start_year > end_year:
                return jsonify({"error": "start_date cannot be after end_date"}), 400

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            experiences = master.get('experience', master.pop('experiences', []))
            master['experience'] = experiences
            if action == 'delete':
                exp_id = (req.get('id') or '').strip()
                if not exp_id:
                    return jsonify({"error": "id is required for delete"}), 400
                original_len = len(experiences)
                master['experience'] = [e for e in experiences if e.get('id') != exp_id]
                if len(master['experience']) == original_len:
                    return jsonify({"ok": False, "error": "Experience not found"}), 404
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            exp_data = req.get('experience') or {}
            if not exp_data.get('title') or not exp_data.get('company'):
                return jsonify({"error": "title and company are required"}), 400
            if action == 'add':
                new_id  = 'exp_' + str(int(datetime.now().timestamp() * 1000))
                loc: Dict[str, Any] = {}
                if exp_data.get('city'):
                    loc['city'] = exp_data['city']
                if exp_data.get('state'):
                    loc['state'] = exp_data['state']
                new_exp: Dict[str, Any] = {
                    'id':               new_id,
                    'title':            exp_data['title'],
                    'company':          exp_data['company'],
                    'location':         loc,
                    'start_date':       exp_data.get('start_date', ''),
                    'end_date':         exp_data.get('end_date', ''),
                    'employment_type':  exp_data.get('employment_type', 'full_time'),
                    'importance':       int(exp_data.get('importance') or 5),
                    'tags':             exp_data.get('tags') or [],
                    'domain_relevance': exp_data.get('domain_relevance') or [],
                    'achievements':     [],
                }
                experiences.append(new_exp)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "id": new_id})
            exp_id = (req.get('id') or exp_data.get('id') or '').strip()
            if not exp_id:
                return jsonify({"error": "id is required for update"}), 400
            existing_exp = next((e for e in experiences if e.get('id') == exp_id), None)
            if not existing_exp:
                return jsonify({"ok": False, "error": "Experience not found"}), 404
            for field in ('title', 'company', 'start_date', 'end_date', 'employment_type'):
                if field in exp_data:
                    existing_exp[field] = exp_data[field]
            if 'importance' in exp_data:
                existing_exp['importance'] = int(exp_data['importance'])
            if 'city' in exp_data or 'state' in exp_data:
                exp_loc = existing_exp.setdefault('location', {})
                if 'city' in exp_data:
                    exp_loc['city'] = exp_data['city']
                if 'state' in exp_data:
                    exp_loc['state'] = exp_data['state']
            for field in ('tags', 'domain_relevance'):
                if field in exp_data:
                    existing_exp[field] = exp_data[field]
            save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "id": exp_id})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/skill")
    def master_data_update_skill():
        """Add, update, or delete a skill, or manage skill categories."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete', 'add_category', 'delete_category'):
            return jsonify({"error": "action must be add, update, delete, add_category, or delete_category"}), 400

        category_key_pattern = re.compile(r'^[A-Za-z0-9_-]+$')

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            skills = master.get('skills', [])

            valid_exp_ids = {
                (exp.get('id') or '').strip()
                for exp in (master.get('experience') or master.get('experiences') or [])
                if isinstance(exp, dict) and (exp.get('id') or '').strip()
            }

            def _sanitize_experience_ids(raw_ids: Any) -> List[str]:
                if raw_ids is None:
                    return []
                if isinstance(raw_ids, str):
                    raw_ids = [x.strip() for x in raw_ids.split(',') if x.strip()]
                if not isinstance(raw_ids, list):
                    return []
                cleaned: List[str] = []
                seen = set()
                for item in raw_ids:
                    if not isinstance(item, str):
                        continue
                    exp_id = item.strip()
                    if not exp_id or exp_id not in valid_exp_ids or exp_id in seen:
                        continue
                    cleaned.append(exp_id)
                    seen.add(exp_id)
                return cleaned

            def _skill_name(item: Any) -> str:
                if isinstance(item, str):
                    return item
                if isinstance(item, dict):
                    return str(item.get('name') or '')
                return ''

            def _skill_experiences(item: Any) -> List[str]:
                if isinstance(item, dict) and isinstance(item.get('experiences'), list):
                    return _sanitize_experience_ids(item.get('experiences'))
                return []

            def _skill_group(item: Any) -> Optional[str]:
                if isinstance(item, dict):
                    grp = item.get('group')
                    if grp is None:
                        return None
                    g = str(grp).strip()
                    return g if g else None
                return None

            def _skill_payload(name: str, experience_ids: List[str], group: Optional[str]) -> Any:
                if experience_ids or group:
                    d: Dict[str, Any] = {'name': name}
                    if experience_ids:
                        d['experiences'] = experience_ids
                    if group:
                        d['group'] = group
                    return d
                return name

            if action == 'add_category':
                cat_key = (req.get('category_key') or '').strip()
                cat_name = (req.get('category_name') or cat_key).strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                if not category_key_pattern.match(cat_key):
                    return jsonify({"error": "category_key must contain only letters, numbers, underscores, or hyphens"}), 400
                if not isinstance(skills, dict):
                    return jsonify({"ok": False, "error": "Skills field is not a category dict"}), 400
                if cat_key in skills:
                    return jsonify({"ok": False, "error": "Category already exists"}), 409
                skills[cat_key] = {"category": cat_name, "skills": []}
                master['skills'] = skills
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "category_added"})

            if action == 'delete_category':
                cat_key = (req.get('category_key') or '').strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                if not isinstance(skills, dict):
                    return jsonify({"ok": False, "error": "Skills field is not a category dict"}), 400
                if cat_key not in skills:
                    return jsonify({"ok": False, "error": "Category not found"}), 404
                del skills[cat_key]
                master['skills'] = skills
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "category_deleted"})

            skill_name = (req.get('skill') or '').strip()
            if not skill_name:
                return jsonify({"error": "skill is required"}), 400
            if len(skill_name) > 100:
                return jsonify({"error": "skill must be 100 characters or fewer"}), 400
            skill_lower = skill_name.lower()
            new_skill_name = (req.get('skill_new') or skill_name).strip()
            has_experience_field = 'experiences' in req
            has_group_field = 'group' in req
            requested_experience_ids = _sanitize_experience_ids(req.get('experiences'))
            requested_group = (req.get('group') or '').strip() if has_group_field else None
            if requested_group == '':
                requested_group = None

            if isinstance(skills, list):
                existing_lower = {_skill_name(s).strip().lower() for s in skills}
                if action == 'add':
                    if skill_lower in existing_lower:
                        return jsonify({"ok": False, "error": "Skill already exists"}), 409
                    skills.append(_skill_payload(skill_name, requested_experience_ids, requested_group))
                    master['skills'] = skills
                    save_master(master, master_path)
                    return jsonify({"ok": True, "action": "added"})
                if action == 'update':
                    idx = next((i for i, s in enumerate(skills) if _skill_name(s) == skill_name), -1)
                    if idx < 0:
                        return jsonify({"ok": False, "error": "Skill not found"}), 404
                    if new_skill_name != skill_name and new_skill_name.lower() in existing_lower:
                        return jsonify({"ok": False, "error": "Updated skill name already exists"}), 409
                    effective_experience_ids = (
                        requested_experience_ids if has_experience_field else _skill_experiences(skills[idx])
                    )
                    effective_group = requested_group if has_group_field else _skill_group(skills[idx])
                    skills[idx] = _skill_payload(new_skill_name, effective_experience_ids, effective_group)
                    master['skills'] = skills
                    save_master(master, master_path)
                    return jsonify({"ok": True, "action": "updated"})
                idx = next((i for i, s in enumerate(skills) if _skill_name(s) == skill_name), -1)
                if idx < 0:
                    return jsonify({"ok": False, "error": "Skill not found"}), 404
                del skills[idx]
                master['skills'] = skills
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})

            if isinstance(skills, dict):
                cat_key = (req.get('category') or '').strip()
                if not cat_key and action in ('add', 'update', 'delete'):
                    return jsonify({"error": "category is required for categorized skills"}), 400
                if cat_key and cat_key not in skills:
                    return jsonify({"ok": False, "error": "Category not found"}), 404
                cat_val = skills.get(cat_key)
                if isinstance(cat_val, list):
                    cat_list = cat_val
                elif isinstance(cat_val, dict):
                    raw_list = cat_val.setdefault('skills', [])
                    if not isinstance(raw_list, list):
                        return jsonify({"ok": False, "error": "Category skills must be a list"}), 400
                    cat_list = raw_list
                else:
                    cat_list = []

                cat_existing_lower = {_skill_name(s).strip().lower() for s in cat_list}
                if action == 'add':
                    if skill_lower in cat_existing_lower:
                        return jsonify({"ok": False, "error": "Skill already exists in category"}), 409
                    cat_list.append(_skill_payload(skill_name, requested_experience_ids, requested_group))
                    save_master(master, master_path)
                    return jsonify({"ok": True, "action": "added"})
                if action == 'update':
                    idx = next((i for i, s in enumerate(cat_list) if _skill_name(s) == skill_name), -1)
                    if idx < 0:
                        return jsonify({"ok": False, "error": "Skill not found in category"}), 404
                    if new_skill_name != skill_name and new_skill_name.lower() in cat_existing_lower:
                        return jsonify({"ok": False, "error": "Updated skill name already exists in category"}), 409
                    effective_experience_ids = (
                        requested_experience_ids if has_experience_field else _skill_experiences(cat_list[idx])
                    )
                    effective_group = requested_group if has_group_field else _skill_group(cat_list[idx])
                    cat_list[idx] = _skill_payload(new_skill_name, effective_experience_ids, effective_group)
                    save_master(master, master_path)
                    return jsonify({"ok": True, "action": "updated"})
                idx = next((i for i, s in enumerate(cat_list) if _skill_name(s) == skill_name), -1)
                if idx < 0:
                    return jsonify({"ok": False, "error": "Skill not found in category"}), 404
                del cat_list[idx]
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})

            return jsonify({"ok": False, "error": "Unexpected skills format"}), 400
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/education")
    def master_data_update_education():
        """Add, update, or delete an education entry in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        start_year = None
        end_year = None
        if action in ('add', 'update'):
            if req.get('start_year') not in (None, ''):
                try:
                    start_year = int(req.get('start_year'))
                except (TypeError, ValueError):
                    return jsonify({"error": "start_year must be an integer"}), 400
                if start_year < 1900 or start_year > 2100:
                    return jsonify({"error": "start_year must be between 1900 and 2100"}), 400
            if req.get('end_year') not in (None, ''):
                try:
                    end_year = int(req.get('end_year'))
                except (TypeError, ValueError):
                    return jsonify({"error": "end_year must be an integer"}), 400
                if end_year < 1900 or end_year > 2100:
                    return jsonify({"error": "end_year must be between 1900 and 2100"}), 400
            if start_year is not None and end_year is not None and start_year > end_year:
                return jsonify({"error": "start_year cannot be greater than end_year"}), 400

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            education = master.setdefault('education', [])
            if action == 'delete':
                idx = req.get('idx')
                if not isinstance(idx, int):
                    return jsonify({"error": "idx (int) is required for delete"}), 400
                if idx < 0 or idx >= len(education):
                    return jsonify({"ok": False, "error": "Index out of range"}), 404
                education.pop(idx)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            edu_data: Dict[str, Any] = {}
            for field in ('degree', 'field', 'institution'):
                if field in req:
                    edu_data[field] = req[field]
            if req.get('city') or req.get('state'):
                loc: Dict[str, Any] = {}
                if req.get('city'):
                    loc['city'] = req['city']
                if req.get('state'):
                    loc['state'] = req['state']
                edu_data['location'] = loc
            if start_year is not None:
                edu_data['start_year'] = start_year
            if end_year is not None:
                edu_data['end_year'] = end_year
            if action == 'add':
                if not edu_data.get('degree') or not edu_data.get('institution'):
                    return jsonify({"error": "degree and institution are required"}), 400
                education.append(edu_data)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "idx": len(education) - 1})
            idx = req.get('idx')
            if not isinstance(idx, int):
                return jsonify({"error": "idx (int) is required for update"}), 400
            if idx < 0 or idx >= len(education):
                return jsonify({"ok": False, "error": "Index out of range"}), 404
            education[idx].update(edu_data)
            save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "idx": idx})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/award")
    def master_data_update_award():
        """Add, update, or delete an award entry in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        parsed_year = None
        if action in ('add', 'update') and req.get('year') not in (None, ''):
            try:
                parsed_year = int(req.get('year'))
            except (TypeError, ValueError):
                return jsonify({"error": "year must be an integer"}), 400
            if parsed_year < 1900 or parsed_year > 2100:
                return jsonify({"error": "year must be between 1900 and 2100"}), 400

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            awards = master.setdefault('awards', [])
            if action == 'delete':
                idx = req.get('idx')
                if not isinstance(idx, int):
                    return jsonify({"error": "idx (int) is required for delete"}), 400
                if idx < 0 or idx >= len(awards):
                    return jsonify({"ok": False, "error": "Index out of range"}), 404
                awards.pop(idx)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            award_data: Dict[str, Any] = {}
            for field in ('title', 'description'):
                if field in req:
                    award_data[field] = req[field]
            if parsed_year is not None:
                award_data['year'] = parsed_year
            if 'relevant_for' in req:
                award_data['relevant_for'] = req['relevant_for']
            if action == 'add':
                if not award_data.get('title'):
                    return jsonify({"error": "title is required"}), 400
                awards.append(award_data)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "idx": len(awards) - 1})
            idx = req.get('idx')
            if not isinstance(idx, int):
                return jsonify({"error": "idx (int) is required for update"}), 400
            if idx < 0 or idx >= len(awards):
                return jsonify({"ok": False, "error": "Index out of range"}), 404
            awards[idx].update(award_data)
            save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "idx": idx})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    # ------------------------------------------------------------------
    # Certifications CRUD
    # ------------------------------------------------------------------

    @bp.post("/api/master-data/certification")
    def master_data_update_certification():
        """Add, update, or delete a certification entry in the master CV."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        parsed_year = None
        if action in ('add', 'update') and req.get('year') not in (None, ''):
            try:
                parsed_year = int(req.get('year'))
            except (TypeError, ValueError):
                return jsonify({"error": "year must be an integer"}), 400
            if parsed_year < 1900 or parsed_year > 2100:
                return jsonify({"error": "year must be between 1900 and 2100"}), 400

        try:
            master, master_path = load_master(orchestrator.master_data_path)
            certs = master.setdefault('certifications', [])
            if action == 'delete':
                idx = req.get('idx')
                if not isinstance(idx, int):
                    return jsonify({"error": "idx (int) is required for delete"}), 400
                if idx < 0 or idx >= len(certs):
                    return jsonify({"ok": False, "error": "Index out of range"}), 404
                certs.pop(idx)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            cert_data: Dict[str, Any] = {}
            for field in ('name', 'issuer'):
                if field in req:
                    cert_data[field] = req[field]
            if parsed_year is not None:
                cert_data['year'] = parsed_year
            if action == 'add':
                if not cert_data.get('name'):
                    return jsonify({"error": "name is required"}), 400
                certs.append(cert_data)
                save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "idx": len(certs) - 1})
            idx = req.get('idx')
            if not isinstance(idx, int):
                return jsonify({"error": "idx (int) is required for update"}), 400
            if idx < 0 or idx >= len(certs):
                return jsonify({"ok": False, "error": "Index out of range"}), 404
            certs[idx].update(cert_data)
            save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "idx": idx})
        except Exception:
            logger.exception("Failed to update certifications")
            return jsonify({
                "ok": False,
                "error": "Failed to update certifications.",
            }), 500

    # ------------------------------------------------------------------
    # Master-data history, restore, and export
    # ------------------------------------------------------------------

    @bp.get("/api/master-data/history")
    def master_data_history():
        """List timestamped backup snapshots of Master_CV_Data.json."""
        entry = get_session()
        orchestrator = entry.orchestrator
        master_path = Path(orchestrator.master_data_path)
        backup_dir = master_path.parent / "backups"
        snapshots: List[Dict[str, Any]] = []
        if backup_dir.exists():
            for p in sorted(backup_dir.glob("Master_CV_*.json"), reverse=True):
                stat = p.stat()
                snapshots.append({
                    "filename": p.name,
                    "size":     stat.st_size,
                    "mtime":    stat.st_mtime,
                })
        return jsonify({"ok": True, "snapshots": snapshots})

    @bp.post("/api/master-data/restore")
    def master_data_restore():
        """Restore master CV from a named backup snapshot."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        filename = (req.get('filename') or '').strip()
        if not filename:
            return jsonify({"ok": False, "error": "filename is required"}), 400
        # Validate filename format to avoid path traversal.
        # Accepts both backup formats:
        #   web_app._save_master  → Master_CV_Data.YYYYMMDD_HHMMSS_ffffff.bak.json
        #   routes._save_master   → Master_CV_YYYYMMDDTHHMMSSZ.json
        _BACKUP_NAME_RE = (
            r'Master_CV_Data\.\d{8}_\d{6}_\d+\.bak\.json'
            r'|Master_CV_\d{8}T\d{6}Z\.json'
        )
        if not re.fullmatch(_BACKUP_NAME_RE, filename):
            return jsonify({"ok": False, "error": "Invalid backup filename format"}), 400
        master_path = Path(orchestrator.master_data_path)
        backup_dir = master_path.parent / "backups"
        source = _resolve_backup_path(backup_dir, filename)
        if source is None:
            return jsonify({"ok": False, "error": "Invalid backup filename format"}), 400
        if not source.exists():
            return jsonify({"ok": False, "error": "Backup not found"}), 404
        try:
            # Create a safety backup of the current master before overwriting
            ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            safety_path = backup_dir / f"Master_CV_{ts}.json"
            if master_path.exists():
                shutil.copy2(master_path, safety_path)
            shutil.copy2(source, master_path)
            subprocess.run(
                ['git', '-C', str(master_path.parent), 'add', master_path.name],
                capture_output=True, check=False,
            )
            # Reload orchestrator in-memory master data from the restored file
            restored_data, _ = load_master(str(master_path))
            orchestrator.master_data = restored_data
            return jsonify({"ok": True, "restored_from": filename, "safety_backup": safety_path.name})
        except Exception:
            logger.exception("Failed to restore master data backup")
            return jsonify({
                "ok": False,
                "error": "Failed to restore the selected backup.",
            }), 500

    @bp.get("/api/master-data/export")
    def master_data_export():
        """Download the current Master_CV_Data.json as a file attachment."""
        entry = get_session()
        orchestrator = entry.orchestrator
        try:
            master, _ = load_master(orchestrator.master_data_path)
            payload = json.dumps(master, indent=2)
            from flask import Response
            return Response(
                payload,
                mimetype='application/json',
                headers={
                    'Content-Disposition': 'attachment; filename="Master_CV_Data.json"',
                },
            )
        except Exception:
            logger.exception("Failed to export master data")
            return jsonify({
                "ok": False,
                "error": "Failed to export master data.",
            }), 500

    # ------------------------------------------------------------------
    # Generate professional summary
    # ------------------------------------------------------------------

    @bp.post("/api/generate-summary")
    def generate_professional_summary():
        """Generate (or refine) a custom LLM professional summary for this session."""
        # duckflow:
        #   id: summary_api_generate_live
        #   kind: api
        #   timestamp: "2026-03-27T01:23:28Z"
        #   status: live
        #   handles:
        #     - "POST /api/generate-summary"
        #   calls:
        #     - "llm:generate_professional_summary"
        #   reads:
        #     - "state:job_analysis"
        #     - "state:experience_decisions"
        #     - "state:customizations.recommended_experiences"
        #   writes:
        #     - "state:session_summaries.ai_generated"
        #     - "state:summary_focus_override"
        #   returns:
        #     - "response:POST /api/generate-summary.summary"
        #   notes: "Live summary-generation route writes the generated summary into session state and sets it active."
        entry = get_session()
        validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            req = request.get_json() or {}
            refinement_prompt = (req.get('refinement_prompt') or '').strip() or None
            previous_summary  = (req.get('previous_summary')  or '').strip() or None

            job_analysis = conversation.state.get('job_analysis')
            if not job_analysis:
                return jsonify({"ok": False, "error": "No job analysis found. Analyse a job description first."}), 400

            all_experiences = orchestrator.master_data.get('experience', []) if orchestrator and orchestrator.master_data else []
            experience_decisions = conversation.state.get('experience_decisions') or {}
            if experience_decisions:
                selected_exp_ids = {
                    eid for eid, action in experience_decisions.items()
                    if action in ('emphasize', 'include')
                }
                selected_experiences = [e for e in all_experiences if e.get('id') in selected_exp_ids]
            else:
                customizations = conversation.state.get('customizations') or {}
                recommended_ids = set(customizations.get('recommended_experiences') or [])
                selected_experiences = (
                    [e for e in all_experiences if e.get('id') in recommended_ids]
                    if recommended_ids else all_experiences
                )

            with entry.lock:
                summary = llm_client_ref['value'].generate_professional_summary(
                    job_analysis=job_analysis,
                    master_data=orchestrator.master_data if orchestrator else {},
                    selected_experiences=selected_experiences,
                    refinement_prompt=refinement_prompt,
                    previous_summary=previous_summary,
                )

                if not summary:
                    return jsonify({"ok": False, "error": "LLM returned an empty summary. Please try again."}), 500

                session_summaries = conversation.state.get('session_summaries') or {}
                session_summaries['ai_generated'] = summary
                conversation.state['session_summaries'] = session_summaries
                conversation.state['summary_focus_override'] = 'ai_generated'
                conversation._save_session()
            session_registry.touch(sid)

            return jsonify({"ok": True, "summary": summary})

        except Exception as exc:
            import traceback
            traceback.print_exc()
            return jsonify({"ok": False, "error": str(exc)}), 500

    # ------------------------------------------------------------------
    # Publication CRUD  (master-data / publications.bib)
    # ------------------------------------------------------------------

    @bp.get("/api/master-data/publications")
    def master_data_get_publications():
        """Return all publications stored in publications.bib."""
        entry = get_session()
        validate_owner(entry)
        orchestrator = entry.orchestrator
        pubs = orchestrator.publications or {}
        result = []
        for key, pub in pubs.items():
            item: Dict[str, Any] = {
                "key": key,
                "type": pub.get("type", ""),
                "fields": pub.get("fields", {}),
            }
            try:
                item["formatted_citation"] = format_publication(pub, style="apa")
            except Exception:
                item["formatted_citation"] = ""
            result.append(item)
        try:
            bib_path = orchestrator.publications_path
            content = bib_path.read_text(encoding="utf-8") if bib_path.exists() else ""
        except Exception:
            content = ""
            bib_path = orchestrator.publications_path
        return jsonify({
            "ok": True,
            "publications": result,
            "content": content,
            "path": str(bib_path),
            "count": len(pubs),
        })

    @bp.put("/api/master-data/publications")
    def master_data_save_raw_publications():
        """Overwrite publications.bib with raw BibTeX text."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        content = req.get("content", "")

        try:
            parsed = bibtex_text_to_publications(content)
        except Exception as e:
            return jsonify({"ok": False, "error": f"BibTeX parse error: {e}"}), 400

        if content.strip() and not parsed:
            return jsonify({
                "ok": False,
                "error": "No valid BibTeX entries found — file not saved.",
            }), 400

        bib_path = orchestrator.publications_path
        backup_path = None
        try:
            if bib_path.exists():
                backup_dir = bib_path.parent / "backups"
                backup_dir.mkdir(parents=True, exist_ok=True)
                ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                backup_path = backup_dir / f"{bib_path.stem}.{ts}{bib_path.suffix}"
                shutil.copy2(bib_path, backup_path)
        except Exception as e:
            return jsonify({"ok": False, "error": f"Backup failed: {e}"}), 500

        try:
            bib_path.write_text(content, encoding="utf-8")
            orchestrator.publications = parsed
            return jsonify({"ok": True, "count": len(parsed)})
        except Exception as e:
            if backup_path and backup_path.exists():
                try:
                    shutil.copy2(backup_path, bib_path)
                except Exception:
                    pass
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/publications/validate")
    def master_data_validate_publications():
        """Parse BibTeX text and report errors without saving anything."""
        entry = get_session()
        validate_owner(entry)
        req = request.get_json() or {}
        bibtex_text = req.get("bibtex_text", "")

        if not bibtex_text or not bibtex_text.strip():
            return jsonify({"ok": True, "count": 0, "entries": []})

        try:
            parsed = bibtex_text_to_publications(bibtex_text)
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 400

        if not parsed:
            return jsonify({
                "ok": False,
                "error": "No valid BibTeX entries found in the supplied text.",
            }), 400

        entries = [
            {"key": key, "type": value.get("type", "")}
            for key, value in parsed.items()
        ]
        return jsonify({"ok": True, "count": len(entries), "entries": entries})

    @bp.post("/api/master-data/publication")
    def master_data_update_publication():
        """Add, update, or delete a single publication in publications.bib."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        action = (req.get("action") or "").strip()
        if action not in ("add", "update", "delete"):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        key = (req.get("key") or "").strip()
        if not key:
            return jsonify({"error": "key is required"}), 400

        pubs = dict(orchestrator.publications or {})

        try:
            if action == "delete":
                if key not in pubs:
                    return jsonify({"ok": False, "error": f"Key '{key}' not found"}), 404
                del pubs[key]
                orchestrator.publications_path.write_text(
                    serialize_publications_to_bibtex(pubs), encoding="utf-8"
                )
                orchestrator.publications = parse_bibtex_file(
                    str(orchestrator.publications_path)
                )
                return jsonify({"ok": True, "action": "deleted"})

            fields = req.get("fields")
            if not isinstance(fields, dict):
                return jsonify({"error": "fields (dict) is required for add/update"}), 400
            entry_type = (req.get("type") or "").strip()
            if not entry_type:
                return jsonify({"error": "type is required for add/update"}), 400

            if not fields.get("title"):
                return jsonify({"error": "fields.title is required"}), 400
            if not fields.get("year"):
                return jsonify({"error": "fields.year is required"}), 400
            if not fields.get("author") and not fields.get("editor"):
                return jsonify({"error": "fields.author or fields.editor is required"}), 400

            if action == "add" and key in pubs:
                return jsonify({"error": f"Key '{key}' already exists; use action=update"}), 409

            pubs[key] = {"key": key, "type": entry_type, "fields": fields}
            orchestrator.publications_path.write_text(
                serialize_publications_to_bibtex(pubs), encoding="utf-8"
            )
            orchestrator.publications = parse_bibtex_file(
                str(orchestrator.publications_path)
            )
            return jsonify({"ok": True, "action": action, "key": key})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @bp.post("/api/master-data/publications/import")
    def master_data_import_publications():
        """Parse a BibTeX string and merge entries into publications.bib."""
        entry = get_session()
        validate_owner(entry)
        phase_error = _require_master_data_write_phase(entry)
        if phase_error is not None:
            return phase_error
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        bibtex_text = req.get("bibtex_text", "")
        overwrite = bool(req.get("overwrite", False))

        if not bibtex_text or not bibtex_text.strip():
            return jsonify({"error": "bibtex_text is required"}), 400

        try:
            imported = bibtex_text_to_publications(bibtex_text)
        except Exception as e:
            return jsonify({"ok": False, "error": f"BibTeX parse error: {e}"}), 400

        if not imported:
            return jsonify({"ok": False, "error": "No valid BibTeX entries found"}), 400

        pubs = dict(orchestrator.publications or {})
        added = 0
        updated = 0
        skipped = 0
        for key, pub in imported.items():
            if key in pubs:
                if overwrite:
                    pubs[key] = pub
                    updated += 1
                else:
                    skipped += 1
            else:
                pubs[key] = pub
                added += 1

        try:
            orchestrator.publications_path.write_text(
                serialize_publications_to_bibtex(pubs), encoding="utf-8"
            )
            orchestrator.publications = parse_bibtex_file(
                str(orchestrator.publications_path)
            )
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

        return jsonify({
            "ok": True,
            "added": added,
            "updated": updated,
            "skipped": skipped,
            "total": len(pubs),
        })

    @bp.post("/api/master-data/publications/convert")
    def master_data_convert_publications():
        """Use the LLM to convert free-form citation text to BibTeX."""
        entry = get_session()
        validate_owner(entry)
        orchestrator = entry.orchestrator

        if not getattr(orchestrator, "llm", None):
            return jsonify({"ok": False, "error": "No LLM provider configured for this session"}), 503

        req = request.get_json() or {}
        text = (req.get("text") or "").strip()
        if not text:
            return jsonify({"error": "text is required"}), 400

        try:
            bibtex = orchestrator.llm.convert_text_to_bibtex(text)
        except LLMError as e:
            return jsonify({"ok": False, "error": str(e)}), 500
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

        return jsonify({"ok": True, "bibtex": bibtex})

    # ------------------------------------------------------------------
    # Cover letter
    # ------------------------------------------------------------------

    @bp.get("/api/cover-letter/prior")
    def cover_letter_prior():
        """Scan prior sessions for saved cover letters and return previews."""
        try:
            from utils.config import get_config
            cfg         = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            results     = []
            if output_base.exists():
                for session_file in sorted(output_base.rglob('session.json'), reverse=True)[:30]:
                    try:
                        with open(session_file, encoding='utf-8') as f:
                            data = json.load(f)
                        state = data.get('state', {})
                        cl_text = state.get('cover_letter_text')
                        if not cl_text:
                            continue
                        params = state.get('cover_letter_params') or {}
                        job_analysis = state.get('job_analysis') or {}
                        results.append({
                            'session_path': str(session_file),
                            'company':      job_analysis.get('company', ''),
                            'role':         job_analysis.get('title', ''),
                            'date':         data.get('timestamp', '')[:10],
                            'tone':         params.get('tone', ''),
                            'preview':      cl_text[:200],
                            'full_text':    cl_text,
                        })
                    except Exception:
                        pass
            return jsonify({'ok': True, 'sessions': results})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @bp.post("/api/cover-letter/generate")
    def cover_letter_generate():
        """Generate a cover letter using the LLM and current session context."""
        entry = get_session()
        validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        with entry.lock:
            body            = request.get_json(silent=True) or {}
            tone            = body.get('tone', 'startup/tech')
            hiring_manager  = (body.get('hiring_manager') or 'Hiring Manager').strip()
            company_address = (body.get('company_address') or '').strip()
            highlight       = (body.get('highlight') or '').strip()
            reuse_body      = (body.get('reuse_body') or '').strip()

            job_analysis  = conversation.state.get('job_analysis') or {}
            master        = orchestrator.master_data or {}
            personal_info = master.get('personal_info', {})

            skills_raw = master.get('skills', [])
            all_skills = conversation.normalize_skills_data(skills_raw)
            skill_names = [
                s.get('name', str(s)) if isinstance(s, dict) else str(s)
                for s in all_skills
            ]
            top_skills = ', '.join(skill_names[:12]) if skill_names else '(see attached CV)'

            summaries = master.get('professional_summaries', {})
            if isinstance(summaries, dict) and summaries:
                summary_text = next(iter(summaries.values()))
            else:
                summary_text = master.get('summary', '')

            achievements   = master.get('selected_achievements', [])
            top_ach_titles = '\n'.join(f'- {a.get("title", "")}' for a in achievements[:4]) or '(see CV)'

            answers_snippet = ''
            answers = conversation.state.get('post_analysis_answers') or {}
            if answers:
                answers_snippet = 'Candidate context:\n' + '\n'.join(
                    f'- {q}: {a}' for q, a in list(answers.items())[:6]
                )

            tone_hint = _TONE_GUIDANCE.get(tone, '')
            company   = job_analysis.get('company', 'the company')
            role      = job_analysis.get('title', 'the position')
            keywords  = ', '.join((job_analysis.get('ats_keywords') or [])[:12])
            req_skills = ', '.join((job_analysis.get('required_skills') or [])[:10])

            today = datetime.now().strftime('%B %d, %Y')
            pieces   = [f'Date: {today}']
            if company_address:
                pieces.append(company_address)
            header_block = '\n'.join(pieces) + '\n\n'

            reuse_instruction = (
                f'\nUse the following prior cover letter as a starting point, '
                f'adapting it to the new role:\n\n"""\n{reuse_body}\n"""\n'
            ) if reuse_body else ''

            prompt = f"""\
You are a professional career coach writing a tailored cover letter.

Tone style: {tone} — {tone_hint}

TARGET ROLE
  Company: {company}
  Position: {role}
  Key requirements: {req_skills or keywords or '(see job description)'}

CANDIDATE PROFILE
  Name: {personal_info.get('name', 'The candidate')}
  Summary: {summary_text[:400] if summary_text else '(see CV)'}
  Top skills: {top_skills}
  Key achievements:
{top_ach_titles}

{answers_snippet}
{reuse_instruction}
{'Please especially highlight: ' + highlight if highlight else ''}

Write a compelling, personalised cover letter (3–4 paragraphs, ~300–400 words).
Start directly with the salutation line: "Dear {hiring_manager},"
Do NOT include a date, address block, or subject line — return only the letter body starting with the salutation.
Reference concrete skills and achievements from the candidate profile.
Close professionally with a call to action.
"""

            try:
                response = llm_client_ref['value'].chat(
                    messages=[
                        {'role': 'system', 'content': 'You write tailored, professional cover letters. Return only the letter body text.'},
                        {'role': 'user',   'content': prompt},
                    ],
                    temperature=0.7,
                )
            except Exception as e:
                return jsonify({'ok': False, 'error': f'LLM error: {e}'}), 500

            letter_text = header_block + response.strip()
            conversation.state['cover_letter_text']   = letter_text
            conversation.state['cover_letter_params'] = {
                'tone': tone, 'hiring_manager': hiring_manager,
                'company_address': company_address, 'highlight': highlight,
            }
        session_registry.touch(sid)
        return jsonify({'ok': True, 'text': letter_text})

    @bp.post("/api/cover-letter/save")
    def cover_letter_save():
        """Save cover letter text to DOCX in the output directory and update metadata.json."""
        # duckflow:
        #   id: cover_letter_api_save_live
        #   kind: api
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   handles:
        #     - "POST /api/cover-letter/save"
        #   reads:
        #     - "request:POST /api/cover-letter/save.text"
        #     - "state:generated_files.output_dir"
        #     - "state:cover_letter_reused_from"
        #   writes:
        #     - "state:cover_letter_text"
        #     - "file:metadata.cover_letter_text"
        #     - "file:metadata.cover_letter_reused_from"
        #     - "file:artifact.cover_letter_docx"
        #   returns:
        #     - "response:POST /api/cover-letter/save.filename"
        #   notes: "Saves the finalized cover-letter body to session state, writes a DOCX artifact in the application output directory, and appends the reusable text metadata."
        entry = get_session()
        validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        with entry.lock:
            text = (request.get_json(silent=True) or {}).get('text', '').strip()
            if not text:
                return jsonify({'error': 'text is required'}), 400

            generated = conversation.state.get('generated_files')
            if not generated or not generated.get('output_dir'):
                return jsonify({'error': 'No generated CV found — please generate your CV first.'}), 400

            try:
                from docx import Document
                from docx.shared import Pt

                output_dir   = Path(generated['output_dir'])
                job_analysis = conversation.state.get('job_analysis') or {}
                company      = (job_analysis.get('company') or 'Company').replace(' ', '_')
                date_str     = datetime.now().strftime('%Y-%m-%d')
                filename     = f'CoverLetter_{company}_{date_str}.docx'
                docx_path    = output_dir / filename

                doc = Document()
                for para_text in text.split('\n'):
                    p = doc.add_paragraph(para_text)
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
                doc.save(str(docx_path))

                metadata_path = output_dir / 'metadata.json'
                if metadata_path.exists():
                    with open(metadata_path, encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {}
                metadata['cover_letter_text']        = text
                metadata['cover_letter_reused_from'] = conversation.state.get('cover_letter_reused_from')
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)

                conversation.state['cover_letter_text'] = text
                session_registry.touch(sid)
                return jsonify({'ok': True, 'filename': filename})
            except Exception as e:
                return jsonify({'ok': False, 'error': str(e)}), 500

    # ------------------------------------------------------------------
    # Screening responses
    # ------------------------------------------------------------------

    @bp.post("/api/screening/search")
    def screening_search():
        """For a single question, find the best prior library match and top 3 relevant experiences."""
        try:
            body     = request.get_json(force=True) or {}
            question = (body.get('question') or '').strip()
            if not question:
                return jsonify({'ok': False, 'error': 'question required'}), 400

            from utils.config import get_config
            config   = get_config()
            lib_path = Path(config.master_cv_path).parent / 'response_library.json'
            library: list = []
            if lib_path.exists():
                with open(lib_path, encoding='utf-8') as f:
                    library = json.load(f)

            scored_prior = sorted(
                [
                    {
                        'score': _text_similarity(
                            question,
                            (e.get('question') or '') + ' ' + (e.get('response_text') or ''),
                        ),
                        'entry': e,
                    }
                    for e in library
                ],
                key=lambda x: x['score'],
                reverse=True,
            )
            best_prior = scored_prior[0] if scored_prior and scored_prior[0]['score'] >= 0.25 else None

            with open(config.master_cv_path, encoding='utf-8') as f:
                master = json.load(f)
            exps = master.get('experience', [])

            def _exp_text(e: dict) -> str:
                return ' '.join([
                    e.get('title', ''), e.get('company', ''), e.get('summary', ''),
                    ' '.join(e.get('achievements', [])[:5]),
                ])

            scored_exps = sorted(
                [{'score': _text_similarity(question, _exp_text(e)), 'exp': e, 'idx': i}
                 for i, e in enumerate(exps)],
                key=lambda x: x['score'],
                reverse=True,
            )[:3]

            return jsonify({
                'ok':     True,
                'prior':  best_prior['entry'] if best_prior else None,
                'experiences': [
                    {
                        'idx':        x['idx'],
                        'score':      round(x['score'], 2),
                        'title':      x['exp'].get('title', ''),
                        'company':    x['exp'].get('company', ''),
                        'date_range': x['exp'].get('date_range', ''),
                        'summary':    (x['exp'].get('summary') or '')[:200],
                    }
                    for x in scored_exps
                ],
            })
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @bp.post("/api/screening/generate")
    def screening_generate():
        """Generate a draft screening-question response via LLM."""
        try:
            body            = request.get_json(force=True) or {}
            question        = (body.get('question') or '').strip()
            fmt             = body.get('format', 'direct')
            exp_indices     = body.get('experience_indices') or []
            prior_response  = (body.get('prior_response') or '').strip()

            if not question:
                return jsonify({'ok': False, 'error': 'question required'}), 400

            from utils.config import get_config
            config = get_config()
            with open(config.master_cv_path, encoding='utf-8') as f:
                master = json.load(f)
            exps     = master.get('experience', [])
            selected = [exps[i] for i in exp_indices if isinstance(i, int) and 0 <= i < len(exps)]

            fmt_name, word_range, fmt_instructions = _SCREENING_FORMAT_GUIDANCE.get(
                fmt, _SCREENING_FORMAT_GUIDANCE['direct']
            )

            exp_blocks = '\n\n'.join(
                f"Role: {e.get('title', '')} at {e.get('company', '')}\n"
                f"Summary: {(e.get('summary') or '')[:300]}\n"
                f"Key achievements: {'; '.join((e.get('achievements') or [])[:5])}"
                for e in selected
            ) or 'No specific experience provided.'

            entry = get_session()
            conversation = entry.manager
            answers              = conversation.state.get('post_analysis_answers') or {}
            cover_letter_snippet = (conversation.state.get('cover_letter_text') or '')[:400]

            cl_context = (
                '\n'.join(f'- {k}: {v}' for k, v in list(answers.items())[:8])
                if answers else 'None provided.'
            )

            prior_block = (
                f'\nUse the following prior response as a starting point, adapting it as needed:\n'
                f'"""\n{prior_response}\n"""\n'
            ) if prior_response else ''

            prompt = (
                f'You are drafting a screening-question response for a job application.\n'
                f'Response format: {fmt_name} (~{word_range}). {fmt_instructions}\n\n'
                f'Question:\n"{question}"\n\n'
                f'Relevant experience:\n{exp_blocks}\n\n'
                f'Applicant preferences / context:\n{cl_context}\n'
                + (f'\nCover letter excerpt (for tone/context):\n{cover_letter_snippet}\n' if cover_letter_snippet else '')
                + prior_block
                + '\nWrite only the response text. No preamble, labels, or meta-commentary.'
            )

            try:
                response_text = llm_client_ref['value'].chat(
                    messages=[
                        {'role': 'system', 'content': 'You write concise, tailored screening-question responses for job applications.'},
                        {'role': 'user',   'content': prompt},
                    ],
                    temperature=0.7,
                )
            except Exception as e:
                return jsonify({'ok': False, 'error': f'LLM error: {e}'}), 500

            return jsonify({'ok': True, 'text': response_text.strip()})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @bp.post("/api/screening/save")
    def screening_save():
        """Save screening responses to DOCX, update metadata.json, upsert response_library.json."""
        # duckflow:
        #   id: screening_api_save_live
        #   kind: api
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   handles:
        #     - "POST /api/screening/save"
        #   reads:
        #     - "request:POST /api/screening/save.responses"
        #     - "state:job_analysis.company"
        #   writes:
        #     - "state:screening_responses"
        #     - "file:metadata.screening_responses"
        #     - "file:artifact.screening_docx"
        #     - "file:response_library.json"
        #   returns:
        #     - "response:POST /api/screening/save.filename"
        #     - "response:POST /api/screening/save.count"
        #   notes: "Persists saved screening responses in session state, writes the archive DOCX and metadata entry, and upserts the reusable response library."
        entry = get_session()
        validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            body         = request.get_json(force=True) or {}
            responses_in = body.get('responses') or []
            if not responses_in:
                return jsonify({'ok': False, 'error': 'No responses to save.'}), 400

            with entry.lock:
                output_dir = Path(orchestrator.output_dir)
                if not output_dir.exists():
                    return jsonify({'ok': False, 'error': 'Output directory not found. Generate a CV first.'}), 400

                from docx import Document as _DocxDoc
                from docx.shared import Pt as _Pt
                doc = _DocxDoc()
                doc.add_heading('Screening Question Responses', 0)
                for item in responses_in:
                    doc.add_heading((item.get('question') or '')[:120], level=2)
                    para = doc.add_paragraph(item.get('response_text') or '')
                    para.style.font.size = _Pt(11)
                    doc.add_paragraph()

                date_str = datetime.now().strftime('%Y-%m-%d')
                filename = f'Screening_Responses_{date_str}.docx'
                doc_path = output_dir / filename
                doc.save(str(doc_path))

                # Update metadata.json
                metadata_path = output_dir / 'metadata.json'
                if metadata_path.exists():
                    with open(metadata_path, encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {}
                metadata['screening_responses'] = responses_in
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)

                # Upsert response_library.json
                from utils.config import get_config
                config   = get_config()
                lib_path = Path(config.master_cv_path).parent / 'response_library.json'
                library: list = []
                if lib_path.exists():
                    with open(lib_path, encoding='utf-8') as f:
                        library = json.load(f)

                job_analysis = conversation.state.get('job_analysis') or {}
                company      = job_analysis.get('company', '') or ''
                session_p    = str(output_dir)
                for item in responses_in:
                    library.append({
                        'question':      item.get('question', ''),
                        'topic_tag':     item.get('topic_tag', ''),
                        'response_text': item.get('response_text', ''),
                        'format':        item.get('format', ''),
                        'company':       company,
                        'date':          date_str,
                        'session_path':  session_p,
                    })
                with open(lib_path, 'w', encoding='utf-8') as f:
                    json.dump(library, f, indent=2)

                conversation.state['screening_responses'] = responses_in
                session_registry.touch(sid)
                return jsonify({'ok': True, 'filename': filename, 'count': len(responses_in)})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    return bp
