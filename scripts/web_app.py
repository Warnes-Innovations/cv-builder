#!/usr/bin/env python3
# Copyright (C) 2026 Gregory R. Warnes
# SPDX-License-Identifier: AGPL-3.0-or-later
#
# This file is part of CV-Builder.
# For commercial licensing, contact greg@warnes-innovations.com

"""
Minimal Web UI for LLM-Driven CV Generator

Serves a single-page web app with endpoints to:
- Submit a job description
- Send chat messages to the assistant
- View current state
- Save session

Run:
    python scripts/web_app.py

Then open http://localhost:5001
"""

import argparse
import copy
import dataclasses
import json
import logging
import os
import shutil
import subprocess
import sys
import threading
import traceback
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
from dotenv import load_dotenv

logger = logging.getLogger(__name__)

from flask import Flask, jsonify, redirect, request, send_file, send_from_directory, url_for
import requests
from urllib.parse import urlparse
import re
from bs4 import BeautifulSoup

# Load environment variables from .env file
env_path = Path(__file__).parent.parent / '.env'
if env_path.exists():
    load_dotenv(env_path)

# Ensure scripts are importable
sys.path.insert(0, str(Path(__file__).parent))

from utils.config import get_config, validate_config, ConfigurationError, setup_logging
from utils.llm_client import get_llm_provider, PROVIDER_MODELS, PROVIDER_BILLING, MODEL_INFO
from utils.cv_orchestrator import CVOrchestrator, validate_ats_report
from utils.conversation_manager import ConversationManager, Phase
from utils.llm_client import LLMError, LLMAuthError, LLMRateLimitError, LLMContextLengthError
from utils.copilot_auth import CopilotAuthManager
from utils.pricing_cache import (
    get_cached_pricing, get_pricing_updated_at, get_pricing_source,
    refresh_pricing_cache, maybe_refresh_in_background, lookup_runtime_pricing_bulk, STATIC_PRICING,
)
from utils.spell_checker import SpellChecker
from utils.master_data_validator import validate_master_data_file
from utils.bibtex_parser import (
    parse_bibtex_file,
    format_publication,
    serialize_publications_to_bibtex,
    bibtex_text_to_publications,
)
from utils.session_registry import (
    SessionRegistry, SessionNotFoundError, SessionOwnedError
)


# ---------------------------------------------------------------------------
# API response DTOs — typed dataclasses for the most critical endpoints.
#
# Using dataclasses ensures:
#   • All fields are always present in the JSON response (no silent omissions).
#   • Shape changes (field additions / renames) surface as TypeError at the call site.
#   • The expected response contract is documented in one place.
#
# JS mirror validators live in web/app.js (parseStatusResponse, etc.).
# Update both sides together when adding/removing fields.
# ---------------------------------------------------------------------------

@dataclass
class StatusResponse:
    """Response shape for GET /api/status."""
    position_name: Optional[str]
    phase: Optional[str]
    llm_provider: str
    llm_model: Optional[str]
    job_description: bool
    job_description_text: Optional[str]
    job_analysis: Optional[Dict[str, Any]]
    post_analysis_questions: List[Any]
    post_analysis_answers: Dict[str, Any]
    customizations: Optional[Dict[str, Any]]
    generated_files: Optional[Dict[str, Any]]
    generation_progress: List[Any]
    persuasion_warnings: List[Any]
    all_experience_ids: List[str]
    all_experiences: List[Dict[str, Any]]
    all_skills: List[Any]
    all_achievements: List[Any]
    professional_summaries: Dict[str, Any]
    copilot_auth: Dict[str, Any]
    iterating: bool
    reentry_phase: Optional[str]
    experience_decisions: Dict[str, Any]
    skill_decisions: Dict[str, Any]
    achievement_decisions: Dict[str, Any]
    publication_decisions: Dict[str, Any]
    summary_focus_override: Optional[str]
    extra_skills: List[Any]
    extra_skill_matches: Dict[str, List[str]]
    session_file: str
    max_skills: int
    achievement_edits: Dict[str, Any]
    intake: Dict[str, Any]


@dataclass
class SessionItem:
    """One entry in the SessionListResponse."""
    path: str
    position_name: str
    timestamp: str
    phase: str
    has_job: bool
    has_analysis: bool
    has_customizations: bool


@dataclass
class SessionListResponse:
    """Response shape for GET /api/sessions."""
    sessions: List[SessionItem]

    def to_dict(self) -> Dict[str, Any]:
        return {"sessions": [dataclasses.asdict(s) for s in self.sessions]}


@dataclass
class RewritesResponse:
    """Response shape for GET /api/rewrites."""
    ok: bool
    rewrites: List[Any]
    persuasion_warnings: List[Any]
    phase: str


@dataclass
class MessageResponse:
    """Response shape for POST /api/message."""
    ok: bool
    response: Any
    phase: Optional[str]


@dataclass
class ActionResponse:
    """Response shape for POST /api/action."""
    ok: bool
    result: Any
    phase: Optional[str]


# ---------------------------------------------------------------------------
# Model-catalog helpers (module-level for testability)
# ---------------------------------------------------------------------------

_CATALOG_LIST_MODELS_CAPABLE: set[str] = {"openai", "anthropic", "gemini", "groq"}
_CATALOG_STATIC_ONLY: set[str] = {"copilot-oauth", "copilot", "github", "local"}


def _frontend_project_root() -> Path:
    """Return the repository root for frontend bundle checks."""
    return Path(__file__).resolve().parent.parent


_FRONTEND_IMPORT_RE = re.compile(
    r'["\'](?P<spec>[^"\']+)["\']',
)


def _iter_frontend_import_specs(source: str) -> List[str]:
    """Extract relative import specs from single-line import/export statements."""
    specs: List[str] = []
    for line in source.splitlines():
        stripped = line.strip()
        if not stripped.startswith(('import ', 'export ')):
            continue
        if stripped.startswith('export ') and ' from ' not in stripped:
            continue

        match = _FRONTEND_IMPORT_RE.search(stripped)
        if match:
            specs.append(match.group('spec'))

    return specs


def _resolve_frontend_bundle_import(importer: Path, spec: str) -> Optional[Path]:
    """Resolve a relative JS import used by the bundle entrypoint graph."""
    if not spec.startswith('.'):
        return None

    candidate = (importer.parent / spec).resolve()
    if candidate.is_file():
        return candidate

    if candidate.suffix:
        return None

    for resolved in (candidate.with_suffix('.js'), candidate / 'index.js'):
        if resolved.is_file():
            return resolved

    return None


def _frontend_bundle_inputs(project_root: Path) -> List[Path]:
    """Return source files whose mtimes determine whether the bundle is stale."""
    resolved_project_root = project_root.resolve()
    build_script = project_root / 'scripts' / 'build.mjs'
    entrypoint = project_root / 'web' / 'src' / 'main.js'
    inputs: List[Path] = []

    if build_script.exists():
        inputs.append(build_script)

    if not entrypoint.exists():
        return inputs

    pending = [entrypoint]
    visited: set[Path] = set()

    while pending:
        path = pending.pop()
        if path in visited or not path.exists():
            continue

        visited.add(path)
        inputs.append(path)

        try:
            source = path.read_text(encoding='utf-8')
        except OSError:
            continue

        for spec in _iter_frontend_import_specs(source):
            resolved = _resolve_frontend_bundle_import(path, spec)
            if not resolved:
                continue
            try:
                resolved.relative_to(resolved_project_root)
            except ValueError:
                continue
            pending.append(resolved)

    return inputs


def _frontend_bundle_is_outdated(project_root: Optional[Path] = None) -> bool:
    """Return True when web/bundle.js is missing or older than its inputs."""
    root = project_root or _frontend_project_root()
    bundle_path = root / 'web' / 'bundle.js'
    if not bundle_path.exists():
        return True

    inputs = _frontend_bundle_inputs(root)
    if not inputs:
        return False

    latest_input_mtime = max(path.stat().st_mtime for path in inputs)
    return latest_input_mtime > bundle_path.stat().st_mtime


def _ensure_frontend_bundle_current(project_root: Optional[Path] = None) -> bool:
    """Rebuild web/bundle.js when frontend sources are newer than the bundle."""
    root = project_root or _frontend_project_root()
    if not _frontend_bundle_is_outdated(root):
        return False

    node_bin = shutil.which('node')
    if not node_bin:
        raise RuntimeError(
            'Frontend bundle is outdated, but Node.js is not available to rebuild web/bundle.js.'
        )

    logger.info('Frontend bundle is stale; rebuilding web/bundle.js')
    result = subprocess.run(
        [node_bin, 'scripts/build.mjs'],
        cwd=root,
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        details = (result.stderr or result.stdout or '').strip()
        raise RuntimeError(
            'Failed to rebuild frontend bundle: '
            f'{details or "node scripts/build.mjs exited non-zero"}'
        )

    return True


def _frontend_bundle_built_at(
    project_root: Optional[Path] = None,
) -> Optional[str]:
    """Return the bundle build timestamp derived from web/bundle.js mtime."""
    root = project_root or _frontend_project_root()
    bundle_path = root / 'web' / 'bundle.js'
    if not bundle_path.exists():
        return None

    built_at = datetime.fromtimestamp(
        bundle_path.stat().st_mtime,
        tz=timezone.utc,
    ).astimezone()
    return built_at.isoformat(timespec='seconds')


def _catalog_provider_api_key(provider: str) -> Optional[str]:
    if provider == "openai":
        return os.getenv("OPENAI_API_KEY")
    if provider == "anthropic":
        return os.getenv("ANTHROPIC_API_KEY")
    if provider == "gemini":
        return os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if provider == "groq":
        return os.getenv("GROQ_API_KEY")
    return None


def _catalog_provider_api_base(provider: str) -> Optional[str]:
    if provider == "openai":
        return os.getenv("OPENAI_BASE_URL")
    if provider == "anthropic":
        return os.getenv("ANTHROPIC_BASE_URL")
    if provider == "gemini":
        return os.getenv("GOOGLE_GEMINI_BASE_URL")
    if provider == "groq":
        return os.getenv("GROQ_BASE_URL")
    return None


def _catalog_normalize_model_id(model_id: str) -> str:
    """Normalize Gemini-style 'models/gemini-2.5-flash' IDs to bare model names."""
    if model_id.startswith("models/"):
        return model_id.split("/", 1)[1]
    return model_id


def _catalog_discover_provider_models(provider: str) -> Optional[List[str]]:
    """Return runtime model list for providers that support list_models, or None."""
    if provider not in _CATALOG_LIST_MODELS_CAPABLE or provider in _CATALOG_STATIC_ONLY:
        return None

    api_key = _catalog_provider_api_key(provider)
    if not api_key:
        return None

    try:
        from any_llm import list_models as anyllm_list_models
    except Exception:
        return None

    try:
        kwargs: Dict[str, Any] = {
            "provider": provider,
            "api_key":  api_key,
        }
        api_base = _catalog_provider_api_base(provider)
        if api_base:
            kwargs["api_base"] = api_base

        models = anyllm_list_models(**kwargs)
        names: List[str] = []
        for item in models:
            model_id = getattr(item, "id", None)
            if model_id is None:
                model_id = str(item)
            model_id = _catalog_normalize_model_id(str(model_id))
            if model_id:
                names.append(model_id)

        # Keep order stable and deduplicate.
        unique: List[str] = []
        seen: set[str] = set()
        for name in names:
            if name in seen:
                continue
            unique.append(name)
            seen.add(name)
        return unique
    except Exception:
        return None


# ── Dynamic model catalog cache ───────────────────────────────────────────────
# Populated at startup (background) and on demand; keyed by provider name.
# Allows /api/model and /api/model-catalog to return the full live model list
# without a blocking API call on every request.

_dynamic_model_cache: Dict[str, List[str]] = {}
_dynamic_model_cache_lock = threading.Lock()
 

def _get_available_models(provider: str, current_model: Optional[str] = None) -> List[str]:
    """Return the best-available model list for a provider.

    Preference order: cached live list > static PROVIDER_MODELS fallback.
    If *current_model* is provided and not already in the list it is prepended
    so the active model is always visible in the selector.
    """
    cached = _dynamic_model_cache.get(provider)
    models: List[str] = list(cached) if cached is not None else list(PROVIDER_MODELS.get(provider, []))
    if current_model and current_model not in models:
        models = [current_model] + models
    return models


def _refresh_dynamic_model_cache(provider: str) -> None:
    """Fetch the live model list for *provider* and store it in the cache (blocking)."""
    discovered = _catalog_discover_provider_models(provider)
    if discovered:
        with _dynamic_model_cache_lock:
            _dynamic_model_cache[provider] = discovered


def _maybe_refresh_dynamic_cache_in_background(provider: str) -> None:
    """Kick off a background refresh for *provider* if not already cached."""
    if provider not in _CATALOG_LIST_MODELS_CAPABLE:
        return
    if provider in _dynamic_model_cache:
        return
    t = threading.Thread(
        target=_refresh_dynamic_model_cache,
        args=(provider,),
        daemon=True,
        name=f"model-catalog-{provider}",
    )
    t.start()


def _text_similarity(query: str, target: str) -> float:
    """Simple word-overlap similarity score (0–1) for response library search."""
    _STOP = {
        'a', 'an', 'the', 'and', 'or', 'for', 'in', 'of', 'to', 'is',
        'are', 'was', 'were', 'i', 'my', 'your', 'we', 'our', 'this',
        'that', 'it', 'with', 'as', 'by', 'at', 'on', 'be', 'have',
        'has', 'had', 'do', 'does', 'did', 'will', 'would', 'can',
        'could', 'should', 'may', 'might', 'from', 'into', 'about',
    }
    def _tok(s: str) -> set:
        return {w.lower() for w in re.findall(r'\w+', s) if w.lower() not in _STOP and len(w) > 2}
    q_tok = _tok(query)
    t_tok = _tok(target)
    if not q_tok or not t_tok:
        return 0.0
    return len(q_tok & t_tok) / max(len(q_tok), len(t_tok))


_SCREENING_FORMAT_GUIDANCE: dict = {
    'direct':    ('Direct/Concise',    '150–200 words',
                  'Be clear and direct. State the answer, give one concrete example, close concisely.'),
    'star':      ('STAR',              '250–350 words',
                  'Use the STAR framework: Situation, Task, Action, Result. 1–2 sentences each.'),
    'technical': ('Technical Detail', '400–500 words',
                  'Provide full technical depth: context, methodology, tools/technologies, outcomes with metrics.'),
}


def _web_app_build_objects(args, auth_manager):
    """
    Instantiate LLMClient, CVOrchestrator, and ConversationManager from CLI args.

    Defined at module level so tests can patch ``CVOrchestrator`` and
    ``ConversationManager`` in ``scripts.web_app`` and have the patches
    take effect when ``session_registry.create()`` calls this factory.
    """
    llm_client = get_llm_provider(
        provider=args.llm_provider,
        model=args.model,
        auth_manager=auth_manager,
    )
    orchestrator = CVOrchestrator(
        master_data_path=args.master_data,
        publications_path=args.publications,
        output_dir=args.output_dir,
        llm_client=llm_client,
    )
    conversation = ConversationManager(orchestrator=orchestrator, llm_client=llm_client)
    return conversation, orchestrator


def create_app(args) -> Flask:
    app = Flask(__name__, static_folder=None)

    # Validate configuration before initializing dependencies.
    # Raises ConfigurationError with a clear message if no LLM provider is set.
    validate_config(provider=args.llm_provider)

    # Ensure the served frontend bundle matches the current web sources.
    bundle_rebuilt = _ensure_frontend_bundle_current()
    bundle_status = 'rebuilt' if bundle_rebuilt else 'already current'
    bundle_built_at = _frontend_bundle_built_at() or 'unknown'
    app.config['FRONTEND_BUNDLE_STATUS'] = bundle_status
    app.config['FRONTEND_BUNDLE_BUILT_AT'] = bundle_built_at
    logger.info(
        'Frontend bundle status: %s',
        bundle_status,
    )
    logger.info(
        'Frontend bundle built at: %s',
        bundle_built_at,
    )

    # Kick off a background pricing-cache refresh if the cache is stale
    maybe_refresh_in_background()

    # Kick off a background dynamic model catalog refresh for the active provider
    _maybe_refresh_dynamic_cache_in_background(args.llm_provider)

    # Copilot OAuth auth manager (shared across all requests)
    auth_manager = CopilotAuthManager()
    _auth_poll: dict = {"polling": False, "error": None, "device_code": None, "interval": 5}

    # ── Provider / model state ───────────────────────────────────────────────
    _provider_name: str = args.llm_provider
    _current_model: Optional[str] = args.model  # short form; updated by set_model()

    # One shared LLM client used by model-catalog / test endpoints.
    # Per-session LLM clients are created inside each SessionEntry.
    llm_client = get_llm_provider(
        provider=_provider_name, model=args.model, auth_manager=auth_manager
    )

    # ── Session registry ─────────────────────────────────────────────────────
    # Replaces the single global conversation + orchestrator pair.
    # Each browser tab gets its own session identified by a short UUID.
    _app_config = get_config()

    def _build_objects_for_registry(_config_ignored):
        """Factory passed to SessionRegistry; uses the app's CLI args."""
        return _web_app_build_objects(args, auth_manager)

    session_registry = SessionRegistry(
        idle_timeout_minutes=_app_config.idle_timeout_minutes,
        build_objects=_build_objects_for_registry,
    )
    app.session_registry = session_registry  # exposed for test access

    @app.before_request
    def _evict_idle_sessions():
        """Lazily evict stale sessions on every request."""
        session_registry.evict_idle()

    # ── Session helpers ──────────────────────────────────────────────────────

    def _get_session(required: bool = True, allow_missing: bool = False):
        """Extract session_id from the request and return the SessionEntry.

        For GET requests reads from query string.
        For POST/PUT/DELETE reads from query string OR JSON body.
        Returns None (does not raise) when required=False and session_id absent.
        Returns None when allow_missing=True and the supplied session_id is stale.
        Aborts 400 when required=True and session_id absent.
        Aborts 404 when session_id present but not in registry.
        """
        from flask import abort as _abort
        sid = request.args.get('session_id')
        if not sid and request.is_json:
            sid = (request.get_json(silent=True) or {}).get('session_id')
        
        source = "query" if request.args.get('session_id') else ("json_body" if sid else "none")
        logger.debug(
            "_get_session: session_id=%s (source=%s, method=%s, path=%s)",
            sid or "<missing>", source, request.method, request.path
        )
        
        if not sid:
            if required:
                _abort(400, description='session_id is required')
            return None
        try:
            return session_registry.get_or_404(sid)
        except SessionNotFoundError:
            if allow_missing and not required:
                logger.debug(
                    "_get_session: optional stale session_id=%s ignored for %s",
                    sid,
                    request.path,
                )
                return None
            _abort(404, description=f'Session not found: {sid}')

    def _validate_owner(entry) -> None:
        """Validate that the request's owner_token matches the session's owner.

        Reads owner_token from the JSON body or query string (GET requests).
        Aborts 403 if the token does not match.
        Skips validation if the session has no owner set yet (unclaimed).
        """
        from flask import abort as _abort
        if entry.owner_token is None:
            logger.debug(
                "_validate_owner: session %s is unclaimed (skipping validation)",
                entry.session_id
            )
            return  # unclaimed — allow any caller
        token = (request.get_json(silent=True) or {}).get('owner_token')
        if token is None:
            token = request.args.get('owner_token')
        token_match = (token == entry.owner_token)
        logger.debug(
            "_validate_owner: session %s (claimed=%s, token_match=%s)",
            entry.session_id, entry.owner_token is not None, token_match
        )
        if not token_match:
            _abort(403, description='Not the session owner')

    def _infer_position_name(job_text: str) -> Optional[str]:
        """Infer a concise position label from job text."""
        if not job_text:
            return None

        lines = [line.strip() for line in job_text.splitlines() if line.strip()]
        if not lines:
            return None

        title = lines[0]
        company = lines[1] if len(lines) > 1 else ""

        if " at " in title.lower() and not company:
            parts = title.split(" at ", 1)
            if len(parts) == 2:
                title, company = parts[0].strip(), parts[1].strip()

        if title and company:
            label = f"{title} at {company}"
        else:
            label = title or company

        return label[:120] if label else None

    def _coerce_to_dict(value: Any) -> Dict[str, Any]:
        """Return value as a dict, parsing from JSON string if necessary.

        Defensive helper for session state fields that older sessions may have
        persisted as a JSON string rather than a nested object.
        """
        if isinstance(value, dict):
            return value
        if isinstance(value, str):
            try:
                parsed = json.loads(value)
                if isinstance(parsed, dict):
                    return parsed
            except json.JSONDecodeError:
                pass
        return {}

    def _extract_json_payload(text: str) -> Any:
        """Best-effort extraction of JSON array/object from model output.

        Uses bracket-depth counting to locate the outermost JSON container
        rather than greedy regexes, so it handles multiple code blocks,
        backticks inside strings, and JSON embedded in prose correctly.
        """
        if not text:
            return None

        # Fast path: already clean JSON.
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        # Find the first JSON container character.
        start = -1
        open_char = ''
        for i, ch in enumerate(text):
            if ch in ('{', '['):
                start = i
                open_char = ch
                break
        if start == -1:
            return None

        close_char = '}' if open_char == '{' else ']'
        depth = 0
        in_string = False
        escape_next = False
        for j in range(start, len(text)):
            ch = text[j]
            if escape_next:
                escape_next = False
                continue
            if ch == '\\' and in_string:
                escape_next = True
                continue
            if ch == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == open_char:
                depth += 1
            elif ch == close_char:
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start:j + 1])
                    except json.JSONDecodeError:
                        return None
        return None

    def _fallback_post_analysis_questions(analysis: Dict[str, Any]) -> List[Dict[str, str]]:
        """Deterministic fallback if LLM question generation fails."""
        questions: List[Dict] = []

        role_level = analysis.get("role_level")
        if role_level:
            questions.append({
                "type": "experience_level",
                "question": f"This role appears to be at {role_level} level. Should I emphasize your most senior experiences or include a broader range to show career progression?",
                "choices": ["Emphasize most senior", "Broader career progression", "Let you decide based on analysis"],
            })

        required_skills = analysis.get("required_skills") or []
        if isinstance(required_skills, list):
            skill_text = " ".join(str(s).lower() for s in required_skills)
            if any(token in skill_text for token in ("leadership", "management", "team")):
                questions.append({
                    "type": "leadership_focus",
                    "question": "This role has leadership components. Would you prefer me to emphasize your management experience or focus more on your technical contributions?",
                    "choices": ["Emphasize management", "Focus on technical", "Balance both equally"],
                })

        domain = analysis.get("domain")
        if domain:
            questions.append({
                "type": "domain_expertise",
                "question": f"The role is in {domain}. Do you have particular projects or achievements in this domain that you'd like me to highlight?",
                "choices": ["Highlight domain-specific achievements", "Use all available experience", "Prioritize most recent work"],
            })

        company = analysis.get("company")
        if company:
            questions.append({
                "type": "company_culture",
                "question": f"For {company}, would you like me to tailor emphasis toward their culture and values? If so, what should I prioritize?",
                "choices": ["Research-driven / academic", "Industry / commercial impact", "Innovation / startup", "Use cultural indicators from job description"],
            })

        return questions[:4]

    def _generate_post_analysis_questions(
        analysis: Dict[str, Any],
        job_text: Optional[str],
        prior_qa: Optional[Dict[str, Any]] = None,
    ) -> List[Dict]:
        """Generate clarifying questions from the LLM in JSON format.

        ``prior_qa`` is a dict mapping previous question keys to the candidate's
        answers; when supplied the LLM is instructed to avoid redundant questions
        and focus only on gaps not yet covered.
        """
        prior_section = ""
        if prior_qa:
            lines = "\n".join(f"- {k}: {v}" for k, v in prior_qa.items())
            prior_section = f"""
Previously answered questions (do NOT repeat these topics):
{lines}

"""
        prompt = f"""You are helping tailor a CV to a specific job.

Create 2-4 concise, high-value clarifying questions for the candidate before generating customization recommendations.

Requirements:
- Questions must be specific to this role, company, and analysis.
- Focus on tradeoffs that affect selection/emphasis of experiences and skills.
- Avoid generic or repetitive questions.
- Keep each question under 220 characters.
- For each question, provide 2-4 button-answer choices covering the most likely responses.
- Return ONLY valid JSON as an array of objects.
{prior_section}
Schema:
[
  {{"type": "short_snake_case", "question": "...", "choices": ["Option A", "Option B", "Option C"]}}
]

Job Analysis:
{json.dumps(analysis, indent=2)}

Job Description (excerpt):
{(job_text or '')[:2500]}
"""

        response = llm_client.chat(
            messages=[
                {"role": "system", "content": "You generate targeted CV-optimization clarification questions and respond with strict JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
        )

        payload = _extract_json_payload(response)
        if isinstance(payload, dict) and isinstance(payload.get("questions"), list):
            payload = payload.get("questions")

        if not isinstance(payload, list):
            return []

        cleaned: List[Dict] = []
        for item in payload:
            if not isinstance(item, dict):
                continue
            question = str(item.get("question", "")).strip()
            qtype = str(item.get("type", "clarification")).strip().lower().replace(" ", "_")
            choices = item.get("choices")
            if not isinstance(choices, list):
                choices = []
            choices = [str(c).strip() for c in choices if str(c).strip()][:4]
            if not question:
                continue
            entry: Dict = {
                "type": qtype[:40] or "clarification",
                "question": question[:220],
            }
            if choices:
                entry["choices"] = choices
            cleaned.append(entry)

        return cleaned[:4]

    # Preload job description if provided — create a session for it
    _preload_session_id: Optional[str] = None
    if args.job_file:
        job_file_path = Path(args.job_file)
        if job_file_path.exists():
            job_text = job_file_path.read_text(encoding="utf-8")
            _preload_sid, _preload_entry = session_registry.create(_app_config)
            _preload_session_id = _preload_sid
            _preload_conv = _preload_entry.manager
            _preload_conv.add_job_description(job_text)
            # Extract position name from filename (remove date suffix and extension)
            position_name = job_file_path.stem
            position_name = re.sub(r'_\d{4}-\d{2}-\d{2}$', '', position_name)
            _preload_conv.state["position_name"] = position_name
            print(f"✓ Position name set to: {position_name}")
            print(f"✓ Pre-loaded session ID: {_preload_sid}")

            # Try to load the most recent session for this position
            try:
                loaded = _preload_conv._load_latest_session_for_position(position_name)
                if loaded:
                    print(f"✓ Restored previous session for: {position_name}")
                else:
                    _preload_conv.conversation_history.append({
                        "role": "system",
                        "content": (
                            f"Job description loaded: {job_text.split(chr(10))[0]} at "
                            + (job_text.split(chr(10))[1] if len(job_text.split(chr(10))) > 1 else 'Company')
                        ),
                    })
            except Exception as e:
                print(f"⚠ Could not load previous session: {e}")
                _preload_conv.conversation_history.append({
                    "role": "system",
                    "content": (
                        f"Job description loaded: {job_text.split(chr(10))[0]} at "
                        + (job_text.split(chr(10))[1] if len(job_text.split(chr(10))) > 1 else 'Company')
                    ),
                })
            session_registry.touch(_preload_sid)

    @app.get("/")
    def index():
        if _preload_session_id and not request.args.get("session"):
            return redirect(url_for("index", session=_preload_session_id))
        page_path = Path(__file__).parent.parent / "web" / "index.html"
        return send_file(page_path)

    @app.get("/favicon.ico")
    def favicon():
        return "", 204

    @app.get("/<path:filename>")
    def static_web(filename):
        web_dir = Path(__file__).parent.parent / "web"
        return send_from_directory(web_dir, filename)

    @app.get("/logo")
    def logo():
        # Serve white on transparent logo from web/media
        logo_path = Path(__file__).parent.parent / "web" / "media" / "logo_white_transparent.png"
        if logo_path.exists():
            return send_file(logo_path)
        else:
            # Return 404 if logo not found
            return "", 404

    @app.get("/api/status")
    def status():
        # Read-only: no lock, no touch.
        # When no session_id is present return a minimal liveness response so
        # health-check probes (test harness, conftest) get HTTP 200 without
        # needing an active session.
        entry = _get_session(required=False)
        if entry is None:
            return jsonify({
                "ok": True,
                "alive": True,
                "phase": None,
                "llm_provider": _provider_name,
                "llm_model": _current_model,
            })
        conversation = entry.manager
        orchestrator = entry.orchestrator
        # Get all experience IDs from master data
        all_experience_ids = []
        all_experiences = []
        all_skills = []
        all_achievements = []
        professional_summaries = {}
        if orchestrator and orchestrator.master_data:
            experiences = orchestrator.master_data.get('experience', [])
            all_experience_ids = [exp.get('id') for exp in experiences if exp.get('id')]
            for exp in experiences:
                if not isinstance(exp, dict):
                    continue
                ach = exp.get('achievements') or []
                ach_text = []
                for item in ach:
                    if isinstance(item, dict):
                        txt = (item.get('text') or item.get('description') or '').strip()
                    else:
                        txt = str(item).strip()
                    if txt:
                        ach_text.append(txt)
                all_experiences.append({
                    'id': exp.get('id', ''),
                    'title': exp.get('title', ''),
                    'company': exp.get('company', ''),
                    'achievements': ach_text,
                })
            all_achievements = orchestrator.master_data.get('selected_achievements', [])
            professional_summaries = dict(orchestrator.master_data.get('professional_summaries', {}))
            # Merge in any LLM-generated summaries saved this session (e.g. 'ai_recommended')
            session_summaries = conversation.state.get('session_summaries') or {}
            professional_summaries.update(session_summaries)

            # Get all skills - use canonical flat list format
            skills_data = orchestrator.master_data.get('skills', [])
            all_skills = conversation.normalize_skills_data(skills_data)
        return jsonify(dataclasses.asdict(StatusResponse(
            position_name=conversation.state.get("position_name"),
            phase=conversation.state.get("phase"),
            llm_provider=_provider_name,
            llm_model=_current_model,
            job_description=bool(conversation.state.get("job_description")),
            job_description_text=conversation.state.get("job_description"),
            job_analysis=conversation.state.get("job_analysis"),  # full data, not just bool
            post_analysis_questions=conversation.state.get("post_analysis_questions") or [],
            post_analysis_answers=conversation.state.get("post_analysis_answers") or {},
            customizations=conversation.state.get("customizations"),  # full data, not just bool
            generated_files=conversation.state.get("generated_files"),
            generation_progress=conversation.state.get("generation_progress") or [],
            persuasion_warnings=conversation.state.get("persuasion_warnings") or [],
            all_experience_ids=all_experience_ids,
            all_experiences=all_experiences,
            all_skills=all_skills,
            all_achievements=all_achievements,
            professional_summaries=professional_summaries,
            copilot_auth=auth_manager.status,
            iterating=bool(conversation.state.get("iterating")),
            reentry_phase=conversation.state.get("reentry_phase"),
            experience_decisions=conversation.state.get("experience_decisions")   or {},
            skill_decisions=conversation.state.get("skill_decisions")         or {},
            achievement_decisions=conversation.state.get("achievement_decisions")   or {},
            publication_decisions=conversation.state.get("publication_decisions")   or {},
            summary_focus_override=conversation.state.get("summary_focus_override"),
            extra_skills=conversation.state.get("extra_skills")            or [],
            extra_skill_matches=conversation.state.get("extra_skill_matches") or {},
            session_file=str(getattr(conversation, "session_file", "") or ""),
            max_skills=int(conversation.state.get("max_skills") or get_config().get("generation.max_skills", 20)),
            achievement_edits=conversation.state.get("achievement_edits")       or {},
            intake=conversation.state.get("intake")                             or {},
        )))

    @app.get("/api/master-fields")
    def master_fields():
        """Return selected_achievements and professional_summaries directly from the master CV file.

        This endpoint re-reads directly from disk to bypass any in-memory state
        issues, ensuring the data is always fresh and available.
        """
        entry = _get_session()
        orchestrator = entry.orchestrator
        try:
            data, _ = _load_master(orchestrator.master_data_path)
            return jsonify({
                "ok": True,
                "selected_achievements":   data.get('selected_achievements', []),
                "professional_summaries":  data.get('professional_summaries', {}),
                "experiences":             data.get('experience', []),
            })
        except Exception as e:
            return jsonify({
                "ok": False,
                "error": str(e),
                "selected_achievements":  [],
                "professional_summaries": {},
                "experiences":            [],
            }), 500

    # ── Master data management endpoints ────────────────────────────────────

    @app.get("/api/master-data/overview")
    def master_data_overview():
        """Return a profile summary (counts + personal info) from the master CV file."""
        entry = _get_session()
        orchestrator = entry.orchestrator
        try:
            data, _ = _load_master(orchestrator.master_data_path)
            personal  = data.get('personal_info', {})
            skills    = data.get('skills', [])
            skill_count = (
                sum(len(v) if isinstance(v, list) else 0 for v in skills.values())
                if isinstance(skills, dict) else len(skills)
            )
            summaries = data.get('professional_summaries', {})
            return jsonify({
                "ok":                True,
                "name":              personal.get('name', ''),
                "headline":          personal.get('headline', personal.get('title', '')),
                "email":             (personal.get('contact') or {}).get('email',
                                     personal.get('email', '')),
                "experience_count":  len(data.get('experience', [])),
                "skill_count":       skill_count,
                "achievement_count": len(data.get('selected_achievements', [])),
                "summary_count":     len(summaries) if isinstance(summaries, dict)
                                     else len(summaries),
                "education_count":   len(data.get('education', [])),
                "publication_count": len(data.get('publications', [])),
            })
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/update-achievement")
    def master_data_update_achievement():
        """Update or delete a selected achievement, or add a new one to the master CV.

        Body:
            id     — achievement id (required)
            action — 'delete' removes the entry; omit or any other value upserts
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        ach_id = (req.get('id') or '').strip()
        if not ach_id:
            return jsonify({"error": "id is required"}), 400

        if 'importance' in req:
            try:
                importance = int(req.get('importance'))
            except (TypeError, ValueError):
                return jsonify({"error": "importance must be an integer"}), 400
            if importance < 1 or importance > 10:
                return jsonify({"error": "importance must be between 1 and 10"}), 400

        if 'relevant_for' in req:
            relevant_for = req.get('relevant_for')
            if isinstance(relevant_for, str):
                req['relevant_for'] = [x.strip() for x in relevant_for.split(',') if x.strip()]
            elif isinstance(relevant_for, list):
                if not all(isinstance(x, str) for x in relevant_for):
                    return jsonify({"error": "relevant_for must contain only strings"}), 400
            else:
                return jsonify({"error": "relevant_for must be a list or comma-separated string"}), 400

        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            achievements = master.setdefault('selected_achievements', [])
            if req.get('action') == 'delete':
                original_len = len(achievements)
                master['selected_achievements'] = [a for a in achievements if a.get('id') != ach_id]
                if len(master['selected_achievements']) == original_len:
                    return jsonify({"ok": False, "error": "Achievement not found"}), 404
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted", "id": ach_id})
            existing = next((a for a in achievements if a.get('id') == ach_id), None)
            if existing:
                for field in ('title', 'description', 'relevant_for', 'importance'):
                    if field in req:
                        existing[field] = req[field]
                action = 'updated'
            else:
                new_ach = {'id': ach_id}
                for field in ('title', 'description', 'relevant_for', 'importance'):
                    if field in req:
                        new_ach[field] = req[field]
                achievements.append(new_ach)
                action = 'added'
            _save_master(master, master_path)
            return jsonify({"ok": True, "action": action, "id": ach_id})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/update-summary")
    def master_data_update_summary():
        """Update, add, or delete a named professional summary variant in the master CV.

        Body:
            key    — variant name/slug (required)
            text   — summary text (required unless action is 'delete')
            action — 'delete' removes the variant; omit or any other value upserts
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req  = request.get_json() or {}
        key  = (req.get('key') or '').strip()
        text = (req.get('text') or '').strip()
        if not key:
            return jsonify({"error": "key is required"}), 400
        action = req.get('action')
        if action != 'delete' and not text:
            return jsonify({"error": "text is required for add/update"}), 400
        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            summaries = master.get('professional_summaries', {})
            if isinstance(summaries, list):
                summaries = {str(i): v for i, v in enumerate(summaries)}
                master['professional_summaries'] = summaries
            if action == 'delete':
                if key not in summaries:
                    return jsonify({"ok": False, "error": "Summary not found"}), 404
                del summaries[key]
                master['professional_summaries'] = summaries
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted", "key": key})
            is_new = key not in summaries
            summaries[key] = text
            master['professional_summaries'] = summaries
            _save_master(master, master_path)
            return jsonify({"ok": True, "action": "added" if is_new else "updated", "key": key})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.get("/api/master-data/full")
    def master_data_full():
        """Return all editable sections of the master CV for the structured editor."""
        entry = _get_session()
        orchestrator = entry.orchestrator
        try:
            master, _ = _load_master(orchestrator.master_data_path)
            return jsonify({
                "ok":                     True,
                "personal_info":          master.get('personal_info', {}),
                "experience":             master.get('experience', master.get('experiences', [])),
                "skills":                 master.get('skills', []),
                "education":              master.get('education', []),
                "awards":                 master.get('awards', []),
                "selected_achievements":  master.get('selected_achievements', []),
                "professional_summaries": master.get('professional_summaries', {}),
            })
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.get("/api/master-data/validate")
    def master_data_validate():
        """Validate the master data file structure and optional JSON schema."""
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator

        use_schema_param = (request.args.get('use_schema') or 'true').strip().lower()
        use_schema = use_schema_param in ('1', 'true', 'yes', 'on')
        schema_path = request.args.get('schema_path')

        result = validate_master_data_file(
            str(orchestrator.master_data_path),
            use_schema=use_schema,
            schema_path=schema_path,
        )

        return jsonify({
            'ok': result.valid,
            **result.to_dict(),
        })

    @app.post("/api/master-data/preview-diff")
    def master_data_preview_diff():
        """Return a read-only before/after diff preview for master-data edits.

        Supported sections:
            - personal_info
            - skill
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req = request.get_json() or {}
        section = (req.get('section') or '').strip()
        if section not in ('personal_info', 'skill'):
            return jsonify({"error": "section must be one of: personal_info, skill"}), 400

        try:
            master, _ = _load_master(orchestrator.master_data_path)

            if section == 'personal_info':
                pi = master.get('personal_info', {})
                contact = pi.get('contact', {}) if isinstance(pi, dict) else {}
                address = contact.get('address', {}) if isinstance(contact, dict) else {}

                existing = {
                    'name': pi.get('name', ''),
                    'title': pi.get('title', ''),
                    'email': contact.get('email', pi.get('email', '')),
                    'phone': contact.get('phone', ''),
                    'linkedin': contact.get('linkedin', ''),
                    'website': contact.get('website', ''),
                    'city': address.get('city', ''),
                    'state': address.get('state', ''),
                }

                changes = []
                for key in ('name', 'title', 'email', 'phone', 'linkedin', 'website', 'city', 'state'):
                    if key not in req:
                        continue
                    new_val = req.get(key)
                    old_val = existing.get(key, '')
                    old_norm = str(old_val or '').strip()
                    new_norm = str(new_val or '').strip()
                    if old_norm == new_norm:
                        continue
                    changes.append({
                        'field': key,
                        'old': old_val,
                        'new': new_val,
                    })

                return jsonify({
                    'ok': True,
                    'section': section,
                    'changed': bool(changes),
                    'changes': changes,
                })

            # section == 'skill'
            action = (req.get('action') or '').strip()
            if action not in ('add', 'delete', 'add_category', 'delete_category'):
                return jsonify({"error": "action must be add, delete, add_category, or delete_category"}), 400

            skills = master.get('skills', [])
            changes = []

            if action in ('add_category', 'delete_category'):
                cat_key = (req.get('category_key') or '').strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                exists = isinstance(skills, dict) and cat_key in skills
                if action == 'add_category' and not exists:
                    changes.append({'field': f'skills.category.{cat_key}', 'old': None, 'new': 'created'})
                if action == 'delete_category' and exists:
                    changes.append({'field': f'skills.category.{cat_key}', 'old': 'exists', 'new': None})
            else:
                skill_name = (req.get('skill') or '').strip()
                if not skill_name:
                    return jsonify({"error": "skill is required"}), 400

                if isinstance(skills, list):
                    existing_lower = {str(s).strip().lower() for s in skills if isinstance(s, str)}
                    exists = skill_name.lower() in existing_lower
                    if action == 'add' and not exists:
                        changes.append({'field': 'skills', 'old': None, 'new': skill_name})
                    if action == 'delete' and exists:
                        changes.append({'field': 'skills', 'old': skill_name, 'new': None})
                elif isinstance(skills, dict):
                    cat_key = (req.get('category') or '').strip()
                    if not cat_key:
                        return jsonify({"error": "category is required for categorized skills"}), 400
                    cat_val = skills.get(cat_key)
                    if isinstance(cat_val, list):
                        cat_list = cat_val
                    elif isinstance(cat_val, dict):
                        cat_list = cat_val.get('skills', []) if isinstance(cat_val.get('skills', []), list) else []
                    else:
                        cat_list = []
                    existing_lower = {str(s).strip().lower() for s in cat_list if isinstance(s, str)}
                    exists = skill_name.lower() in existing_lower
                    if action == 'add' and not exists:
                        changes.append({'field': f'skills.{cat_key}', 'old': None, 'new': skill_name})
                    if action == 'delete' and exists:
                        changes.append({'field': f'skills.{cat_key}', 'old': skill_name, 'new': None})

            return jsonify({
                'ok': True,
                'section': section,
                'changed': bool(changes),
                'changes': changes,
            })
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/personal-info")
    def master_data_update_personal_info():
        """Update personal_info fields in the master CV.

        Body (all optional — only provided keys are updated):
            name, title, email, phone, linkedin, website, city, state
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req = request.get_json() or {}

        email = str(req.get('email') or '').strip()
        if email and '@' not in email:
            return jsonify({"error": "email must be a valid email address"}), 400

        for url_field in ('linkedin', 'website'):
            url_val = str(req.get(url_field) or '').strip()
            if url_val and not re.match(r'^https?://', url_val):
                return jsonify({"error": f"{url_field} must start with http:// or https://"}), 400

        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            pi = master.setdefault('personal_info', {})
            for field in ('name', 'title'):
                if field in req:
                    pi[field] = req[field]
            contact = pi.setdefault('contact', {})
            for field in ('email', 'phone', 'linkedin', 'website'):
                if field in req:
                    contact[field] = req[field]
            if 'city' in req or 'state' in req:
                address = contact.setdefault('address', {})
                if 'city' in req:
                    address['city'] = req['city']
                if 'state' in req:
                    address['state'] = req['state']
            _save_master(master, master_path)
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/experience")
    def master_data_update_experience():
        """Add, update, or delete an experience entry in the master CV.

        Body:
            action     — 'add' | 'update' | 'delete'
            id         — experience id (required for update/delete)
            experience — dict with fields: title, company, city, state,
                         start_date, end_date, employment_type, importance,
                         tags, domain_relevance
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        if action in ('add', 'update'):
            exp_data = req.get('experience') or {}
            if not exp_data.get('title') or not exp_data.get('company'):
                return jsonify({"error": "title and company are required"}), 400

            try:
                importance_val = int(exp_data.get('importance') or 5)
            except (TypeError, ValueError):
                return jsonify({"error": "importance must be an integer"}), 400
            if importance_val < 1 or importance_val > 10:
                return jsonify({"error": "importance must be between 1 and 10"}), 400

            employment_type = str(exp_data.get('employment_type') or 'full_time').strip()
            allowed_types = {
                'full_time', 'part_time', 'contract', 'consulting',
                'internship', 'self_employed',
            }
            if employment_type not in allowed_types:
                return jsonify({"error": "employment_type is invalid"}), 400

            start_year = _extract_year(exp_data.get('start_date'))
            end_year = _extract_year(exp_data.get('end_date'))
            if start_year is not None and end_year is not None and start_year > end_year:
                return jsonify({"error": "start_date cannot be after end_date"}), 400

        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            experiences = master.get('experience', master.pop('experiences', []))
            master['experience'] = experiences
            if action == 'delete':
                exp_id = (req.get('id') or '').strip()
                if not exp_id:
                    return jsonify({"error": "id is required for delete"}), 400
                original_len = len(experiences)
                master['experience'] = [e for e in experiences if e.get('id') != exp_id]
                if len(master['experience']) == original_len:
                    return jsonify({"ok": False, "error": "Experience not found"}), 404
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            exp_data = req.get('experience') or {}
            if action == 'add':
                new_id  = 'exp_' + str(int(datetime.now().timestamp() * 1000))
                loc: Dict[str, Any] = {}
                if exp_data.get('city'):
                    loc['city'] = exp_data['city']
                if exp_data.get('state'):
                    loc['state'] = exp_data['state']
                new_exp: Dict[str, Any] = {
                    'id':               new_id,
                    'title':            exp_data['title'],
                    'company':          exp_data['company'],
                    'location':         loc,
                    'start_date':       exp_data.get('start_date', ''),
                    'end_date':         exp_data.get('end_date', ''),
                    'employment_type':  exp_data.get('employment_type', 'full_time'),
                    'importance':       int(exp_data.get('importance') or 5),
                    'tags':             exp_data.get('tags') or [],
                    'domain_relevance': exp_data.get('domain_relevance') or [],
                    'achievements':     [],
                }
                experiences.append(new_exp)
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "id": new_id})
            exp_id = (req.get('id') or exp_data.get('id') or '').strip()
            if not exp_id:
                return jsonify({"error": "id is required for update"}), 400
            existing_exp = next((e for e in experiences if e.get('id') == exp_id), None)
            if not existing_exp:
                return jsonify({"ok": False, "error": "Experience not found"}), 404
            for field in ('title', 'company', 'start_date', 'end_date', 'employment_type'):
                if field in exp_data:
                    existing_exp[field] = exp_data[field]
            if 'importance' in exp_data:
                existing_exp['importance'] = int(exp_data['importance'])
            if 'city' in exp_data or 'state' in exp_data:
                exp_loc = existing_exp.setdefault('location', {})
                if 'city' in exp_data:
                    exp_loc['city'] = exp_data['city']
                if 'state' in exp_data:
                    exp_loc['state'] = exp_data['state']
            for field in ('tags', 'domain_relevance'):
                if field in exp_data:
                    existing_exp[field] = exp_data[field]
            _save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "id": exp_id})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/skill")
    def master_data_update_skill():
        """Add, update, or delete a skill, or manage skill categories.

        skill add/update/delete body:
            action   — 'add' | 'update' | 'delete'
            skill    — skill name string
            skill_new — updated skill name (for update)
            category — category key (required when skills is a dict)
            experiences — optional list of experience IDs associated with skill

        category management body:
            action        — 'add_category' | 'delete_category'
            category_key  — slug key
            category_name — display name (for add_category only)
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete', 'add_category', 'delete_category'):
            return jsonify({"error": "action must be add, update, delete, add_category, or delete_category"}), 400

        category_key_pattern = re.compile(r'^[A-Za-z0-9_-]+$')
        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            skills = master.get('skills', [])

            valid_exp_ids = {
                (exp.get('id') or '').strip()
                for exp in (master.get('experience') or master.get('experiences') or [])
                if isinstance(exp, dict) and (exp.get('id') or '').strip()
            }

            def _sanitize_experience_ids(raw_ids: Any) -> List[str]:
                if raw_ids is None:
                    return []
                if isinstance(raw_ids, str):
                    raw_ids = [x.strip() for x in raw_ids.split(',') if x.strip()]
                if not isinstance(raw_ids, list):
                    return []
                cleaned: List[str] = []
                seen = set()
                for item in raw_ids:
                    if not isinstance(item, str):
                        continue
                    exp_id = item.strip()
                    if not exp_id or exp_id not in valid_exp_ids or exp_id in seen:
                        continue
                    cleaned.append(exp_id)
                    seen.add(exp_id)
                return cleaned

            def _skill_name(item: Any) -> str:
                if isinstance(item, str):
                    return item
                if isinstance(item, dict):
                    return str(item.get('name') or '')
                return ''

            def _skill_payload(name: str, experience_ids: List[str], group: Optional[str] = None) -> Any:
                if experience_ids or group:
                    d: Dict[str, Any] = {'name': name}
                    if experience_ids:
                        d['experiences'] = experience_ids
                    if group:
                        d['group'] = group
                    return d
                return name

            def _skill_experiences(item: Any) -> List[str]:
                if isinstance(item, dict) and isinstance(item.get('experiences'), list):
                    return _sanitize_experience_ids(item.get('experiences'))
                return []

            def _skill_group(item: Any) -> Optional[str]:
                if isinstance(item, dict):
                    g = (item.get('group') or '').strip()
                    return g if g else None
                return None

            if action == 'add_category':
                cat_key  = (req.get('category_key') or '').strip()
                cat_name = (req.get('category_name') or cat_key).strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                if not category_key_pattern.match(cat_key):
                    return jsonify({"error": "category_key must contain only letters, numbers, underscores, or hyphens"}), 400
                if not cat_name:
                    return jsonify({"error": "category_name is required"}), 400
                if not isinstance(skills, dict):
                    return jsonify({"ok": False, "error": "Skills field is not a category dict"}), 400
                if cat_key in skills:
                    return jsonify({"ok": False, "error": "Category already exists"}), 409
                skills[cat_key] = {"category": cat_name, "skills": []}
                master['skills'] = skills
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "category_added"})
            if action == 'delete_category':
                cat_key = (req.get('category_key') or '').strip()
                if not cat_key:
                    return jsonify({"error": "category_key is required"}), 400
                if not isinstance(skills, dict):
                    return jsonify({"ok": False, "error": "Skills field is not a category dict"}), 400
                if cat_key not in skills:
                    return jsonify({"ok": False, "error": "Category not found"}), 404
                del skills[cat_key]
                master['skills'] = skills
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "category_deleted"})
            skill_name = (req.get('skill') or '').strip()
            if not skill_name:
                return jsonify({"error": "skill is required"}), 400
            if len(skill_name) > 100:
                return jsonify({"error": "skill must be 100 characters or fewer"}), 400

            skill_lower = skill_name.lower()
            new_skill_name = (req.get('skill_new') or skill_name).strip()
            has_experience_field = 'experiences' in req
            has_group_field = 'group' in req
            requested_experience_ids = _sanitize_experience_ids(req.get('experiences'))
            requested_group = (req.get('group') or '').strip() or None
            if isinstance(skills, list):
                if action == 'add':
                    existing_lower = {_skill_name(s).lower() for s in skills}
                    if skill_lower in existing_lower:
                        return jsonify({"ok": False, "error": "Skill already exists"}), 409
                    skills.append(_skill_payload(skill_name, requested_experience_ids, requested_group))
                    master['skills'] = skills
                    _save_master(master, master_path)
                    return jsonify({"ok": True, "action": "added"})
                if action == 'update':
                    idx = next((i for i, s in enumerate(skills) if _skill_name(s) == skill_name), -1)
                    if idx < 0:
                        return jsonify({"ok": False, "error": "Skill not found"}), 404
                    if new_skill_name != skill_name and any(_skill_name(s) == new_skill_name for s in skills):
                        return jsonify({"ok": False, "error": "Updated skill name already exists"}), 409
                    effective_experience_ids = (
                        requested_experience_ids if has_experience_field else _skill_experiences(skills[idx])
                    )
                    effective_group = requested_group if has_group_field else _skill_group(skills[idx])
                    skills[idx] = _skill_payload(new_skill_name, effective_experience_ids, effective_group)
                    master['skills'] = skills
                    _save_master(master, master_path)
                    return jsonify({"ok": True, "action": "updated"})
                idx = next((i for i, s in enumerate(skills) if _skill_name(s) == skill_name), -1)
                if idx < 0:
                    return jsonify({"ok": False, "error": "Skill not found"}), 404
                del skills[idx]
                master['skills'] = skills
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            if isinstance(skills, dict):
                cat_key = (req.get('category') or '').strip()
                if not cat_key:
                    return jsonify({"error": "category is required for categorized skills"}), 400
                if cat_key not in skills:
                    return jsonify({"ok": False, "error": "Category not found"}), 404
                cat_val  = skills[cat_key]
                if isinstance(cat_val, list):
                    cat_list = cat_val
                elif isinstance(cat_val, dict):
                    raw_list = cat_val.setdefault('skills', [])
                    if not isinstance(raw_list, list):
                        return jsonify({"ok": False, "error": "Category skills must be a list"}), 400
                    cat_list = raw_list
                else:
                    return jsonify({"ok": False, "error": "Category value must be a list or object"}), 400

                cat_existing_lower = {
                    str(s).strip().lower(): s for s in cat_list if isinstance(s, str)
                }
                if action == 'add':
                    if any(_skill_name(s).lower() == skill_lower for s in cat_list):
                        return jsonify({"ok": False, "error": "Skill already exists in category"}), 409
                    cat_list.append(_skill_payload(skill_name, requested_experience_ids, requested_group))
                    _save_master(master, master_path)
                    return jsonify({"ok": True, "action": "added"})
                if action == 'update':
                    idx = next((i for i, s in enumerate(cat_list) if _skill_name(s) == skill_name), -1)
                    if idx < 0:
                        return jsonify({"ok": False, "error": "Skill not found in category"}), 404
                    if new_skill_name != skill_name and any(_skill_name(s) == new_skill_name for s in cat_list):
                        return jsonify({"ok": False, "error": "Updated skill name already exists in category"}), 409
                    effective_experience_ids = (
                        requested_experience_ids if has_experience_field else _skill_experiences(cat_list[idx])
                    )
                    effective_group = requested_group if has_group_field else _skill_group(cat_list[idx])
                    cat_list[idx] = _skill_payload(new_skill_name, effective_experience_ids, effective_group)
                    _save_master(master, master_path)
                    return jsonify({"ok": True, "action": "updated"})
                idx = next((i for i, s in enumerate(cat_list) if _skill_name(s) == skill_name), -1)
                if idx < 0:
                    return jsonify({"ok": False, "error": "Skill not found in category"}), 404
                del cat_list[idx]
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            return jsonify({"ok": False, "error": "Unexpected skills format"}), 400
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/education")
    def master_data_update_education():
        """Add, update, or delete an education entry in the master CV.

        Body:
            action      — 'add' | 'update' | 'delete'
            idx         — list index (int, required for update/delete)
            degree, field, institution, city, state, start_year, end_year
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        start_year = None
        end_year = None
        if action in ('add', 'update'):
            if req.get('start_year') not in (None, ''):
                try:
                    start_year = int(req.get('start_year'))
                except (TypeError, ValueError):
                    return jsonify({"error": "start_year must be an integer"}), 400
                if start_year < 1900 or start_year > 2100:
                    return jsonify({"error": "start_year must be between 1900 and 2100"}), 400

            if req.get('end_year') not in (None, ''):
                try:
                    end_year = int(req.get('end_year'))
                except (TypeError, ValueError):
                    return jsonify({"error": "end_year must be an integer"}), 400
                if end_year < 1900 or end_year > 2100:
                    return jsonify({"error": "end_year must be between 1900 and 2100"}), 400

            if start_year is not None and end_year is not None and start_year > end_year:
                return jsonify({"error": "start_year cannot be greater than end_year"}), 400

        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            education = master.setdefault('education', [])
            if action == 'delete':
                idx = req.get('idx')
                if not isinstance(idx, int):
                    return jsonify({"error": "idx (int) is required for delete"}), 400
                if idx < 0 or idx >= len(education):
                    return jsonify({"ok": False, "error": "Index out of range"}), 404
                education.pop(idx)
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            edu_data: Dict[str, Any] = {}
            for field in ('degree', 'field', 'institution'):
                if field in req:
                    edu_data[field] = req[field]
            if req.get('city') or req.get('state'):
                loc: Dict[str, Any] = {}
                if req.get('city'):
                    loc['city'] = req['city']
                if req.get('state'):
                    loc['state'] = req['state']
                edu_data['location'] = loc
            if start_year is not None:
                edu_data['start_year'] = start_year
            if end_year is not None:
                edu_data['end_year'] = end_year
            if action == 'add':
                if not edu_data.get('degree') or not edu_data.get('institution'):
                    return jsonify({"error": "degree and institution are required"}), 400
                education.append(edu_data)
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "idx": len(education) - 1})
            idx = req.get('idx')
            if not isinstance(idx, int):
                return jsonify({"error": "idx (int) is required for update"}), 400
            if idx < 0 or idx >= len(education):
                return jsonify({"ok": False, "error": "Index out of range"}), 404
            education[idx].update(edu_data)
            _save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "idx": idx})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/award")
    def master_data_update_award():
        """Add, update, or delete an award entry in the master CV.

        Body:
            action      — 'add' | 'update' | 'delete'
            idx         — list index (int, required for update/delete)
            title, year, description, relevant_for
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get('action') or '').strip()
        if action not in ('add', 'update', 'delete'):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        parsed_year = None
        if action in ('add', 'update') and req.get('year') not in (None, ''):
            try:
                parsed_year = int(req.get('year'))
            except (TypeError, ValueError):
                return jsonify({"error": "year must be an integer"}), 400
            if parsed_year < 1900 or parsed_year > 2100:
                return jsonify({"error": "year must be between 1900 and 2100"}), 400

        try:
            master, master_path = _load_master(orchestrator.master_data_path)
            awards = master.setdefault('awards', [])
            if action == 'delete':
                idx = req.get('idx')
                if not isinstance(idx, int):
                    return jsonify({"error": "idx (int) is required for delete"}), 400
                if idx < 0 or idx >= len(awards):
                    return jsonify({"ok": False, "error": "Index out of range"}), 404
                awards.pop(idx)
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "deleted"})
            award_data: Dict[str, Any] = {}
            for field in ('title', 'description'):
                if field in req:
                    award_data[field] = req[field]
            if parsed_year is not None:
                award_data['year'] = parsed_year
            if 'relevant_for' in req:
                award_data['relevant_for'] = req['relevant_for']
            if action == 'add':
                if not award_data.get('title'):
                    return jsonify({"error": "title is required"}), 400
                awards.append(award_data)
                _save_master(master, master_path)
                return jsonify({"ok": True, "action": "added", "idx": len(awards) - 1})
            idx = req.get('idx')
            if not isinstance(idx, int):
                return jsonify({"error": "idx (int) is required for update"}), 400
            if idx < 0 or idx >= len(awards):
                return jsonify({"ok": False, "error": "Index out of range"}), 404
            awards[idx].update(award_data)
            _save_master(master, master_path)
            return jsonify({"ok": True, "action": "updated", "idx": idx})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    # ------------------------------------------------------------------
    # Publication CRUD  (master-data / publications.bib)
    # ------------------------------------------------------------------

    @app.get("/api/master-data/publications")
    def master_data_get_publications():
        """Return all publications stored in publications.bib.

        Returns both the parsed structured list and the raw BibTeX file content
        so the UI can render either the CRUD table or the raw text editor from
        a single request.

        Response fields:
            ok                 — True
            publications       — list of {key, type, fields, formatted_citation}
            content            — raw .bib file text (empty string if file missing)
            path               — absolute path to the .bib file
            count              — number of parsed entries in memory
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        pubs = orchestrator.publications or {}
        result = []
        for key, pub in pubs.items():
            item: Dict[str, Any] = {
                "key":   key,
                "type":  pub.get("type", ""),
                "fields": pub.get("fields", {}),
            }
            try:
                item["formatted_citation"] = format_publication(pub, style="apa")
            except Exception:
                item["formatted_citation"] = ""
            result.append(item)
        try:
            bib_path = orchestrator.publications_path
            content  = bib_path.read_text(encoding="utf-8") if bib_path.exists() else ""
        except Exception:
            content  = ""
            bib_path = orchestrator.publications_path
        return jsonify({
            "ok":           True,
            "publications": result,
            "content":      content,
            "path":         str(bib_path),
            "count":        len(pubs),
        })

    @app.put("/api/master-data/publications")
    def master_data_save_raw_publications():
        """Overwrite publications.bib with raw BibTeX text.

        Safety contract:
          1. The incoming content is parsed FIRST.  If parsing fails or the
             content is non-empty but yields zero entries, the request is
             rejected (400) and the file is never touched.
          2. A timestamped backup is created before writing.
          3. If the file write itself fails, the backup is automatically
             restored so the live .bib is never left in a broken state.

        Body:
            content  — raw BibTeX text to write (empty string clears the file,
                       which is accepted because an empty file is valid)

        Response:
            ok      — True
            count   — number of entries parsed from the saved content
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req     = request.get_json() or {}
        content = req.get("content", "")

        # --- Step 1: parse before touching the file ---
        try:
            parsed = bibtex_text_to_publications(content)
        except Exception as e:
            return jsonify({"ok": False, "error": f"BibTeX parse error: {e}"}), 400

        if content.strip() and not parsed:
            return jsonify({
                "ok":    False,
                "error": "No valid BibTeX entries found — file not saved.",
            }), 400

        # --- Step 2: backup ---
        bib_path    = orchestrator.publications_path
        backup_path = None
        try:
            if bib_path.exists():
                backup_dir  = bib_path.parent / "backups"
                backup_dir.mkdir(parents=True, exist_ok=True)
                ts          = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                backup_path = backup_dir / f"{bib_path.stem}.{ts}{bib_path.suffix}"
                shutil.copy2(bib_path, backup_path)
        except Exception as e:
            return jsonify({"ok": False, "error": f"Backup failed: {e}"}), 500

        # --- Step 3: write — restore backup on failure ---
        try:
            bib_path.write_text(content, encoding="utf-8")
            orchestrator.publications = parsed
            return jsonify({"ok": True, "count": len(parsed)})
        except Exception as e:
            if backup_path and backup_path.exists():
                try:
                    shutil.copy2(backup_path, bib_path)
                except Exception:
                    pass  # best-effort restore; backup is still on disk
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/publications/validate")
    def master_data_validate_publications():
        """Parse BibTeX text and report errors without saving anything.

        Use this to give the user instant feedback before they commit a save.

        Body:
            bibtex_text  — raw BibTeX string to validate

        Response on success:
            ok      — True
            count   — number of entries parsed
            entries — list of {key, type} for each valid entry
        Response on parse failure:
            ok      — False
            error   — description of the problem
        """
        entry = _get_session()
        _validate_owner(entry)
        req         = request.get_json() or {}
        bibtex_text = req.get("bibtex_text", "")

        if not bibtex_text or not bibtex_text.strip():
            return jsonify({"ok": True, "count": 0, "entries": []})

        try:
            parsed = bibtex_text_to_publications(bibtex_text)
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 400

        if not parsed:
            return jsonify({
                "ok":    False,
                "error": "No valid BibTeX entries found in the supplied text.",
            }), 400

        entries = [
            {"key": k, "type": v.get("type", "")}
            for k, v in parsed.items()
        ]
        return jsonify({"ok": True, "count": len(entries), "entries": entries})

    @app.post("/api/master-data/publication")
    def master_data_update_publication():
        """Add, update, or delete a single publication in publications.bib.

        Body:
            action  — 'add' | 'update' | 'delete'
            key     — BibTeX cite key (required for all actions)
            type    — entry type, e.g. 'article' (required for add/update)
            fields  — dict of BibTeX fields (required for add/update;
                      must include at least title and year, and one of
                      author or editor)
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req    = request.get_json() or {}
        action = (req.get("action") or "").strip()
        if action not in ("add", "update", "delete"):
            return jsonify({"error": "action must be add, update, or delete"}), 400

        key = (req.get("key") or "").strip()
        if not key:
            return jsonify({"error": "key is required"}), 400

        pubs = dict(orchestrator.publications or {})

        try:
            if action == "delete":
                if key not in pubs:
                    return jsonify({"ok": False, "error": f"Key '{key}' not found"}), 404
                del pubs[key]
                orchestrator.publications_path.write_text(
                    serialize_publications_to_bibtex(pubs), encoding="utf-8"
                )
                orchestrator.publications = parse_bibtex_file(
                    str(orchestrator.publications_path)
                )
                return jsonify({"ok": True, "action": "deleted"})

            # add / update
            fields = req.get("fields")
            if not isinstance(fields, dict):
                return jsonify({"error": "fields (dict) is required for add/update"}), 400
            entry_type = (req.get("type") or "").strip()
            if not entry_type:
                return jsonify({"error": "type is required for add/update"}), 400

            if not fields.get("title"):
                return jsonify({"error": "fields.title is required"}), 400
            if not fields.get("year"):
                return jsonify({"error": "fields.year is required"}), 400
            if not fields.get("author") and not fields.get("editor"):
                return jsonify({"error": "fields.author or fields.editor is required"}), 400

            if action == "add" and key in pubs:
                return jsonify({"error": f"Key '{key}' already exists; use action=update"}), 409

            pubs[key] = {"key": key, "type": entry_type, "fields": fields}
            orchestrator.publications_path.write_text(
                serialize_publications_to_bibtex(pubs), encoding="utf-8"
            )
            orchestrator.publications = parse_bibtex_file(
                str(orchestrator.publications_path)
            )
            return jsonify({"ok": True, "action": action, "key": key})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/master-data/publications/import")
    def master_data_import_publications():
        """Parse a BibTeX string and merge entries into publications.bib.

        Body:
            bibtex_text  — raw BibTeX string to import
            overwrite    — bool (default false); if true, existing keys are
                           replaced by the imported entries
        Returns:
            added, updated, skipped counts and the full updated publication list.
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator
        req         = request.get_json() or {}
        bibtex_text = req.get("bibtex_text", "")
        overwrite   = bool(req.get("overwrite", False))

        if not bibtex_text or not bibtex_text.strip():
            return jsonify({"error": "bibtex_text is required"}), 400

        try:
            imported = bibtex_text_to_publications(bibtex_text)
        except Exception as e:
            return jsonify({"ok": False, "error": f"BibTeX parse error: {e}"}), 400

        if not imported:
            return jsonify({"ok": False, "error": "No valid BibTeX entries found"}), 400

        pubs   = dict(orchestrator.publications or {})
        added  = 0
        updated = 0
        skipped = 0
        for key, pub in imported.items():
            if key in pubs:
                if overwrite:
                    pubs[key] = pub
                    updated += 1
                else:
                    skipped += 1
            else:
                pubs[key] = pub
                added += 1

        try:
            orchestrator.publications_path.write_text(
                serialize_publications_to_bibtex(pubs), encoding="utf-8"
            )
            orchestrator.publications = parse_bibtex_file(
                str(orchestrator.publications_path)
            )
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

        return jsonify({
            "ok": True,
            "added": added,
            "updated": updated,
            "skipped": skipped,
            "total": len(pubs),
        })

    @app.post("/api/master-data/publications/convert")
    def master_data_convert_publications():
        """Use the LLM to convert free-form citation text to BibTeX.

        This endpoint returns a BibTeX preview string and does NOT save anything.

        Body:
            text  — free-form citation text (plain text, DOI, APA string, etc.)
        Returns:
            bibtex  — generated BibTeX string
        """
        entry = _get_session()
        _validate_owner(entry)
        orchestrator = entry.orchestrator

        if not getattr(orchestrator, "llm", None):
            return jsonify({"ok": False, "error": "No LLM provider configured for this session"}), 503

        req  = request.get_json() or {}
        text = (req.get("text") or "").strip()
        if not text:
            return jsonify({"error": "text is required"}), 400

        try:
            bibtex = orchestrator.llm.convert_text_to_bibtex(text)
        except LLMError as e:
            return jsonify({"ok": False, "error": str(e)}), 500
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

        return jsonify({"ok": True, "bibtex": bibtex})

    @app.post("/api/generate-summary")
    def generate_professional_summary():
        """Generate (or refine) a custom LLM professional summary for this session.

        Body (JSON, all optional):
            refinement_prompt: str   — user instructions for iterative refinement
            previous_summary:  str   — existing text to refine (required when
                                       refinement_prompt is provided)

        On success stores the result as ``session_summaries['ai_generated']`` and
        auto-selects it via ``summary_focus_override``.

        Returns:
            {"ok": true, "summary": "<text>"}
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            req = request.get_json() or {}
            refinement_prompt = (req.get('refinement_prompt') or '').strip() or None
            previous_summary  = (req.get('previous_summary')  or '').strip() or None

            job_analysis = conversation.state.get('job_analysis')
            if not job_analysis:
                return jsonify({"ok": False, "error": "No job analysis found. Analyse a job description first."}), 400

            # Build a compact list of selected experiences.
            # Prefer the user's submitted decisions (experience_decisions) over the
            # LLM's original recommendations, so the summary reflects what the user
            # actually chose to include.
            all_experiences = orchestrator.master_data.get('experience', []) if orchestrator and orchestrator.master_data else []
            experience_decisions = conversation.state.get('experience_decisions') or {}
            if experience_decisions:
                # Include experiences the user accepted (emphasize or include)
                selected_exp_ids = {
                    eid for eid, action in experience_decisions.items()
                    if action in ('emphasize', 'include')
                }
                selected_experiences = [e for e in all_experiences if e.get('id') in selected_exp_ids]
            else:
                # Fall back to LLM recommendations if user hasn't submitted yet
                customizations = conversation.state.get('customizations') or {}
                recommended_ids = set(customizations.get('recommended_experiences') or [])
                selected_experiences = (
                    [e for e in all_experiences if e.get('id') in recommended_ids]
                    if recommended_ids else all_experiences
                )

            with entry.lock:
                summary = llm_client.generate_professional_summary(
                    job_analysis=job_analysis,
                    master_data=orchestrator.master_data if orchestrator else {},
                    selected_experiences=selected_experiences,
                    refinement_prompt=refinement_prompt,
                    previous_summary=previous_summary,
                )

                if not summary:
                    return jsonify({"ok": False, "error": "LLM returned an empty summary. Please try again."}), 500

                # Persist in session_summaries so the orchestrator can resolve it
                session_summaries = conversation.state.get('session_summaries') or {}
                session_summaries['ai_generated'] = summary
                conversation.state['session_summaries'] = session_summaries
                # Auto-select the generated summary
                conversation.state['summary_focus_override'] = 'ai_generated'
                conversation._save_session()
            session_registry.touch(sid)

            return jsonify({"ok": True, "summary": summary})

        except Exception as exc:
            traceback.print_exc()
            return jsonify({"ok": False, "error": str(exc)}), 500

    # ── Cover Letter endpoints (Phase 14) ───────────────────────────────────

    _TONE_GUIDANCE: Dict[str, str] = {
        'startup/tech':   'Energetic, direct, outcome-focused.  Emphasise velocity, impact, and technical depth.',
        'pharma/biotech': 'Precise, methodical, compliance-aware.  Reference domain credentials and regulatory rigour.',
        'academia':       'Scholarly, collaborative.  Highlight publications, teaching experience, and departmental service.',
        'financial':      'Professional, quantitative, risk-aware.  Emphasise fiduciary responsibility and data-driven decisions.',
        'leadership':     'Strategic, vision-focused, people-first.  Highlight team-building and organisational impact.',
    }

    @app.get("/api/cover-letter/prior")
    def cover_letter_prior():
        """Scan prior sessions for saved cover letters and return previews."""
        try:
            from utils.config import get_config
            cfg         = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            results     = []
            if output_base.exists():
                for session_file in sorted(output_base.rglob('session.json'), reverse=True)[:30]:
                    try:
                        with open(session_file, encoding='utf-8') as f:
                            data = json.load(f)
                        state = data.get('state', {})
                        cl_text = state.get('cover_letter_text')
                        if not cl_text:
                            continue
                        params = state.get('cover_letter_params') or {}
                        job_analysis = state.get('job_analysis') or {}
                        results.append({
                            'session_path': str(session_file),
                            'company':      job_analysis.get('company', ''),
                            'role':         job_analysis.get('title', ''),
                            'date':         data.get('timestamp', '')[:10],
                            'tone':         params.get('tone', ''),
                            'preview':      cl_text[:200],
                            'full_text':    cl_text,
                        })
                    except Exception:
                        pass
            return jsonify({'ok': True, 'sessions': results})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @app.post("/api/cover-letter/generate")
    def cover_letter_generate():
        """Generate a cover letter using the LLM and current session context."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        with entry.lock:
            body            = request.get_json(silent=True) or {}
            tone            = body.get('tone', 'startup/tech')
            hiring_manager  = (body.get('hiring_manager') or 'Hiring Manager').strip()
            company_address = (body.get('company_address') or '').strip()
            highlight       = (body.get('highlight') or '').strip()
            reuse_body      = (body.get('reuse_body') or '').strip()

            job_analysis  = conversation.state.get('job_analysis') or {}
            master        = orchestrator.master_data or {}
            personal_info = master.get('personal_info', {})

            # Build skills/summary snippet from master data
            skills_raw = master.get('skills', [])
            all_skills = conversation.normalize_skills_data(skills_raw)
            skill_names = [
                s.get('name', str(s)) if isinstance(s, dict) else str(s)
                for s in all_skills
            ]
            top_skills = ', '.join(skill_names[:12]) if skill_names else '(see attached CV)'

            summaries = master.get('professional_summaries', {})
            if isinstance(summaries, dict) and summaries:
                summary_text = next(iter(summaries.values()))
            else:
                summary_text = master.get('summary', '')

            achievements   = master.get('selected_achievements', [])
            top_ach_titles = '\n'.join(f'- {a.get("title", "")}' for a in achievements[:4]) or '(see CV)'

            answers_snippet = ''
            answers = conversation.state.get('post_analysis_answers') or {}
            if answers:
                answers_snippet = 'Candidate context:\n' + '\n'.join(
                    f'- {q}: {a}' for q, a in list(answers.items())[:6]
                )

            tone_hint = _TONE_GUIDANCE.get(tone, '')
            company   = job_analysis.get('company', 'the company')
            role      = job_analysis.get('title', 'the position')
            keywords  = ', '.join((job_analysis.get('ats_keywords') or [])[:12])
            req_skills = ', '.join((job_analysis.get('required_skills') or [])[:10])

            today = datetime.now().strftime('%B %d, %Y')
            pieces   = [f'Date: {today}']
            if company_address:
                pieces.append(company_address)
            header_block = '\n'.join(pieces) + '\n\n'  # LLM provides the salutation

            reuse_instruction = (
                f'\nUse the following prior cover letter as a starting point, '
                f'adapting it to the new role:\n\n"""\n{reuse_body}\n"""\n'
            ) if reuse_body else ''

            prompt = f"""\
You are a professional career coach writing a tailored cover letter.

Tone style: {tone} — {tone_hint}

TARGET ROLE
  Company: {company}
  Position: {role}
  Key requirements: {req_skills or keywords or '(see job description)'}

CANDIDATE PROFILE
  Name: {personal_info.get('name', 'The candidate')}
  Summary: {summary_text[:400] if summary_text else '(see CV)'}
  Top skills: {top_skills}
  Key achievements:
{top_ach_titles}

{answers_snippet}
{reuse_instruction}
{'Please especially highlight: ' + highlight if highlight else ''}

Write a compelling, personalised cover letter (3–4 paragraphs, ~300–400 words).
Start directly with the salutation line: "Dear {hiring_manager},"
Do NOT include a date, address block, or subject line — return only the letter body starting with the salutation.
Reference concrete skills and achievements from the candidate profile.
Close professionally with a call to action.
"""

            try:
                response = llm_client.chat(
                    messages=[
                        {'role': 'system', 'content': 'You write tailored, professional cover letters. Return only the letter body text.'},
                        {'role': 'user',   'content': prompt},
                    ],
                    temperature=0.7,
                )
            except Exception as e:
                return jsonify({'ok': False, 'error': f'LLM error: {e}'}), 500

            # Prepend date + address block; LLM response already contains the salutation
            letter_text = header_block + response.strip()
            conversation.state['cover_letter_text']   = letter_text
            conversation.state['cover_letter_params'] = {
                'tone': tone, 'hiring_manager': hiring_manager,
                'company_address': company_address, 'highlight': highlight,
            }
        session_registry.touch(sid)
        return jsonify({'ok': True, 'text': letter_text})

    @app.post("/api/cover-letter/save")
    def cover_letter_save():
        """Save cover letter text to DOCX in the output directory and update metadata.json."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        with entry.lock:
            text = (request.get_json(silent=True) or {}).get('text', '').strip()
            if not text:
                return jsonify({'error': 'text is required'}), 400

            generated = conversation.state.get('generated_files')
            if not generated or not generated.get('output_dir'):
                return jsonify({'error': 'No generated CV found — please generate your CV first.'}), 400

            try:
                from docx import Document
                from docx.shared import Pt

                output_dir   = Path(generated['output_dir'])
                job_analysis = conversation.state.get('job_analysis') or {}
                company      = (job_analysis.get('company') or 'Company').replace(' ', '_')
                date_str     = datetime.now().strftime('%Y-%m-%d')
                filename     = f'CoverLetter_{company}_{date_str}.docx'
                docx_path    = output_dir / filename

                doc = Document()
                # Use Normal style and write paragraphs
                for para_text in text.split('\n'):
                    p = doc.add_paragraph(para_text)
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
                doc.save(str(docx_path))

                # Update metadata.json
                metadata_path = output_dir / 'metadata.json'
                if metadata_path.exists():
                    with open(metadata_path, encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {}
                metadata['cover_letter_text']        = text
                metadata['cover_letter_reused_from'] = conversation.state.get('cover_letter_reused_from')
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)

                conversation.state['cover_letter_text'] = text
                session_registry.touch(sid)
                return jsonify({'ok': True, 'filename': filename})
            except Exception as e:
                return jsonify({'ok': False, 'error': str(e)}), 500

    # ── Phase 15: Screening Response endpoints ──────────────────────────────

    @app.post("/api/screening/search")
    def screening_search():
        """For a single question, find the best prior library match and top 3 relevant experiences."""
        try:
            body     = request.get_json(force=True) or {}
            question = (body.get('question') or '').strip()
            if not question:
                return jsonify({'ok': False, 'error': 'question required'}), 400

            from utils.config import get_config
            config   = get_config()
            lib_path = Path(config.master_cv_path).parent / 'response_library.json'
            library: list = []
            if lib_path.exists():
                with open(lib_path, encoding='utf-8') as f:
                    library = json.load(f)

            # Score every library entry against this question
            scored_prior = sorted(
                [
                    {
                        'score': _text_similarity(
                            question,
                            (e.get('question') or '') + ' ' + (e.get('response_text') or ''),
                        ),
                        'entry': e,
                    }
                    for e in library
                ],
                key=lambda x: x['score'],
                reverse=True,
            )
            best_prior = scored_prior[0] if scored_prior and scored_prior[0]['score'] >= 0.25 else None

            # Score experiences from master CV
            with open(config.master_cv_path, encoding='utf-8') as f:
                master = json.load(f)
            exps = master.get('experience', [])

            def _exp_text(e: dict) -> str:
                return ' '.join([
                    e.get('title', ''), e.get('company', ''), e.get('summary', ''),
                    ' '.join(e.get('achievements', [])[:5]),
                ])

            scored_exps = sorted(
                [{'score': _text_similarity(question, _exp_text(e)), 'exp': e, 'idx': i}
                 for i, e in enumerate(exps)],
                key=lambda x: x['score'],
                reverse=True,
            )[:3]

            return jsonify({
                'ok':     True,
                'prior':  best_prior['entry'] if best_prior else None,
                'experiences': [
                    {
                        'idx':        x['idx'],
                        'score':      round(x['score'], 2),
                        'title':      x['exp'].get('title', ''),
                        'company':    x['exp'].get('company', ''),
                        'date_range': x['exp'].get('date_range', ''),
                        'summary':    (x['exp'].get('summary') or '')[:200],
                    }
                    for x in scored_exps
                ],
            })
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @app.post("/api/screening/generate")
    def screening_generate():
        """Generate a draft screening-question response via LLM."""
        try:
            body            = request.get_json(force=True) or {}
            question        = (body.get('question') or '').strip()
            fmt             = body.get('format', 'direct')
            exp_indices     = body.get('experience_indices') or []
            prior_response  = (body.get('prior_response') or '').strip()

            if not question:
                return jsonify({'ok': False, 'error': 'question required'}), 400

            from utils.config import get_config
            config = get_config()
            with open(config.master_cv_path, encoding='utf-8') as f:
                master = json.load(f)
            exps     = master.get('experience', [])
            selected = [exps[i] for i in exp_indices if isinstance(i, int) and 0 <= i < len(exps)]

            fmt_name, word_range, fmt_instructions = _SCREENING_FORMAT_GUIDANCE.get(
                fmt, _SCREENING_FORMAT_GUIDANCE['direct']
            )

            exp_blocks = '\n\n'.join(
                f"Role: {e.get('title', '')} at {e.get('company', '')}\n"
                f"Summary: {(e.get('summary') or '')[:300]}\n"
                f"Key achievements: {'; '.join((e.get('achievements') or [])[:5])}"
                for e in selected
            ) or 'No specific experience provided.'

            entry = _get_session()
            conversation = entry.manager
            answers              = conversation.state.get('post_analysis_answers') or {}
            cover_letter_snippet = (conversation.state.get('cover_letter_text') or '')[:400]

            cl_context = (
                '\n'.join(f'- {k}: {v}' for k, v in list(answers.items())[:8])
                if answers else 'None provided.'
            )

            prior_block = (
                f'\nUse the following prior response as a starting point, adapting it as needed:\n'
                f'"""\n{prior_response}\n"""\n'
            ) if prior_response else ''

            prompt = (
                f'You are drafting a screening-question response for a job application.\n'
                f'Response format: {fmt_name} (~{word_range}). {fmt_instructions}\n\n'
                f'Question:\n"{question}"\n\n'
                f'Relevant experience:\n{exp_blocks}\n\n'
                f'Applicant preferences / context:\n{cl_context}\n'
                + (f'\nCover letter excerpt (for tone/context):\n{cover_letter_snippet}\n' if cover_letter_snippet else '')
                + prior_block
                + '\nWrite only the response text. No preamble, labels, or meta-commentary.'
            )

            try:
                response_text = llm_client.chat(
                    messages=[
                        {'role': 'system', 'content': 'You write concise, tailored screening-question responses for job applications.'},
                        {'role': 'user',   'content': prompt},
                    ],
                    temperature=0.7,
                )
            except Exception as e:
                return jsonify({'ok': False, 'error': f'LLM error: {e}'}), 500

            return jsonify({'ok': True, 'text': response_text.strip()})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    @app.post("/api/screening/save")
    def screening_save():
        """Save screening responses to DOCX, update metadata.json, upsert response_library.json."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            body         = request.get_json(force=True) or {}
            responses_in = body.get('responses') or []
            if not responses_in:
                return jsonify({'ok': False, 'error': 'No responses to save.'}), 400

            with entry.lock:
                output_dir = Path(orchestrator.output_dir)
                if not output_dir.exists():
                    return jsonify({'ok': False, 'error': 'Output directory not found. Generate a CV first.'}), 400

                from docx import Document as _DocxDoc
                from docx.shared import Pt as _Pt
                doc = _DocxDoc()
                doc.add_heading('Screening Question Responses', 0)
                for item in responses_in:
                    doc.add_heading((item.get('question') or '')[:120], level=2)
                    para = doc.add_paragraph(item.get('response_text') or '')
                    para.style.font.size = _Pt(11)
                    doc.add_paragraph()

                date_str = datetime.now().strftime('%Y-%m-%d')
                filename = f'Screening_Responses_{date_str}.docx'
                doc_path = output_dir / filename
                doc.save(str(doc_path))

                # Update metadata.json
                metadata_path = output_dir / 'metadata.json'
                if metadata_path.exists():
                    with open(metadata_path, encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {}
                metadata['screening_responses'] = responses_in
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)

                # Upsert response_library.json
                from utils.config import get_config
                config   = get_config()
                lib_path = Path(config.master_cv_path).parent / 'response_library.json'
                library: list = []
                if lib_path.exists():
                    with open(lib_path, encoding='utf-8') as f:
                        library = json.load(f)

                job_analysis = conversation.state.get('job_analysis') or {}
                company      = job_analysis.get('company', '') or ''
                session_p    = str(output_dir)
                for item in responses_in:
                    library.append({
                        'question':      item.get('question', ''),
                        'topic_tag':     item.get('topic_tag', ''),
                        'response_text': item.get('response_text', ''),
                        'format':        item.get('format', ''),
                        'company':       company,
                        'date':          date_str,
                        'session_path':  session_p,
                    })
                with open(lib_path, 'w', encoding='utf-8') as f:
                    json.dump(library, f, indent=2)

                conversation.state['screening_responses'] = responses_in
                session_registry.touch(sid)
                return jsonify({'ok': True, 'filename': filename, 'count': len(responses_in)})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    # ── Copilot OAuth endpoints ──────────────────────────────────────────────

    @app.post("/api/copilot-auth/start")
    def copilot_auth_start():
        """Begin Device Flow: returns user_code + verification_uri for the user to open."""
        try:
            flow = auth_manager.start_device_flow()
            _auth_poll["device_code"] = flow["device_code"]
            _auth_poll["interval"]    = flow.get("interval", 5)
            _auth_poll["error"]       = None
            return jsonify({
                "user_code":        flow["user_code"],
                "verification_uri": flow["verification_uri"],
                "interval":         flow.get("interval", 5),
                "expires_in":       flow.get("expires_in", 900),
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/copilot-auth/poll")
    def copilot_auth_poll():
        """Start a background thread that polls GitHub until the user approves the device flow."""
        import threading
        if _auth_poll["polling"]:
            return jsonify({"ok": True, "message": "Already polling"})
        device_code = _auth_poll.get("device_code")
        interval    = _auth_poll.get("interval", 5)
        if not device_code:
            return jsonify({"error": "No device flow in progress — call /start first"}), 400

        def _do_poll():
            _auth_poll["polling"] = True
            _auth_poll["error"]   = None
            try:
                auth_manager.complete_device_flow(device_code, interval)
            except Exception as exc:
                _auth_poll["error"] = str(exc)
            finally:
                _auth_poll["polling"] = False

        threading.Thread(target=_do_poll, daemon=True).start()
        return jsonify({"ok": True})

    @app.get("/api/copilot-auth/status")
    def copilot_auth_status():
        """Return current auth state (authenticated, polling, error)."""
        return jsonify({
            **auth_manager.status,
            "polling": _auth_poll["polling"],
            "error":   _auth_poll["error"],
        })

    @app.post("/api/copilot-auth/logout")
    def copilot_auth_logout():
        """Clear stored credentials."""
        auth_manager.logout()
        return jsonify({"ok": True})

    # ── Model selection endpoints ─────────────────────────────────────────────

    @app.get("/api/model")
    def get_model():
        """Return current model, all provider models, and pricing metadata.

        If a `session_id` is supplied (query or JSON body), prefer any
        session-scoped provider/model persisted in that session's state.
        """
        entry = _get_session(required=False, allow_missing=True)
        session_provider = None
        session_model = None
        if entry:
            conversation = entry.manager
            session_provider = conversation.state.get("provider")
            session_model = conversation.state.get("model")

        provider_for_view = session_provider or _provider_name
        current = session_model or _current_model or (llm_client.model if hasattr(llm_client, "model") else None)
        available = _get_available_models(provider_for_view, current_model=current)
        billing = PROVIDER_BILLING.get(provider_for_view, {"type": "per_token", "note": ""})
        live      = get_cached_pricing()
        models_with_info = [
            {
                "model":              m,
                "context_window":     MODEL_INFO.get(m, {}).get("context_window"),
                "cost_input":         (live.get(m) or MODEL_INFO.get(m, {})).get("cost_input"),
                "cost_output":        (live.get(m) or MODEL_INFO.get(m, {})).get("cost_output"),
                "copilot_multiplier": MODEL_INFO.get(m, {}).get("copilot_multiplier"),
                "notes":              MODEL_INFO.get(m, {}).get("notes", ""),
            }
            for m in available
        ]
        # Cross-provider model list for the model-selection UI
        all_models = []
        for prov in PROVIDER_MODELS:
            prov_models       = _get_available_models(prov)
            prov_billing_type = PROVIDER_BILLING.get(prov, {}).get("type", "per_token")
            for m in prov_models:
                pricing = live.get(m) or MODEL_INFO.get(m, {})
                price_source = "static_baseline"
                all_models.append({
                    "provider":           prov,
                    "model":              m,
                    "source":             "list_models" if prov in _dynamic_model_cache else "fallback_static",
                    "price_source":       price_source,
                    "billing_type":       prov_billing_type,
                    "context_window":     MODEL_INFO.get(m, {}).get("context_window"),
                    "cost_input":         pricing.get("cost_input"),
                    "cost_output":        pricing.get("cost_output"),
                    "copilot_multiplier": MODEL_INFO.get(m, {}).get("copilot_multiplier"),
                    "notes":              MODEL_INFO.get(m, {}).get("notes", ""),
                })
        return jsonify({
            "provider":           provider_for_view,
            "providers":          sorted(PROVIDER_MODELS.keys()),
            "list_models_capable": ["openai", "anthropic", "gemini", "groq"],
            "billing_type":       billing["type"],
            "billing_note":       billing["note"],
            "model":              current,
            "available":          models_with_info,
            "all_models":         all_models,
            "pricing_updated_at": get_pricing_updated_at(),
            "pricing_source":     get_pricing_source(),
        })

    @app.get("/api/model-catalog")
    def get_model_catalog():
        """Return model rows for selected providers.

        For providers that support list_models and have credentials configured,
        build rows from runtime-discovered models. Otherwise, fall back to the
        static provider/model catalog.
        """
        list_models_capable = _CATALOG_LIST_MODELS_CAPABLE

        selected_param = (request.args.get("providers") or "").strip()
        if selected_param:
            selected = [p.strip() for p in selected_param.split(",") if p.strip()]
        else:
            selected = [_provider_name]

        selected = [p for p in selected if p in PROVIDER_MODELS]
        if not selected:
            selected = [_provider_name]

        live = get_cached_pricing()
        rows: List[Dict[str, Any]] = []
        provider_sources: Dict[str, str] = {}
        provider_models: Dict[str, List[str]] = {}
        runtime_candidates: List[tuple[str, str]] = []

        for provider in selected:
            # Use the in-process cache when available; otherwise do a blocking fetch
            # and store the result so subsequent calls are fast.
            if provider in _dynamic_model_cache:
                discovered = _dynamic_model_cache[provider]
            else:
                discovered = _catalog_discover_provider_models(provider)
                if discovered:
                    with _dynamic_model_cache_lock:
                        _dynamic_model_cache[provider] = discovered
            if discovered:
                model_list = discovered
                provider_sources[provider] = "list_models"
                runtime_candidates.extend((provider, name) for name in model_list)
            else:
                model_list = PROVIDER_MODELS.get(provider, [])
                provider_sources[provider] = "fallback_static"
            provider_models[provider] = model_list

        runtime_prices = lookup_runtime_pricing_bulk(runtime_candidates, cached_pricing=live)

        for provider in selected:
            model_list = provider_models.get(provider, [])
            prov_billing_type = PROVIDER_BILLING.get(provider, {}).get("type", "per_token")
            for model_name in model_list:
                pricing = (
                    live.get(model_name)
                    or runtime_prices.get(model_name)
                    or MODEL_INFO.get(model_name, {})
                )
                if model_name in runtime_prices and model_name not in STATIC_PRICING:
                    price_source = "runtime_cache"
                else:
                    price_source = "static_baseline"
                base_notes = MODEL_INFO.get(model_name, {}).get("notes", "")
                if provider_sources[provider] == "list_models" and not base_notes:
                    base_notes = "Discovered via list_models"
                rows.append({
                    "provider":           provider,
                    "model":              model_name,
                    "source":             provider_sources[provider],
                    "price_source":       price_source,
                    "billing_type":       prov_billing_type,
                    "context_window":     MODEL_INFO.get(model_name, {}).get("context_window"),
                    "cost_input":         pricing.get("cost_input"),
                    "cost_output":        pricing.get("cost_output"),
                    "copilot_multiplier": MODEL_INFO.get(model_name, {}).get("copilot_multiplier"),
                    "notes":              base_notes,
                })

        return jsonify({
            "providers":            sorted(PROVIDER_MODELS.keys()),
            "selected_providers":   selected,
            "list_models_capable":  sorted(list(list_models_capable)),
            "provider_sources":     provider_sources,
            "all_models":           rows,
            "pricing_updated_at":   get_pricing_updated_at(),
            "pricing_source":       get_pricing_source(),
        })

    @app.post("/api/model-pricing/refresh")
    def refresh_model_pricing():
        """Refresh the pricing cache. Fetches live prices from OpenRouter; falls back to static."""
        pricing = refresh_pricing_cache()
        return jsonify({
            "ok":          True,
            "updated_at":  get_pricing_updated_at(),
            "source":      get_pricing_source(),
            "model_count": len(pricing),
        })

    @app.post("/api/model")
    def set_model():
        """Switch the active model and optionally the provider."""
        nonlocal llm_client, _provider_name, _current_model

        def _format_probe_error(provider_name: str, probe_error: Optional[str]) -> str:
            if not probe_error:
                return "Model probe failed."

            friendly = probe_error.strip()
            if provider_name == "github":
                friendly = friendly.replace("with OpenAI", "with GitHub Models")
                friendly = friendly.replace("by OpenAI", "by GitHub Models")
                friendly = friendly.replace("(openai)", "(github)")
            return friendly

        def _probe_client(candidate_client) -> tuple[bool, Optional[str]]:
            """Run a minimal chat call to ensure model/provider is reachable."""
            try:
                candidate_client.chat(
                    messages=[{"role": "user", "content": "Reply with one word: ready"}],
                    temperature=0,
                    max_tokens=8,
                )
                return True, None
            except Exception as exc:
                return False, str(exc)

        data     = request.get_json(silent=True) or {}
        model    = data.get("model", "").strip()
        provider = (data.get("provider") or _provider_name).strip()
        if not model:
            return jsonify({"error": "Missing model"}), 400
        available = PROVIDER_MODELS.get(provider, [])
        # Providers with dynamic model catalogs (e.g., list_models-backed) can
        # accept models beyond the static catalog. Enforce strict validation only
        # for providers that are intentionally static.
        static_only = {"copilot-oauth", "copilot", "github", "local"}
        if provider in static_only and available and model not in available:
            return jsonify({"error": f"Unknown model '{model}' for provider '{provider}'"}), 400
        try:
            candidate_client = get_llm_provider(provider=provider, model=model, auth_manager=auth_manager)
            ok, probe_error = _probe_client(candidate_client)
            if not ok:
                formatted_error = _format_probe_error(provider, probe_error)
                return jsonify({
                    "error": f"Model '{model}' is not currently available for provider '{provider}'. {formatted_error}",
                    "provider": provider,
                    "model": model,
                }), 400

            llm_client     = candidate_client
            _provider_name = provider
            _current_model = model
            # Update all active sessions with the new LLM client
            for _entry in session_registry.all_active():
                _entry.orchestrator.llm = llm_client
                _entry.manager.llm = llm_client

            # If the caller supplied a session_id, persist the chosen
            # provider/model into that session's state so it survives reloads.
            entry = _get_session(required=False, allow_missing=True)
            if entry:
                try:
                    _validate_owner(entry)
                    conv = entry.manager
                    conv.state["provider"] = provider
                    conv.state["model"] = model
                    conv._save_session()
                    session_registry.touch(entry.session_id)
                except Exception:
                    # Preserve original behavior even if session write fails
                    pass

            return jsonify({"ok": True, "provider": provider, "model": model})
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    @app.post("/api/model/test")
    def test_model():
        """Smoke-test the active LLM: send a minimal 1-token prompt.

        Returns {ok, provider, model, latency_ms} on success or
        {ok: false, error, provider, model} on failure.
        """
        import time
        t0 = time.monotonic()
        try:
            llm_client.chat(
                messages=[{"role": "user", "content": "Reply with one word: ready"}],
            )
            latency_ms = round((time.monotonic() - t0) * 1000)
            return jsonify({
                "ok":         True,
                "provider":   _provider_name,
                "model":      _current_model,
                "latency_ms": latency_ms,
            })
        except Exception as exc:
            return jsonify({
                "ok":       False,
                "error":    str(exc),
                "provider": _provider_name,
                "model":    _current_model,
            }), 200   # 200 so JS can read the body regardless

    # ── Job / chat endpoints ──────────────────────────────────────────────────

    @app.post("/api/job")
    def submit_job():
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        job_text: Optional[str] = data.get("job_text")
        if not job_text:
            return jsonify({"error": "Missing job_text"}), 400
        with entry.lock:
            # Store job description in state and also add to conversation history
            conversation.add_job_description(job_text)
            conversation.state["position_name"] = _infer_position_name(job_text)
            conversation.conversation_history.append({
                "role": "user",
                "content": job_text,
            })
        session_registry.touch(sid)
        return jsonify({"ok": True, "message": "Job description added."})
    
    @app.post("/api/fetch-job-url")
    def fetch_job_url():
        """Fetch job description from URL with enhanced error handling"""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        url = data.get("url")

        if not url:
            return jsonify({"error": "Missing URL"}), 400

        try:
            # Validate URL format
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                return jsonify({"error": "Invalid URL format"}), 400
            
            domain = parsed.netloc.lower()
            
            # Check for protected job boards that require special handling
            protected_sites = {
                'linkedin.com': {
                    'name': 'LinkedIn',
                    'message': 'LinkedIn requires login to view job descriptions. Please copy the job text manually from your browser.',
                    'instructions': [
                        '1. Open the LinkedIn job posting in your browser',
                        '2. Log in if needed and scroll to view the full job description', 
                        '3. Select and copy the job description text',
                        '4. Use the "Paste Text" tab to submit it directly'
                    ]
                },
                'indeed.com': {
                    'name': 'Indeed', 
                    'message': 'Indeed has anti-bot protection. Please copy the job text manually.',
                    'instructions': [
                        '1. Open the Indeed job posting in your browser',
                        '2. Copy the job description text',
                        '3. Use the "Paste Text" tab to submit it'
                    ]
                },
                'glassdoor.com': {
                    'name': 'Glassdoor',
                    'message': 'Glassdoor requires authentication. Please copy the job text manually.',
                    'instructions': ['Copy job text from browser and use "Paste Text" tab']
                }
            }
            
            # Check if this is a protected site
            for site_domain, site_info in protected_sites.items():
                if site_domain in domain:
                    return jsonify({
                        "error": f"{site_info['name']} Protection Detected",
                        "message": site_info['message'],
                        "instructions": site_info['instructions'],
                        "site_name": site_info['name'],
                        "protected_site": True
                    }), 400
            
            # Enhanced headers to mimic real browser
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Cache-Control': 'max-age=0',
                'sec-ch-ua': '"Not A(Brand";v="99", "Google Chrome";v="121", "Chromium";v="121"',
                'sec-ch-ua-mobile': '?0',
                'sec-ch-ua-platform': '"macOS"',
                'Sec-Fetch-Dest': 'document',
                'Sec-Fetch-Mode': 'navigate',
                'Sec-Fetch-Site': 'none',
                'Sec-Fetch-User': '?1'
            }
            
            # Fetch the URL with timeout and proper error handling
            print(f"📡 Fetching URL: {url}")
            response = requests.get(url, timeout=30, headers=headers, allow_redirects=True)
            
            # Check response status
            if response.status_code != 200:
                if response.status_code == 403:
                    return jsonify({
                        "error": "Access Forbidden (403)",
                        "message": "The website is blocking automated access. Please try copying the job description manually.",
                        "instructions": [
                            "1. Open the URL in your browser",
                            "2. Copy the job description text", 
                            "3. Use the 'Paste Text' tab to submit it"
                        ],
                        "status_code": response.status_code
                    }), 400
                elif response.status_code == 404:
                    return jsonify({
                        "error": "Page Not Found (404)",
                        "message": "The job posting may have been removed or the URL is incorrect.",
                        "status_code": response.status_code
                    }), 400
                else:
                    response.raise_for_status()
            
            # Extract text content
            content_type = response.headers.get('content-type', '').lower()
            print(f"📄 Content type: {content_type}")
            
            if 'text/plain' in content_type:
                job_text = response.text
            elif 'text/html' in content_type or 'html' in content_type:
                # Parse HTML and extract text
                soup = BeautifulSoup(response.text, 'html.parser')

                # ── Pre-extract structured data BEFORE stripping script tags ──
                # JS-rendered SPAs (e.g. Workday) have an empty body but rich
                # JSON-LD structured data and og:description meta tags in <head>.
                import json as _json
                json_ld_text = None
                for script_tag in soup.find_all('script', type='application/ld+json'):
                    try:
                        ld_data = _json.loads(script_tag.string or '')
                        desc = ld_data.get('description') if isinstance(ld_data, dict) else None
                        if desc and len(desc) > 100:
                            json_ld_text = desc
                            print(f"📋 Found JSON-LD job description ({len(json_ld_text)} chars)")
                            break
                    except Exception:
                        pass

                meta_desc_text = None
                for meta in soup.find_all('meta'):
                    prop = meta.get('property', '') or meta.get('name', '')
                    if prop in ('og:description', 'description'):
                        content = meta.get('content', '')
                        if len(content) > 100:
                            meta_desc_text = content
                            print(f"📋 Found meta description ({len(meta_desc_text)} chars)")
                            break

                # Remove script and style elements
                for script in soup(["script", "style", "nav", "header", "footer"]):
                    script.decompose()
                
                # Try to find job-specific content first
                job_selectors = [
                    '.job-description',
                    '.job-content', 
                    '.posting-description',
                    '.description',
                    '[data-testid="job-description"]',
                    '.job-details'
                ]
                
                job_content = None
                for selector in job_selectors:
                    elements = soup.select(selector)
                    if elements:
                        job_content = elements[0]
                        break
                
                # If no specific job content found, get main content or body
                if not job_content:
                    job_content = soup.find('main') or soup.find('article') or soup.find('body') or soup
                
                # Get text content
                job_text = job_content.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in job_text.splitlines())
                job_text = '\n'.join(line for line in lines if line)

                # For JS-rendered SPAs body text may be nearly empty — fall back
                # to structured data extracted before stripping script tags.
                if len(job_text.strip()) < 200:
                    if json_ld_text:
                        job_text = json_ld_text
                        print(f"ℹ️ Using JSON-LD structured data (body text was too short)")
                    elif meta_desc_text:
                        job_text = meta_desc_text
                        print(f"ℹ️ Using meta description (body text was too short)")

                # Basic validation - check if we got meaningful content
                if len(job_text.strip()) < 100:
                    return jsonify({
                        "error": "Insufficient Content",
                        "message": "The fetched content appears to be too short or may not contain the job description.",
                        "instructions": [
                            "1. Check if the URL is correct",
                            "2. Try opening the URL in your browser first",
                            "3. Copy the job description manually and use 'Paste Text' tab"
                        ],
                        "content_length": len(job_text)
                    }), 400
            else:
                return jsonify({
                    "error": f"Unsupported content type: {content_type}",
                    "message": "The URL does not contain text or HTML content that can be processed."
                }), 400
            
            # Store job description in state
            with entry.lock:
                conversation.add_job_description(job_text)
                conversation.state["position_name"] = _infer_position_name(job_text)
            session_registry.touch(sid)
            print(f"✅ Successfully fetched {len(job_text)} characters from {domain}")
            
            return jsonify({
                "ok": True,
                "job_text": job_text,
                "message": f"Job description fetched from {domain}",
                "source_url": url,
                "content_length": len(job_text)
            })
            
        except requests.Timeout:
            return jsonify({
                "error": "Request Timeout",
                "message": "The website took too long to respond. Please try again or use manual text input.",
                "instructions": ["Try copying the job description manually and use the 'Paste Text' tab"]
            }), 500
        except requests.ConnectionError:
            return jsonify({
                "error": "Connection Error", 
                "message": "Unable to connect to the website. Please check the URL or your internet connection.",
                "instructions": ["Verify the URL is correct and accessible in your browser"]
            }), 500
        except requests.RequestException as e:
            return jsonify({
                "error": "Network Error",
                "message": f"Failed to fetch URL: {str(e)}",
                "instructions": ["Try copying the job description manually and use the 'Paste Text' tab"],
                "technical_details": str(e)
            }), 500
        except Exception as e:
            print(f"❌ Error processing URL {url}: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({
                "error": "Processing Error",
                "message": f"Error processing content: {str(e)}",
                "instructions": ["Try copying the job description manually and use the 'Paste Text' tab"],
                "technical_details": str(e)
            }), 500
    
    @app.post("/api/upload-file")
    def upload_file():
        """Extract text from an uploaded file (txt, md, html, pdf, docx, etc.)."""
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400

        f = request.files['file']
        if not f.filename:
            return jsonify({"error": "Empty filename"}), 400

        filename_lower = f.filename.lower()
        raw = f.read()

        try:
            # ── Plain text / Markdown ─────────────────────────────────────────
            if any(filename_lower.endswith(ext) for ext in ('.txt', '.md', '.rst', '.text')):
                text = raw.decode('utf-8', errors='replace')

            # ── HTML ──────────────────────────────────────────────────────────
            elif any(filename_lower.endswith(ext) for ext in ('.html', '.htm')):
                soup = BeautifulSoup(raw, 'html.parser')
                for tag in soup(['script', 'style', 'head', 'nav', 'footer']):
                    tag.decompose()
                text = soup.get_text(separator='\n')

            # ── PDF ───────────────────────────────────────────────────────────
            elif filename_lower.endswith('.pdf'):
                import io
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(io.BytesIO(raw))
                    pages = [page.extract_text() or '' for page in reader.pages]
                    text = '\n\n'.join(pages)
                except ImportError:
                    return jsonify({"error": "PDF support not available. Run: pip install pypdf"}), 500

            # ── DOCX ──────────────────────────────────────────────────────────
            elif filename_lower.endswith('.docx'):
                import io
                try:
                    import mammoth
                    result = mammoth.extract_raw_text(io.BytesIO(raw))
                    text = result.value
                except ImportError:
                    try:
                        from docx import Document
                        doc = Document(io.BytesIO(raw))
                        text = '\n'.join(p.text for p in doc.paragraphs)
                    except ImportError:
                        return jsonify({"error": "DOCX support not available. Run: pip install python-docx"}), 500

            # ── DOC (legacy Word) ─────────────────────────────────────────────
            elif filename_lower.endswith('.doc'):
                return jsonify({
                    "error": "Legacy .doc format not supported",
                    "message": "Please save the file as .docx or copy-paste the content."
                }), 400

            # ── RTF ───────────────────────────────────────────────────────────
            elif filename_lower.endswith('.rtf'):
                # Strip RTF control words crudely — good enough for job descriptions
                import re as _re
                text_bytes = raw.decode('latin-1', errors='replace')
                text = _re.sub(r'\\[a-z]+\d*\s?|[{}]', ' ', text_bytes)

            else:
                # Try decoding as UTF-8 fallback for unknown extensions
                try:
                    text = raw.decode('utf-8', errors='replace')
                except Exception:
                    return jsonify({
                        "error": f"Unsupported file type: {filename_lower.rsplit('.',1)[-1]}",
                        "message": "Supported formats: txt, md, html, pdf, docx, rtf"
                    }), 400

            # Clean up whitespace
            import re as _re
            text = _re.sub(r'\n{3,}', '\n\n', text).strip()

            if len(text) < 50:
                return jsonify({
                    "error": "Insufficient Content",
                    "message": "The file appears to be empty or contains no readable text.",
                    "content_length": len(text)
                }), 400

            print(f"📎 Uploaded file '{f.filename}': extracted {len(text)} characters")
            return jsonify({
                "ok":             True,
                "text":           text,
                "filename":       f.filename,
                "content_length": len(text),
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": f"Error reading file: {str(e)}"}), 500

    @app.post("/api/load-job-file")
    def load_job_file():
        """Load a job description from a file."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        filename = data.get("filename")
        if not filename:
            return jsonify({"error": "Missing filename"}), 400

        # Look for the file in sample_jobs directory
        job_file_path = Path(__file__).parent.parent / "sample_jobs" / filename

        # Also check CV directory for other job files
        if not job_file_path.exists():
            cv_path = Path.home() / "CV" / "files" / filename
            if cv_path.exists():
                job_file_path = cv_path

        if not job_file_path.exists():
            return jsonify({"error": f"File not found: {filename}"}), 404

        try:
            with open(job_file_path, 'r', encoding='utf-8') as f:
                job_text = f.read()

            with entry.lock:
                # Store job description in state
                conversation.add_job_description(job_text)
                conversation.state["position_name"] = _infer_position_name(job_text)
            session_registry.touch(sid)

            return jsonify({
                "ok": True,
                "job_text": job_text,
                "message": f"Loaded job description from {filename}"
            })
        except Exception as e:
            return jsonify({"error": f"Failed to load file: {str(e)}"}), 500

    @app.get("/api/positions")
    def positions():
        # Reuse ConversationManager helper to list positions
        entry = _get_session()
        conversation = entry.manager
        try:
            names = conversation._list_positions()
            return jsonify({"positions": names})
        except Exception as e:
            return jsonify({"error": str(e), "positions": []}), 500

    @app.post("/api/position")
    def set_position():
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        name = data.get("name")
        open_latest = bool(data.get("open_latest"))
        if not name:
            return jsonify({"error": "Missing name"}), 400
        with entry.lock:
            conversation.state["position_name"] = name
            loaded = False
            if open_latest:
                loaded = conversation._load_latest_session_for_position(name)
        session_registry.touch(sid)
        return jsonify({"ok": True, "loaded": loaded, "position_name": name})

    @app.post("/api/message")
    def send_message():
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        msg: Optional[str] = data.get("message")
        if not msg:
            return jsonify({"error": "Missing message"}), 400
        try:
            with entry.lock:
                response = conversation._process_message(msg)
            session_registry.touch(sid)
            return jsonify(dataclasses.asdict(MessageResponse(
                ok=True,
                response=response,
                phase=conversation.state.get("phase"),
            )))
        except LLMAuthError as e:
            logger.warning("LLM auth error in /api/message: %s", e)
            return jsonify({"error": str(e), "error_type": "auth"}), 401
        except LLMRateLimitError as e:
            logger.warning("LLM rate limit in /api/message: %s", e)
            return jsonify({"error": str(e), "error_type": "rate_limit"}), 429
        except LLMContextLengthError as e:
            logger.warning("LLM context length in /api/message: %s", e)
            return jsonify({"error": str(e), "error_type": "context_length"}), 400
        except LLMError as e:
            logger.error("LLM provider error in /api/message: %s", e)
            return jsonify({"error": str(e), "error_type": "provider"}), 502
        except Exception as e:
            # Comprehensive error logging
            logger.error(
                "ERROR in /api/message endpoint: %s (msg: %s)",
                type(e).__name__,
                msg[:100] if len(msg) > 100 else msg,
                exc_info=True
            )
            return jsonify({"error": str(e)}), 500

    @app.post("/api/rename-current-session")
    def rename_current_session():
        """Rename the currently active session (no path needed)."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        new_name = (data.get("new_name") or "").strip()[:200]
        if not new_name:
            return jsonify({"error": "Missing new_name"}), 400
        with entry.lock:
            conversation.state["position_name"] = new_name
            conversation._save_session()
        session_registry.touch(sid)
        return jsonify({"ok": True, "new_name": new_name})

    @app.post("/api/rename-session")
    def rename_session():
        """Rename a session's position_name in its session.json file."""
        # This is a disk-level rename; it does NOT require an active session_id.
        data = request.get_json(silent=True) or {}
        path     = data.get("path")
        new_name = (data.get("new_name") or "").strip()[:200]
        if not path:
            return jsonify({"error": "Missing path"}), 400
        if not new_name:
            return jsonify({"error": "Missing new_name"}), 400
        session_file = Path(path)
        if not session_file.exists():
            return jsonify({"error": f"Session not found: {path}"}), 404
        # Safety: must be inside the configured output base
        try:
            cfg = get_config()
            output_base = Path(cfg.get("data.output_dir", "~/CV/files")).expanduser()
            if not session_file.resolve().is_relative_to(output_base.resolve()):
                return jsonify({"error": "Path is outside the output directory"}), 400
        except Exception:
            pass  # if we can't verify, proceed (path existence check above is sufficient)
        try:
            with open(session_file, "r", encoding="utf-8") as f:
                session_data = json.load(f)
            session_data.setdefault("state", {})["position_name"] = new_name
            with open(session_file, "w", encoding="utf-8") as f:
                json.dump(session_data, f, indent=2, default=str)
            # If this is an active in-memory session, sync in-memory state too
            for _entry in session_registry.all_active():
                _active = getattr(_entry.manager, "session_file", None)
                if _active and str(Path(_active).resolve()) == str(session_file.resolve()):
                    _entry.manager.state["position_name"] = new_name
                    _entry.manager._save_session()
            return jsonify({"ok": True, "new_name": new_name})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    def _usage_prompt_tokens(usage) -> int | None:
        """Extract prompt/input token count from a provider usage object or dict."""
        if usage is None:
            return None
        if isinstance(usage, dict):
            return usage.get("prompt_tokens") or usage.get("input_tokens")
        return (
            getattr(usage, "prompt_tokens", None)
            or getattr(usage, "input_tokens", None)
        )

    @app.get("/api/context-stats")
    def context_stats():
        """Return a rough token-usage estimate for the current session."""
        entry = _get_session()
        conversation = entry.manager
        orchestrator = entry.orchestrator
        MODEL_CONTEXT_WINDOWS = {
            "gemini-2.5-pro":    2_000_000,
            "gemini-2.5-flash":  1_048_576,
            "gemini-2.0-flash":  1_048_576,
            "gemini-1.5-pro":    2_000_000,
            "gemini-1.5-flash":  1_048_576,
            "gpt-4o":              128_000,
            "gpt-4o-mini":         128_000,
            "gpt-4-turbo":         128_000,
            "o1":                  200_000,
            "o3-mini":             200_000,
            "claude-3-5-sonnet":   200_000,
            "claude-3-5-haiku":    200_000,
            "claude-3-opus":       200_000,
            "llama":               128_000,
        }
        try:
            model_name     = getattr(orchestrator.llm, "model", "") or ""
            context_window = 128_000
            for key, size in MODEL_CONTEXT_WINDOWS.items():
                if key.lower() in model_name.lower():
                    context_window = size
                    break

            # Prefer exact prompt-token count from the last API response; fall back to
            # a char-based estimate (chars / 4) when no real usage is available yet.
            real_tokens  = _usage_prompt_tokens(getattr(orchestrator.llm, "last_usage", None))
            if real_tokens is not None:
                token_count  = real_tokens
                token_source = "exact"
            else:
                state_chars   = len(json.dumps(conversation.state, default=str))
                history_chars = sum(len(str(m.get("content", ""))) for m in conversation.conversation_history)
                base_overhead = 4_000  # boilerplate in the system prompt
                token_count   = (state_chars + history_chars + base_overhead) // 4
                token_source  = "estimated"

            return jsonify({
                "ok":               True,
                "estimated_tokens": token_count,
                "token_source":     token_source,
                "context_window":   context_window,
                "model":            model_name,
                "history_messages": len(conversation.conversation_history),
            })
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.post("/api/generation-settings")
    def update_generation_settings():
        """Update per-session generation settings (max_skills, etc.).

        Request body: ``{"max_skills": int}``
        All fields are optional; only provided fields are updated.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        if "max_skills" in data:
            v = data["max_skills"]
            if not isinstance(v, int) or not (1 <= v <= 100):
                return jsonify({"error": "max_skills must be an integer between 1 and 100"}), 400
            with entry.lock:
                conversation.state["max_skills"] = v
                conversation._save_session()
        else:
            with entry.lock:
                conversation._save_session()
        session_registry.touch(sid)
        cfg_default = get_config().get("generation.max_skills", 20)
        return jsonify({
            "ok": True,
            "max_skills": int(conversation.state.get("max_skills") or cfg_default),
        })

    @app.post("/api/post-analysis-responses")
    def post_analysis_responses():
        """Persist generated post-analysis questions and user answers into session state."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        questions = data.get("questions") or []
        answers = data.get("answers") or {}

        if not isinstance(questions, list):
            return jsonify({"error": "questions must be a list"}), 400
        if not isinstance(answers, dict):
            return jsonify({"error": "answers must be an object"}), 400

        cleaned_questions = []
        for q in questions[:8]:
            if not isinstance(q, dict):
                continue
            question_text = str(q.get("question", "")).strip()
            qtype = str(q.get("type", "clarification")).strip().lower().replace(" ", "_")
            if question_text:
                cleaned_questions.append({
                    "type": qtype[:40] or "clarification",
                    "question": question_text[:2000],
                })

        cleaned_answers = {}
        for key, value in answers.items():
            clean_key = str(key).strip().lower().replace(" ", "_")[:40]
            if not clean_key:
                continue
            clean_value = str(value).strip()
            if clean_value:
                cleaned_answers[clean_key] = clean_value[:1000]

        with entry.lock:
            conversation.state["post_analysis_questions"] = cleaned_questions
            conversation.state["post_analysis_answers"] = cleaned_answers
            conversation._save_session()
        session_registry.touch(sid)

        return jsonify({
            "ok": True,
            "questions_count": len(cleaned_questions),
            "answers_count": len(cleaned_answers),
        })

    @app.post("/api/save")
    def save():
        entry = _get_session()
        conversation = entry.manager
        try:
            with entry.lock:
                conversation._save_session()
            session_file = str(conversation.session_dir / "session.json") if conversation.session_dir else None
            return jsonify({"ok": True, "session_file": session_file})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.get("/api/sessions")
    def list_sessions():
        """List saved sessions, most recent first."""
        try:
            from utils.config import get_config
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            sessions = []
            if output_base.exists():
                for session_file in sorted(output_base.rglob("session.json"), reverse=True):
                    # Exclude anything inside .trash
                    if trash_dir in session_file.parents:
                        continue
                    try:
                        import json as _json
                        with open(session_file) as f:
                            data = _json.load(f)
                        state = data.get('state', {})
                        sessions.append(SessionItem(
                            path=str(session_file),
                            position_name=state.get('position_name') or session_file.parent.name,
                            timestamp=data.get('timestamp', ''),
                            phase=state.get('phase', ''),
                            has_job=bool(state.get('job_description')),
                            has_analysis=bool(state.get('job_analysis')),
                            has_customizations=bool(state.get('customizations')),
                        ))
                    except Exception:
                        pass
            return jsonify(SessionListResponse(sessions=sessions[:20]).to_dict())  # cap at 20 most recent
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.get("/api/load-items")
    def load_items():
        """Merged list of saved sessions and server-side job files for the Load Job panel."""
        items = []

        # ── Sessions ──────────────────────────────────────────────────────────
        try:
            from utils.config import get_config as _cfg
            cfg = _cfg()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            if output_base.exists():
                for session_file in sorted(output_base.rglob("session.json"), reverse=True):
                    if trash_dir in session_file.parents:
                        continue
                    try:
                        import json as _json
                        with open(session_file) as f:
                            data = _json.load(f)
                        state = data.get('state', {})
                        items.append({
                            "kind":         "session",
                            "path":         str(session_file),
                            "label":        state.get('position_name') or session_file.parent.name,
                            "timestamp":    data.get('timestamp', ''),
                            "phase":        state.get('phase', ''),
                            "has_job":      bool(state.get('job_description')),
                            "has_analysis": bool(state.get('job_analysis')),
                            "has_cv":       bool(state.get('generated_files')),
                        })
                    except Exception:
                        pass
        except Exception:
            pass

        items = items[:20]  # cap sessions at 20

        # ── Server-side job files ──────────────────────────────────────────────
        try:
            sample_jobs_dir = Path(__file__).parent.parent / "sample_jobs"
            if sample_jobs_dir.exists():
                for f in sorted(sample_jobs_dir.iterdir()):
                    if f.suffix.lower() in {'.txt', '.md', '.html', '.pdf', '.docx', '.rtf'}:
                        label = f.stem.replace('_', ' ').replace('-', ' ').title()
                        items.append({
                            "kind":      "file",
                            "path":      str(f),
                            "filename":  f.name,
                            "label":     label,
                            "timestamp": "",
                            "phase":     "",
                        })
        except Exception:
            pass

        return jsonify({"items": items})

    @app.post("/api/load-session")
    def load_session_endpoint():
        """Load a saved session file and register it in the session registry."""
        data = request.get_json(silent=True) or {}
        path = data.get("path")
        if not path:
            return jsonify({"error": "Missing path"}), 400
        session_file = Path(path)
        if not session_file.exists():
            return jsonify({"error": f"Session file not found: {path}"}), 404
        try:
            sid, entry = session_registry.load_from_file(str(session_file), _app_config)
            conversation = entry.manager
            return jsonify({
                "ok":            True,
                "session_id":    sid,
                "redirect_url":  f"/?session={sid}",
                "session_file":  str(session_file),
                "position_name": conversation.state.get("position_name"),
                "phase":         conversation.state.get("phase"),
                "has_job":       bool(conversation.state.get("job_description")),
                "has_analysis":  bool(conversation.state.get("job_analysis")),
                "history_count": len(conversation.conversation_history),
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/delete-session")
    def delete_session_endpoint():
        """Move a session directory to the .trash folder (recoverable).

        Accepts a JSON body with either:
        - ``path``: full path to a ``session.json`` file → moves its parent
          directory into ``<output_base>/.trash/``.
        - ``session_id``: legacy positional identifier (directory name or
          suffix) → falls back to a name-matching search for backward compat.
        """
        data = request.get_json(silent=True) or {}
        path_param  = data.get("path") or data.get("session_id")
        if not path_param:
            return jsonify({"error": "Missing path or session_id"}), 400

        try:
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'

            def _move_to_trash(job_dir: Path):
                """Move job_dir into .trash, handling name collisions."""
                trash_dir.mkdir(parents=True, exist_ok=True)
                dest = trash_dir / job_dir.name
                if dest.exists():
                    dest = trash_dir / f"{job_dir.name}_{int(datetime.now().timestamp())}"
                import shutil as _shutil
                _shutil.move(str(job_dir), str(dest))
                print(f"Trashed: {job_dir} → {dest}")

            # ── Preferred: caller supplies the full session.json path ──────────
            candidate = Path(path_param)
            if candidate.exists() and candidate.name == 'session.json':
                job_dir = candidate.parent
                if job_dir.resolve().is_relative_to(output_base.resolve()):
                    _move_to_trash(job_dir)
                    return jsonify({"success": True, "message": "Session moved to Trash"})
                else:
                    return jsonify({"error": "Path is outside the output directory"}), 400

            # ── Fallback: match by directory name or position name ────────────
            deleted = False
            for session_file in output_base.rglob('session.json'):
                if trash_dir in session_file.parents:
                    continue
                job_dir = session_file.parent
                if path_param in job_dir.name or job_dir.name == path_param:
                    _move_to_trash(job_dir)
                    deleted = True
                    break

            if deleted:
                return jsonify({"success": True, "message": "Session moved to Trash"})
            return jsonify({"error": f"Session not found: {path_param}"}), 404

        except Exception as e:
            print(f"ERROR in delete_session: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.get("/api/trash")
    def list_trash():
        """List sessions in the .trash folder."""
        try:
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            items = []
            if trash_dir.exists():
                for session_file in sorted(trash_dir.rglob("session.json"), reverse=True):
                    try:
                        import json as _json
                        with open(session_file) as f:
                            data = _json.load(f)
                        state = data.get('state', {})
                        items.append({
                            "path":          str(session_file),
                            "position_name": state.get('position_name') or session_file.parent.name,
                            "timestamp":     data.get('timestamp', ''),
                            "phase":         state.get('phase', ''),
                        })
                    except Exception:
                        pass
            return jsonify({"items": items, "count": len(items)})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/trash/restore")
    def trash_restore():
        """Restore a trashed session back to the output directory."""
        data = request.get_json(silent=True) or {}
        path_param = data.get("path")
        if not path_param:
            return jsonify({"error": "Missing path"}), 400
        try:
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            candidate   = Path(path_param)
            if not candidate.exists() or candidate.name != 'session.json':
                return jsonify({"error": "Session file not found"}), 404
            job_dir = candidate.parent
            if not job_dir.resolve().is_relative_to(trash_dir.resolve()):
                return jsonify({"error": "Path is not inside trash"}), 400
            dest = output_base / job_dir.name
            if dest.exists():
                dest = output_base / f"{job_dir.name}_restored_{int(datetime.now().timestamp())}"
            import shutil as _shutil
            _shutil.move(str(job_dir), str(dest))
            print(f"Restored: {job_dir} → {dest}")
            return jsonify({"success": True, "message": "Session restored"})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/trash/delete")
    def trash_delete_one():
        """Permanently delete a single item from trash."""
        data = request.get_json(silent=True) or {}
        path_param = data.get("path")
        if not path_param:
            return jsonify({"error": "Missing path"}), 400
        try:
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            candidate   = Path(path_param)
            if not candidate.exists() or candidate.name != 'session.json':
                return jsonify({"error": "Session file not found"}), 404
            job_dir = candidate.parent
            if not job_dir.resolve().is_relative_to(trash_dir.resolve()):
                return jsonify({"error": "Path is not inside trash"}), 400
            import shutil as _shutil
            _shutil.rmtree(job_dir)
            print(f"Permanently deleted: {job_dir}")
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/trash/empty")
    def trash_empty():
        """Permanently delete everything in the .trash folder."""
        try:
            cfg = get_config()
            output_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash_dir   = output_base / '.trash'
            if trash_dir.exists():
                import shutil as _shutil
                _shutil.rmtree(trash_dir)
                trash_dir.mkdir()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/action")
    def do_action():
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        action = data.get("action")
        if not action:
            return jsonify({"error": "Missing action"}), 400

        # Format the payload correctly for _execute_action
        payload = {"action": action}
        if data.get("job_text"):
            payload["job_text"] = data["job_text"]
        if data.get("user_preferences"):
            payload["user_preferences"] = data["user_preferences"]

        try:
            with entry.lock:
                result = conversation._execute_action(payload)
            if not result:
                return jsonify({"error": "Invalid or unsupported action"}), 400
            session_registry.touch(sid)
            return jsonify(dataclasses.asdict(ActionResponse(
                ok=True,
                result=result,
                phase=conversation.state.get("phase"),
            )))
        except LLMAuthError as e:
            logger.warning("LLM auth error in /api/action: %s", e)
            return jsonify({"error": str(e), "error_type": "auth"}), 401
        except LLMRateLimitError as e:
            logger.warning("LLM rate limit in /api/action: %s", e)
            return jsonify({"error": str(e), "error_type": "rate_limit"}), 429
        except LLMContextLengthError as e:
            logger.warning("LLM context length in /api/action: %s", e)
            return jsonify({"error": str(e), "error_type": "context_length"}), 400
        except LLMError as e:
            logger.error("LLM provider error in /api/action: %s", e)
            return jsonify({"error": str(e), "error_type": "provider"}), 502
        except Exception as e:
            logger.error("Action execution error: %s", e, exc_info=True)
            return jsonify({"error": str(e)}), 500

    @app.post("/api/back-to-phase")
    def back_to_phase():
        """Navigate back to a prior phase without clearing downstream state.

        Body: ``{"phase": "analysis"|"customizations"|"rewrite"|"spell"|...,
                 "feedback": "optional refinement note"}``
        Resolves frontend step labels to internal phase strings automatically.
        If ``feedback`` is provided it is injected as a user message so the
        next LLM call sees it as context.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        target = data.get("phase")
        if not target:
            return jsonify({"error": "Missing phase"}), 400
        try:
            with entry.lock:
                result = conversation.back_to_phase(target)
                feedback = (data.get("feedback") or "").strip()
                if feedback:
                    conversation.conversation_history.append({
                        "role": "user",
                        "content": f"[Refinement feedback for {target}]: {feedback}",
                    })
            session_registry.touch(sid)
            return jsonify(result)
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/re-run-phase")
    def re_run_phase():
        """Re-execute the LLM call for a phase with downstream context preserved.

        Body: ``{"phase": "analysis"|"customizations"|"rewrite"}``
        Returns ``{ok, phase, prior_output, new_output}``.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        target = data.get("phase")
        if not target:
            return jsonify({"error": "Missing phase"}), 400
        try:
            with entry.lock:
                result = conversation.re_run_phase(target)
            if not result.get("ok"):
                return jsonify(result), 400
            session_registry.touch(sid)
            return jsonify(result)
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.get("/api/intake-metadata")
    def intake_metadata():
        """Return extracted or confirmed intake metadata for the current session.

        If intake has already been confirmed (``state.intake.confirmed == True``),
        the stored values are returned as-is.  Otherwise the fields are extracted
        heuristically from the stored job description text.

        Response: ``{role, company, date_applied, confirmed}``
        """
        entry = _get_session()
        conversation = entry.manager
        intake = conversation.state.get('intake') or {}
        if intake.get('confirmed'):
            return jsonify({
                'role':         intake.get('role'),
                'company':      intake.get('company'),
                'date_applied': intake.get('date_applied'),
                'confirmed':    True,
            })
        extracted = conversation.extract_intake_metadata()
        return jsonify({
            'role':         extracted.get('role'),
            'company':      extracted.get('company'),
            'date_applied': extracted.get('date_applied'),
            'confirmed':    False,
        })

    @app.post("/api/confirm-intake")
    def confirm_intake():
        """Persist user-confirmed intake metadata and immediately save the session.

        Body: ``{company, role, date_applied}``
        All fields are optional; missing/blank fields are stored as ``None``.
        Session is persisted synchronously so the intake record survives a page reload.

        Response: ``{ok, intake: {role, company, date_applied, confirmed}}``
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        intake = {
            'role':         (data.get('role') or '').strip() or None,
            'company':      (data.get('company') or '').strip() or None,
            'date_applied': (data.get('date_applied') or '').strip() or None,
            'confirmed':    True,
        }
        with entry.lock:
            conversation.state['intake'] = intake
            if intake.get('role') and intake.get('company'):
                conversation.state['position_name'] = f"{intake['role']} at {intake['company']}"
            elif intake.get('role'):
                conversation.state['position_name'] = intake['role']
            elif intake.get('company'):
                conversation.state['position_name'] = intake['company']
            conversation._save_session()
        session_registry.touch(sid)
        return jsonify({'ok': True, 'intake': intake})

    @app.get("/api/prior-clarifications")
    def prior_clarifications():
        """Return prior post-analysis answers from the most recent session with a similar role.

        Searches persisted sessions by keyword overlap between the current
        session's ``intake.role`` and each candidate session's ``intake.role``.
        Sessions lacking ``post_analysis_answers`` are skipped.

        Query param: ``?limit=3`` (default 3, max 10)
        Response: ``{found, matches: [{position_name, role, company, date, answers, overlap}]}``
        """
        entry = _get_session()
        conversation = entry.manager
        current_intake = conversation.state.get('intake') or {}
        current_role   = (current_intake.get('role') or '').lower()

        limit = min(int(request.args.get('limit', 3)), 10)

        _STOP = {
            'a', 'an', 'the', 'of', 'in', 'at', 'for', 'to', 'and', 'or',
            'is', 'be', 'with', 'on', 'by', 'senior', 'junior', 'lead',
        }

        def _kw(role: str) -> set:
            return {w for w in role.split() if w not in _STOP and len(w) > 2}

        current_kw = _kw(current_role)

        try:
            from utils.config import get_config as _get_cfg
            cfg      = _get_cfg()
            out_base = Path(cfg.get('data.output_dir', '~/CV/files')).expanduser()
            trash    = out_base / '.trash'
            matches  = []
            for sf in sorted(out_base.rglob('session.json'), reverse=True):
                if trash in sf.parents:
                    continue
                try:
                    with open(sf, 'r', encoding='utf-8') as fh:
                        sd = json.load(fh)
                    st      = sd.get('state', {})
                    answers = st.get('post_analysis_answers') or {}
                    if not answers:
                        continue
                    prior_intake = st.get('intake') or {}
                    prior_role   = (prior_intake.get('role') or st.get('position_name') or '').lower()
                    overlap      = current_kw & _kw(prior_role)
                    if not overlap:
                        continue
                    matches.append({
                        'position_name': st.get('position_name'),
                        'role':          prior_intake.get('role'),
                        'company':       prior_intake.get('company'),
                        'date':          prior_intake.get('date_applied') or sd.get('timestamp', '')[:10],
                        'answers':       answers,
                        'overlap':       sorted(overlap),
                    })
                    if len(matches) >= limit:
                        break
                except Exception:
                    pass
            return jsonify({'found': len(matches) > 0, 'matches': matches})
        except Exception as e:
            return jsonify({'found': False, 'matches': [], 'error': str(e)})

    @app.get("/api/synonym-lookup")
    def synonym_lookup():
        """Look up the canonical form of a skill or keyword via the synonym map.

        Query param: ``?term=ML``
        Returns ``{term, canonical, found}`` — ``found`` is False when no
        mapping exists (canonical == term in that case).
        """
        entry = _get_session()
        conversation = entry.manager
        term = request.args.get("term", "").strip()
        if not term:
            return jsonify({"error": "Missing term query parameter"}), 400
        canonical = conversation.orchestrator.canonical_skill_name(term)
        return jsonify({"term": term, "canonical": canonical, "found": canonical != term})

    @app.get("/api/synonym-map")
    def synonym_map():
        """Return the full synonym map as ``{alias: canonical}``."""
        entry = _get_session()
        conversation = entry.manager
        return jsonify(conversation.orchestrator._synonym_map)

    @app.post("/api/reorder-bullets")
    def reorder_bullets():
        """Persist a user-defined bullet ordering for one experience.

        Body: ``{"experience_id": "exp_001", "order": [2, 0, 1]}``
        ``order`` is a list of original achievement indices in the desired
        display order.  Pass an empty list to reset to relevance-sorted order.
        Returns ``{ok: true}``.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        exp_id = data.get("experience_id")
        order  = data.get("order")
        if not exp_id:
            return jsonify({"error": "Missing experience_id"}), 400
        if order is None:
            return jsonify({"error": "Missing order list"}), 400
        if not isinstance(order, list):
            return jsonify({"error": "order must be a list of integers"}), 400
        try:
            with entry.lock:
                achievement_orders = conversation.state.setdefault("achievement_orders", {})
                if order:
                    achievement_orders[exp_id] = [int(i) for i in order]
                else:
                    achievement_orders.pop(exp_id, None)  # reset → use auto relevance sort
                conversation._save_session()
            session_registry.touch(sid)
            return jsonify({"ok": True, "experience_id": exp_id, "order": order})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.get("/api/proposed-bullet-order")
    def proposed_bullet_order():
        """Return relevance-ranked bullet order for one experience based on job keywords.

        Query param: ``?experience_id=exp_001``
        Returns ``{"proposed_order": [2, 0, 1], "has_job_analysis": true}``.
        Indices are original achievement positions sorted by keyword relevance (highest first).
        Returns natural order when no job analysis is available.
        """
        entry = _get_session()
        conversation = entry.manager
        exp_id = request.args.get("experience_id")
        if not exp_id:
            return jsonify({"error": "Missing experience_id"}), 400
        try:
            master_data = conversation.orchestrator.master_data
            experiences_list = master_data.get("experiences") or master_data.get("experience", [])
            experience = next((e for e in experiences_list if e.get("id") == exp_id), None)
            if not experience:
                return jsonify({"error": f"Experience {exp_id} not found"}), 404

            achievements = list(experience.get("achievements") or [])
            if not achievements:
                return jsonify({"proposed_order": [], "has_job_analysis": False})

            job_analysis = conversation.state.get("job_analysis") or {}
            job_keywords = {kw.lower() for kw in job_analysis.get("ats_keywords", [])}

            if not job_keywords:
                return jsonify({
                    "proposed_order": list(range(len(achievements))),
                    "has_job_analysis": False,
                })

            def ach_score(ach):
                text = (ach.get("text", "") if isinstance(ach, dict) else str(ach)).lower()
                tokens = set(re.findall(r'\b\w+\b', text))
                return len(tokens & job_keywords)

            proposed = sorted(range(len(achievements)), key=lambda i: ach_score(achievements[i]), reverse=True)
            return jsonify({"proposed_order": proposed, "has_job_analysis": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/reorder-rows")
    def reorder_rows():
        """Persist a user-defined row ordering for experiences or skills.

        Body: ``{"type": "experience"|"skill", "ordered_ids": ["exp_001", ...]}``
        Pass an empty list to reset to relevance-scored order.
        Returns ``{ok: true}``.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        row_type   = data.get("type")
        ordered_ids = data.get("ordered_ids")
        if row_type not in ("experience", "skill"):
            return jsonify({"error": "type must be 'experience' or 'skill'"}), 400
        if ordered_ids is None:
            return jsonify({"error": "Missing ordered_ids"}), 400
        if not isinstance(ordered_ids, list):
            return jsonify({"error": "ordered_ids must be a list"}), 400
        state_key = "experience_row_order" if row_type == "experience" else "skill_row_order"
        try:
            with entry.lock:
                if ordered_ids:
                    conversation.state[state_key] = [str(i) for i in ordered_ids]
                else:
                    conversation.state.pop(state_key, None)
                conversation._save_session()
            session_registry.touch(sid)
            return jsonify({"ok": True, "type": row_type, "ordered_ids": ordered_ids})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.post("/api/post-analysis-draft-response")
    def post_analysis_draft_response():
        """Use the LLM to draft an answer for a single clarification question."""
        entry = _get_session()
        conversation = entry.manager
        try:
            body          = request.get_json(force=True) or {}
            question      = (body.get('question') or '').strip()
            question_type = (body.get('question_type') or '').strip()
            analysis      = body.get('analysis') or {}

            if not question:
                return jsonify({'ok': False, 'error': 'question required'}), 400

            existing_answers = conversation.state.get('post_analysis_answers') or {}

            context_items: List[str] = []
            if isinstance(analysis, dict):
                for key, label in [('job_title', 'Job Title'), ('company_name', 'Company'),
                                   ('role_level', 'Role Level'), ('domain', 'Domain')]:
                    if analysis.get(key):
                        context_items.append(f"{label}: {analysis[key]}")
            context = '\n'.join(context_items) or 'Not available'

            prior_answers_block = (
                '\n'.join(f'- {k}: {v}' for k, v in list(existing_answers.items())[:6])
                if existing_answers else 'None yet.'
            )

            prompt = (
                'You are helping a job applicant answer a clarifying question about their CV preferences.\n'
                'The question was asked to better tailor their CV for a specific job.\n\n'
                f'Job context:\n{context}\n\n'
                f'Previously answered questions:\n{prior_answers_block}\n\n'
                f'Question to answer:\n"{question}"\n\n'
                'Write a concise, first-person DRAFT answer (1–3 sentences) the applicant could give.\n'
                'Base it on the most sensible choice for someone applying to this role.\n'
                'Write only the answer text. No preamble, no labels.'
            )

            try:
                draft = llm_client.chat(
                    messages=[
                        {'role': 'system', 'content': 'You write concise draft answers for job application clarifying questions.'},
                        {'role': 'user',   'content': prompt},
                    ],
                    temperature=0.7,
                    # No max_tokens cap: thinking-capable models (e.g. Gemini 2.0 Flash)
                    # count reasoning tokens against max_output_tokens, so a 200-token
                    # limit leaves only a handful of tokens for the visible response.
                    # The prompt already constrains output to 1-3 sentences.
                )
            except Exception as e:
                err_str = str(e)
                if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str or 'quota' in err_str.lower() or 'rate' in err_str.lower():
                    return jsonify({'ok': False, 'error': 'Rate limit reached — please wait a moment and try again.', 'rate_limited': True}), 429
                return jsonify({'ok': False, 'error': f'LLM error: {e}'}), 500

            return jsonify({'ok': True, 'text': draft.strip()})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500


    @app.post("/api/post-analysis-questions")
    def post_analysis_questions():
        """Generate post-analysis clarifying questions, preferably via LLM."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}

        analysis = _coerce_to_dict(
            data.get("analysis") or conversation.state.get("job_analysis")
        )

        if not analysis:
            return jsonify({"ok": True, "questions": []})

        # Pass prior answers so the LLM avoids re-asking already-covered topics.
        prior_qa = conversation.state.get("post_analysis_answers") or None

        questions: List[Dict[str, str]] = []
        source = "fallback"
        try:
            questions = _generate_post_analysis_questions(
                analysis=analysis,
                job_text=conversation.state.get("job_description"),
                prior_qa=prior_qa,
            )
            if questions:
                source = "llm"
        except Exception as e:
            print(f"Question generation failed, using fallback: {e}")

        if not questions:
            questions = _fallback_post_analysis_questions(analysis)

        with entry.lock:
            conversation.state["post_analysis_questions"] = questions
            conversation._save_session()
        session_registry.touch(sid)

        return jsonify({"ok": True, "questions": questions, "source": source})

    @app.get("/api/history")
    def history():
        # Return the conversation history for chat-style rendering
        entry = _get_session()
        conversation = entry.manager
        return jsonify({
            "history": conversation.conversation_history,
            "phase": conversation.state.get("phase"),
        })

    def normalize_experience_id(exp_id):
        """Return the experience ID unchanged for direct lookup against master data."""
        return exp_id
    
    @app.post("/api/experience-details")
    def get_experience_details():
        entry = _get_session()
        conversation = entry.manager
        data = request.get_json(silent=True) or {}
        experience_id = data.get("experience_id")
        if not experience_id:
            return jsonify({"error": "Missing experience_id"}), 400

        try:
            # Normalize the ID to match master data format
            normalized_id = normalize_experience_id(experience_id)

            # Look up experience details in master data
            master_data = conversation.orchestrator.master_data
            experience = None
            
            # Search through experiences in master data (check both 'experience' and 'experiences')
            experiences_list = master_data.get("experiences") or master_data.get("experience", [])
            if experiences_list:
                for exp in experiences_list:
                    if exp.get("id") == normalized_id or exp.get("id") == experience_id:
                        experience = exp
                        break
            
            if experience:
                return jsonify({"experience": experience})
            else:
                # Log available IDs for debugging
                available_ids = [exp.get("id") for exp in experiences_list] if experiences_list else []
                print(f"DEBUG: Experience '{experience_id}' (normalized: '{normalized_id}') not found")
                print(f"DEBUG: Available IDs: {available_ids[:10]}")
                return jsonify({"experience": None, "message": f"Experience {experience_id} not found"})
                
        except Exception as e:
            print(f"ERROR in get_experience_details: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.route('/api/review-decisions', methods=['POST'])
    def save_review_decisions():
        """Save user's review decisions for experiences/skills"""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.json

        if not data:
            return jsonify({"error": "No data provided"}), 400

        decision_type = data.get('type')  # 'experiences' or 'skills'
        decisions = data.get('decisions', {})

        if not decision_type or not decisions:
            return jsonify({"error": "Missing type or decisions"}), 400

        try:
            # Store decisions in conversation state
            if decision_type == 'experiences':
                conversation.state['experience_decisions'] = decisions
                message = f"Saved decisions for {len(decisions)} experiences"
            elif decision_type == 'skills':
                conversation.state['skill_decisions'] = decisions
                extra_skills = data.get('extra_skills', [])
                if extra_skills:
                    conversation.state['extra_skills'] = extra_skills

                raw_matches = data.get('extra_skill_matches') or {}
                sanitized_matches: Dict[str, List[str]] = {}
                valid_ids = {
                    (exp.get('id') or '').strip()
                    for exp in (conversation.orchestrator.master_data.get('experience') or [])
                    if isinstance(exp, dict) and (exp.get('id') or '').strip()
                }
                if isinstance(raw_matches, dict):
                    for skill_name, exp_ids in raw_matches.items():
                        if not isinstance(skill_name, str) or not skill_name.strip():
                            continue
                        if isinstance(exp_ids, str):
                            exp_ids = [x.strip() for x in exp_ids.split(',') if x.strip()]
                        if not isinstance(exp_ids, list):
                            continue
                        cleaned = []
                        seen = set()
                        for exp_id in exp_ids:
                            if not isinstance(exp_id, str):
                                continue
                            exp_id = exp_id.strip()
                            if not exp_id or exp_id not in valid_ids or exp_id in seen:
                                continue
                            cleaned.append(exp_id)
                            seen.add(exp_id)
                        if cleaned:
                            sanitized_matches[skill_name.strip()] = cleaned
                conversation.state['extra_skill_matches'] = sanitized_matches

                # Mirror extra-skill decisions into customizations so render/generation
                # paths can use them without relying on separate top-level state.
                customizations = dict(conversation.state.get('customizations') or {})
                customizations['extra_skills'] = conversation.state.get('extra_skills') or []
                customizations['extra_skill_matches'] = sanitized_matches
                conversation.state['customizations'] = customizations
                message = f"Saved decisions for {len(decisions)} skills"
            elif decision_type == 'achievements':
                conversation.state['achievement_decisions'] = decisions
                accepted_suggestions = data.get('accepted_suggestions', [])
                if accepted_suggestions:
                    conversation.state['accepted_suggested_achievements'] = accepted_suggestions
                message = f"Saved decisions for {len(decisions)} achievements" + (
                    f" (+{len(accepted_suggestions)} AI suggestions accepted)" if accepted_suggestions else ""
                )
            elif decision_type == 'publications':
                conversation.state['publication_decisions'] = decisions
                message = f"Saved decisions for {len(decisions)} publications"
            elif decision_type == 'summary_focus':
                # decisions is a single string key here
                conversation.state['summary_focus_override'] = decisions
                message = "Saved summary focus preference"
            else:
                return jsonify({"error": f"Invalid type: {decision_type}"}), 400
            
            # Save the updated state
            conversation._save_session()

            print(f"Saved {decision_type} decisions: {decisions}")
            session_registry.touch(sid)
            return jsonify({"success": True, "message": message})

        except Exception as e:
            print(f"ERROR in save_review_decisions: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.route('/api/save-achievement-edits', methods=['POST'])
    def save_achievement_edits():
        """Save per-experience achievement edits from the achievements editor tab."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.json or {}
        edits = data.get('edits', {})
        if not edits:
            return jsonify({"error": "No edits provided"}), 400
        # Convert string keys (from JSON) to int if needed
        normalized = {int(k): v for k, v in edits.items() if str(k).lstrip("-").isdigit()}
        with entry.lock:
            conversation.state['achievement_edits'] = normalized
            conversation._save_session()
        session_registry.touch(sid)
        total = sum(len(v) for v in normalized.values())
        return jsonify({"success": True, "message": f"Saved edits for {len(normalized)} experiences ({total} achievements)"})

    @app.route('/api/rewrite-achievement', methods=['POST'])
    def rewrite_achievement():
        """Ask the LLM to rewrite a single achievement bullet."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        data = request.json or {}
        achievement_text = data.get('achievement_text', '').strip()
        experience_index = data.get('experience_index')
        user_instructions    = data.get('user_instructions', '').strip()
        previous_suggestions = data.get('previous_suggestions') or []
        if not isinstance(previous_suggestions, list):
            previous_suggestions = []

        if not achievement_text:
            return jsonify({"error": "achievement_text is required"}), 400

        # Gather context: experience title + job description
        experience_context = ''
        exp_idx = None
        if experience_index is not None:
            try:
                exp_idx = int(experience_index)
                _master, _ = _load_master(orchestrator.master_data_path)
                experiences = _master.get('experience', [])
                if 0 <= exp_idx < len(experiences):
                    exp = experiences[exp_idx]
                    title   = exp.get('title', exp.get('position', ''))
                    company = exp.get('company', exp.get('organization', ''))
                    experience_context = f"{title} at {company}".strip(' at')
            except (ValueError, TypeError, OSError):
                pass

        achievement_index = data.get('achievement_index')
        ach_idx = None
        if achievement_index is not None:
            try:
                ach_idx = int(achievement_index)
            except (ValueError, TypeError):
                ach_idx = None

        job_description = conversation.state.get('job_description') or ''

        try:
            rewritten = llm_client.rewrite_achievement(
                achievement_text=achievement_text,
                experience_context=experience_context,
                job_description=job_description,
                user_instructions=user_instructions,
                previous_suggestions=previous_suggestions,
            )
            log_id = conversation.log_achievement_rewrite(
                original_text=achievement_text,
                experience_context=experience_context,
                user_instructions=user_instructions,
                previous_suggestions=previous_suggestions,
                suggested_text=rewritten,
                experience_index=exp_idx,
                achievement_index=ach_idx,
            )
            return jsonify({"rewritten": rewritten, "log_id": log_id})
        except Exception as e:
            print(f"ERROR rewriting achievement: {e}")
            return jsonify({"error": str(e)}), 500

    @app.route('/api/rewrite-achievement-outcome', methods=['POST'])
    def rewrite_achievement_outcome():
        """Record the user's accept/reject decision for an AI rewrite suggestion."""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        data = request.json or {}
        log_id = data.get('log_id', '').strip()
        outcome = data.get('outcome', '').strip()
        accepted_text = data.get('accepted_text') or None

        if not log_id:
            return jsonify({"error": "log_id is required"}), 400
        if outcome not in ('accepted', 'rejected'):
            return jsonify({"error": "outcome must be 'accepted' or 'rejected'"}), 400

        found = conversation.update_achievement_rewrite_outcome(
            log_id=log_id,
            outcome=outcome,
            accepted_text=accepted_text,
        )
        if not found:
            return jsonify({"error": "log_id not found"}), 404
        return jsonify({"ok": True})

    @app.route('/api/cv-data', methods=['GET'])
    def get_cv_data():
        """Get current CV data for editing"""
        entry = _get_session()
        conversation = entry.manager
        orchestrator = entry.orchestrator
        try:
            # Get CV data from orchestrator's master data
            cv_data = {
                'personal_info': {},
                'summary': '',
                'experiences': [],
                'skills': []
            }

            if orchestrator and orchestrator.master_data:
                master_data = orchestrator.master_data
                
                # Get personal info
                personal_info = master_data.get('personal_info', {})
                cv_data['personal_info'] = {
                    'name': personal_info.get('name', ''),
                    'email': personal_info.get('email', ''),
                    'phone': personal_info.get('phone', ''),
                    'location': personal_info.get('location', '')
                }
                
                # Get summary
                cv_data['summary'] = master_data.get('summary', '')
                
                # Get experiences
                experiences = master_data.get('experience', [])
                cv_data['experiences'] = []
                for exp in experiences:
                    exp_data = {
                        'title': exp.get('title', ''),
                        'company': exp.get('company', ''),
                        'start_date': exp.get('start_date', ''),
                        'end_date': exp.get('end_date', ''),
                        'current': exp.get('current', False),
                        'location': exp.get('location', ''),
                        'achievements': exp.get('achievements', [])
                    }
                    cv_data['experiences'].append(exp_data)
                
                # Get skills
                skills_data = master_data.get('skills', [])
                all_skills = conversation.normalize_skills_data(skills_data)
                cv_data['skills'] = all_skills
            
            return jsonify(cv_data)
            
        except Exception as e:
            print(f"ERROR in get_cv_data: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.route('/api/cv-data', methods=['POST'])
    def save_cv_data():
        """Save edited CV data"""
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        try:
            data = request.json
            if not data:
                return jsonify({"error": "No data provided"}), 400

            # Store the edited CV data in the conversation state for now
            # In a full implementation, you'd want to update the master data file
            with entry.lock:
                conversation.state['edited_cv_data'] = data
                conversation._save_session()
            session_registry.touch(sid)
            
            # Log the changes
            print(f"CV data updated:")
            if 'personal_info' in data:
                print(f"  - Personal info: {data['personal_info']}")
            if 'summary' in data:
                print(f"  - Summary: {len(data.get('summary', ''))} chars")
            if 'experiences' in data:
                print(f"  - Experiences: {len(data['experiences'])} items")
            if 'skills' in data:
                print(f"  - Skills: {len(data['skills'])} items")
            
            return jsonify({"success": True, "message": "CV data saved successfully"})
            
        except Exception as e:
            print(f"ERROR in save_cv_data: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.get("/api/rewrites")
    def get_rewrites():
        """Propose LLM text rewrites aligned with the target job description.

        Calls ``orchestrator.propose_rewrites`` with the master CV data and the
        stored job analysis, stores the proposals in session state, and returns
        them for the frontend rewrite-review panel.

        Returns ``phase: 'generation'`` when no proposals are produced (either
        no LLM is configured or the LLM found nothing to rewrite) so the
        frontend can fall through gracefully.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            # Return cached rewrites on restore rather than re-running the LLM
            stored_rewrites = conversation.state.get('pending_rewrites')
            if stored_rewrites:
                return jsonify(dataclasses.asdict(RewritesResponse(
                    ok=True,
                    rewrites=stored_rewrites,
                    persuasion_warnings=conversation.state.get('persuasion_warnings', []),
                    phase=Phase.REWRITE_REVIEW,
                )))

            job_analysis = conversation.state.get('job_analysis')
            if not job_analysis:
                return jsonify({"error": "Job analysis not available. Analyse the job first."}), 400

            content = orchestrator.master_data
            if not content:
                return jsonify({"error": "CV master data not loaded."}), 400

            rewrites = orchestrator.propose_rewrites(
                content,
                job_analysis,
                conversation_history=conversation.conversation_history,
                user_preferences=conversation.state.get('post_analysis_answers'),
            )
            conversation.state['pending_rewrites'] = rewrites

            # Run persuasion quality checks (Phase 10)
            if rewrites:
                persuasion_warnings = conversation.run_persuasion_checks(
                    rewrites,
                    job_analysis,
                    orchestrator.master_data
                )
                conversation.state['persuasion_warnings'] = persuasion_warnings
            else:
                conversation.state['persuasion_warnings'] = []

            if rewrites:
                conversation.state['phase'] = Phase.REWRITE_REVIEW
                phase = Phase.REWRITE_REVIEW
            else:
                # No proposals (no LLM or nothing to rewrite) — skip review step
                phase = Phase.GENERATION

            conversation._save_session()
            session_registry.touch(sid)
            return jsonify(dataclasses.asdict(RewritesResponse(
                ok=True,
                rewrites=rewrites,
                persuasion_warnings=conversation.state.get('persuasion_warnings', []),
                phase=phase,
            )))

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.post("/api/rewrites/approve")
    def approve_rewrites():
        """Submit accept / edit / reject decisions for pending rewrite proposals.

        Request body::

            {"decisions": [{"id": str, "outcome": "accept"|"reject"|"edit",
                            "final_text": str|null}, ...]}

        Delegates to :meth:`ConversationManager.submit_rewrite_decisions` which
        builds ``approved_rewrites``, ``rewrite_audit``, advances the phase to
        ``'generation'``, and persists the session.
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        data = request.get_json(silent=True) or {}
        decisions = data.get('decisions')
        if decisions is None:
            return jsonify({"error": "Missing decisions"}), 400
        if not isinstance(decisions, list):
            return jsonify({"error": "decisions must be a list"}), 400

        try:
            with entry.lock:
                summary = conversation.submit_rewrite_decisions(decisions)
            session_registry.touch(sid)
            return jsonify({
                "ok":             True,
                "approved_count": summary['approved_count'],
                "rejected_count": summary['rejected_count'],
                "phase":          summary['phase'],
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.get("/api/publication-recommendations")
    def publication_recommendations():
        """Return LLM-ranked publication recommendations for the current job.

        Reads `session.publication_recommendations` if already computed, or
        computes them from `orchestrator.publications` + `session.job_analysis`.
        Computation runs at most once per session (cached in state).
        """
        entry = _get_session()
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sid = entry.session_id
        try:
            # Return cached recommendations if available.
            cached = conversation.state.get('publication_recommendations')
            if cached is not None:
                return jsonify({"ok": True, "recommendations": cached, "source": "cache",
                                "total_count": len(cached)})

            job_analysis = conversation.state.get('job_analysis')
            if not job_analysis:
                return jsonify({"ok": True, "recommendations": [], "source": "no_analysis"})

            if not orchestrator.publications:
                return jsonify({"ok": True, "recommendations": [], "source": "no_publications"})

            candidate_name = ''
            if orchestrator.master_data:
                candidate_name = orchestrator.master_data.get('personal_info', {}).get('name', '')

            # Convert bibtex_parser dict-of-dicts to list for the LLM ranker.
            pubs_list = list(orchestrator.publications.values())

            try:
                recommendations = llm_client.rank_publications_for_job(
                    publications=pubs_list,
                    job_analysis=job_analysis,
                    candidate_name=candidate_name,
                    max_results=15,
                )
                if not recommendations:
                    raise RuntimeError("LLM returned no publication recommendations")
                source = "llm"
            except Exception as rank_err:
                print(f"Publication ranking failed, using score-based fallback: {rank_err}")
                # Fallback: use the existing score-based _select_publications.
                selected = orchestrator._select_publications(job_analysis, max_count=15)
                recommendations = []
                for pub in selected:
                    recommendations.append({
                        'cite_key':          pub.get('key', ''),
                        'title':             pub.get('title', ''),
                        'venue':             pub.get('journal') or pub.get('booktitle') or '',
                        'year':              pub.get('year', ''),
                        'is_first_author':   False,
                        'relevance_score':   pub.get('relevance_score', 5),
                        'confidence':        'Medium',
                        'rationale':         '',
                        'authority_signals': [],
                        'venue_warning':     '' if (pub.get('journal') or pub.get('booktitle')) else 'No venue found',
                        'formatted_citation': pub.get('formatted', ''),
                    })
                source = "fallback"

            # Mark LLM-recommended publications and add any remaining pubs as not recommended.
            recommended_keys = {r['cite_key'] for r in recommendations}
            for r in recommendations:
                r['is_recommended'] = True

            if orchestrator.publications:
                try:
                    from utils.bibtex_parser import format_publication as _fmt_pub
                except ImportError:
                    _fmt_pub = None
                not_recommended = []
                for key, pub in orchestrator.publications.items():
                    if key in recommended_keys:
                        continue
                    if _fmt_pub:
                        try:
                            formatted = _fmt_pub(pub, style='apa')
                        except Exception:
                            formatted = ''
                    else:
                        formatted = ''
                    if not formatted:
                        formatted = f"{pub.get('authors', '')} ({pub.get('year', '')}). {pub.get('title', '')}".strip('. ')
                    not_recommended.append({
                        'cite_key':          key,
                        'title':             pub.get('title', ''),
                        'venue':             pub.get('journal') or pub.get('booktitle') or '',
                        'year':              pub.get('year', ''),
                        'is_first_author':   False,
                        'relevance_score':   0,
                        'confidence':        '',
                        'rationale':         '',
                        'authority_signals': [],
                        'venue_warning':     '' if (pub.get('journal') or pub.get('booktitle')) else 'No venue found',
                        'formatted_citation': formatted,
                        'is_recommended':    False,
                    })
                not_recommended.sort(key=lambda p: -int(str(p['year']).strip() or '0'))
                recommendations.extend(not_recommended)

            conversation.state['publication_recommendations'] = recommendations
            conversation._save_session()
            session_registry.touch(sid)

            total_count = len(recommendations)
            return jsonify({"ok": True, "recommendations": recommendations, "source": source, "total_count": total_count})

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.get("/api/download/<filename>")
    def download_file(filename):
        """Download generated CV files"""
        entry = _get_session()
        conversation = entry.manager
        try:
            # Get generated files from conversation state
            generated_files = conversation.state.get('generated_files', {})
            
            # Find the requested file
            file_path = None
            
            # Check if generated_files is the dictionary structure returned by orchestrator
            if isinstance(generated_files, dict) and 'files' in generated_files:
                # This is the structure: {'output_dir': 'path', 'files': ['file1', 'file2'], 'metadata': {}}
                output_dir = Path(generated_files['output_dir'])
                for file_name in generated_files['files']:
                    if file_name == filename:
                        file_path = output_dir / filename
                        break
            else:
                # Legacy structure or other format - search in different ways
                for file_type, file_data in generated_files.items():
                    if isinstance(file_data, dict):
                        # File data is a dict with path info
                        check_filename = file_data.get('filename') if hasattr(file_data, 'get') else None
                        if check_filename == filename:
                            file_path = Path(file_data.get('path', file_data))
                            break
                    elif isinstance(file_data, (str, Path)):
                        # File data is a direct path
                        if Path(file_data).name == filename:
                            file_path = Path(file_data)
                            break
            
            if not file_path or not file_path.exists():
                return jsonify({"error": "File not found on disk"}), 404
            
            # Determine MIME type
            mime_type = 'application/octet-stream'
            if filename.endswith('.pdf'):
                mime_type = 'application/pdf'
            elif filename.endswith('.docx'):
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filename.endswith('.html'):
                mime_type = 'text/html'
            
            return send_file(
                str(file_path),
                as_attachment=True,
                download_name=filename,
                mimetype=mime_type
            )
            
        except Exception as e:
            print(f"ERROR in download_file: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    # ------------------------------------------------------------------ #
    # Spell / Grammar Check endpoints  (Phase 6)                          #
    # ------------------------------------------------------------------ #

    # Lazy singleton — the LanguageTool JVM starts on first call.
    _spell_checker: SpellChecker = SpellChecker()

    def _prepopulate_spell_dict(orchestrator_inst) -> None:
        """Load domain terms and proper nouns from master data into the custom dictionary."""
        try:
            master = orchestrator_inst.master_data or {}
            skills = master.get('skills', {})
            all_names: list = []

            def _collect_names(values) -> None:
                for value in values or []:
                    if isinstance(value, dict):
                        name = (
                            value.get('name', '')
                            or value.get('title', '')
                            or value.get('degree', '')
                            or value.get('institution', '')
                            or value.get('company', '')
                            or value.get('issuer', '')
                            or value.get('language', '')
                            or value.get('proficiency', '')
                        )
                        if name:
                            all_names.append(str(name))
                    elif value:
                        all_names.append(str(value))

            if isinstance(skills, dict):
                for cat_skills in skills.values():
                    if isinstance(cat_skills, list):
                        _collect_names(cat_skills)
            elif isinstance(skills, list):
                _collect_names(skills)

            pinfo = master.get('personal_info', {})
            all_names.extend(filter(None, [
                pinfo.get('name', ''),
                pinfo.get('title', ''),
            ]))

            for exp in master.get('experience', []) or []:
                if isinstance(exp, dict):
                    all_names.extend(filter(None, [
                        exp.get('company', ''),
                        exp.get('title', ''),
                    ]))

            for edu in master.get('education', []) or []:
                if isinstance(edu, dict):
                    all_names.extend(filter(None, [
                        edu.get('institution', ''),
                        edu.get('degree', ''),
                        edu.get('field', ''),
                    ]))

            for award in master.get('awards', []) or []:
                if isinstance(award, dict):
                    all_names.extend(filter(None, [
                        award.get('degree', ''),
                        award.get('title', ''),
                    ]))

            for cert in master.get('certifications', []) or []:
                if isinstance(cert, dict):
                    all_names.extend(filter(None, [
                        cert.get('name', ''),
                        cert.get('issuer', ''),
                    ]))

            for lang in pinfo.get('languages', []) or []:
                if isinstance(lang, dict):
                    all_names.extend(filter(None, [
                        lang.get('language', ''),
                        lang.get('proficiency', ''),
                    ]))
                else:
                    all_names.append(str(lang))

            _spell_checker.prepopulate_from_skills(all_names)
        except Exception:
            pass

    @app.get("/api/spell-check-sections")
    def spell_check_sections():
        """Return the text sections that need spell checking for the current session.

        Covers the professional summary and every non-omitted achievement bullet
        from the selected work experiences, with approved rewrites already applied
        so the spell-checked text matches what the CV will actually render.
        """
        entry = _get_session()
        conversation = entry.manager
        orchestrator = entry.orchestrator
        sections = []

        def _append_section(section_id: str, label: str, text: str, context: str = 'skill') -> None:
            text = (text or '').strip()
            if not text:
                return
            sections.append({
                'id':      section_id,
                'label':   label,
                'text':    text,
                'context': context,
            })
        try:
            state             = conversation.state
            job_analysis      = state.get('job_analysis') or {}
            customizations    = state.get('customizations') or {}
            approved_rewrites = state.get('approved_rewrites') or []
            spell_audit       = state.get('spell_audit') or []

            _prepopulate_spell_dict(orchestrator)
            selected_content = orchestrator.build_render_ready_content(
                job_analysis,
                customizations,
                approved_rewrites=approved_rewrites,
                spell_audit=spell_audit,
                max_skills=state.get('max_skills'),
                use_semantic_match=False,
            )

            cv_data = orchestrator._prepare_cv_data_for_template(selected_content, job_analysis)

            _append_section('summary', 'Professional Summary', cv_data.get('professional_summary', ''), 'summary')

            for i, ach in enumerate(selected_content.get('achievements', []) or []):
                text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                _append_section(f'selected_ach_{i}', f'Selected Achievement {i + 1}', text, 'bullet')

            for idx, skill in enumerate(selected_content.get('skills', []) or []):
                if isinstance(skill, dict):
                    rendered_skill = skill.get('name', '')
                    if skill.get('years'):
                        rendered_skill += f" ({skill['years']} yrs)"
                else:
                    rendered_skill = str(skill)
                _append_section(f'skill_{idx}', f'Skill {idx + 1}', rendered_skill, 'skill')

            for idx, edu in enumerate(selected_content.get('education', []) or []):
                if not isinstance(edu, dict):
                    continue
                _append_section(f'edu_{idx}_degree', f'Education {idx + 1} — Degree', edu.get('degree', ''), 'skill')
                _append_section(f'edu_{idx}_field', f'Education {idx + 1} — Field', edu.get('field', ''), 'skill')
                _append_section(f'edu_{idx}_institution', f'Education {idx + 1} — Institution', edu.get('institution', ''), 'skill')

            for idx, award in enumerate(selected_content.get('awards', []) or []):
                if not isinstance(award, dict):
                    continue
                _append_section(
                    f'award_{idx}_title',
                    f'Award {idx + 1}',
                    award.get('degree', '') or award.get('title', ''),
                    'skill',
                )

            for idx, cert in enumerate(selected_content.get('certifications', []) or []):
                if not isinstance(cert, dict):
                    continue
                _append_section(f'cert_{idx}_name', f'Certification {idx + 1} — Name', cert.get('name', ''), 'skill')
                _append_section(f'cert_{idx}_issuer', f'Certification {idx + 1} — Issuer', cert.get('issuer', ''), 'skill')

            for idx, lang in enumerate(selected_content.get('personal_info', {}).get('languages', []) or []):
                if isinstance(lang, dict):
                    _append_section(f'lang_{idx}_language', f'Language {idx + 1}', lang.get('language', ''), 'skill')
                    _append_section(f'lang_{idx}_proficiency', f'Language {idx + 1} — Proficiency', lang.get('proficiency', ''), 'skill')
                else:
                    _append_section(f'lang_{idx}', f'Language {idx + 1}', str(lang), 'skill')

            for idx, pub in enumerate(selected_content.get('publications', []) or []):
                if not isinstance(pub, dict):
                    continue
                if pub.get('formatted'):
                    _append_section(f'pub_{idx}_formatted', f'Publication {idx + 1}', pub.get('formatted', ''), 'skill')
                else:
                    _append_section(f'pub_{idx}_title', f'Publication {idx + 1} — Title', pub.get('title', ''), 'skill')
                    _append_section(f'pub_{idx}_authors', f'Publication {idx + 1} — Authors', pub.get('authors', ''), 'skill')
                    _append_section(f'pub_{idx}_journal', f'Publication {idx + 1} — Venue', pub.get('journal', '') or pub.get('booktitle', ''), 'skill')

            # Experience bullets from every selected experience (rewrites already applied)
            for exp in cv_data.get('experiences', []) or []:
                exp_id  = exp.get('id', '')
                company = exp.get('company', 'Experience')
                role    = exp.get('title', '')
                label   = f"{company} \u2014 {role}" if role else company

                ach_list = exp.get('ordered_achievements') or exp.get('achievements') or []
                for i, ach in enumerate(ach_list):
                    text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                    _append_section(f"exp_{exp_id}_ach_{i}", f"{label} (bullet {i + 1})", text, 'bullet')

        except Exception as e:
            return jsonify({'error': str(e)}), 500

        aggregate_stats = _spell_checker.aggregate_stats([s['text'] for s in sections])
        return jsonify({
            'ok':               True,
            'sections':         sections,
            'aggregate_stats':  aggregate_stats,
            'custom_dict_size': len(_spell_checker.get_custom_dict()),
        })

    @app.post("/api/spell-check")
    def spell_check_text():
        """Check a single text fragment.

        Body: ``{"text": "...", "context": "bullet"|"summary"|"skill"}``
        Returns: ``{"ok": true, "suggestions": [...]}``
        """
        entry = _get_session()
        orchestrator = entry.orchestrator
        try:
            body    = request.get_json(force=True) or {}
            text    = body.get('text', '')
            context = body.get('context', 'bullet')
            if context not in ('bullet', 'summary', 'skill'):
                context = 'bullet'

            _prepopulate_spell_dict(orchestrator)
            result = _spell_checker.check(text, context=context)
            custom_dict_size = len(_spell_checker.get_custom_dict())
            return jsonify({
                'ok':              True,
                'suggestions':     result['suggestions'],
                'stats':           result['stats'],
                'custom_dict_size': custom_dict_size,
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.get("/api/custom-dictionary")
    def custom_dictionary_get():
        """Return the current custom dictionary word list."""
        try:
            words = _spell_checker.get_custom_dict()
            return jsonify({'ok': True, 'words': words})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.post("/api/custom-dictionary")
    def custom_dictionary_add():
        """Add a word to the custom dictionary.

        Body: ``{"word": "MyTechTerm"}``
        Returns: ``{"ok": true, "added": true|false}``
        """
        try:
            body  = request.get_json(force=True) or {}
            word  = body.get('word', '').strip()
            if not word:
                return jsonify({'error': 'word is required'}), 400
            added = _spell_checker.add_word(word)
            return jsonify({'ok': True, 'added': added})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.post("/api/spell-check-complete")
    def spell_check_complete():
        """Record spell-check audit and advance phase to generation.

        Body: ``{"spell_audit": [...]}``
        Each audit entry: ``{context_type, location, original, suggestion, rule, outcome, final}``
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        try:
            body        = request.get_json(force=True) or {}
            spell_audit = body.get('spell_audit', [])
            with entry.lock:
                result = conversation.complete_spell_check(spell_audit)
            session_registry.touch(sid)
            return jsonify({'ok': True, **result})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ------------------------------------------------------------------ #
    # Layout Instructions (Phase 12)                                      #
    # ------------------------------------------------------------------ #

    @app.get("/api/layout-html")
    def get_layout_html():
        """Return the HTML content of the most recently generated CV.

        Reads the ``.html`` file from the session's generated output directory.
        Used by the layout-preview panel on first load when tabData.cv has no
        inline HTML.

        Returns: ``{"ok": true, "html": "..."}`` or ``{"error": "..."}``
        """
        entry = _get_session()
        conversation = entry.manager
        try:
            generated = conversation.state.get('generated_files')
            if not generated or not isinstance(generated, dict):
                return jsonify({'error': 'No generated CV found — generate your CV first.'}), 404
            output_dir = Path(generated.get('output_dir', ''))
            if not output_dir.is_dir():
                return jsonify({'error': f'Output directory not found: {output_dir}'}), 404
            html_files = sorted(output_dir.glob('*.html'))
            if not html_files:
                return jsonify({'error': 'No HTML file found in output directory.'}), 404
            html_content = html_files[0].read_text(encoding='utf-8', errors='replace')
            return jsonify({'ok': True, 'html': html_content})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.post("/api/layout-instruction")
    def apply_layout_instruction():
        """Apply a natural-language layout instruction to the current HTML.

        Body: ``{"instruction": "Move Publications after Skills", "current_html": "..."}``
            Optional: ``"prior_instructions": [...]``

        Returns:
            Success: ``{"ok": true, "html": "...", "summary": "...", "confidence": 0.95}``
            Clarification: ``{"ok": false, "error": "clarify", "question": "..."}``
            Error: ``{"ok": false, "error": "error_type", "details": "..."}``
        """
        entry = _get_session()
        conversation = entry.manager
        try:
            body = request.get_json(force=True) or {}
            instruction_text = body.get('instruction', '').strip()
            current_html = body.get('current_html', '')
            prior_instructions = body.get('prior_instructions', [])

            if not instruction_text:
                return jsonify({'error': 'Missing instruction text'}), 400
            if not current_html:
                return jsonify({'error': 'Missing current HTML'}), 400

            # Call orchestrator to apply instruction
            result = conversation.orchestrator.apply_layout_instruction(
                instruction_text=instruction_text,
                current_html=current_html,
                prior_instructions=prior_instructions
            )

            if result.get('error'):
                return jsonify({
                    'ok':           False,
                    'error':        result['error'],
                    'question':     result.get('question'),
                    'details':      result.get('details'),
                    'confidence':   result.get('confidence'),
                    'raw_response': result.get('raw_response'),
                })

            return jsonify({
                'ok': True,
                'html': result['html'],
                'summary': result['summary'],
                'confidence': result['confidence']
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.get("/api/layout-history")
    def get_layout_history():
        """Return the current session's applied layout instruction history.

        Returns:
            ``{"instructions": [...], "count": int}``
        """
        entry = _get_session()
        conversation = entry.manager
        try:
            instructions = conversation.state.get('layout_instructions')
            if not instructions:
                instructions = (
                    conversation.state.get('generation_state', {}).get('layout_instructions', [])
                )
            return jsonify({
                'instructions': instructions,
                'count': len(instructions)
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.post("/api/layout-complete")
    def complete_layout_review():
        """Record layout instruction outcomes and advance phase to refinement (finalise).

        Body: ``{"layout_instructions": [...]}``
        Each instruction should have: ``{timestamp, instruction_text, change_summary, confirmation}``

        Returns:
            ``{"ok": true, "instructions_applied": int, "phase": "refinement"}``
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        try:
            body = request.get_json(force=True) or {}
            layout_instructions = body.get('layout_instructions', [])
            if not layout_instructions:
                layout_instructions = (
                    conversation.state.get('layout_instructions')
                    or conversation.state.get('generation_state', {}).get('layout_instructions', [])
                )
            with entry.lock:
                result = conversation.complete_layout_review(layout_instructions)
            session_registry.touch(sid)
            return jsonify({'ok': True, **result})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.post("/api/layout-settings")
    def update_layout_settings():
        """Persist layout display settings to session state.

        Supported keys:
          ``base_font_size`` — CSS font size string, e.g. ``"10px"`` or ``"11px"``.

        Returns:
            ``{"ok": true}``
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid          = entry.session_id
        try:
            body = request.get_json(force=True) or {}
            with entry.lock:
                if 'base_font_size' in body:
                    raw = str(body['base_font_size']).strip()
                    # Accept bare numbers ("10") or values with unit ("10px")
                    if not raw.endswith('px'):
                        raw = raw + 'px'
                    conversation.state['base_font_size'] = raw
                    # Keep customizations dict in sync if it already exists
                    if 'customizations' in conversation.state:
                        conversation.state['customizations']['base_font_size'] = raw
                conversation._save_session()
            session_registry.touch(sid)
            return jsonify({'ok': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ------------------------------------------------------------------ #
    # ATS Validation + Page Count  (Phase 7)                              #
    # ------------------------------------------------------------------ #

    @app.get("/api/ats-validate")
    def ats_validate():
        """Run 16-check ATS validation on the latest generated CV files.

        Returns:
            ``{"ok": true, "checks": [...], "page_count": int|null,
               "summary": {"pass": N, "warn": N, "fail": N}}``
        """
        entry = _get_session()
        conversation = entry.manager
        sid = entry.session_id
        try:
            generated = conversation.state.get('generated_files')
            if not generated or not isinstance(generated, dict):
                return jsonify({'ok': False, 'error': 'No CV files generated yet'}), 400

            output_dir  = Path(generated.get('output_dir', ''))
            if not output_dir.is_dir():
                return jsonify({'ok': False, 'error': f'Output directory not found: {output_dir}'}), 404

            job_analysis = _coerce_to_dict(conversation.state.get('job_analysis'))

            checks, page_count = validate_ats_report(output_dir, job_analysis)

            # Cache page_count in session state
            if page_count is not None:
                conversation.state['page_count'] = page_count

            summary = {
                'pass': sum(1 for c in checks if c['status'] == 'pass'),
                'warn': sum(1 for c in checks if c['status'] == 'warn'),
                'fail': sum(1 for c in checks if c['status'] == 'fail'),
            }

            # Cache validation results for finalise (includes page_count and all checks)
            conversation.state['validation_results'] = {
                'page_count': page_count,
                'checks': checks,
                'summary': summary,
                'validation_date': datetime.now().isoformat(),
            }
            session_registry.touch(sid)

            return jsonify({
                'ok':         True,
                'checks':     checks,
                'page_count': page_count,
                'summary':    summary,
            })
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': str(e)}), 500

    @app.get("/api/persuasion-check")
    def persuasion_check():
        """Run rule-based persuasion checks on selected experience bullets.

        Returns
        -------
        ``{"ok": true, "findings": [...], "summary": {total_bullets, flagged,
           strong_count}}``

        Each finding has: ``exp_id``, ``bullet_index``, ``text``, ``severity``,
        ``issues`` (list of ``{type, severity, suggestion}``).
        """
        entry = _get_session()
        conversation = entry.manager
        try:
            experiences = None

            # Use selected_content stored during generation if available
            generated = conversation.state.get('generated_files')
            if generated:
                # Re-derive selected content from current state
                job_analysis   = _coerce_to_dict(conversation.state.get('job_analysis'))
                customizations = conversation.state.get('customizations') or {}
                try:
                    selected = conversation.orchestrator._select_content_hybrid(
                        job_analysis, customizations
                    )
                    experiences = selected.get('experiences', [])
                except Exception:
                    experiences = None

            # Fallback: analyse raw master data experiences
            if experiences is None:
                experiences = (
                    conversation.orchestrator.master_data.get('experience')
                    or conversation.orchestrator.master_data.get('experiences')
                    or []
                )

            result = conversation.orchestrator.check_persuasion(experiences)
            return jsonify({'ok': True, **result})
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({'error': str(e)}), 500

    # ── Phase 11: Finalise & Archive ────────────────────────────────────────

    # ── Staged Generation Routes (GAP-20 Phase 1) ─────────────────────────────
    # Contract: tasks/contracts/phase0-contract.md

    @app.get("/api/cv/generation-state")
    def get_generation_state():
        """Return staged generation phase and metadata (no raw HTML)."""
        entry = _get_session()
        gen   = entry.manager.state.get("generation_state") or {}
        return jsonify({
            "ok":                        True,
            "phase":                     gen.get("phase", "idle"),
            "preview_available":         bool(gen.get("preview_html")),
            "layout_confirmed":          gen.get("layout_confirmed", False),
            "page_count_estimate":       gen.get("page_count_estimate"),
            "page_length_warning":       gen.get("page_length_warning", False),
            "layout_instructions_count": len(gen.get("layout_instructions", [])),
            "ats_score":                 gen.get("ats_score"),
            "final_generated_at":        gen.get("final_generated_at"),
        })

    @app.post("/api/cv/generate-preview")
    def generate_cv_preview():
        """Generate an HTML preview of the CV and store it in generation_state.

        Renders the HTML template from current CV state (job_analysis +
        customizations + approved rewrites + spell audit).  Falls back to the
        most recent HTML file on disk when the customisation data is incomplete.

        Returns the rendered HTML so the frontend can display it immediately
        without a second round-trip to /api/layout-html.
        """
        import uuid as _u
        entry = _get_session()
        conv  = entry.manager
        if not conv.state.get("job_analysis"):
            return jsonify({"error": "Run job analysis first."}), 400

        html_str = None

        # ── Try to render fresh HTML from current CV state ──────────────────
        customizations = conv.state.get("customizations")
        if customizations:
            try:
                approved_rewrites = conv.state.get("approved_rewrites") or []
                spell_audit       = conv.state.get("spell_audit")
                if spell_audit is None:
                    legacy_spell = conv.state.get("spell_check") or {}
                    spell_audit = legacy_spell.get("audit") or [] if isinstance(legacy_spell, dict) else []
                html_str = conv.orchestrator.render_html_preview(
                    job_analysis=conv.state["job_analysis"],
                    customizations=customizations,
                    approved_rewrites=approved_rewrites,
                    spell_audit=spell_audit,
                )
            except Exception as _exc:
                app.logger.warning("render_html_preview failed: %s", _exc)

        # ── Fallback: load most recent HTML file from output directory ───────
        if not html_str:
            generated      = conv.state.get("generated_files") or {}
            output_dir_str = generated.get("output_dir", "")
            if output_dir_str:
                output_dir = Path(output_dir_str)
                if output_dir.is_dir():
                    for p in sorted(output_dir.glob("*.html")):
                        html_str = p.read_text(encoding="utf-8")

        if not html_str:
            return jsonify({"error": "No CV content available — complete customisation first."}), 404

        now     = datetime.now().isoformat()
        prev_id = str(_u.uuid4())
        gen = conv.state.setdefault("generation_state", {})
        gen.update({
            "phase":                "layout_review",
            "preview_html":         html_str,
            "preview_request_id":   prev_id,
            "preview_generated_at": now,
            "layout_confirmed":     False,
        })
        if "layout_instructions" not in gen:
            gen["layout_instructions"] = []
        conv._save_session()
        return jsonify({
            "ok":                  True,
            "html":                html_str,
            "preview_request_id":  prev_id,
            "page_count_estimate": gen.get("page_count_estimate"),
            "page_length_warning": gen.get("page_length_warning", False),
        })

    @app.post("/api/cv/layout-refine")
    def refine_cv_layout():
        """Apply a layout instruction to the stored preview and return updated HTML.

        Bridges the existing ``apply_layout_instruction`` orchestrator method
        into the staged generation contract.  The frontend does not need to
        pass the current HTML — it is read from and written back to
        ``generation_state.preview_html`` in the session.

        Body: ``{"instruction": str, "session_id": str}``

        Returns:
            ``{"ok": true, "html": str, "summary": str, "confidence": float,
               "preview_request_id": str}``
        """
        import uuid as _u
        entry = _get_session()
        conv  = entry.manager
        gen   = conv.state.get("generation_state") or {}

        phase = gen.get("phase", "idle")
        if phase not in ("preview", "layout_review"):
            return jsonify({
                "error": "Call /api/cv/generate-preview first before refining layout."
            }), 400

        body = request.get_json(force=True) or {}
        instruction_text = (body.get("instruction") or "").strip()
        if not instruction_text:
            return jsonify({"error": "Missing instruction text."}), 400

        current_html = gen.get("preview_html", "")
        if not current_html:
            return jsonify({"error": "No preview HTML in session — call generate-preview first."}), 400

        prior_instructions = gen.get("layout_instructions", [])

        result = conv.orchestrator.apply_layout_instruction(
            instruction_text=instruction_text,
            current_html=current_html,
            prior_instructions=prior_instructions,
        )

        if result.get("error"):
            return jsonify({
                "ok":           False,
                "error":        result["error"],
                "question":     result.get("question"),
                "details":      result.get("details"),
                "raw_response": result.get("raw_response"),
            })

        updated_html = result["html"]
        now     = datetime.now().isoformat()
        prev_id = str(_u.uuid4())

        instruction_record = {
            "id":           prev_id,
            "text":         instruction_text,
            "submitted_at": now,
            "applied":      True,
            "summary":      result.get("summary", ""),
            "confidence":   result.get("confidence"),
        }

        gen = conv.state.setdefault("generation_state", {})
        gen["preview_html"]        = updated_html
        gen["preview_request_id"]  = prev_id
        gen["preview_generated_at"] = now
        gen["phase"]               = "layout_review"
        gen["layout_confirmed"]    = False
        gen.setdefault("layout_instructions", []).append(instruction_record)
        conv._save_session()

        return jsonify({
            "ok":                 True,
            "html":               updated_html,
            "summary":            result.get("summary", ""),
            "confidence":         result.get("confidence"),
            "preview_request_id": prev_id,
        })

    @app.post("/api/cv/confirm-layout")
    def confirm_cv_layout():
        """Lock current preview; enables /api/cv/generate-final."""
        entry = _get_session()
        conv  = entry.manager
        gen   = conv.state.get("generation_state") or {}
        if not gen.get("preview_html"):
            return jsonify({"error": "No preview — call /api/cv/generate-preview first."}), 400
        if gen.get("layout_confirmed") or gen.get("phase") == "confirmed":
            return jsonify({"error": "Layout already confirmed."}), 400
        import hashlib as _hl
        now   = datetime.now().isoformat()
        chash = _hl.sha256(gen["preview_html"].encode()).hexdigest()[:16]
        gen   = conv.state.setdefault("generation_state", {})
        gen.update({
            "phase": "confirmed", "layout_confirmed": True,
            "confirmed_at": now, "confirmed_preview_hash": chash,
        })
        conv._save_session()
        return jsonify({"ok": True, "confirmed": True, "confirmed_at": now, "hash": chash})

    @app.post("/api/cv/ats-score")
    def compute_cv_ats_score():
        """Return ATS match score for current session state (GAP-21).

        Computes the score from job_analysis and customizations stored in
        the session.  Optionally re-persists the score in generation_state.

        Request body:
            { "session_id": str, "basis": str | None }

        Response:
            { "ok": True, "ats_score": { ...Phase-0 contract schema... } }
        """
        from utils.scoring import compute_ats_score as _compute_ats_score
        entry = _get_session()
        conv  = entry.manager
        job_analysis   = conv.state.get("job_analysis") or {}
        customizations = dict(conv.state.get("customizations") or {})
        body  = request.get_json(silent=True) or {}
        basis = body.get("basis", "review_checkpoint")

        # Enrich customizations with session-level decision fields so the ATS
        # scorer sees the actual content the user has approved/edited.
        #
        # skill_decisions: {skill_name: 'keep'|'exclude'} → approved_skills list
        skill_decisions = conv.state.get("skill_decisions") or {}
        extra_skills    = conv.state.get("extra_skills") or []
        kept_skills = [k for k, v in skill_decisions.items() if v != "exclude"]
        kept_skills += [s for s in extra_skills if s not in kept_skills]
        if kept_skills:
            existing = [
                (s.get("name") if isinstance(s, dict) else s)
                for s in customizations.get("approved_skills", [])
            ]
            customizations["approved_skills"] = list(
                customizations.get("approved_skills", [])
            ) + [s for s in kept_skills if s not in existing]

        # approved_rewrites stored at top-level state (not nested in customizations)
        if not customizations.get("approved_rewrites"):
            state_rewrites = conv.state.get("approved_rewrites") or []
            if state_rewrites:
                customizations["approved_rewrites"] = state_rewrites

        # achievement_edits: {expIdx: [bullet, ...]} — fold into experience text
        achievement_edits = conv.state.get("achievement_edits") or {}
        if achievement_edits and not customizations.get("approved_rewrites"):
            bullet_rewrites = []
            for bullets in achievement_edits.values():
                if isinstance(bullets, list):
                    bullet_rewrites.extend(
                        {"rewritten": b, "section": "experience"}
                        for b in bullets if isinstance(b, str) and b.strip()
                    )
            if bullet_rewrites:
                customizations.setdefault("approved_rewrites", [])
                customizations["approved_rewrites"] = (
                    customizations["approved_rewrites"] + bullet_rewrites
                )

        # session_summaries override master professional_summaries
        session_summaries = conv.state.get("session_summaries") or {}
        if session_summaries and not customizations.get("selected_summary"):
            focus = conv.state.get("summary_focus_override", "ai_recommended")
            chosen = session_summaries.get(focus) or next(iter(session_summaries.values()), "")
            if chosen:
                customizations["selected_summary"] = chosen

        score = _compute_ats_score(job_analysis, customizations, basis=basis)
        # Persist latest score into generation_state for client polling
        gen = conv.state.setdefault("generation_state", {})
        gen["ats_score"] = score
        conv._save_session()
        return jsonify({"ok": True, "ats_score": score})

    @app.post("/api/cv/generate-final")
    def generate_cv_final():
        """Regenerate human-readable HTML+PDF from the confirmed preview; mark final_complete.

        Requires layout_confirmed == true and an existing output_dir (from the
        earlier generate_cv run).  Takes the confirmed HTML out of generation_state,
        writes it to the output directory, and reconverts to PDF so that any layout
        instructions applied during the preview loop are reflected in the final files.

        ATS DOCX is derived from structured data (not HTML) and is left unchanged.
        """
        entry = _get_session()
        conv  = entry.manager
        gen   = conv.state.get("generation_state") or {}
        if not gen.get("layout_confirmed"):
            return jsonify({"error": "Confirm layout first via /api/cv/confirm-layout."}), 400

        confirmed_html = gen.get("preview_html")
        if not confirmed_html:
            return jsonify({"error": "No confirmed preview HTML in session."}), 400

        generated = conv.state.get("generated_files") or {}
        if not generated.get("output_dir"):
            return jsonify({"error": "No generated files — complete workflow first."}), 404

        output_dir = Path(generated["output_dir"])

        # Build descriptive filename: CV_{company}_{position}_{date}
        _analysis = _coerce_to_dict(conv.state.get("job_analysis"))
        _company  = re.sub(r'[^\w]', '_', (_analysis.get('company') or '').strip())[:30]
        _position = re.sub(r'[^\w]', '_', (conv.state.get('position_name') or
                                            _analysis.get('job_title') or '').strip())[:40]
        _date     = datetime.now().strftime('%Y-%m-%d')
        _parts    = [p for p in ['CV', _company, _position, _date] if p and p != '_']
        filename_base = '_'.join(_parts) if len(_parts) > 2 else 'CV_final'

        try:
            final_paths = conv.orchestrator.generate_final_from_confirmed_html(
                confirmed_html=confirmed_html,
                output_dir=output_dir,
                filename_base=filename_base,
            )
        except Exception as exc:
            app.logger.error("generate_final_from_confirmed_html failed: %s", exc)
            return jsonify({"error": f"Final generation failed: {exc}"}), 500

        now = datetime.now().isoformat()
        gen = conv.state.setdefault("generation_state", {})
        gen.update({
            "phase": "final_complete",
            "final_generated_at": now,
            "final_output_paths": final_paths,
        })
        generated.update({
            "final_html": final_paths["html"],
            "final_pdf": final_paths["pdf"],
            "files": [
                final_paths["html"],
                final_paths["pdf"],
            ],
        })
        conv._save_session()

        outputs = dict(generated)
        return jsonify({"ok": True, "generated_at": now, "outputs": outputs})

    @app.post("/api/finalise")
    def finalise_application():
        """Finalise the application: update metadata, upsert response library, git commit.

        Request body (all optional):
            status  — "draft" | "ready" | "sent"  (default "ready")
            notes   — free-text notes string

        Returns:
            { ok, commit_hash, summary } on success
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        with entry.lock:
            generated = conversation.state.get('generated_files')
            if not generated or not generated.get('output_dir'):
                return jsonify({'error': 'No generated CV to finalise. Please generate first.'}), 400

            try:
                body        = request.get_json(silent=True) or {}
                app_status  = body.get('status', 'ready')
                notes       = body.get('notes', '')

                if app_status not in ('draft', 'ready', 'sent'):
                    return jsonify({'error': "status must be 'draft', 'ready', or 'sent'"}), 400

                output_dir   = Path(generated['output_dir'])
                metadata_path = output_dir / 'metadata.json'

                # Load existing metadata (created during generate_cv)
                if metadata_path.exists():
                    with open(metadata_path, encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {}

                # Augment metadata with finalise data
                metadata['application_status'] = app_status
                metadata['notes']              = notes
                metadata['finalised_at']       = datetime.now().isoformat()
                metadata['clarification_answers'] = conversation.state.get('post_analysis_answers') or {}
                metadata['spell_audit']           = conversation.state.get('spell_audit') or []
                metadata['layout_instructions']   = conversation.state.get('layout_instructions') or []
                metadata['validation_results']    = conversation.state.get('validation_results') or {}
                ats_score = ((conversation.state.get('generation_state') or {}).get('ats_score'))
                if ats_score is not None:
                    metadata['ats_score'] = ats_score

                # Upsert screening responses into response_library.json (if any)
                screening = metadata.get('screening_responses') or []
                if screening:
                    library_path = Path(conversation.orchestrator.master_data_path).parent / 'response_library.json'
                    if library_path.exists():
                        with open(library_path, encoding='utf-8') as f:
                            library = json.load(f)
                    else:
                        library = {}
                    for resp in screening:
                        tag = resp.get('topic_tag') or resp.get('question', '')[:40]
                        if tag:
                            library[tag] = resp
                    library_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(library_path, 'w', encoding='utf-8') as f:
                        json.dump(library, f, indent=2)

                # Write updated metadata
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)

                # Git commit
                company  = (metadata.get('company') or 'Unknown').replace(' ', '_')
                role     = (metadata.get('role') or 'Role').replace(' ', '_')
                date_str = datetime.now().strftime('%Y-%m-%d')
                commit_msg = f"feat: Add {company}_{role}_{date_str} application"

                commit_hash = None
                git_error   = None
                try:
                    repo_root = Path(__file__).parent.parent
                    subprocess.run(
                        ['git', 'add', str(output_dir)],
                        cwd=str(repo_root), check=True, capture_output=True
                    )
                    result = subprocess.run(
                        ['git', 'commit', '-m', commit_msg],
                        cwd=str(repo_root), capture_output=True, text=True
                    )
                    if result.returncode == 0:
                        # Extract short hash from the commit output
                        m = re.search(r'\b([0-9a-f]{7,40})\b', result.stdout)
                        commit_hash = m.group(1) if m else None
                    else:
                        git_error = result.stderr.strip() or result.stdout.strip()
                except Exception as git_exc:
                    git_error = str(git_exc)

                # Advance phase
                conversation.state['phase'] = Phase.REFINEMENT
                conversation.save_session()
                session_registry.touch(sid)

                # Build keyword match summary
                job_analysis   = conversation.state.get('job_analysis') or {}
                ats_keywords   = job_analysis.get('ats_keywords') or []
                approved_count = len(conversation.state.get('approved_rewrites') or [])

                summary = {
                    'files':          generated.get('files', []),
                    'output_dir':     str(output_dir),
                    'ats_keywords':   ats_keywords,
                    'ats_score':      ats_score,
                    'approved_rewrites': approved_count,
                    'application_status': app_status,
                }

                return jsonify({
                    'ok':          True,
                    'commit_hash': commit_hash,
                    'git_error':   git_error,
                    'summary':     summary,
                })
            except Exception as e:
                traceback.print_exc()
                return jsonify({'error': str(e)}), 500

    # ── Phase 11: Master Data Harvest ───────────────────────────────────────

    @app.get("/api/harvest/candidates")
    def harvest_candidates():
        """Compile candidate write-back items from the current session.

        Returns:
            { candidates: List[{id, type, label, original, proposed, rationale}] }
        """
        entry = _get_session()
        conversation = entry.manager
        try:
            candidates = _compile_harvest_candidates(conversation)
            return jsonify({'ok': True, 'candidates': candidates})
        except Exception as e:
            traceback.print_exc()
            return jsonify({'error': str(e)}), 500

    @app.post("/api/harvest/apply")
    def harvest_apply():
        """Write selected harvest candidates back to Master_CV_Data.json and git commit.

        Request body:
            selected_ids — list of candidate ids to apply

        Returns:
            { ok, written_count, diff_summary, commit_hash }
        """
        entry = _get_session()
        _validate_owner(entry)
        conversation = entry.manager
        sid = entry.session_id
        with entry.lock:
            try:
                body         = request.get_json(silent=True) or {}
                selected_ids = body.get('selected_ids') or []

                if not selected_ids:
                    return jsonify({'ok': True, 'written_count': 0, 'diff_summary': [], 'commit_hash': None})

                candidates_by_id = {c['id']: c for c in _compile_harvest_candidates(conversation)}
                selected = [candidates_by_id[sid] for sid in selected_ids if sid in candidates_by_id]
                if not selected:
                    return jsonify({'ok': True, 'written_count': 0, 'diff_summary': [], 'commit_hash': None})

                master_path = Path(conversation.orchestrator.master_data_path)
                with open(master_path, encoding='utf-8') as f:
                    master = json.load(f)

                diff_summary: List[Dict[str, Any]] = []

                for cand in selected:
                    ctype = cand['type']
                    if ctype == 'improved_bullet':
                        # Update the matching bullet in master experience data
                        applied = _harvest_apply_bullet(master, cand['original'], cand['proposed'])
                        diff_summary.append({
                            'id':      cand['id'],
                            'type':    ctype,
                            'applied': applied,
                            'label':   cand['label'],
                        })
                    elif ctype in ('new_skill', 'skill_gap_confirmed'):
                        skill_name = cand['proposed']
                        applied    = _harvest_add_skill(master, skill_name)
                        diff_summary.append({
                            'id':      cand['id'],
                            'type':    ctype,
                            'applied': applied,
                            'label':   cand['label'],
                        })
                    elif ctype == 'summary_variant':
                        applied = _harvest_add_summary_variant(master, cand['proposed'])
                        diff_summary.append({
                            'id':      cand['id'],
                            'type':    ctype,
                            'applied': applied,
                            'label':   cand['label'],
                        })

                # Write updated master data via shared helper (backup + validation)
                _save_master(master, master_path)

                # Reload in orchestrator
                conversation.orchestrator.master_data = master

                # Git commit
                job_analysis = conversation.state.get('job_analysis') or {}
                company  = (job_analysis.get('company') or 'Unknown').replace(' ', '_')
                role     = (job_analysis.get('title') or 'Role').replace(' ', '_')
                date_str = datetime.now().strftime('%Y-%m-%d')
                commit_msg = f"chore: Update master CV data from {company}_{role}_{date_str} session"

                commit_hash = None
                git_error   = None
                try:
                    repo_root = Path(__file__).parent.parent
                    subprocess.run(
                        ['git', 'add', str(master_path)],
                        cwd=str(repo_root), check=True, capture_output=True
                    )
                    result = subprocess.run(
                        ['git', 'commit', '-m', commit_msg],
                        cwd=str(repo_root), capture_output=True, text=True
                    )
                    if result.returncode == 0:
                        m = re.search(r'\b([0-9a-f]{7,40})\b', result.stdout)
                        commit_hash = m.group(1) if m else None
                    else:
                        git_error = result.stderr.strip() or result.stdout.strip()
                except Exception as git_exc:
                    git_error = str(git_exc)

                written_count = sum(1 for d in diff_summary if d.get('applied'))
                session_registry.touch(sid)
                return jsonify({
                    'ok':           True,
                    'written_count': written_count,
                    'diff_summary': diff_summary,
                    'commit_hash':  commit_hash,
                    'git_error':    git_error,
                })
            except Exception as e:
                traceback.print_exc()
                return jsonify({'error': str(e)}), 500

    # ── Session management endpoints ─────────────────────────────────────────

    @app.post("/api/sessions/new")
    def sessions_new():
        """Create a new session and return its ID.

        Returns ``{"ok": true, "session_id": "abcd1234", "redirect_url": "/?session=abcd1234"}``.
        """
        sid, _entry = session_registry.create(_app_config)
        return jsonify({"ok": True, "session_id": sid, "redirect_url": f"/?session={sid}"})

    @app.post("/api/sessions/claim")
    def sessions_claim():
        """Claim ownership of a session with a tab token.

        Body: ``{"session_id": "...", "owner_token": "..."}``
        Returns ``{"ok": true}`` or 409 if already owned by another token.
        """
        body = request.get_json(silent=True) or {}
        sid = body.get("session_id")
        token = body.get("owner_token")
        if not sid or not token:
            return jsonify({"error": "session_id and owner_token required"}), 400
        try:
            session_registry.claim(sid, token)
            return jsonify({"ok": True})
        except SessionNotFoundError as e:
            return jsonify({
                "ok": False,
                "error": "session_not_found",
                "message": str(e),
            })
        except SessionOwnedError as e:
            return jsonify({"error": "session_owned", "message": str(e)}), 409

    @app.post("/api/sessions/takeover")
    def sessions_takeover():
        """Forcibly take over a session (e.g. after page reload).

        Body: ``{"session_id": "...", "owner_token": "..."}``
        Returns ``{"ok": true}``.
        """
        body = request.get_json(silent=True) or {}
        sid = body.get("session_id")
        token = body.get("owner_token")
        if not sid or not token:
            return jsonify({"error": "session_id and owner_token required"}), 400
        try:
            session_registry.takeover(sid, token)
            return jsonify({"ok": True})
        except SessionNotFoundError as e:
            return jsonify({"error": str(e)}), 404

    @app.get("/api/sessions/active")
    def sessions_active():
        """Return a list of all active in-memory sessions.

        Returns ``{"sessions": [{"session_id": "...", "created": "...",
        "last_modified": "...", "position_name": "...", "phase": "...",
        "owner_token": "..."}]}``.
        """
        entries = session_registry.all_active()
        requester_token = request.args.get("owner_token")
        return jsonify({
            "sessions": [
                {
                    "session_id":          e.session_id,
                    "position_name":       (e.manager.state or {}).get("position_name"),
                    "phase":               (e.manager.state or {}).get("phase"),
                    "created":             e.created.isoformat(),
                    "last_modified":       e.last_modified.isoformat(),
                    "claimed":             e.owner_token is not None,
                    "owned_by_requester":  bool(
                        requester_token
                        and e.owner_token
                        and requester_token == e.owner_token
                    ),
                }
                for e in entries
            ]
        })

    @app.delete("/api/sessions/<session_id>/evict")
    def sessions_evict(session_id):
        """Save and remove a specific session from the registry.

        Returns ``{"ok": true}``.
        """
        entry = session_registry.get(session_id)
        if entry is None:
            return jsonify({"error": f"Session not found: {session_id}"}), 404
        _validate_owner(entry)
        try:
            entry.manager._save_session()
        except Exception:
            pass
        session_registry.remove(session_id)
        return jsonify({"ok": True})

    return app


# ---------------------------------------------------------------------------
# Master CV IO helpers (module-level for testability)
# ---------------------------------------------------------------------------

def _load_master(master_data_path: str) -> "tuple[dict, Path]":
    """Read master CV JSON from disk and return (data, path)."""
    p = Path(master_data_path).expanduser()

    validation = validate_master_data_file(str(p), use_schema=True)
    if not validation.valid:
        msg = "; ".join(validation.errors) or "master data validation failed"
        raise ValueError(f"Master data validation failed before load: {msg}")

    with open(p, 'r', encoding='utf-8') as f:
        return json.load(f), p


def _extract_year(value: Any) -> Optional[int]:
    """Extract a 4-digit year from a free-form date string."""
    raw = str(value or '').strip()
    if not raw:
        return None
    m = re.search(r'\b(19|20)\d{2}\b', raw)
    if not m:
        return None
    return int(m.group(0))


def _validate_master_for_persistence(master: Dict[str, Any]) -> None:
    """Validate top-level master CV structure before writing to disk."""
    if not isinstance(master, dict):
        raise ValueError("master data must be a JSON object")

    errors: List[str] = []

    if 'personal_info' in master and not isinstance(master.get('personal_info'), dict):
        errors.append("personal_info must be an object")

    for key in ('experience', 'education', 'awards', 'selected_achievements'):
        if key in master and not isinstance(master.get(key), list):
            errors.append(f"{key} must be a list")

    if 'skills' in master and not isinstance(master.get('skills'), (list, dict)):
        errors.append("skills must be a list or object")

    if 'professional_summaries' in master and not isinstance(master.get('professional_summaries'), (dict, list)):
        errors.append("professional_summaries must be an object or list")

    if errors:
        raise ValueError("Invalid master data: " + "; ".join(errors))


def _save_master(master: Dict[str, Any], master_path: Path) -> None:
    """Write master CV data to disk and stage the file in git.

    Creates a timestamped backup before overwrite when the target file exists.
    """
    _validate_master_for_persistence(master)

    backup_path: Optional[Path] = None
    if master_path.exists():
        backup_dir = master_path.parent / "backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        backup_path = backup_dir / f"{master_path.stem}.{ts}.bak{master_path.suffix}"
        shutil.copy2(master_path, backup_path)

    with open(master_path, 'w', encoding='utf-8') as f:
        json.dump(master, f, indent=2)

    validation = validate_master_data_file(str(master_path), use_schema=True)
    if not validation.valid:
        if backup_path is not None and backup_path.exists():
            shutil.copy2(backup_path, master_path)
        msg = "; ".join(validation.errors) or "master data validation failed"
        raise ValueError(f"Master data validation failed after write: {msg}")

    subprocess.run(
        ['git', '-C', str(master_path.parent), 'add', master_path.name],
        capture_output=True, check=False,
    )


# ---------------------------------------------------------------------------
# Phase 11 harvest helpers (module-level for testability)
# ---------------------------------------------------------------------------

def _compile_harvest_candidates(conversation) -> List[Dict[str, Any]]:
    """Return candidate write-back items for the current session.

    Extracted from the harvest_candidates view so harvest_apply can call it
    directly without round-tripping through the Flask response layer.
    """
    candidates: List[Dict[str, Any]] = []

    approved_rewrites = conversation.state.get('approved_rewrites') or []
    customizations    = conversation.state.get('customizations') or {}
    post_answers      = conversation.state.get('post_analysis_answers') or {}

    # 1. Improved bullets — approved rewrites where content differs from master
    for rw in approved_rewrites:
        if rw.get('section') == 'summary':  # handled separately as summary_variant
            continue
        proposed = rw.get('proposed', '')
        original = rw.get('original', '')
        if not proposed or not original:
            continue
        if proposed.strip() == original.strip():
            continue
        candidates.append({
            'id':        f"rewrite_{rw.get('id', len(candidates))}",
            'type':      'improved_bullet',
            'label':     f"Improved bullet — {rw.get('context', rw.get('id', 'unknown'))}",
            'original':  original,
            'proposed':  proposed,
            'rationale': rw.get('rationale') or 'Approved rewrite improves ATS-keyword coverage or adds a quantified metric.',
        })

    # 2. New / renamed skills added during the session
    for skill in customizations.get('new_skills_added') or []:
        if not skill:
            continue
        candidates.append({
            'id':        f"skill_{skill.replace(' ', '_')}",
            'type':      'new_skill',
            'label':     f"New skill — {skill}",
            'original':  '(not in master data)',
            'proposed':  skill,
            'rationale': 'Skill was added during the skills review step.',
        })

    # 3. Summary variant — if summary was rewritten and approved
    summary_rewrite = next(
        (rw for rw in approved_rewrites if rw.get('section') == 'summary'), None
    )
    if summary_rewrite and summary_rewrite.get('proposed'):
        cand_id = 'summary_variant'
        if not any(c['id'] == cand_id for c in candidates):
            candidates.append({
                'id':        cand_id,
                'type':      'summary_variant',
                'label':     'Professional summary variant',
                'original':  summary_rewrite.get('original', ''),
                'proposed':  summary_rewrite.get('proposed', ''),
                'rationale': 'Rewritten summary could be stored as a named variant for future reuse.',
            })

    # 4. Clarification-answer-revealed skills (user confirmed yes to a skill gap)
    for key, val in post_answers.items():
        if not isinstance(val, str):
            continue
        if key.startswith('skill_gap_') and val.lower() in ('yes', 'true', '1'):
            skill_name = key[len('skill_gap_'):]
            cand_id    = f'skill_gap_{skill_name}'
            if not any(c['id'] == cand_id for c in candidates):
                candidates.append({
                    'id':        cand_id,
                    'type':      'skill_gap_confirmed',
                    'label':     f"Confirmed skill — {skill_name}",
                    'original':  '(not in master data)',
                    'proposed':  skill_name,
                    'rationale': 'You confirmed this skill in response to a clarifying question.',
                })

    return candidates


def _harvest_apply_bullet(master: Dict, original: str, proposed: str) -> bool:
    """Replace ``original`` bullet text with ``proposed`` in master experience data."""
    experiences = (
        master.get('experience')
        or master.get('experiences')
        or []
    )
    for exp in experiences:
        achievements = exp.get('achievements') or exp.get('bullets') or []
        for i, bullet in enumerate(achievements):
            text = bullet if isinstance(bullet, str) else bullet.get('text', '')
            if text.strip() == original.strip():
                if isinstance(bullet, str):
                    achievements[i] = proposed
                else:
                    bullet['text'] = proposed
                return True
    return False


def _harvest_add_skill(master: Dict, skill_name: str) -> bool:
    """Add ``skill_name`` to master skills data; returns True if actually added."""
    skills = master.get('skills')
    if isinstance(skills, list):
        if skill_name not in skills:
            skills.append(skill_name)
            return True
        return False
    elif isinstance(skills, dict):
        # Category dict — add to a named 'other'/'general'/'additional' bucket;
        # never auto-add to an arbitrary category.
        for cat_key, cat_val in skills.items():
            cat_list: List[str] = []
            if isinstance(cat_val, list):
                cat_list = cat_val
            elif isinstance(cat_val, dict) and isinstance(cat_val.get('skills'), list):
                cat_list = cat_val['skills']
            if cat_key.lower() in ('other', 'general', 'additional'):
                if skill_name not in cat_list:
                    cat_list.append(skill_name)
                    return True
                return False
        # No suitable category — add a new 'Other' category
        skills['Other'] = [skill_name]
        return True
    # skills field is missing — create as list
    master['skills'] = [skill_name]
    return True


def _harvest_add_summary_variant(master: Dict, new_summary: str) -> bool:
    """Store ``new_summary`` as a named variant in master data.

    If ``professional_summaries`` already exists as a list, appends.
    Otherwise creates it.
    """
    variants = master.get('professional_summaries')
    if isinstance(variants, list):
        if new_summary not in variants:
            variants.append(new_summary)
            return True
        return False
    master['professional_summaries'] = [new_summary]
    return True


def parse_args():
    config = get_config()
    
    parser = argparse.ArgumentParser(description="Minimal Web UI for CV Generator")
    parser.add_argument("--job-file", help="Path to job description text file")
    parser.add_argument("--master-data", default=config.master_cv_path,
                       help=f"Path to Master_CV_Data.json")
    parser.add_argument("--publications", default=config.publications_path,
                       help=f"Path to publications.bib")
    parser.add_argument("--output-dir", default=config.output_dir,
                       help=f"Output directory")
    parser.add_argument("--llm-provider", choices=["copilot-oauth", "copilot", "github", "openai", "anthropic", "gemini", "groq", "local", "copilot-sdk", "stub"],
                       default=config.llm_provider,
                       help=f"LLM provider (default: {config.llm_provider})")
    parser.add_argument("--model", default=config.llm_model, help="Specific model to use")
    parser.add_argument("--port", type=int, default=config.web_port,
                       help=f"Port to run on (default: {config.web_port})")
    parser.add_argument("--debug", action="store_true", help="Run Flask in debug mode")
    return parser.parse_args()


def main():
    args = parse_args()
    config = get_config()
    
    # Set up logging before anything else
    setup_logging(config)

    app = create_app(args)
    bundle_status = app.config.get('FRONTEND_BUNDLE_STATUS', 'unknown')
    bundle_built_at = app.config.get('FRONTEND_BUNDLE_BUILT_AT', 'unknown')

    # ── Startup banner ────────────────────────────────────────────────────────
    model_source = (
        "env var"      if os.getenv("CV_LLM_MODEL")
        else ".env"    if (Path(__file__).parent.parent / ".env").exists()
                          and _env_file_has_value("CV_LLM_MODEL")
        else "config.yaml"
    )
    provider_source = (
        "env var"      if os.getenv("CV_LLM_PROVIDER")
        else "CLI arg" if args.llm_provider != config.llm_provider
        else ".env"    if (Path(__file__).parent.parent / ".env").exists()
                          and _env_file_has_value("CV_LLM_PROVIDER")
        else "config.yaml"
    )
    print(
        f"\n"
        f"  ┌─────────────────────────────────────────────────────┐\n"
        f"  │                    CV Builder                       │\n"
        f"  ├──────────┬──────────────────────────────────────────┤\n"
        f"  │ provider │ {args.llm_provider:<24}  [{provider_source}]\n"
        f"  │ model    │ {args.model or '(provider default)':<24}  [{model_source}]\n"
        f"  │ port     │ {args.port}\n"
        f"  │ data     │ {args.master_data}\n"
        f"  │ output   │ {args.output_dir}\n"
        f"  │ bundle   │ {bundle_status}\n"
        f"  │ built at │ {bundle_built_at}\n"
        f"  └──────────┴──────────────────────────────────────────┘\n",
        flush=True,
    )

    app.run(host=config.web_host, port=args.port, debug=args.debug)


def _env_file_has_value(key: str) -> bool:
    """Return True if *key* is set (uncommented) in the project .env file."""
    env_path = Path(__file__).parent.parent / ".env"
    try:
        for line in env_path.read_text().splitlines():
            line = line.strip()
            if line.startswith("#") or "=" not in line:
                continue
            if line.split("=", 1)[0].strip() == key:
                return True
    except OSError:
        pass
    return False


if __name__ == "__main__":
    main()
