# Copyright (C) 2026 Gregory R. Warnes
# SPDX-License-Identifier: AGPL-3.0-or-later
#
# This file is part of CV-Builder.
# For commercial licensing, contact greg@warnes-innovations.com

"""
CV Orchestrator - Bridges LLM intelligence with document generation utilities.

This module coordinates between:
- LLM-driven content selection
- Traditional utility functions (scoring, formatting)
- Document generation (DOCX/PDF)
"""

import copy
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import logging
import re
import sys
import time
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timezone, date as _date
import subprocess
import weasyprint  # noqa: F401  -- kept for test mock path (patch cv_orchestrator.weasyprint.HTML)
from collections import Counter, defaultdict
from bs4 import BeautifulSoup, Comment

from .scoring import (
    calculate_relevance_score,
    calculate_skill_score,
)
from .bibtex_parser import parse_bibtex_file, format_publication
from .config import get_config
from .llm_client import LLMClient
from .master_data_mutations import (
    ALLOWED_OPS as _MDU_ALLOWED_OPS,
    ALLOWED_SECTIONS as _MDU_ALLOWED_SECTIONS,
    _apply_master_data_change as _mdu_apply_change,
)
from .master_data_validator import validate_master_data, validate_master_data_file
from .session_data_view import SessionDataView
from .prompt_safety import sanitize_instruction_text, scan_text_for_injection
from .template_renderer import safe_css_size, safe_url

logger = logging.getLogger(__name__)


class PDFRendererNotFoundError(OSError):
    """Raised when no supported PDF renderer (Chrome, WeasyPrint) is found."""


class PDFRenderingError(RuntimeError):
    """Raised when a PDF renderer is found but fails to produce output."""


_LAYOUT_URL_ATTRS = ('href', 'src', 'srcset', 'poster', 'xlink:href')
_LAYOUT_PRESERVED_HEAD_TAGS = {'link', 'script', 'meta', 'base'}
_SCHEMA_ORG_CONTEXTS = {'https://schema.org', 'http://schema.org'}

def _append_layout_finding(
    findings: List[Dict[str, Any]],
    issue: str,
    detail: str,
    fragment: Optional[str] = None,
) -> None:
    """Append a normalized layout safety finding."""
    entry: Dict[str, Any] = {'issue': issue, 'detail': detail}
    if fragment:
        entry['fragment'] = fragment[:500]
    findings.append(entry)


def _summarize_layout_findings(
    *finding_groups: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Flatten finding groups while preserving order."""
    merged: List[Dict[str, Any]] = []
    for group in finding_groups:
        merged.extend(group or [])
    return merged


def _is_exact_schema_org_context(value: Any) -> bool:
    """Return true when a JSON-LD @context is exactly schema.org."""
    if isinstance(value, str):
        return value in _SCHEMA_ORG_CONTEXTS
    if isinstance(value, list):
        return any(
            isinstance(item, str) and item in _SCHEMA_ORG_CONTEXTS
            for item in value
        )
    return False


class CVOrchestrator:
    """Orchestrates CV generation with LLM + utilities."""

    _TEMPLATES_DIR: Path = Path(__file__).parent.parent.parent / 'templates'
    _CV_TEMPLATE_FILE: Path = _TEMPLATES_DIR / 'cv-template.html'

    def __init__(
        self,
        master_data_path: str,
        publications_path: str,
        output_dir: str,
        llm_client: LLMClient
    ):
        self.master_data_path = Path(master_data_path).expanduser()
        self.publications_path = Path(publications_path).expanduser()
        self.output_dir = Path(output_dir).expanduser()
        self.llm = llm_client
        
        # Load master data
        self.master_data = self._load_master_data()

        # Load publications if available
        self.publications = {}
        if self.publications_path.exists():
            self.publications = parse_bibtex_file(str(self.publications_path))

        # Load synonym map for ATS skill normalisation
        self._synonym_map: Dict[str, str] = self._load_synonym_map()
        # Full expansion index: any form (lower) -> canonical
        self._expansion_index: Dict[str, str] = {}
        for alias, canonical in self._synonym_map.items():
            if alias.startswith('_'):  # skip comment keys
                continue
            self._expansion_index[alias.lower()] = canonical
            self._expansion_index[canonical.lower()] = canonical
    
    def _load_master_data(self) -> Dict:
        """Load Master_CV_Data.json."""
        if not self.master_data_path.exists():
            raise FileNotFoundError(
                f"Master data file not found: {self.master_data_path}\n"
                "Please create Master_CV_Data.json first."
            )

        validation = validate_master_data_file(str(self.master_data_path), use_schema=True)
        if not validation.valid:
            msg = "; ".join(validation.errors) or "master data validation failed"
            raise ValueError(f"Master data validation failed before load: {msg}")
        
        with open(self.master_data_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _load_synonym_map(self) -> Dict[str, str]:
        """Load scripts/data/synonym_map.json, returning {} gracefully if missing."""
        map_path = Path(__file__).parent.parent / 'data' / 'synonym_map.json'
        if not map_path.exists():
            return {}
        try:
            with open(map_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return {k: v for k, v in data.items() if not k.startswith('_')}
        except Exception:
            return {}

    def canonical_skill_name(self, name: str) -> str:
        """Return the canonical form of a skill name using the synonym map.

        Examples: 'ML' -> 'Machine Learning', 'sklearn' -> 'scikit-learn'.
        Unknown names are returned unchanged.
        """
        return self._expansion_index.get(name.lower(), name)

    def _prepare_cv_data_for_template(
        self,
        selected_content: Dict,
        job_analysis: Dict,
        template_variant: str = 'standard',
        base_font_size: str | None = None,
        customizations: Optional[Dict] = None,
    ) -> Dict:
        """Prepare CV data in the format expected by the HTML resume template."""

        # Work on copies so template-specific normalization never mutates the
        # session's selected content in place.
        personal_info = copy.deepcopy(selected_content.get('personal_info', {}))

        # Validate contact information
        contact = personal_info.get('contact', {})
        address = contact.get('address', {})
        if address:
            address_display = f"{address.get('city', '')}, {address.get('state', '')}"
            address_display = address_display.strip(', ')
            contact['address_display'] = address_display
        contact['linkedin_href'] = safe_url(contact.get('linkedin', ''))
        contact['website_href'] = safe_url(contact.get('website', ''))
        
        # Ensure languages key exists (template expects it)
        if 'languages' not in personal_info:
            personal_info['languages'] = []
        else:
            personal_info['languages'] = self._normalize_language_entries(
                personal_info.get('languages', [])
            )

        # Get professional summary
        professional_summary = selected_content.get('summary', '')
        if not professional_summary.strip():
            professional_summary = f"Experienced professional applying for {job_analysis.get('title', 'position')}"
        
        # Format skills by category
        _show_proficiency = True
        if isinstance(customizations, dict):
            raw_show_prof = customizations.get('skills_show_proficiency', True)
            if isinstance(raw_show_prof, bool):
                _show_proficiency = raw_show_prof
            elif isinstance(raw_show_prof, str):
                _show_proficiency = raw_show_prof.strip().lower() not in {'false', '0', 'no', 'never'}
        skills_by_category = self._organize_skills_by_category(
            selected_content.get('skills', []),
            template_variant,
            selected_content.get('skill_category_order', []),
            show_proficiency=_show_proficiency,
        )
        
        # Format publications
        publications = self._format_publications(selected_content.get('publications', []))

        experiences = self._normalize_experiences_for_template(
            copy.deepcopy(selected_content.get('experiences', []))
        )
        achievements = self._normalize_achievement_entries(
            copy.deepcopy(selected_content.get('achievements', []))
        )

        publications_start_new_page = False
        if isinstance(customizations, dict):
            raw_publications_page_break = customizations.get(
                'publications_start_new_page',
                customizations.get(
                    'publications_page_break',
                    customizations.get('start_publications_on_new_page', False),
                ),
            )
            if isinstance(raw_publications_page_break, str):
                publications_start_new_page = raw_publications_page_break.strip().lower() in {
                    '1',
                    'true',
                    'yes',
                    'on',
                }
            else:
                publications_start_new_page = bool(raw_publications_page_break)

        skills_show_experience = 'individual'
        if isinstance(customizations, dict):
            raw_skills_exp = customizations.get(
                'skills_show_experience',
                customizations.get('show_skill_experience', 'individual'),
            )
            if (
                isinstance(raw_skills_exp, str)
                and raw_skills_exp.strip().lower() in {'always', 'never', 'individual'}
            ):
                skills_show_experience = raw_skills_exp.strip().lower()

        # Get awards and certifications
        awards = selected_content.get('awards', [])
        certifications = selected_content.get('certifications', [])
        applicant_tagline = self._resolve_applicant_tagline(
            customizations=customizations,
            personal_info=personal_info,
            job_analysis=job_analysis,
        )
        
        # Add template metadata
        # duckflow:
        #   id: cv_render.scripts_utils_cv_orchestrator.L224
        #   kind: artifact
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   writes:
        #     - "artifact:template_metadata[\"skills_section_title\"]"
        #   notes: "Seeds the render metadata with the default skills section title before preview or final-generation overrides are applied."
        template_metadata = {
            'variant': template_variant,
            'generated_date': datetime.now().isoformat(),
            'job_title': job_analysis.get('title', ''),
            'applicant_tagline': applicant_tagline,
            'company': job_analysis.get('company', ''),
            'total_publications_count': len(self.publications) if self.publications else 0,
            'skills_section_title': 'Skills',
            'publications_start_new_page': publications_start_new_page,
            'skills_show_experience': skills_show_experience,
        }
        human_skills_title = self._resolve_human_skills_title(customizations)
        
        cv_data = {
            'personal_info': personal_info,
            'professional_summary': professional_summary,
            'experiences': experiences,
            'achievements': achievements,
            'education': selected_content.get('education', []),
            'skills_by_category': skills_by_category,
            'awards': awards,
            'certifications': certifications,
            'publications': publications,
            'template_metadata': template_metadata,
            'base_font_size': safe_css_size(base_font_size, default='10px'),
            'human_skills_title': human_skills_title,
        }

        return cv_data

    @staticmethod
    def _resolve_applicant_tagline(
        customizations: Optional[Dict],
        personal_info: Optional[Dict],
        job_analysis: Optional[Dict],
    ) -> str:
        """Resolve a non-placeholder applicant tagline for resume headers."""
        job_title = str((job_analysis or {}).get('title') or '').strip().lower()

        candidates: List[Any] = []
        if isinstance(customizations, dict):
            candidates.extend([
                customizations.get('tagline_override'),
                customizations.get('tagline'),
                customizations.get('applicant_tagline'),
            ])

        if isinstance(personal_info, dict):
            candidates.extend([
                personal_info.get('applicant_tagline'),
                personal_info.get('tagline'),
                personal_info.get('headline'),
                personal_info.get('professional_headline'),
                personal_info.get('professional_title'),
                personal_info.get('title'),
            ])

        for candidate in candidates:
            text = str(candidate or '').strip()
            if not text:
                continue

            normalized = re.sub(r'\s+', ' ', text).strip()
            lowered = normalized.lower()
            if lowered == job_title:
                continue
            if lowered == 'debug resume render':
                continue
            return normalized

        return ''

    @staticmethod
    def _resolve_human_skills_title(customizations: Optional[Dict]) -> str:
        """Return the human-facing skills heading from session customizations."""
        if not isinstance(customizations, dict):
            return 'Technical Skills'

        candidate_values = [
            customizations.get('skills_section_title'),
            customizations.get('skills_title'),
        ]

        section_titles = customizations.get('section_titles')
        if isinstance(section_titles, dict):
            candidate_values.extend([
                section_titles.get('human_skills'),
                section_titles.get('skills_human'),
                section_titles.get('skills'),
            ])

        for candidate in candidate_values:
            if isinstance(candidate, str) and candidate.strip():
                return candidate.strip()

        return 'Technical Skills'

    @staticmethod
    def _extract_display_text(item: Any, preferred_fields: Optional[List[str]] = None) -> str:
        """Extract the best human-readable text from a template item."""
        if item is None:
            return ''
        if isinstance(item, str):
            return item.strip()
        if not isinstance(item, dict):
            return str(item).strip()

        field_order = preferred_fields or []
        fallback_fields = [
            'text',
            'description',
            'summary',
            'formatted',
            'formatted_citation',
            'title',
            'name',
            'degree',
            'institution',
            'language',
            'value',
        ]

        for field_name in field_order + fallback_fields:
            value = item.get(field_name)
            if isinstance(value, str) and value.strip():
                return value.strip()

        return ''

    def _normalize_achievement_entries(self, achievements: List[Any]) -> List[Any]:
        """Ensure achievement-like entries always expose a human-readable ``text``."""
        normalized = []
        for achievement in achievements or []:
            if isinstance(achievement, dict):
                entry = dict(achievement)
                entry['text'] = self._extract_display_text(
                    entry,
                    preferred_fields=['text', 'description', 'summary', 'title', 'name'],
                )
                normalized.append(entry)
            else:
                normalized.append(achievement)
        return normalized

    def _normalize_experiences_for_template(self, experiences: List[Any]) -> List[Any]:
        """Normalize experience achievement payloads before template rendering."""
        normalized = []
        for experience in experiences or []:
            if not isinstance(experience, dict):
                normalized.append(experience)
                continue

            entry = dict(experience)
            for key in ('ordered_achievements', 'achievements'):
                if isinstance(entry.get(key), list):
                    entry[key] = self._normalize_achievement_entries(entry.get(key, []))
            normalized.append(entry)
        return normalized

    def _apply_session_achievement_edits(
        self,
        selected_content: Dict,
        achievement_edits: Dict[Any, Any],
    ) -> Dict:
        """Overlay session-edited bullets onto the selected experience list."""
        if not achievement_edits:
            return selected_content

        updated = copy.deepcopy(selected_content)
        master_experiences = self.master_data.get('experience') or []
        selected_experiences = updated.get('experiences') or []
        experience_by_id = {
            str(exp.get('id') or '').strip(): exp
            for exp in selected_experiences
            if isinstance(exp, dict) and str(exp.get('id') or '').strip()
        }

        for raw_key, raw_items in achievement_edits.items():
            try:
                exp_idx = int(raw_key)
            except (TypeError, ValueError):
                continue
            if not (0 <= exp_idx < len(master_experiences)):
                continue

            master_exp = master_experiences[exp_idx] or {}
            exp_id = str(master_exp.get('id') or '').strip()
            if not exp_id or exp_id not in experience_by_id:
                continue

            visible_achievements = []
            items = raw_items if isinstance(raw_items, list) else [raw_items]
            for item in items:
                if isinstance(item, dict):
                    text = str(item.get('text') or item.get('description') or item.get('content') or '').strip()
                    hidden = bool(item.get('hidden'))
                else:
                    text = str(item or '').strip()
                    hidden = False

                if text and not hidden:
                    visible_achievements.append({'text': text})

            selected_exp = experience_by_id[exp_id]
            selected_exp['achievements'] = visible_achievements
            if 'ordered_achievements' in selected_exp:
                selected_exp['ordered_achievements'] = copy.deepcopy(visible_achievements)

        return updated

    def _normalize_language_entries(self, languages: List[Any]) -> List[str]:
        """Convert language records to simple display strings for the template."""
        normalized = []
        for language in languages or []:
            if isinstance(language, dict):
                name = self._extract_display_text(
                    language,
                    preferred_fields=['language', 'name'],
                )
                proficiency = str(language.get('proficiency', '')).strip()
                if name and proficiency:
                    normalized.append(f"{name} ({proficiency})")
                elif name:
                    normalized.append(name)
            else:
                text = str(language).strip()
                if text:
                    normalized.append(text)
        return normalized

    def _deduplicate_skills(self, skills: List[Dict]) -> List[Dict]:
        """Deduplicate skills by canonical synonym name, merging aliases."""
        canonical_seen: Dict[str, Dict] = {}  # canonical_lower -> merged skill dict
        for skill in skills:
            name = skill.get('name', '')
            canonical = self.canonical_skill_name(name)
            key = canonical.lower()
            if key not in canonical_seen:
                merged = dict(skill)
                merged['name'] = canonical if canonical != name else name
                merged.setdefault('aliases', list(skill.get('aliases') or []))
                if canonical != name and name not in merged['aliases']:
                    merged['aliases'].append(name)
                canonical_seen[key] = merged
            else:
                existing = canonical_seen[key]
                # Keep the entry with more years; add the other name as alias
                if skill.get('years', 0) > existing.get('years', 0):
                    alias_name = existing.get('name', '')
                    existing.update({k: v for k, v in skill.items() if k != 'aliases'})
                    existing['name'] = canonical
                    existing.setdefault('aliases', [])
                    if alias_name and alias_name not in existing['aliases']:
                        existing['aliases'].append(alias_name)
                else:
                    existing.setdefault('aliases', [])
                    if name and name not in existing['aliases'] and name != existing['name']:
                        existing['aliases'].append(name)
        return list(canonical_seen.values())

    def _group_skills_by_category(self, skills: List[Dict]) -> Dict[str, List[Dict]]:
        """Group a flat list of skills by their category field."""
        category_skills: Dict[str, List[Dict]] = defaultdict(list)
        for skill in skills:
            category = skill.get('category', 'General')
            category_skills[category].append(skill)
        return category_skills

    def _sort_categories(
        self,
        category_skills: Dict[str, List[Dict]],
        variant: str,
        category_order: Optional[List[str]],
        show_proficiency: bool,
    ) -> List[Dict]:
        """Sort skill categories by priority order and build the final sorted list."""
        custom_order = []
        for category in category_order or []:
            label = str(category or '').strip()
            if label and label not in custom_order:
                custom_order.append(label)

        priority_orders = {
            'standard': ['Core Expertise', 'Programming', 'Technical', 'Tools', 'General'],
            'technical': ['Programming', 'Technical', 'Tools', 'Core Expertise', 'General'],
            'academic': ['Research', 'Technical', 'Programming', 'Core Expertise', 'General']
        }
        priority_order = custom_order or priority_orders.get(variant, priority_orders['standard'])

        sorted_categories = []
        for category in priority_order:
            if category in category_skills:
                skills_list = sorted(category_skills[category],
                                     key=lambda x: (-x.get('years', 0), x.get('name', '')))
                sorted_categories.append({
                    'category': category,
                    'skills': self._group_inline_skills(skills_list, show_proficiency=show_proficiency)
                })

        remaining_categories = sorted(set(category_skills.keys()) - set(priority_order))
        for category in remaining_categories:
            skills_list = sorted(category_skills[category],
                                 key=lambda x: (-x.get('years', 0), x.get('name', '')))
            sorted_categories.append({
                'category': category,
                'skills': self._group_inline_skills(skills_list, show_proficiency=show_proficiency)
            })

        return sorted_categories

    def _organize_skills_by_category(
        self,
        skills: List[Dict],
        variant: str,
        category_order: Optional[List[str]] = None,
        show_proficiency: bool = True,
    ) -> List[Dict]:
        """Organize skills by category, deduplicating by canonical synonym name."""
        if not skills:
            return []
        deduped = self._deduplicate_skills(skills)
        grouped  = self._group_skills_by_category(deduped)
        return self._sort_categories(grouped, variant, category_order, show_proficiency)

    def _group_inline_skills(
        self,
        skills_list: List[Dict],
        show_proficiency: bool = True,
    ) -> List[Dict]:
        """Combine skills that share the same non-empty `group` key into a
        single inline entry.  The first member becomes the representative entry
        with an added `group_names` list.  Ungrouped skills pass through unchanged."""
        groups: Dict[str, List[Dict]] = {}
        group_insertion_idx: Dict[str, int] = {}
        result: List[Any] = []

        for skill in skills_list:
            g = (skill.get('group') or '').strip()
            if g:
                if g not in groups:
                    groups[g] = []
                    group_insertion_idx[g] = len(result)
                    result.append(None)  # placeholder
                groups[g].append(skill)
            else:
                result.append(skill)

        for g, members in groups.items():
            primary = dict(members[0])
            primary['group_names'] = [m['name'] for m in members]
            primary['group_display_names'] = [self._skill_inline_label(m, show_proficiency=show_proficiency) for m in members]
            primary['display_name'] = self._skill_inline_label(primary, show_proficiency=show_proficiency)
            result[group_insertion_idx[g]] = primary

        finalized = [s for s in result if s is not None]
        for skill in finalized:
            if isinstance(skill, dict) and 'display_name' not in skill:
                skill['display_name'] = self._skill_inline_label(skill, show_proficiency=show_proficiency)
        return finalized

    @staticmethod
    def _skill_inline_label(
        skill: Dict[str, Any],
        show_proficiency: bool = True,
    ) -> str:
        """Return a human-readable inline label for a skill entry."""
        name = str(skill.get('name') or '').strip()
        if not name:
            return ''

        parenthetical = str(skill.get('parenthetical') or '').strip()
        if parenthetical:
            return f"{name} ({parenthetical})"

        qualifier_parts = []
        if show_proficiency:
            proficiency = str(skill.get('proficiency') or '').strip()
            if proficiency:
                qualifier_parts.append(proficiency[:1].upper() + proficiency[1:])

        raw_subskills = skill.get('subskills', skill.get('sub_skills', []))
        if isinstance(raw_subskills, str):
            raw_subskills = [item.strip() for item in raw_subskills.split(',')]
        subskills = [
            str(item).strip()
            for item in raw_subskills or []
            if str(item).strip()
        ]
        qualifier_parts.extend(subskills)

        if qualifier_parts:
            return f"{name} ({', '.join(qualifier_parts)})"

        years = skill.get('years')
        if years:
            return f"{name} ({years} yrs)"

        return name

    @staticmethod
    def _normalize_extra_skill_entry(raw_skill: Any) -> Optional[Dict[str, Any]]:
        """Return a dict form for session-only extra skills."""
        if isinstance(raw_skill, str):
            name = raw_skill.strip()
            return {'name': name} if name else None
        if not isinstance(raw_skill, dict):
            return None

        name = str(raw_skill.get('name') or '').strip()
        if not name:
            return None

        normalized = dict(raw_skill)
        normalized['name'] = name
        raw_subskills = normalized.get('subskills', normalized.get('sub_skills'))
        if isinstance(raw_subskills, str):
            raw_subskills = [item.strip() for item in raw_subskills.split(',')]
        if isinstance(raw_subskills, list):
            cleaned = []
            seen = set()
            for item in raw_subskills:
                label = str(item or '').strip()
                if not label or label in seen:
                    continue
                cleaned.append(label)
                seen.add(label)
            if cleaned:
                normalized['subskills'] = cleaned
            else:
                normalized.pop('subskills', None)
                normalized.pop('sub_skills', None)
        return normalized

    @staticmethod
    def _publication_year_value(pub: Dict[str, Any]) -> Optional[int]:
        """Return a parsed publication year from explicit fields or cite key."""
        year_sources = [
            pub.get('year'),
            (pub.get('fields') or {}).get('year') if isinstance(pub.get('fields'), dict) else None,
        ]
        for source in year_sources:
            text = str(source or '').strip()
            if not text:
                continue
            match = re.search(r'(19|20)\d{2}', text)
            if match:
                return int(match.group(0))

        key_text = str(pub.get('key') or '').strip()
        key_match = re.search(r'(19|20)\d{2}', key_text)
        if key_match:
            return int(key_match.group(0))

        return None

    def _sort_selected_publications(
        self,
        publications: List[Dict[str, Any]],
        customizations: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """Sort selected publications unless the user has explicitly ordered them."""
        if not publications:
            return []

        # Explicit row order from UI/session takes precedence.
        explicit_order = (
            customizations.get('publication_row_order')
            or customizations.get('publication_order')
            or []
        )
        if explicit_order:
            order_map = {str(key): idx for idx, key in enumerate(explicit_order)}
            return sorted(
                publications,
                key=lambda pub: order_map.get(str(pub.get('key') or ''), len(order_map)),
            )

        sort_pref = str(
            customizations.get('publication_sort_order')
            or customizations.get('publication_sort')
            or ''
        ).strip().lower()

        # Default: newest first (descending).
        descending = sort_pref not in {'asc', 'ascending', 'oldest', 'oldest_first', 'chronological_asc'}

        def _sort_key(pub: Dict[str, Any]) -> tuple:
            parsed_year = self._publication_year_value(pub)
            normalized_year = parsed_year if parsed_year is not None else -1
            title = str(pub.get('title') or '').lower()
            return (normalized_year, title)

        return sorted(publications, key=_sort_key, reverse=descending)

    def _format_publications(self, publications: List) -> List[Dict]:
        """Format publications for template consumption."""
        owner_name = self.master_data.get('personal_info', {}).get('name', '') if self.master_data else ''
        # Extract last name: handle "Last, First" (BibTeX/comma style) and "First Last" (natural)
        if ',' in owner_name:
            owner_last = owner_name.split(',')[0].strip().lower()
        else:
            owner_last = owner_name.strip().split()[-1].lower() if owner_name.strip() else ''

        formatted_pubs = []
        for pub in publications:
            if isinstance(pub, dict):
                entry: Dict[str, Any] = {}
                fields = pub.get('fields', {}) if isinstance(pub.get('fields'), dict) else {}
                note_text = str(pub.get('note') or fields.get('note') or '').strip()
                url_text = str(pub.get('url') or fields.get('url') or '').strip()
                formatted_text = str(pub.get('formatted') or '').strip()
                title_text = str(pub.get('title') or '').strip()

                combined_text = ' '.join(
                    [
                        note_text,
                        url_text,
                        formatted_text,
                        title_text,
                    ]
                ).lower()
                is_r_package = (
                    'r package' in combined_text
                    or 'cran.r-project.org' in combined_text
                    or ' bioconductor' in combined_text
                    or 'bioconductor.org' in combined_text
                )

                entry_type = str(pub.get('type') or fields.get('ENTRYTYPE') or '').lower()
                is_software_entry = (
                    fields.get('type') == 'software'
                    or pub.get('type') == 'software'
                    or is_r_package
                )
                is_patent = (
                    entry_type in ('patent', 'patents')
                    or 'patent' in str(pub.get('title', '')).lower()
                    or 'patent' in str(fields.get('note', '')).lower()
                )

                venue_text = (
                    pub.get('journal')
                    or pub.get('booktitle')
                    or pub.get('institution')
                    or pub.get('school')
                    or pub.get('publisher')
                    or pub.get('organization')
                    or fields.get('journal')
                    or fields.get('booktitle')
                    or fields.get('institution')
                    or fields.get('school')
                    or fields.get('publisher')
                    or fields.get('organization')
                    or fields.get('howpublished')
                    or fields.get('series')
                    or (fields.get('note') if not is_software_entry else '')
                    or ''
                )
                venue_text = str(venue_text).strip()
                if not venue_text and is_r_package:
                    venue_text = (
                        'Bioconductor R package'
                        if 'bioconductor' in combined_text
                        else 'CRAN R package'
                    )

                publication_url = ''
                doi_value = str(pub.get('doi') or fields.get('doi') or '').strip()
                if doi_value:
                    doi_clean = re.sub(r'^https?://(dx\.)?doi\.org/', '', doi_value, flags=re.I)
                    doi_clean = re.sub(r'^doi:\s*', '', doi_clean, flags=re.I)
                    publication_url = safe_url(f"https://doi.org/{doi_clean.lstrip('/')}" )

                if not publication_url:
                    raw_url = str(pub.get('url') or fields.get('url') or '').strip()
                    if raw_url and not re.match(r'^[a-z][a-z0-9+.-]*://', raw_url, flags=re.I):
                        if raw_url.lower().startswith('doi.org/'):
                            raw_url = f'https://{raw_url}'
                        elif raw_url.startswith('www.'):
                            raw_url = f'https://{raw_url}'
                    publication_url = safe_url(raw_url)

                if 'formatted' in pub:
                    formatted = str(pub.get('formatted', '')).strip()
                    if (
                        venue_text
                        and venue_text.lower() not in formatted.lower()
                        and not is_software_entry
                    ):
                        formatted = f"{formatted.rstrip('.')} {venue_text}.".strip()
                    entry['formatted_citation'] = formatted
                elif 'title' in pub:
                    authors = pub.get('authors', 'Unknown')
                    title = pub.get('title', '')
                    year = pub.get('year', '')
                    citation = f"{authors}. {title}. {venue_text} ({year}).".strip()
                    entry['formatted_citation'] = citation
                else:
                    continue

                citation_title = title_text.replace('{', '').replace('}', '').strip()
                entry['title'] = citation_title
                entry['citation_prefix'] = ''
                entry['citation_title'] = ''
                entry['citation_suffix'] = ''
                if citation_title and citation_title in entry['formatted_citation']:
                    prefix, suffix = entry['formatted_citation'].split(citation_title, 1)
                    entry['citation_prefix'] = prefix
                    entry['citation_title'] = citation_title
                    entry['citation_suffix'] = suffix

                entry['publication_url'] = publication_url

                # Detect first authorship: compare owner last name against leading author token
                if owner_last:
                    raw_authors = pub.get('authors', '')
                    first_token = raw_authors.split(',')[0].strip().lower() if raw_authors else ''
                    entry['is_first_author'] = bool(first_token and owner_last in first_token)
                else:
                    entry['is_first_author'] = False

                # Flag entries with no venue so the template can render a warning icon
                has_venue = bool(venue_text or is_software_entry or is_patent)
                entry['venue_warning'] = '' if has_venue else 'No journal or conference name found in BibTeX entry'

                formatted_pubs.append(entry)
        return formatted_pubs
    
    def render_html_preview(
        self,
        job_analysis: Dict,
        customizations: Dict,
        approved_rewrites: Optional[List[Dict]] = None,
        spell_audit: Optional[List[Dict]] = None,
        max_skills: Optional[int] = None,
        template_variant: str = 'standard',
        use_semantic_match: bool = True,
    ) -> str:
        """Render CV as HTML for preview without generating PDF or DOCX.

        Called by the staged generation workflow (GAP-20) to produce the
        preview artifact that the layout-review loop works against.  Does not
        write any files; returns the raw HTML string.

        Parameters mirror ``generate_cv`` but only the HTML rendering path is
        executed, so this is significantly faster than a full generation run.
        """
        # duckflow:
        #   id: summary_orchestrator_preview_html
        #   kind: artifact
        #   timestamp: "2026-03-27T01:23:28Z"
        #   status: shared
        #   reads: ["cv:selected_content.summary"]
        #   returns: ["artifact:generation_state.preview_html"]
        #   notes: "Renders the selected summary into preview HTML without writing files yet."
        selected_content = self.build_render_ready_content(
            job_analysis,
            customizations,
            approved_rewrites=approved_rewrites,
            spell_audit=spell_audit,
            max_skills=max_skills,
            use_semantic_match=use_semantic_match,
        )
        cv_data = self._prepare_cv_data_for_template(
            selected_content,
            job_analysis,
            template_variant,
            customizations.get('base_font_size'),
            customizations=customizations,
        )
        cv_data['json_ld_str']    = self._build_json_ld(cv_data, job_analysis)
        # duckflow:
        #   id: cv_render.scripts_utils_cv_orchestrator.L599
        #   kind: artifact
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   reads:
        #     - "customizations:skills_section_title"
        #   writes:
        #     - "artifact:template_metadata[\"skills_section_title\"]"
        #   notes: "Copies the user-selected skills title into the preview HTML render metadata."
        cv_data['template_metadata']['skills_section_title'] = customizations.get('skills_section_title', 'Skills')
        cv_data['base_font_size'] = customizations.get(
            'base_font_size',
                get_config().get('generation.base_font_size', cv_data.get('base_font_size', '13px')),
        )
        cv_data['page_margin']    = customizations.get(
            'page_margin',
            get_config().get('generation.page_margin', '0.5in'),
        )

        template_file = self._CV_TEMPLATE_FILE
        if not template_file.exists():
            raise FileNotFoundError(f"HTML template not found: {template_file}")

        from .template_renderer import load_template, render_template  # noqa: PLC0415
        template = load_template(str(template_file))
        return render_template(template, cv_data)

    def generate_final_from_confirmed_html(
        self,
        confirmed_html: str,
        output_dir: Path,
        filename_base: str = "CV_final",
        preferred_renderer: str = 'auto',
    ) -> Dict:
        """Write confirmed HTML to disk and regenerate the human-readable PDF.

        Called by ``POST /api/cv/generate-final`` after layout confirmation.
        The confirmed HTML (which may have had layout instructions applied) is
        written to ``output_dir/{filename_base}.html`` and converted to a PDF
        via WeasyPrint.  ATS DOCX is not regenerated here because it is derived
        from structured data, not from HTML layout.

        Returns:
            dict with keys ``html`` and ``pdf`` (absolute path strings).
        """
        # duckflow:
        #   id: summary_orchestrator_final_files
        #   kind: artifact
        #   timestamp: "2026-03-27T01:23:28Z"
        #   status: shared
        #   reads: ["artifact:generation_state.preview_html"]
        #   writes: ["file:generated_files.final_html", "file:generated_files.final_pdf"]
        #   notes: "Commits the confirmed preview HTML to disk and regenerates the final PDF from that same artifact."
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        html_path = output_dir / f"{filename_base}.html"
        html_path.write_text(confirmed_html, encoding="utf-8")

        pdf_path = output_dir / f"{filename_base}.pdf"
        renderer_info = self._convert_html_to_pdf(
            html_path,
            pdf_path,
            preferred_renderer=preferred_renderer,
        )

        return {
            "html": str(html_path),
            "pdf":  str(pdf_path),
            "renderer": renderer_info["renderer"],
            "renderer_detail": renderer_info.get("detail", ""),
        }

    def generate_pdf_variants_from_html(
        self,
        confirmed_html: str,
        output_dir: Path,
        filename_base: str = "CV_preview",
        renderers: tuple[str, ...] = ('chrome', 'weasyprint'),
    ) -> Dict[str, Any]:
        """Write HTML once and attempt PDF generation for each requested renderer."""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        html_path = output_dir / f"{filename_base}.html"
        html_path.write_text(confirmed_html, encoding="utf-8")

        normalized_renderers = [
            str(renderer_name).strip().lower()
            for renderer_name in renderers
            if str(renderer_name).strip()
        ]
        pdfs: Dict[str, Dict[str, Any]] = {}

        def _render_variant(renderer_key: str) -> tuple[str, Dict[str, Any]]:
            pdf_path = output_dir / f"{filename_base}_{renderer_key}.pdf"
            try:
                renderer_info = self._convert_html_to_pdf(
                    html_path,
                    pdf_path,
                    preferred_renderer=renderer_key,
                )
                return renderer_key, {
                    'ok': True,
                    'pdf': str(pdf_path),
                    'renderer': renderer_info['renderer'],
                    'renderer_detail': renderer_info.get('detail', ''),
                    'error': None,
                }
            except Exception as exc:
                logger.warning(
                    "Preview PDF generation failed for %s (%s)",
                    renderer_key,
                    exc,
                )
                return renderer_key, {
                    'ok': False,
                    'pdf': str(pdf_path),
                    'renderer': renderer_key,
                    'renderer_detail': '',
                    'error': str(exc),
                }

        if normalized_renderers:
            max_workers = min(len(normalized_renderers), 4)
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_map = {
                    executor.submit(_render_variant, renderer_key): renderer_key
                    for renderer_key in normalized_renderers
                }
                for future in as_completed(future_map):
                    renderer_key, result = future.result()
                    pdfs[renderer_key] = result

        return {
            'html': str(html_path),
            'pdfs': pdfs,
        }

    def _render_cv_html_pdf(
        self,
        cv_data: Dict,
        output_dir: Path,
        filename_base: str,
        template_variant: str = 'standard'
    ) -> Path:
        """Render CV using the Jinja2 HTML template and convert to PDF.

        Uses `templates/cv-template.html` with the `cv_data` dictionary
        produced by ``_prepare_cv_data_for_template``. The rendered HTML is
        written to `output_dir` and then converted to PDF via WeasyPrint.

        Returns a 2-tuple ``(html_output, pdf_output)``.
        """
        
        template_file = self._CV_TEMPLATE_FILE
        if not template_file.exists():
            raise FileNotFoundError(f"HTML template not found: {template_file}")

        # Render using Jinja2
        from .template_renderer import load_template, render_template
        template = load_template(str(template_file))

        # Render the template with cv_data context
        rendered_html = render_template(template, cv_data)

        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)

        # Write HTML file to output directory
        html_output = output_dir / f"{filename_base}.html"
        with open(html_output, 'w', encoding='utf-8') as f:
            f.write(rendered_html)

        # Convert HTML to PDF
        pdf_output = output_dir / f"{filename_base}.pdf"
        self._convert_html_to_pdf(html_output, pdf_output)

        return html_output, pdf_output
    
    def _render_with_quarto_engine(self, template_file: Path, work_dir: Path) -> Path:
        """Render template using Quarto engine."""         
        html_output = work_dir / f"{template_file.stem}.html"
        
        try:
            # Render to HTML
            render_cmd = [
                'quarto', 'render', str(template_file),
                '--to', 'html',
                '--output', str(html_output)
            ]
            
            subprocess.run(
                render_cmd, 
                capture_output=True, 
                text=True, 
                check=True, 
                cwd=work_dir,
                timeout=60
            )
            
            if not html_output.exists():
                raise FileNotFoundError(f"Quarto render succeeded but HTML output not found: {html_output}")
            
            logger.info("Quarto render successful: %s", html_output.name)
            return html_output
            
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning("Quarto render failed: %s", e)
            return self._create_fallback_html_file(work_dir, template_file.stem)
    
    def _create_fallback_html_file(self, work_dir: Path, base_name: str) -> Path:
        """Create fallback HTML file when Quarto is unavailable.

        Renders cv-template.html via Jinja2 so the output matches the primary
        generation path visually.  Falls back to _create_fallback_html() only
        when the Jinja2 render itself fails.
        """
        html_output = work_dir / f"{base_name}.html"

        data_file = work_dir / 'temp_cv_data.json'
        if data_file.exists():
            with open(data_file, 'r', encoding='utf-8') as f:
                cv_data = json.load(f)
        else:
            cv_data = {'personal_info': {'name': 'CV Data Error'}, 'professional_summary': 'Data loading failed'}

        template_file = self._CV_TEMPLATE_FILE
        html_content = None
        if template_file.exists():
            try:
                from .template_renderer import load_template, render_template  # noqa: PLC0415
                template = load_template(str(template_file))
                html_content = render_template(template, cv_data)
            except Exception as exc:
                logger.warning("Jinja2 render failed in Quarto fallback (%s); using simple HTML", exc)

        if not html_content:
            html_content = self._create_fallback_html(cv_data)

        html_output.write_text(html_content, encoding='utf-8')
        logger.info("Created fallback HTML: %s", html_output.name)
        return html_output

    def _create_fallback_html(self, cv_data: Dict) -> str:
        """Create basic HTML if Quarto is not available."""
        personal_info = cv_data['personal_info']
        contact = personal_info.get('contact', {})
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html><head>',
            '<meta charset="UTF-8">',
            '<link rel="stylesheet" href="cv-style.css">',
            '<title>CV</title>',
            '</head><body>',
            '<div class="cv-container">',
            '<div class="cv-header">',
            f'<h1>{personal_info.get("name", "")}</h1>',
            f'<h2>{cv_data["professional_summary"]}</h2>',
            '<div class="contact-info">',
            f'{contact.get("email", "")} | {contact.get("phone", "")} | {contact.get("address_display", "")}',
            '</div></div>',
            '<div class="cv-body">',
            '<div class="cv-left-column">',
            '<h2>Professional Experience</h2>'
        ]
        
        # Add experiences
        for exp in cv_data['experiences']:
            location = exp.get('location', {})
            location_str = location.get('city', '')
            if location.get('state'):
                location_str += f", {location['state']}"
                
            html_parts.extend([
                '<div class="experience-item">',
                f'<h3>{exp.get("company", "")} | {exp.get("title", "")}</h3>',
                '<div class="experience-meta">',
                f'{location_str} | {exp.get("start_date", "")} - {exp.get("end_date", "")}',
                '</div>'
            ])
            
            if exp.get('achievements'):
                for achievement in exp['achievements']:
                    html_parts.append(f'<p>• {achievement.get("text", "")}</p>')
            
            html_parts.append('</div>')
        
        # Add education
        html_parts.append('<h2>Education</h2>')
        for edu in cv_data['education']:
            location = edu.get('location', {})
            location_str = f"{location.get('city', '')}, {location.get('state', '')}"
            html_parts.extend([
                '<div class="education-item">',
                f'<h3>{edu.get("degree", "")} {edu.get("field", "")}</h3>',
                f'<p><strong>{edu.get("institution", "")}</strong> | {location_str} | {edu.get("end_year", "")}</p>',
                '</div>'
            ])
        
        html_parts.extend([
            '</div>',  # cv-left-column
            '<div class="cv-right-column">',
            '<h2>Core Skills</h2>'
        ])
        
        # Add skills
        for category_data in cv_data['skills_by_category']:
            html_parts.extend([
                '<div class="skills-category">',
                f'<h3>{category_data["category"]}</h3>'
            ])
            for skill in category_data['skills']:
                years_text = f" ({skill['years']} years)" if skill.get('years') else ""
                html_parts.append(f'<p>• {skill["name"]}{years_text}</p>')
            html_parts.append('</div>')
            
        html_parts.extend([
            '</div>',  # cv-right-column
            '</div>',  # cv-body
            '</div>',  # cv-container
            '</body></html>'
        ])
        
        return '\n'.join(html_parts)
    
    def _convert_html_to_pdf(
        self,
        html_file: Path,
        pdf_output: Path,
        preferred_renderer: str = 'auto',
    ) -> Dict[str, str]:
        """Convert HTML file to PDF.

        Chrome/Chromium headless is the primary renderer (--headless=new mode,
        Chrome 112+).  Supports CSS paged media including @page margin boxes,
        page numbers, and proper print layout.

        Falls back to WeasyPrint in a child subprocess (crash-safe) for
        environments where Chrome is unavailable (e.g. headless Linux servers),
        then to a plain-text instruction file as last resort.
        """
        renderer_mode = (preferred_renderer or 'auto').strip().lower()
        if renderer_mode not in {'auto', 'chrome', 'weasyprint'}:
            raise ValueError(
                "preferred_renderer must be one of: auto, chrome, weasyprint"
            )

        # --- Chrome/Chromium headless (primary) ---
        # Try known binary locations: Linux paths first, then macOS app bundles.
        _chrome_candidates = [
            'google-chrome',
            'chromium',
            'chromium-browser',
            '/Applications/Google Chrome.app/Contents/MacOS/Google Chrome',
            '/Applications/Chromium.app/Contents/MacOS/Chromium',
        ]
        html_url = html_file.as_uri()   # file:///absolute/path/to/file.html
        def _try_chrome() -> Dict[str, str]:
            chrome_err_local = None
            chrome_exc_local = None
            for _chrome_bin in _chrome_candidates:
                try:
                    subprocess.run(
                        [
                            _chrome_bin,
                            '--headless=new',
                            '--disable-gpu',
                            '--no-sandbox',
                            f'--print-to-pdf={pdf_output}',
                            '--print-to-pdf-no-header',
                            '--no-pdf-header-footer',
                            html_url,
                        ],
                        check=True,
                        capture_output=True,
                        timeout=60,
                    )
                    logger.info(
                        "Generated PDF via Chrome (%s): %s",
                        Path(_chrome_bin).name,
                        pdf_output.name,
                    )
                    return {
                        'renderer': 'chrome',
                        'detail': str(_chrome_bin),
                        'success': True,
                        'error': None,
                        'fallback_used': False,
                    }
                except FileNotFoundError:
                    continue
                except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
                    chrome_err_local = str(exc)
                    chrome_exc_local = exc
                    break

            if chrome_err_local:
                raise PDFRenderingError(
                    f"Chrome headless failed: {chrome_err_local}"
                ) from chrome_exc_local
            raise PDFRendererNotFoundError('Chrome/Chromium not found')

        def _try_weasyprint() -> Dict[str, str]:
            wp_render = (
                Path(__file__).parent / 'wp_render.py'
            )
            fonts_dir = Path(__file__).parent.parent.parent / 'web' / 'fonts'
            cmd = [
                sys.executable, str(wp_render),
                str(html_file), str(pdf_output),
            ]
            if fonts_dir.is_dir():
                cmd.append(str(fonts_dir))
            wp_result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=120,
            )
            if wp_result.returncode == 0:
                logger.info("Generated PDF using WeasyPrint: %s", pdf_output.name)
                return {
                    'renderer': 'weasyprint',
                    'detail': sys.executable,
                    'success': True,
                    'error': None,
                    'fallback_used': False,
                }

            wp_error_local = (
                wp_result.stderr.decode(errors='replace').strip()
                or f"exit {wp_result.returncode}"
            )
            raise PDFRenderingError(f"WeasyPrint failed: {wp_error_local}")

        chrome_err = None
        wp_error = None

        if renderer_mode == 'chrome':
            return _try_chrome()

        if renderer_mode == 'weasyprint':
            return _try_weasyprint()

        try:
            return _try_chrome()
        except PDFRendererNotFoundError:
            chrome_err = 'not found'
            logger.warning("Chrome/Chromium not found, trying WeasyPrint...")
        except PDFRenderingError as exc:
            chrome_err = str(exc)
            logger.warning("Chrome headless failed (%s), trying WeasyPrint...", chrome_err)
            logger.debug("Chrome headless full error:", exc_info=True)

        try:
            return _try_weasyprint()
        except PDFRenderingError as exc:
            wp_error = str(exc)
            logger.warning("WeasyPrint also failed (%s)", wp_error)
            logger.debug("WeasyPrint full error:", exc_info=True)

        # --- Plain-text fallback ---
        fallback_content = f"""PDF Generation Failed

The system attempted to generate a PDF but encountered issues:
1. Chrome headless: {chrome_err or 'not found'}
2. WeasyPrint error: {wp_error}

To manually create PDF:
1. Open the HTML file: {html_file}
2. Print to PDF using your browser
3. Save as: {pdf_output}

The HTML file contains your formatted CV ready for conversion.
"""
        pdf_output.write_text(fallback_content.strip(), encoding='utf-8')
        logger.warning("Created fallback instructions: %s", pdf_output.name)
        return {
            'renderer': 'fallback-text',
            'detail': pdf_output.name,
            'success': False,
            'error': f"Chrome: {chrome_err or 'not found'}; WeasyPrint: {wp_error}",
            'fallback_used': True,
        }

    def _generate_human_pdf(
        self,
        cv_data: Dict,
        job_analysis: Dict,
        output_dir: Path,
        template_variant: str = 'standard'
    ) -> tuple:
        """Render the human-readable HTML template and convert to PDF.

        ``cv_data`` must already be prepared via ``_prepare_cv_data_for_template``
        (done once in ``generate_cv`` and shared across all format generators).

        Returns a 2-tuple ``(html_path, pdf_path)``.  On failure the first
        element is ``None`` and the second is an error-log ``.txt`` file.
        """
        company       = job_analysis.get('company', 'Company').replace(' ', '')
        role          = job_analysis.get('title', 'Role').replace(' ', '')[:20]
        timestamp     = datetime.now().strftime("%Y-%m-%d")
        filename_base = f"CV_{company}_{role}_{timestamp}"

        try:
            # HTML is rendered and written once; PDF is converted from that file
            html_path, pdf_path = self._render_cv_html_pdf(
                cv_data, output_dir, filename_base, template_variant
            )

            logger.info("Generated human-readable HTML + PDF (%s): %s", template_variant, pdf_path.name)
            logger.debug(
                "_generate_pdf: completed (template=%s, pdf_size=%d bytes)",
                template_variant, pdf_path.stat().st_size if pdf_path.exists() else 0
            )
            return html_path, pdf_path
            
        except Exception as e:
            logger.warning("PDF generation failed: %s", e)
            # Create enhanced fallback with diagnostic info
            fallback_file = output_dir / f"{filename_base}.txt"
            fallback_content = f"""PDF Generation Error: {e}

This indicates an issue with the document generation pipeline.
Please check:
1. Quarto installation (quarto --version)
2. WeasyPrint dependencies (pip install weasyprint)
3. Template files in templates/ directory
4. Chrome/Chromium for PDF rendering

Template variant attempted: {template_variant}
Content data summary:
- Experiences: {len(cv_data.get('experiences', []))} items
- Skills: {len(cv_data.get('skills', []))} items
- Education: {len(cv_data.get('education', []))} items

For manual generation:
1. Check system requirements
2. Retry with different template variant
3. Use browser Print to PDF as fallback
            """
            fallback_file.write_text(fallback_content.strip(), encoding='utf-8')
            logger.warning("Created error log: %s", fallback_file.name)
            return None, fallback_file

    def _build_json_ld(self, cv_data: Dict, job_analysis: Dict) -> str:
        """Build a Schema.org/Person JSON-LD string from prepared ``cv_data``.

        The result is embedded in the HTML ``<head>`` so that ATS parsers,
        search engines, and other structured-data consumers can extract
        candidate information without parsing the visual layout.
        """
        personal_info = cv_data['personal_info']
        contact       = personal_info.get('contact', {})

        # Work history
        has_occupation: List[Dict] = []
        for exp in cv_data.get('experiences', []):
            entry: Dict[str, Any] = {
                '@type':     'Role',
                'name':      exp.get('company', ''),
                'roleName':  exp.get('title', ''),
                'startDate': exp.get('start_date', ''),
                'endDate':   exp.get('end_date', ''),
            }
            loc = exp.get('location', {})
            if loc:
                entry['locationCreated'] = {
                    '@type': 'Place',
                    'address': {
                        '@type':           'PostalAddress',
                        'addressLocality': loc.get('city', ''),
                        'addressRegion':   loc.get('state', ''),
                    }
                }
            ach_texts = [
                (ac.get('text') if isinstance(ac, dict) else ac)
                for ac in exp.get('achievements', [])
            ]
            if ach_texts:
                entry['description'] = ' '.join(filter(None, ach_texts))
            has_occupation.append(entry)

        # Education
        alumni_of = [
            {
                '@type': 'EducationalOrganization',
                'name':  edu.get('institution', ''),
                'description': (
                    edu.get('degree', '')
                    + (f", {edu.get('field', '')}" if edu.get('field') else '')
                    + (f" ({edu.get('end_year') or edu.get('graduation_date', '')})"
                       if (edu.get('end_year') or edu.get('graduation_date')) else '')
                ),
            }
            for edu in cv_data.get('education', [])
        ]

        all_skill_entries = [
            {
                '@type':          'DefinedTerm',
                'name':           sk.get('name', ''),
                'additionalType': 'HardSkill' if self._classify_skill_type(sk) == 'hard' else 'SoftSkill',
            }
            for cat in cv_data.get('skills_by_category', [])
            for sk in cat.get('skills', [])
            if sk.get('name', '')
        ]

        award_strings = [
            f"{aw.get('degree') or aw.get('title', '')} ({aw.get('year', '')})"
            for aw in cv_data.get('awards', [])
        ]

        same_as = [
            value
            for value in (
                safe_url(contact.get('linkedin_href') or contact.get('linkedin', '')),
                safe_url(contact.get('website_href') or contact.get('website', '')),
            )
            if value
        ]

        json_ld: Dict[str, Any] = {
            '@context':   'https://schema.org',
            '@type':      'Person',
            'name':       personal_info.get('name', ''),
            'jobTitle':   job_analysis.get('title', ''),
            'description': cv_data.get('professional_summary', ''),
        }
        if contact.get('email'):
            json_ld['email'] = contact['email']
        if contact.get('phone'):
            json_ld['telephone'] = contact['phone']
        if same_as:
            json_ld['sameAs'] = same_as
        if contact.get('address_display'):
            json_ld['address'] = {
                '@type': 'PostalAddress',
                'addressLocality': contact['address_display'],
            }
        if alumni_of:
            json_ld['alumniOf'] = alumni_of
        if has_occupation:
            json_ld['hasOccupation'] = has_occupation
        if all_skill_entries:
            json_ld['knowsAbout'] = all_skill_entries
        if award_strings:
            json_ld['award'] = award_strings

        self._validate_json_ld(json_ld)
        return json.dumps(json_ld, indent=2, ensure_ascii=False)

    _JSON_LD_REQUIRED_FIELDS: List[str] = ['@context', '@type', 'name']

    def _validate_json_ld(self, json_ld: Dict[str, Any]) -> None:
        """Warn when *json_ld* is missing or has empty required Schema.org fields.

        Logs a ``WARNING`` for each absent or empty required field so that
        callers can detect silently-invalid structured-data output without
        raising an exception at generation time.
        """
        for field in self._JSON_LD_REQUIRED_FIELDS:
            if not json_ld.get(field):
                logger.warning(
                    "JSON-LD validation: required field %r is absent or empty.", field
                )

    # ── Rewrite pipeline ─────────────────────────────────────────────────────

    def propose_rewrites(
        self,
        content: Dict,
        job_analysis: Dict,
        conversation_history: List = None,
        user_preferences: Dict = None,
    ) -> List[Dict]:
        """Propose targeted text rewrites to align CV terminology with the job.

        Delegates to the LLM provider's ``propose_rewrites`` implementation.
        Returns ``[]`` (with a logged warning) when no LLM client is configured
        so the caller can degrade gracefully.

        Args:
            content:              Selected CV content dict from
                                  :meth:`_select_content_hybrid`.
            job_analysis:         Output of the LLM job-description analysis.
            conversation_history: Full chat history for additional context.
            user_preferences:     Post-analysis Q&A answers.

        Returns:
            List of rewrite proposals (see :meth:`LLMClient.propose_rewrites`
            for the full schema).  Always ``[]`` on failure or missing LLM.
        """
        if not self.llm:
            logger.warning(
                "propose_rewrites called but no LLM client is configured. "
                "Returning empty proposals."
            )
            return []
        return self.llm.propose_rewrites(content, job_analysis, conversation_history, user_preferences)

    def apply_approved_rewrites(
        self, content: Dict, approved: List[Dict]
    ) -> Dict:
        """Apply a list of user-approved rewrite proposals to *content*.

        Each approved item specifies a ``location``, ``original`` text, and
        ``proposed`` replacement.  Items are applied individually; any item
        that fails :func:`LLMClient.apply_rewrite_constraints` is skipped
        (with a logged warning) rather than raising an exception.

        Supported rewrite types
        -----------------------
        ``summary``
            Replaces ``content['summary']``.
        ``bullet``
            Resolves ``location`` of the form ``"exp_ID.achievements[N]"``
            and updates the corresponding achievement's ``text`` field.
        ``skill_rename``
            Finds the skill whose name matches ``original`` and renames it
            to ``proposed``.
        ``skill_add``
            Appends a new skill dict to ``content['skills']``.  When
            ``evidence_strength == "weak"`` the entry is also flagged with
            ``candidate_to_confirm: True``.

        Args:
            content:  CV content dict (not mutated — a deep copy is made).
            approved: List of approved rewrite dicts.

        Returns:
            A new content dict with all valid approved rewrites applied.
        """
        result = copy.deepcopy(content)

        logger.debug(
            "apply_approved_rewrites: processing %d rewrite(s)", len(approved)
        )

        for item in approved:
            loc      = item.get('location', '')
            original = item.get('original', '')
            proposed = item.get('proposed', '')
            kind     = item.get('type', '')
            item_id  = item.get('id', '<unknown>')

            # Guard: validate constraint — skip if numbers/dates/names lost.
            if not LLMClient.apply_rewrite_constraints(original, proposed):
                logger.warning(
                    "apply_approved_rewrites: skipping constraint violation "
                    "(id=%r) — protected tokens would be removed.",
                    item_id
                )
                continue

            if kind == 'summary' or loc == 'summary':
                logger.debug(
                    "apply_approved_rewrites: summary rewrite (id=%s, len_before=%d, len_after=%d)",
                    item_id, len(original), len(proposed)
                )
                result['summary'] = proposed

            elif kind == 'bullet':
                # Parse "exp_001.achievements[2]"
                m = re.match(r'^([^.]+)\.achievements\[(\d+)\]$', loc)
                if not m:
                    logger.warning(
                        "apply_approved_rewrites: cannot parse bullet location %r (id=%r)",
                        loc, item_id
                    )
                    continue
                exp_id  = m.group(1)
                ach_idx = int(m.group(2))
                found   = False
                for exp in result.get('experiences', []):
                    if exp.get('id') == exp_id:
                        found = True
                        achs  = exp.get('achievements', [])
                        if 0 <= ach_idx < len(achs):
                            ach = achs[ach_idx]
                            if isinstance(ach, dict):
                                old_text = ach.get('text', '')
                                ach['text'] = proposed
                                logger.debug(
                                    "apply_approved_rewrites: bullet rewrite "
                                    "(id=%s, exp=%s, idx=%d, len_before=%d, len_after=%d)",
                                    item_id, exp_id, ach_idx, len(old_text), len(proposed)
                                )
                            else:
                                achs[ach_idx] = proposed
                                logger.debug(
                                    "apply_approved_rewrites: bullet rewrite (id=%s, exp=%s, idx=%d)",
                                    item_id, exp_id, ach_idx
                                )
                        else:
                            logger.warning(
                                "apply_approved_rewrites: achievement index %d out of range "
                                "for exp %r (id=%r)", ach_idx, exp_id, item_id
                            )
                        break
                if not found:
                    logger.warning(
                        "apply_approved_rewrites: experience %r not found (id=%r)",
                        exp_id, item_id
                    )

            elif kind == 'skill_rename':
                skills  = result.get('skills', [])
                renamed = False
                if isinstance(skills, list):
                    for i, skill in enumerate(skills):
                        if isinstance(skill, dict) and skill.get('name') == original:
                            skill['name'] = proposed
                            renamed = True
                            break
                        elif isinstance(skill, str) and skill == original:
                            skills[i] = proposed
                            renamed = True
                            break
                elif isinstance(skills, dict):
                    for cat_data in skills.values():
                        cat_list = (
                            cat_data.get('skills', [])
                            if isinstance(cat_data, dict)
                            else cat_data
                            if isinstance(cat_data, list)
                            else []
                        )
                        for i, skill in enumerate(cat_list):
                            if isinstance(skill, dict) and skill.get('name') == original:
                                skill['name'] = proposed
                                renamed = True
                                break
                            elif isinstance(skill, str) and skill == original:
                                cat_list[i] = proposed
                                renamed = True
                                break
                        if renamed:
                            break
                if not renamed:
                    logger.warning(
                        "apply_approved_rewrites: skill_rename: original name %r not found (id=%r)",
                        original, item_id
                    )

            elif kind == 'skill_add':
                new_skill: Dict = {
                    'name':                proposed,
                    'candidate_to_confirm': item.get('evidence_strength') == 'weak',
                    'evidence':            item.get('evidence', ''),
                }
                skills = result.get('skills', [])
                if isinstance(skills, list):
                    skills.append(new_skill)
                elif isinstance(skills, dict):
                    first_cat = next(iter(skills.values()), None)
                    if isinstance(first_cat, dict):
                        first_cat.setdefault('skills', []).append(new_skill)
                    elif isinstance(first_cat, list):
                        first_cat.append(new_skill)

            else:
                logger.warning(
                    "apply_approved_rewrites: unknown rewrite type %r (id=%r), skipping.",
                    kind, item_id
                )

        return result

    def apply_accepted_spell_fixes(
        self, content: Dict, spell_audit: List[Dict]
    ) -> Dict:
        """Apply accepted spell-check fixes to the selected content.

        Accepted fixes are grouped by ``section_id`` and applied against the
        exact span that LanguageTool flagged. Offsets are processed in reverse
        order so multiple fixes in the same section do not shift one another.
        """
        result = copy.deepcopy(content)
        accepted_by_section: Dict[str, List[Dict]] = defaultdict(list)

        for item in spell_audit or []:
            if item.get('outcome') != 'accept':
                continue
            section_id = (item.get('section_id') or '').strip()
            replacement = item.get('final') or item.get('suggestion') or ''
            if not section_id or not replacement:
                continue
            accepted_by_section[section_id].append(item)

        for section_id, fixes in accepted_by_section.items():
            if section_id == 'summary':
                summary_text = result.get('summary', '')
                result['summary'] = self._apply_spell_fixes_to_text(summary_text, fixes)
                continue

            match = re.match(r'^selected_ach_(\d+)$', section_id)
            if match:
                ach_idx = int(match.group(1))
                achievements = result.get('achievements') or []
                self._apply_spell_fixes_to_list_item(achievements, ach_idx, fixes)
                continue

            match = re.match(r'^exp_(.+)_ach_(\d+)$', section_id)
            if not match:
                match = re.match(r'^skill_(\d+)$', section_id)
                if match:
                    skills = result.get('skills') or []
                    self._apply_spell_fixes_to_skill(skills, int(match.group(1)), fixes)
                    continue

                match = re.match(r'^edu_(\d+)_(degree|field|institution)$', section_id)
                if match:
                    self._apply_spell_fixes_to_named_field(
                        result.get('education') or [],
                        int(match.group(1)),
                        match.group(2),
                        fixes,
                    )
                    continue

                match = re.match(r'^award_(\d+)_title$', section_id)
                if match:
                    awards = result.get('awards') or []
                    award_idx = int(match.group(1))
                    if 0 <= award_idx < len(awards):
                        award = awards[award_idx]
                        field_name = 'degree' if isinstance(award, dict) and award.get('degree') else 'title'
                        self._apply_spell_fixes_to_named_field(awards, award_idx, field_name, fixes)
                    continue

                match = re.match(r'^cert_(\d+)_(name|issuer)$', section_id)
                if match:
                    self._apply_spell_fixes_to_named_field(
                        result.get('certifications') or [],
                        int(match.group(1)),
                        match.group(2),
                        fixes,
                    )
                    continue

                match = re.match(r'^lang_(\d+)(?:_(language|proficiency))?$', section_id)
                if match:
                    languages = result.get('personal_info', {}).get('languages') or []
                    lang_idx = int(match.group(1))
                    field_name = match.group(2)
                    self._apply_spell_fixes_to_language(languages, lang_idx, field_name, fixes)
                    continue

                match = re.match(r'^pub_(\d+)_(formatted|title|authors|journal|booktitle)$', section_id)
                if match:
                    publications = result.get('publications') or []
                    pub_idx = int(match.group(1))
                    field_name = match.group(2)
                    self._apply_spell_fixes_to_named_field(publications, pub_idx, field_name, fixes)
                    continue

                logger.warning(
                    "apply_accepted_spell_fixes: cannot parse section id %r",
                    section_id
                )
                continue

            exp_id = match.group(1)
            ach_idx = int(match.group(2))
            exp = next(
                (item for item in result.get('experiences', []) if item.get('id') == exp_id),
                None,
            )
            if exp is None:
                logger.warning(
                    "apply_accepted_spell_fixes: experience %r not found for section %r",
                    exp_id, section_id
                )
                continue

            for key in ('ordered_achievements', 'achievements'):
                achievements = exp.get(key) or []
                if not (0 <= ach_idx < len(achievements)):
                    continue
                achievement = achievements[ach_idx]
                current_text = (
                    achievement.get('text', '')
                    if isinstance(achievement, dict)
                    else str(achievement)
                )
                updated_text = self._apply_spell_fixes_to_text(current_text, fixes)
                if isinstance(achievement, dict):
                    achievement['text'] = updated_text
                else:
                    achievements[ach_idx] = updated_text

        return result

    def _apply_spell_fixes_to_list_item(
        self, items: List[Any], item_idx: int, fixes: List[Dict]
    ) -> None:
        """Apply accepted spell fixes to a list item with optional ``text`` field."""
        if not (0 <= item_idx < len(items)):
            return
        item = items[item_idx]
        current_text = item.get('text', '') if isinstance(item, dict) else str(item)
        updated_text = self._apply_spell_fixes_to_text(current_text, fixes)
        if isinstance(item, dict):
            item['text'] = updated_text
        else:
            items[item_idx] = updated_text

    def _apply_spell_fixes_to_skill(
        self, skills: List[Any], skill_idx: int, fixes: List[Dict]
    ) -> None:
        """Apply accepted spell fixes to a skill name."""
        if not (0 <= skill_idx < len(skills)):
            return
        skill = skills[skill_idx]
        current_text = skill.get('name', '') if isinstance(skill, dict) else str(skill)
        updated_text = self._apply_spell_fixes_to_text(current_text, fixes)
        if isinstance(skill, dict):
            skill['name'] = updated_text
        else:
            skills[skill_idx] = updated_text

    def _apply_spell_fixes_to_named_field(
        self, items: List[Any], item_idx: int, field_name: str, fixes: List[Dict]
    ) -> None:
        """Apply accepted spell fixes to a specific named field on a list item."""
        if not (0 <= item_idx < len(items)):
            return
        item = items[item_idx]
        if not isinstance(item, dict) or field_name not in item:
            return
        item[field_name] = self._apply_spell_fixes_to_text(str(item.get(field_name, '')), fixes)

    def _apply_spell_fixes_to_language(
        self, languages: List[Any], lang_idx: int, field_name: Optional[str], fixes: List[Dict]
    ) -> None:
        """Apply accepted spell fixes to language entries."""
        if not (0 <= lang_idx < len(languages)):
            return
        item = languages[lang_idx]
        if isinstance(item, dict):
            target_field = field_name or ('language' if 'language' in item else None)
            if target_field and target_field in item:
                item[target_field] = self._apply_spell_fixes_to_text(str(item.get(target_field, '')), fixes)
        else:
            languages[lang_idx] = self._apply_spell_fixes_to_text(str(item), fixes)

    @staticmethod
    def _apply_spell_fixes_to_text(text: str, fixes: List[Dict]) -> str:
        """Apply accepted spell fixes to a single text fragment."""
        if not text:
            return text

        updated = text
        sortable_fixes = []
        for item in fixes or []:
            try:
                offset = int(item.get('offset'))
                length = int(item.get('length'))
            except (TypeError, ValueError):
                continue
            sortable_fixes.append((offset, length, item))

        for offset, length, item in sorted(sortable_fixes, key=lambda row: row[0], reverse=True):
            replacement = item.get('final') or item.get('suggestion') or ''
            if not replacement:
                continue
            if offset < 0 or length < 0 or offset + length > len(updated):
                continue

            original = item.get('original', '')
            current_span = updated[offset:offset + length]
            if original and current_span != original:
                continue

            updated = updated[:offset] + replacement + updated[offset + length:]

        return updated

    # ── CV generation ─────────────────────────────────────────────────────────

    def generate_cv(
        self,
        job_analysis: Dict,
        customizations: Dict,
        output_dir: Optional[Path] = None,
        approved_rewrites: Optional[List[Dict]] = None,
        rewrite_audit: Optional[List[Dict]] = None,
        spell_audit: Optional[List[Dict]] = None,
        max_skills: Optional[int] = None,
        max_achievements: Optional[int] = None,
        max_publications: Optional[int] = None,
    ) -> Dict:
        """
        Generate CV files based on LLM analysis and recommendations.

        Parameters
        ----------
        job_analysis:
            Output of :meth:`LLMClient.analyze_job_description`.
        customizations:
            Output of :meth:`LLMClient.recommend_customizations`.
        output_dir:
            When provided (e.g. the already-renamed session directory) the CV
            files are written there.  Otherwise a new
            ``{Company}_{RoleSlug}_{date}`` directory is created under
            ``self.output_dir``.
        approved_rewrites:
            Optional list of user-approved rewrite proposals produced by
            :meth:`propose_rewrites`.  Each item is applied via
            :meth:`apply_approved_rewrites` before content is rendered.
            Defaults to ``[]`` (no rewrites) when ``None``.

        Returns
        -------
        Dict with output_dir, files created, metadata
        """
        company   = job_analysis.get('company', 'Company')
        role      = job_analysis.get('title', 'Role')
        role_slug = role.replace(' ', '')[:20]
        timestamp = datetime.now().strftime("%Y-%m-%d")

        if output_dir is not None:
            job_output_dir = Path(output_dir)
        else:
            output_name    = f"{company}_{role_slug}_{timestamp}"
            job_output_dir = self.output_dir / output_name
        job_output_dir.mkdir(parents=True, exist_ok=True)
        # Read existing run counter so re-generations are numbered sequentially
        _prev_run = 0
        _prev_meta_file = job_output_dir / 'metadata.json'
        if _prev_meta_file.exists():
            try:
                with open(_prev_meta_file, encoding='utf-8') as _f:
                    _prev_run = int(json.load(_f).get('generation_run', 0))
            except Exception:  # noqa: BLE001
                pass
        generation_run = _prev_run + 1

        logger.info("Output directory: %s", job_output_dir)
        logger.debug(
            "generate_cv: entry (company=%s, role=%s, max_skills=%s, "
            "rewrites=%d, spell_audit=%d)",
            company, role, max_skills,
            len(approved_rewrites or []), len(spell_audit or [])
        )
        
        selected_content = self.build_render_ready_content(
            job_analysis,
            customizations,
            approved_rewrites=approved_rewrites,
            spell_audit=spell_audit,
            max_skills=max_skills,
            max_achievements=max_achievements,
            max_publications=max_publications,
        )

        date_overlap_warnings = self._detect_date_overlaps(
            selected_content.get('experiences', [])
        )
        long_bullet_warnings = self._detect_long_bullets(
            selected_content.get('experiences', [])
        )
        sparse_experience_warnings = self._detect_sparse_experiences(
            selected_content.get('experiences', [])
        )
        year_only_date_warnings = self._detect_year_only_dates(
            selected_content.get('experiences', [])
        )
        rewrite_audit_mismatches = self._verify_rewrite_audit_alignment(
            selected_content,
            rewrite_audit or [],
        )
        if date_overlap_warnings:
            logger.warning(
                "Employment date overlaps detected (%d): %s",
                len(date_overlap_warnings),
                '; '.join(
                    f"{w['entry_a']} / {w['entry_b']}" for w in date_overlap_warnings
                ),
            )

        # Prepare template data once — shared by all format generators.
        # JSON-LD is built here and embedded directly in cv-template.html,
        # so the single HTML output is both ATS-compatible and print-ready.
        cv_data = self._prepare_cv_data_for_template(
            selected_content,
            job_analysis,
            base_font_size=customizations.get('base_font_size'),
            customizations=customizations,
        )
        cv_data['achievements']    = selected_content.get('achievements', [])
        cv_data['ai_attribution']  = bool(customizations.get('ai_attribution', False))
        cv_data['json_ld_str']    = self._build_json_ld(cv_data, job_analysis)
        # duckflow:
        #   id: cv_render.scripts_utils_cv_orchestrator.L1684
        #   kind: artifact
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   reads:
        #     - "customizations:skills_section_title"
        #   writes:
        #     - "artifact:template_metadata[\"skills_section_title\"]"
        #   notes: "Copies the user-selected skills title into the final HTML/PDF render metadata."
        cv_data['template_metadata']['skills_section_title'] = customizations.get('skills_section_title', 'Skills')
        cv_data['base_font_size'] = customizations.get(
            'base_font_size',
            get_config().get('generation.base_font_size', cv_data.get('base_font_size', '13px')),
        )
        cv_data['page_margin']    = customizations.get(
            'page_margin',
            get_config().get('generation.page_margin', '0.5in'),
        )

        # Generate documents (Phase 10: Track progress)
        files_created = []
        generation_progress = []

        # 1. ATS-optimized DOCX
        progress_ats = {
            'step': 'generating_docx_ats',
            'status': 'in_progress',
            'start_time': time.time()
        }
        # duckflow:
        #   id: cv_render.scripts_utils_cv_orchestrator.L1703
        #   kind: artifact
        #   timestamp: "2026-03-27T02:07:47Z"
        #   status: live
        #   reads:
        #     - "customizations:skills_section_title"
        #   writes:
        #     - "artifact:selected_content[\"skills_section_title\"]"
        #   notes: "Carries the user-selected skills title into the ATS DOCX generation payload."
        selected_content['skills_section_title'] = customizations.get('skills_section_title', 'Skills')
        selected_content['ai_attribution'] = bool(customizations.get('ai_attribution', False))
        ats_file, ats_score_at_generation = self._generate_ats_docx(
            selected_content,
            job_analysis,
            job_output_dir
        )
        progress_ats['status'] = 'complete'
        progress_ats['elapsed_ms'] = int((time.time() - progress_ats['start_time']) * 1000)
        generation_progress.append(progress_ats)
        files_created.append(ats_file.name)

        # 2. Single HTML (ATS metadata embedded) + PDF both rendered from it
        progress_html = {
            'step': 'rendering_html',
            'status': 'in_progress',
            'start_time': time.time()
        }
        html_path, pdf_path = self._generate_human_pdf(
            cv_data,
            job_analysis,
            job_output_dir
        )
        progress_html['status'] = 'complete'
        progress_html['elapsed_ms'] = int((time.time() - progress_html['start_time']) * 1000)
        generation_progress.append(progress_html)
        if html_path is not None:
            files_created.append(html_path.name)
        files_created.append(pdf_path.name)

        # 3. Human-readable DOCX
        progress_docx_human = {
            'step': 'generating_docx_human',
            'status': 'in_progress',
            'start_time': time.time()
        }
        human_docx = self._generate_human_docx(
            cv_data,
            job_analysis,
            job_output_dir,
            skills_heading=self._resolve_human_skills_title(customizations),
        )
        progress_docx_human['status'] = 'complete'
        progress_docx_human['elapsed_ms'] = int((time.time() - progress_docx_human['start_time']) * 1000)
        generation_progress.append(progress_docx_human)
        files_created.append(human_docx.name)

        # Run full ATS validation report now that all output files exist.
        try:
            _ats_checks, _ats_page_count = validate_ats_report(job_output_dir, job_analysis)
        except Exception:
            _ats_checks, _ats_page_count = [], None

        # Save metadata
        metadata = {
            'generation_date': datetime.now().isoformat(),
            'generation_run':  generation_run,
            'company':         company,
            'role':            role,
            'job_analysis':    job_analysis,
            'customizations':  customizations,
            'approved_rewrites': approved_rewrites or [],
            'rewrite_audit':   rewrite_audit or [],
            'spell_audit':     spell_audit or [],
            'selected_content_summary': {
                'experiences_count': len(selected_content['experiences']),
                'skills_count': len(selected_content['skills']),
                'achievements_count': len(selected_content['achievements'])
            },
            'ats_score': ats_score_at_generation,
            'files_generated': files_created,
            'date_overlap_warnings': date_overlap_warnings,
            'long_bullet_warnings': long_bullet_warnings,
            'sparse_experience_warnings': sparse_experience_warnings,
            'year_only_date_warnings': year_only_date_warnings,
            'rewrite_audit_mismatches': rewrite_audit_mismatches,
            'summary_warnings': selected_content.get('summary_warnings', []),
            'publication_warnings': selected_content.get('publication_warnings', []),
            'ats_validation': {
                'checks': _ats_checks,
                'page_count': _ats_page_count,
                'summary': {
                    'pass': sum(1 for c in _ats_checks if c['status'] == 'pass'),
                    'warn': sum(1 for c in _ats_checks if c['status'] == 'warn'),
                    'fail': sum(1 for c in _ats_checks if c['status'] == 'fail'),
                },
            } if _ats_checks else None,
        }

        metadata_file = job_output_dir / 'metadata.json'
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        files_created.append('metadata.json')
        
        # Save job description
        if job_analysis.get('original_text'):
            job_desc_file = job_output_dir / 'job_description.txt'
            job_desc_file.write_text(job_analysis['original_text'], encoding='utf-8')
            files_created.append('job_description.txt')
        
        return {
            'output_dir': str(job_output_dir),
            'files': files_created,
            'metadata': metadata,
            'generation_progress': generation_progress,
        }

    def generate_preview_html_only(
        self,
        job_analysis: Dict,
        customizations: Dict,
        output_dir: Optional[Path] = None,
        approved_rewrites: Optional[List[Dict]] = None,
        spell_audit: Optional[List[Dict]] = None,
        max_skills: Optional[int] = None,
        max_achievements: Optional[int] = None,
        max_publications: Optional[int] = None,
    ) -> Dict:
        """Generate HTML preview only — no PDF, no DOCX.

        Used by the Alt-A workflow (Phase 6 preview step).  Creates the output
        directory and renders the CV Jinja2 template to an HTML file.  PDF and
        DOCX generation are deferred to :meth:`generate_final_from_confirmed_html`
        (Phase 8, via ``POST /api/cv/generate-final``).

        Returns
        -------
        Dict with output_dir, files (list with one HTML filename),
        html_path, and generation_progress.
        """
        company   = job_analysis.get('company', 'Company')
        role      = job_analysis.get('title', 'Role')
        role_slug = role.replace(' ', '')[:20]
        timestamp = datetime.now().strftime("%Y-%m-%d")

        if output_dir is not None:
            job_output_dir = Path(output_dir)
        else:
            output_name    = f"{company}_{role_slug}_{timestamp}"
            job_output_dir = self.output_dir / output_name
        job_output_dir.mkdir(parents=True, exist_ok=True)

        logger.info("Output directory (preview-only): %s", job_output_dir)

        selected_content = self.build_render_ready_content(
            job_analysis,
            customizations,
            approved_rewrites=approved_rewrites,
            spell_audit=spell_audit,
            max_skills=max_skills,
            max_achievements=max_achievements,
            max_publications=max_publications,
        )

        cv_data = self._prepare_cv_data_for_template(
            selected_content,
            job_analysis,
            base_font_size=customizations.get('base_font_size'),
            customizations=customizations,
        )
        cv_data['achievements']   = selected_content.get('achievements', [])
        cv_data['json_ld_str']    = self._build_json_ld(cv_data, job_analysis)
        cv_data['template_metadata']['skills_section_title'] = customizations.get('skills_section_title', 'Skills')
        cv_data['base_font_size'] = customizations.get(
            'base_font_size',
            get_config().get('generation.base_font_size', cv_data.get('base_font_size', '13px')),
        )
        cv_data['page_margin']    = customizations.get(
            'page_margin',
            get_config().get('generation.page_margin', '0.5in'),
        )

        # Render HTML template (no PDF conversion)
        template_file = self._CV_TEMPLATE_FILE
        if not template_file.exists():
            raise FileNotFoundError(f"HTML template not found: {template_file}")

        from .template_renderer import load_template, render_template
        template      = load_template(str(template_file))
        rendered_html = render_template(template, cv_data)

        filename_base = f"CV_{company}_{role_slug}_{timestamp}_preview"
        html_path     = job_output_dir / f"{filename_base}.html"
        html_path.write_text(rendered_html, encoding='utf-8')

        return {
            'output_dir':          str(job_output_dir),
            'files':               [html_path.name],
            'html_path':           str(html_path),
            'generation_progress': [{'step': 'rendering_html', 'status': 'complete'}],
        }

    def build_render_ready_content(
        self,
        job_analysis: Dict,
        customizations: Dict,
        approved_rewrites: Optional[List[Dict]] = None,
        spell_audit: Optional[List[Dict]] = None,
        max_skills: Optional[int] = None,
        max_achievements: Optional[int] = None,
        max_publications: Optional[int] = None,
        use_semantic_match: bool = True,
    ) -> Dict:
        """Build the selected content exactly as it will be rendered."""
        selected_content = self._select_content_hybrid(
            job_analysis,
            customizations,
            max_skills=max_skills,
            max_achievements=max_achievements,
            max_publications=max_publications,
            use_semantic_match=use_semantic_match,
        )
        selected_content = self._apply_session_achievement_edits(
            selected_content,
            customizations.get('achievement_edits') or {},
        )
        selected_content = self.apply_approved_rewrites(
            selected_content,
            approved_rewrites or [],
        )
        return self.apply_accepted_spell_fixes(
            selected_content,
            spell_audit or [],
        )

    def _serialize_html_for_context(self, html: str) -> str:
        """Convert HTML to human-readable outline for LLM context.

        Parses HTML and extracts section names, nesting, and item counts
        to create a concise structure description. Used to give LLM
        context about current CV layout without sending full HTML.

        Args:
            html: The HTML document to serialize

        Returns:
            Human-readable outline showing section structure and item counts
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, 'html.parser')
        outline = []

        h_tags = soup.find_all(['h2', 'h3'])
        total_li_count = len(soup.find_all('li'))

        for i, heading in enumerate(h_tags, 1):
            heading_text = heading.get_text(' ', strip=True)
            if heading_text:
                outline.append(f"{i}. {heading_text}")

        if total_li_count > 0:
            outline.append(f"\nTotal items: {total_li_count}")

        return '\n'.join(outline) if outline else "[No structured sections found]"

    def _sanitize_layout_instruction_text(self, instruction_text: str) -> Dict[str, Any]:
        """Strip prompt-injection phrases from layout instructions."""
        raw_text = str(instruction_text or '')
        sanitized_text, findings = sanitize_instruction_text(raw_text)
        sanitized_text = re.sub(r'\s+', ' ', sanitized_text).strip(' ,;:-')
        return {
            'flagged': bool(findings),
            'findings': findings,
            'raw_text': raw_text,
            'sanitized_text': sanitized_text,
        }

    def _sanitize_layout_context_html(self, html: str) -> Dict[str, Any]:
        """Remove prompt-payload material from HTML before sending it to the LLM."""
        raw_html = str(html or '')
        soup = BeautifulSoup(raw_html, 'html.parser')
        findings: List[Dict[str, Any]] = []

        for comment in soup.find_all(string=lambda value: isinstance(value, Comment)):
            text = str(comment)
            if scan_text_for_injection(text):
                _append_layout_finding(
                    findings,
                    'unsafe_context_comment',
                    'Removed prompt-like comment from layout context HTML.',
                    text,
                )
                comment.extract()

        for element in list(soup.find_all(True)):
            text = element.get_text(' ', strip=True)
            is_hidden = (
                element.has_attr('hidden')
                or element.get('aria-hidden') == 'true'
                or 'display:none' in element.get('style', '').replace(' ', '').lower()
            )
            if is_hidden and text and scan_text_for_injection(text):
                _append_layout_finding(
                    findings,
                    'unsafe_context_element',
                    'Removed prompt-like element from layout context HTML.',
                    text,
                )
                element.decompose()

        return {
            'flagged': bool(findings),
            'findings': findings,
            'raw_html': raw_html,
            'sanitized_html': str(soup),
        }

    def _sanitize_layout_instruction_html(
        self,
        current_html: str,
        modified_html: str,
    ) -> Dict[str, Any]:
        """Sanitize rewritten layout HTML and preserve safe baseline resources."""
        baseline = BeautifulSoup(str(current_html or ''), 'html.parser')
        soup = BeautifulSoup(str(modified_html or ''), 'html.parser')
        findings: List[Dict[str, Any]] = []

        baseline_head_tags = []
        if baseline.head:
            for node in baseline.head.find_all(True, recursive=False):
                if node.name not in _LAYOUT_PRESERVED_HEAD_TAGS:
                    continue
                if node.name == 'script' and node.get('type') == 'application/ld+json':
                    script_text = node.string or node.get_text('', strip=True)
                    try:
                        parsed = json.loads(script_text) if script_text else {}
                    except json.JSONDecodeError:
                        parsed = {}
                    if not _is_exact_schema_org_context(parsed.get('@context')):
                        continue
                baseline_head_tags.append(copy.deepcopy(node))

        if soup.head:
            for node in list(soup.head.find_all(True, recursive=False)):
                if node.name in _LAYOUT_PRESERVED_HEAD_TAGS:
                    _append_layout_finding(
                        findings,
                        'rewritten_head_replaced',
                        'Replaced rewritten head resources with baseline-safe versions.',
                        str(node)[:500],
                    )
                    node.decompose()
            for node in baseline_head_tags:
                soup.head.append(node)

        baseline_anchor_map: Dict[str, str] = {}
        for anchor in baseline.find_all('a', href=True):
            label = anchor.get_text(' ', strip=True)
            href = anchor.get('href', '').strip()
            if label and href:
                baseline_anchor_map[label] = href

        for anchor in soup.find_all('a'):
            label = anchor.get_text(' ', strip=True)
            baseline_href = baseline_anchor_map.get(label)
            if baseline_href and anchor.get('href') != baseline_href:
                _append_layout_finding(
                    findings,
                    'rewritten_url_reset',
                    'Restored baseline URL for matching anchor text.',
                    label,
                )
                anchor['href'] = baseline_href

        for node in list(soup.find_all('script')):
            if node.get('type') == 'application/ld+json':
                script_text = node.string or node.get_text('', strip=True)
                try:
                    parsed = json.loads(script_text) if script_text else {}
                except json.JSONDecodeError:
                    parsed = {}
                if _is_exact_schema_org_context(parsed.get('@context')):
                    continue
            _append_layout_finding(
                findings,
                'unsafe_rewritten_script',
                'Removed non-schema script from rewritten layout HTML.',
                str(node)[:500],
            )
            node.decompose()

        for comment in soup.find_all(string=lambda value: isinstance(value, Comment)):
            text = str(comment)
            if scan_text_for_injection(text):
                _append_layout_finding(
                    findings,
                    'unsafe_rewritten_comment',
                    'Removed prompt-like comment from rewritten layout HTML.',
                    text,
                )
                comment.extract()

        for element in list(soup.find_all(True)):
            text = element.get_text(' ', strip=True)
            is_hidden = (
                element.has_attr('hidden')
                or element.get('aria-hidden') == 'true'
                or 'display:none' in element.get('style', '').replace(' ', '').lower()
            )
            if is_hidden and text and scan_text_for_injection(text):
                _append_layout_finding(
                    findings,
                    'unsafe_rewritten_element',
                    'Removed prompt-like element from rewritten layout HTML.',
                    text,
                )
                element.decompose()

        return {
            'flagged': bool(findings),
            'findings': findings,
            'html': str(soup),
        }

    def apply_layout_instruction(
        self,
        instruction_text: str,
        current_html: str,
        prior_instructions: Optional[List[Dict]] = None
    ) -> Dict:
        """Apply natural-language layout instruction to HTML via LLM.

        Interprets user's plain-English layout request (e.g., "Move Publications
        after Skills") and modifies HTML structure accordingly without altering
        text content.

        Args:
            instruction_text: Plain-English instruction from user
            current_html: Current HTML document to modify
            prior_instructions: List of previously applied instructions (for context)

        Returns:
            {
                'html': modified HTML (if successful),
                'summary': change description,
                'confidence': score 0.0-1.0,
                'error': error message (if applicable),
                'question': clarification question (if confidence < 0.7 or ambiguous),
                'requires_clarification': bool
            }
        """
        # Build LLM prompt
        instruction_safety = self._sanitize_layout_instruction_text(
            instruction_text
        )
        context_safety = self._sanitize_layout_context_html(current_html)

        sanitized_instruction_text = instruction_safety['sanitized_text']
        if not sanitized_instruction_text:
            return {
                'error': 'unsafe_instruction',
                'details': (
                    'The layout instruction only contained unsafe prompt-like '
                    'directives after sanitization.'
                ),
                'safety': {
                    'flagged': True,
                    'findings': _summarize_layout_findings(
                        instruction_safety['findings'],
                        context_safety['findings'],
                    ),
                    'instruction_text': instruction_safety,
                    'current_html': context_safety,
                },
            }

        sanitized_current_html = context_safety['sanitized_html']
        cv_outline = self._serialize_html_for_context(sanitized_current_html)
        prior_context = ""
        if prior_instructions:
            prior_list = [f"- {inst.get('instruction_text', '')}" for inst in prior_instructions]
            prior_context = "\n\nPRIOR INSTRUCTIONS APPLIED:\n" + "\n".join(prior_list)

        prompt = f"""You are a CV layout assistant. Your job is to interpret user requests and modify CV HTML structure.

CURRENT CV STRUCTURE (outline):
{cv_outline}

CURRENT HTML (modify this):
{sanitized_current_html}
{prior_context}

USER INSTRUCTION:
"{sanitized_instruction_text}"

YOUR TASK:
1. Interpret the user's intent
2. Modify the HTML to reflect the instruction (reorder sections, adjust spacing, etc.)
3. Return ONLY valid JSON (no markdown, no explanations outside the JSON):

{{
  "modified_html": "[complete modified HTML]",
  "change_summary": "[2-3 sentence human-readable summary of what changed]",
  "confidence": 0.95,
  "requires_clarification": false
}}

IMPORTANT CONSTRAINTS:
- Never modify text content (only structure/CSS/order)
- Preserve all existing text exactly
- Return the full HTML document (not a diff or excerpt)
- If unsure of intent, set confidence < 0.7 and include clarification_question

If you need clarification, return:
{{
  "requires_clarification": true,
  "clarification_question": "[your question]",
  "confidence": 0.5
}}
"""

        response = ''
        try:
            # Call LLM to interpret and modify HTML
            response = self.llm.call_llm(
                prompt=prompt,
                system_prompt="You are an expert HTML/CSS layout modifier. You modify CV structure without changing content.",
                temperature=0.3  # Low temperature for precise modifications
            )

            # Guard against empty response before JSON parsing
            if not response or not response.strip():
                return {
                    'error': 'parse_error',
                    'details': 'LLM returned an empty response',
                    'raw_response': response or ''
                }

            try:
                result = json.loads(response)
            except json.JSONDecodeError:
                result = self.llm._parse_json_response(response)

            # Validate response structure
            if result.get('requires_clarification', False):
                return {
                    'error': 'clarify',
                    'clarification_question': result.get('clarification_question', ''),
                    'confidence': result.get('confidence', 0.5)
                }

            # Check confidence before HTML validation so low-confidence responses
            # are surfaced correctly even when modified_html is empty/short.
            confidence = result.get('confidence', 0.7)
            if confidence < 0.7:
                return {
                    'error': 'low_confidence',
                    'question': f"Low confidence ({confidence:.0%}). Could you clarify: {instruction_text}?",
                    'confidence': confidence
                }

            # Extract modified HTML and validate it's not empty
            modified_html = result.get('modified_html', '')
            if not modified_html:
                return {
                    'error': 'parse_failed',
                    'details': 'HTML response was empty'
                }

            safety_result = self._sanitize_layout_instruction_html(
                current_html=sanitized_current_html,
                modified_html=modified_html,
            )

            all_findings = _summarize_layout_findings(
                instruction_safety['findings'],
                context_safety['findings'],
                safety_result['findings'],
            )

            return {
                'html': safety_result['html'],
                'summary': result.get('change_summary', 'Layout updated'),
                'confidence': confidence,
                'requires_clarification': False,
                'safety': {
                    'flagged': bool(all_findings),
                    'findings': all_findings,
                    'instruction_text': instruction_safety,
                    'current_html': context_safety,
                    'rewritten_html': {
                        'flagged': safety_result['flagged'],
                        'findings': safety_result['findings'],
                        'raw_html': modified_html,
                        'sanitized_html': safety_result['html'],
                    },
                },
            }

        except (json.JSONDecodeError, ValueError) as e:
            return {
                'error': 'parse_error',
                'details': f'LLM response was not valid JSON: {str(e)}',
                'raw_response': response
            }
        except Exception as e:
            error_type = type(e).__name__.lower()
            error_text = str(e).lower()
            if (
                isinstance(e, TimeoutError)
                or 'timeout' in error_type
                or 'time out' in error_text
                or 'timed out' in error_text
                or 'readtimeout' in error_type
                or 'apitimeouterror' in error_type
            ):
                return {
                    'error': 'timeout',
                    'details': (
                        'Layout instruction request timed out before the model '
                        'returned a rewrite. Retry, or use a narrower layout '
                        'instruction targeting one section at a time.'
                    ),
                }
            return {
                'error': 'processing_error',
                'details': f'Failed to apply layout instruction: {str(e)}'
            }

    def classify_instruction(self, instruction_text: str) -> str:
        """Classify a user instruction as 'layout' or 'content' using the LLM.

        Args:
            instruction_text: The raw user instruction.

        Returns:
            'layout' if the instruction affects structure/presentation only,
            'content' if it requests text edits. Defaults to 'layout' on failure.
        """
        prompt = (
            f'Classify this CV editing instruction as either "layout" or "content".\n\n'
            f'"layout" = structural changes only: section order, spacing, page breaks, '
            f'font size, margins, moving sections.\n'
            f'"content" = text edits: rewriting bullets, editing summary, changing wording.\n\n'
            f'Instruction: "{instruction_text}"\n\n'
            f'Reply with exactly one word: layout or content'
        )
        try:
            response = self.llm.call_llm(
                prompt=prompt,
                system_prompt='Reply with exactly one word: layout or content',
                temperature=0.0,
            )
            word = (response or '').strip().lower().split()[0] if response else ''
            return 'content' if word == 'content' else 'layout'
        except Exception:
            return 'layout'

    def propose_content_change(
        self,
        instruction_text: str,
        content: Dict,
    ) -> Dict:
        """Propose targeted text content changes based on a natural-language instruction.

        Unlike apply_layout_instruction, this method permits text content modifications.
        Proposals are returned (not applied) in the approved_rewrites format so the user
        can review each change before committing.

        Args:
            instruction_text: Plain-English description of the desired edit,
                e.g. "Shorten the second bullet under Genentech to focus on impact".
            content: Render-ready content dict with 'summary', 'experiences' (each
                having 'id', 'title', 'company', 'achievements'), and 'skills'.

        Returns:
            {
                'proposals': list of {type, location, original, proposed, reason, id},
                'error': str or None,
            }
        """
        import uuid as _u

        # ── Build a structured text view of the content for the LLM ──────────
        lines = []

        summary_text = content.get('summary') or ''
        lines.append(f'SUMMARY:\n"{summary_text}"\n')

        lines.append('EXPERIENCES:')
        for idx, exp in enumerate(content.get('experiences') or [], start=1):
            exp_id   = exp.get('id') or f'exp_{idx:03d}'
            company  = exp.get('company', '')
            title    = exp.get('title', '')
            lines.append(f'{idx}. {title} ({company}, id={exp_id})')
            achievements = exp.get('achievements') or []
            for i, ach in enumerate(achievements):
                text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                lines.append(f'   [{i}] "{text}"')

        content_summary = '\n'.join(lines)

        prompt = f"""You are a CV content editor. A user wants to make a targeted text change to their CV.

CURRENT CV CONTENT:
{content_summary}

USER INSTRUCTION:
"{instruction_text}"

YOUR TASK:
Identify the minimal set of text changes that fulfil the instruction and return them as JSON.
Each change must reference a specific piece of text by its exact location.

Return ONLY valid JSON (no markdown, no extra text):

{{
  "proposals": [
    {{
      "type": "bullet",
      "location": "exp_001.achievements[2]",
      "original": "exact original text",
      "proposed": "new text",
      "reason": "brief rationale"
    }}
  ]
}}

Location format:
- Bullet point: "exp_ID.achievements[N]"  (use the id= value shown above, e.g. exp_001)
- Summary paragraph: "summary"

CONSTRAINTS (must follow all):
- Only modify existing text; do NOT add new achievements, skills, or sections.
- Preserve all proper nouns, numbers, dates, and technical terms exactly.
- Proposed text must be a complete, grammatically correct replacement.
- Return at most 5 proposals.
- If the instruction is ambiguous or impossible to fulfil safely, return an empty proposals list.
"""

        response = ''
        try:
            response = self.llm.call_llm(
                prompt=prompt,
                system_prompt=(
                    'You are an expert CV content editor. You propose precise, minimal text '
                    'improvements without inventing facts or altering meaning.'
                ),
                temperature=0.3,
            )

            if not response or not response.strip():
                return {'proposals': [], 'error': 'LLM returned an empty response'}

            try:
                result = json.loads(response)
            except json.JSONDecodeError:
                result = self.llm._parse_json_response(response)

            raw_proposals = result.get('proposals') or []
            validated: list = []
            for p in raw_proposals:
                if not isinstance(p, dict):
                    continue
                p_type     = str(p.get('type') or '').strip()
                p_location = str(p.get('location') or '').strip()
                p_original = str(p.get('original') or '').strip()
                p_proposed = str(p.get('proposed') or '').strip()
                p_reason   = str(p.get('reason') or '').strip()
                if p_type not in ('bullet', 'summary') or not p_location or not p_original or not p_proposed:
                    continue
                validated.append({
                    'type':     p_type,
                    'location': p_location,
                    'original': p_original,
                    'proposed': p_proposed,
                    'reason':   p_reason,
                    'id':       f'cp_{_u.uuid4().hex[:12]}',
                })

            return {'proposals': validated, 'error': None}

        except (json.JSONDecodeError, ValueError) as e:
            return {'proposals': [], 'error': f'LLM response was not valid JSON: {e}'}
        except Exception as e:
            error_type = type(e).__name__.lower()
            error_text = str(e).lower()
            if (
                isinstance(e, TimeoutError)
                or 'timeout' in error_type
                or 'time out' in error_text
                or 'timed out' in error_text
            ):
                return {
                    'proposals': [],
                    'error': (
                        'Content proposal request timed out. Try a more specific '
                        'instruction targeting a single bullet or section.'
                    ),
                }
            return {'proposals': [], 'error': f'Failed to generate content proposals: {e}'}

    _MDU_DEDUP_THRESHOLD = 85

    def _mdu_build_master_index(self, master: Dict[str, Any]) -> str:
        """Compact, human-readable index of existing master data for prompt context.

        Deliberately not a raw JSON dump of the whole file (which could be
        large) — this gives the LLM enough to resolve "the project at Acme"
        against company/title/date text and to judge duplication, at a
        fraction of the token cost.
        """
        lines: List[str] = []
        experiences = master.get('experience') or master.get('experiences') or []
        lines.append('EXISTING EXPERIENCE ENTRIES:')
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            exp_id = exp.get('id', '')
            title = exp.get('title', '')
            company = exp.get('company', '')
            start = exp.get('start_date', '')
            end = exp.get('end_date', '') or 'present'
            lines.append(f'- id={exp_id}: {title} at {company} ({start} - {end})')

        skills = master.get('skills')
        skill_names: List[str] = []
        if isinstance(skills, list):
            for s in skills:
                skill_names.append(s.get('name', '') if isinstance(s, dict) else str(s))
        elif isinstance(skills, dict):
            for cat_val in skills.values():
                items = cat_val.get('skills', []) if isinstance(cat_val, dict) else (cat_val if isinstance(cat_val, list) else [])
                for s in items:
                    skill_names.append(s.get('name', '') if isinstance(s, dict) else str(s))
        if skill_names:
            lines.append('EXISTING SKILLS: ' + ', '.join(sorted(set(n for n in skill_names if n))))

        summaries = master.get('professional_summaries')
        if isinstance(summaries, dict):
            lines.append('EXISTING SUMMARY VARIANTS: ' + ', '.join(summaries.keys()))
        elif isinstance(summaries, list) and summaries:
            lines.append(f'EXISTING SUMMARY VARIANTS: {len(summaries)} unnamed variant(s)')

        return '\n'.join(lines)

    def _mdu_fuzzy_ratio(self, a: str, b: str) -> float:
        """Similarity ratio 0-100. Uses rapidfuzz if available, else a stdlib fallback."""
        try:
            from rapidfuzz import fuzz  # optional accelerator, not a hard dependency
            return fuzz.token_sort_ratio(a, b)
        except ImportError:
            a_tokens = set(re.findall(r'\w+', a.lower()))
            b_tokens = set(re.findall(r'\w+', b.lower()))
            if not a_tokens or not b_tokens:
                return 0.0
            overlap = len(a_tokens & b_tokens)
            union = len(a_tokens | b_tokens)
            return 100.0 * overlap / union if union else 0.0

    def _mdu_identifying_text(self, change: Dict[str, Any]) -> str:
        """Text used to fuzzy-match a proposed change against other entries."""
        proposed = change.get('proposed')
        if isinstance(proposed, dict):
            if 'text' in proposed:
                return str(proposed.get('text') or '')
            if 'name' in proposed:
                return str(proposed.get('name') or '')
            if change.get('section') == 'experience':
                return f"{proposed.get('company', '')} {proposed.get('title', '')} {proposed.get('start_date', '')}"
            return json.dumps(proposed, sort_keys=True)
        return str(proposed or '')

    def _mdu_existing_entries_text(self, master: Dict[str, Any], section: str) -> List[Tuple[str, str]]:
        """Return [(existing_id_or_label, identifying_text), ...] for a section, for dedup matching."""
        out: List[Tuple[str, str]] = []
        if section == 'experience':
            for exp in (master.get('experience') or master.get('experiences') or []):
                if not isinstance(exp, dict):
                    continue
                exp_id = str(exp.get('id', ''))
                for ach in (exp.get('achievements') or exp.get('bullets') or []):
                    text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                    if text:
                        out.append((exp_id, text))
                label = f"{exp.get('company', '')} {exp.get('title', '')} {exp.get('start_date', '')}"
                out.append((exp_id, label))
        elif section == 'skills':
            skills = master.get('skills')
            items = skills if isinstance(skills, list) else []
            if isinstance(skills, dict):
                for cat_val in skills.values():
                    items = items + (cat_val.get('skills', []) if isinstance(cat_val, dict) else (cat_val if isinstance(cat_val, list) else []))
            for s in items:
                name = s.get('name', '') if isinstance(s, dict) else str(s)
                if name:
                    out.append((name, name))
        return out

    def _mdu_dedup_pass(self, master: Dict[str, Any], changes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Two-stage dedup: within this proposal batch, then against existing master data.

        Stage 1 drops later within-batch near-duplicates outright (the LLM
        proposed the same real-world fact twice in one response — showing the
        user two rows for it helps no one). Stage 2 flags (does not drop)
        near-duplicates of existing master-data entries via
        `possible_duplicate_of`, since a true duplicate-vs-legitimate-update
        judgment call belongs to the human reviewer.
        """
        # Stage 1: within-batch.
        deduped: List[Dict[str, Any]] = []
        seen_texts: List[Tuple[str, str]] = []  # (section, text)
        for change in changes:
            if change.get('op') != 'add':
                deduped.append(change)
                continue
            section = change.get('section', '')
            text = self._mdu_identifying_text(change)
            is_dup = any(
                s == section and text and self._mdu_fuzzy_ratio(text, t) >= self._MDU_DEDUP_THRESHOLD
                for s, t in seen_texts
            )
            if is_dup:
                continue  # drop — same fact already in this batch
            seen_texts.append((section, text))
            deduped.append(change)

        # Stage 2: against existing master data — flag, don't drop.
        for change in deduped:
            if change.get('op') != 'add':
                continue
            section = change.get('section', '')
            text = self._mdu_identifying_text(change)
            if not text:
                continue
            for existing_id, existing_text in self._mdu_existing_entries_text(master, section):
                if self._mdu_fuzzy_ratio(text, existing_text) >= self._MDU_DEDUP_THRESHOLD:
                    change['possible_duplicate_of'] = existing_id
                    break

        return deduped

    def _mdu_persuasion_flags(self, change: Dict[str, Any]) -> List[str]:
        """Advisory (non-blocking) persuasion-quality flags for a proposed change's text.

        Reuses check_persuasion()'s own detection logic rather than
        reimplementing verb/quantification/vague-language heuristics.
        """
        proposed = change.get('proposed')
        text = None
        if change.get('section') == 'experience' and isinstance(proposed, dict):
            text = proposed.get('text')
        elif change.get('section') == 'professional_summaries' and isinstance(proposed, str):
            text = proposed
        if not text or not str(text).strip():
            return []
        try:
            result = self.check_persuasion([{'id': '_mdu_tmp', 'achievements': [{'text': text}]}])
        except Exception:
            return []
        flags: List[str] = []
        for finding in (result.get('findings') or []):
            for issue in (finding.get('issues') or []):
                suggestion = issue.get('suggestion') or issue.get('type')
                if suggestion:
                    flags.append(str(suggestion))
        return flags

    def propose_master_data_update(
        self,
        instruction_or_text: str,
        master: Dict[str, Any],
        *,
        source: str,
        prior_clarifications: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """Propose a structured Master_CV_Data.json diff from an NL instruction or document text.

        Args:
            instruction_or_text: user's NL instruction (source='nl_instruction'),
                or extracted document text (source='document_ingestion').
            master: current in-memory master data dict (for entity-resolution
                context, existing-id validation, and dedup matching).
            source: 'nl_instruction' | 'document_ingestion' — selects the
                prompt variant and default change-count cap.
            prior_clarifications: prior Q&A pairs from earlier turns of the
                same clarification exchange, appended to the prompt on
                resubmission so multi-round clarification is possible.

        Returns:
            On success: {'changes': [MasterDataChange, ...], 'error': None}.
            On ambiguity/low-confidence: {'changes': [], 'error': 'clarify',
                'clarification_question': str, 'confidence': float} — one
                canonical shape regardless of which condition triggered it
                (unlike apply_layout_instruction's two inconsistent branches).
            On unsafe input (fully stripped by sanitization):
                {'changes': [], 'error': 'unsafe_instruction', 'details': str}.
            On any other failure: {'changes': [], 'error': str}.
        """
        raw_text = str(instruction_or_text or '')
        sanitized_text, safety_findings = sanitize_instruction_text(raw_text)
        sanitized_text = re.sub(r'\s+', ' ', sanitized_text).strip(' ,;:-')
        if not sanitized_text:
            return {
                'changes': [],
                'error': 'unsafe_instruction',
                'details': (
                    'The instruction or document text only contained unsafe '
                    'prompt-like directives after sanitization.'
                ),
            }

        master_index = self._mdu_build_master_index(master)

        prior_context = ''
        if prior_clarifications:
            qa_lines = [
                f'- Q: {qa.get("question", "")}\n  A: {qa.get("answer", "")}'
                for qa in prior_clarifications
            ]
            prior_context = '\n\nPRIOR CLARIFICATION EXCHANGE:\n' + '\n'.join(qa_lines)

        if source == 'nl_instruction':
            max_changes = 5
            task_description = (
                'The user gave a plain-English instruction describing a change to make to '
                'their master CV data. Resolve which existing entry (if any) the instruction '
                'refers to by matching company/title/date text against the index above — '
                'never invent or guess an id that is not listed there.'
            )
        else:
            max_changes = 30
            task_description = (
                'The text below is extracted from an uploaded document (an old CV or '
                'LinkedIn export). Extract structured additions to the master data — new '
                'experience entries, achievements, skills, education, awards, or '
                'certifications — that are not already present (see the existing-data index '
                'above). Do not fabricate content not supported by the document text.'
            )

        prompt = f"""You are a master-CV-data update assistant. Your job is to turn the input below into a structured, reviewable set of proposed changes to a candidate's master CV data — never write directly to the data yourself.

EXISTING MASTER DATA (for entity resolution and duplicate-avoidance):
{master_index}

{task_description}

INPUT:
\"\"\"{sanitized_text}\"\"\"
{prior_context}

YOUR TASK:
1. If you cannot confidently identify which existing entry this input refers to (e.g. more than one experience could plausibly match, or the input clearly references something not in the existing data with no clear place to add it), do NOT guess — set "requires_clarification": true and ask a specific question in "clarification_question" that names the ambiguous options by company/title. Otherwise set "requires_clarification": false.
2. If you are unsure for any other reason (vague input, low signal), set "confidence" below 0.7. Otherwise set it to your genuine confidence, up to 1.0.
3. When proposing new or updated text (achievement bullets, summary text), write it with a strong opening verb and quantify impact where the input actually supports it — do NOT invent numbers, dates, or facts not present in the input.
4. Propose at most {max_changes} changes.

Return ONLY valid JSON (no markdown, no extra text):

{{
  "requires_clarification": false,
  "clarification_question": "",
  "confidence": 0.95,
  "changes": [
    {{
      "section": "experience",
      "op": "add",
      "parent_id": "exp_005",
      "field": "achievements",
      "proposed": {{"text": "Delivered a Kubernetes-based deployment pipeline, cutting release time by 40%.", "keywords": ["Kubernetes"]}},
      "label": "New achievement — Senior Engineer @ Acme (exp_005)",
      "rationale": "Input described a Kubernetes project at the company matching exp_005."
    }}
  ]
}}

"section" must be one of: personal_info, experience, skills, education, awards, certifications, selected_achievements, professional_summaries.
"op" must be "add" or "update" — never "delete"; deletions are handled through the existing structured editors, not this path.
"parent_id" is the existing entry's id when adding a nested item (e.g. an achievement) into it, or updating a field on it; omit/null for a brand-new top-level entry (e.g. an entirely new experience) or for sections without nested ids (skills, education, awards, certifications, selected_achievements, professional_summaries).
If requires_clarification is true, return an empty "changes" list.
"""

        response = ''
        try:
            response = self.llm.call_llm(
                prompt=prompt,
                system_prompt=(
                    'You propose precise, minimal, well-supported additions to a '
                    'candidate\'s master CV data record. You never fabricate facts, '
                    'numbers, or dates not present in the input, and you never guess '
                    'at an existing entry\'s identity when genuinely ambiguous.'
                ),
                temperature=0.2 if source == 'nl_instruction' else 0.25,
            )

            if not response or not response.strip():
                return {'changes': [], 'error': 'LLM returned an empty response'}

            try:
                result = json.loads(response)
            except json.JSONDecodeError:
                result = self.llm._parse_json_response(response)

            if not isinstance(result, dict):
                return {'changes': [], 'error': 'LLM response was not a JSON object'}

            requires_clarification = bool(result.get('requires_clarification', False))
            confidence = result.get('confidence', 0.7)
            try:
                confidence = float(confidence)
            except (TypeError, ValueError):
                confidence = 0.7

            if requires_clarification or confidence < 0.7:
                question = str(result.get('clarification_question') or '').strip()
                if not question:
                    question = (
                        f'Could you clarify: "{sanitized_text[:200]}"? '
                        'I could not confidently determine what to update.'
                    )
                return {
                    'changes': [],
                    'error': 'clarify',
                    'clarification_question': question,
                    'confidence': confidence,
                }

            raw_changes = result.get('changes') or []
            if not isinstance(raw_changes, list):
                return {'changes': [], 'error': 'LLM response "changes" was not a list'}

            # Layer 1: shape/allow-list validation, plus parent_id existence check
            # (the concrete anti-hallucination guard for "add it to exp_005" cases).
            existing_ids = {
                str(exp.get('id', ''))
                for exp in (master.get('experience') or master.get('experiences') or [])
                if isinstance(exp, dict) and exp.get('id')
            }
            validated: List[Dict[str, Any]] = []
            for raw in raw_changes:
                if not isinstance(raw, dict):
                    continue
                section = str(raw.get('section') or '').strip()
                op = str(raw.get('op') or '').strip()
                if section not in _MDU_ALLOWED_SECTIONS or op not in _MDU_ALLOWED_OPS:
                    continue
                proposed = raw.get('proposed')
                if proposed is None or proposed == '':
                    continue
                identifying = proposed if isinstance(proposed, str) else json.dumps(proposed)
                if len(identifying) > 2000:
                    continue  # sanity cap against degenerate LLM output
                parent_id = raw.get('parent_id') or None
                if parent_id and str(parent_id) not in existing_ids:
                    continue  # references an id that doesn't exist — drop, don't guess
                validated.append({
                    'id':                  f'mdu_{uuid.uuid4().hex[:12]}',
                    'section':             section,
                    'op':                  op,
                    'parent_id':           str(parent_id) if parent_id else None,
                    'field':               raw.get('field'),
                    'original':            raw.get('original'),
                    'proposed':            proposed,
                    'label':               str(raw.get('label') or '').strip() or f'{op} {section}',
                    'rationale':           str(raw.get('rationale') or '').strip(),
                    'source':              source,
                })

            # Dedup: within-batch drop, then vs-master-data flag.
            validated = self._mdu_dedup_pass(master, validated)

            # Persuasion advisory (non-blocking).
            for change in validated:
                flags = self._mdu_persuasion_flags(change)
                if flags:
                    change['persuasion_flags'] = flags

            # Layer 2: dry-run apply + schema validation. Drop anything that
            # would not actually apply cleanly, so the user never sees a
            # proposal that would fail validation later.
            final_changes: List[Dict[str, Any]] = []
            for change in validated:
                trial = copy.deepcopy(master)
                applied = _mdu_apply_change(trial, change)
                if not applied:
                    continue
                schema_result = validate_master_data(trial)
                if not schema_result.valid:
                    continue
                final_changes.append(change)

            return {'changes': final_changes, 'error': None}

        except (json.JSONDecodeError, ValueError) as e:
            return {'changes': [], 'error': f'LLM response was not valid JSON: {e}'}
        except Exception as e:
            error_type = type(e).__name__.lower()
            error_text = str(e).lower()
            if (
                isinstance(e, TimeoutError)
                or 'timeout' in error_type
                or 'time out' in error_text
                or 'timed out' in error_text
            ):
                return {
                    'changes': [],
                    'error': (
                        'Master data update request timed out. Try a more specific '
                        'instruction, or a shorter document excerpt.'
                    ),
                }
            return {'changes': [], 'error': f'Failed to generate master data update proposal: {e}'}

    def analyze_harvest_candidates(
        self,
        candidates: List[Dict[str, Any]],
        job_analysis: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Use the LLM to evaluate whether each harvest candidate should be promoted to master CV.

        Args:
            candidates: List of harvest candidates from _compile_harvest_candidates().
            job_analysis: The session's job analysis dict (for context).

        Returns:
            {
                'analyses': [{'id', 'recommendation', 'confidence', 'reasoning'}, ...],
                'error': str or None,
            }
        """
        if not self.llm:
            return {'analyses': [], 'error': 'No LLM configured'}

        if not candidates:
            return {'analyses': [], 'error': None}

        job_title   = (job_analysis or {}).get('title', 'Unknown Role')
        job_company = (job_analysis or {}).get('company', 'Unknown Company')
        requirements = (job_analysis or {}).get('required_skills') or []
        requirements_text = ', '.join(str(r) for r in requirements[:20]) if requirements else 'Not specified'

        cand_lines = []
        for c in candidates:
            ctype     = c.get('type', 'unknown')
            cid       = c.get('id', '')
            label     = c.get('label', '')
            original  = c.get('original', '')
            proposed  = c.get('proposed', '')
            rationale = c.get('rationale', '')
            cand_lines.append(
                f'- id: {cid}\n'
                f'  type: {ctype}\n'
                f'  label: {label}\n'
                f'  original: {original!r}\n'
                f'  proposed: {proposed!r}\n'
                f'  rationale: {rationale}'
            )
        candidates_text = '\n'.join(cand_lines)

        prompt = f"""You are a CV master-data curator. You are reviewing candidate improvements generated during a job application for:
  Job: {job_title} at {job_company}
  Key requirements: {requirements_text}

These candidates were generated during the application session and are being considered for permanent promotion to the user's master CV.

CANDIDATES:
{candidates_text}

YOUR TASK:
For each candidate, evaluate whether it should be permanently added to the user's master CV.

Promotion criteria by type:
- improved_bullet: Promote if the rewrite adds job-neutral improvements (metrics, specificity, clearer impact). Skip if it is tailored only to this job or degrades the original.
- new_skill: Promote if the skill likely reflects genuine ongoing expertise. Skip if it was added solely to match this job description with no other evidence.
- skill_gap_confirmed: Promote if the user's confirmation is credible and the skill is absent from master. Be slightly more skeptical — answers to clarifying questions can be aspirational.
- summary_variant: Promote if it is clearly stronger and more broadly applicable. Skip if it is very company-specific or similar to the original.

Return ONLY a valid JSON array, no markdown, no extra text:
[
  {{
    "id": "<candidate id>",
    "recommendation": "promote" or "skip",
    "confidence": "high" or "medium" or "low",
    "reasoning": "One to two sentence explanation focused on long-term CV value, not this specific job."
  }}
]

Include one entry per candidate. Do not omit any candidate."""

        response = ''
        try:
            response = self.llm.call_llm(
                prompt=prompt,
                system_prompt=(
                    'You are a professional CV curator. Evaluate whether generated improvements '
                    'have long-term value for a master CV, not just for one specific job application. '
                    'Return ONLY a valid JSON array.'
                ),
                temperature=0.3,
            )

            if not response or not response.strip():
                return {'analyses': [], 'error': 'LLM returned an empty response'}

            try:
                raw = json.loads(response)
            except json.JSONDecodeError:
                raw = self.llm._parse_json_response(response)

            if not isinstance(raw, list):
                return {'analyses': [], 'error': 'LLM response was not a JSON array'}

            valid_recs  = {'promote', 'skip'}
            valid_confs = {'high', 'medium', 'low'}
            analyses: list = []
            for item in raw:
                if not isinstance(item, dict):
                    continue
                rec       = str(item.get('recommendation') or '').strip().lower()
                conf      = str(item.get('confidence') or '').strip().lower()
                cid       = str(item.get('id') or '').strip()
                reasoning = str(item.get('reasoning') or '').strip()
                if not cid or rec not in valid_recs or conf not in valid_confs:
                    continue
                analyses.append({
                    'id':             cid,
                    'recommendation': rec,
                    'confidence':     conf,
                    'reasoning':      reasoning,
                })
            return {'analyses': analyses, 'error': None}

        except (json.JSONDecodeError, ValueError) as e:
            return {'analyses': [], 'error': f'LLM response was not valid JSON: {e}'}
        except Exception as e:
            error_type = type(e).__name__.lower()
            error_text = str(e).lower()
            if (
                isinstance(e, TimeoutError)
                or 'timeout' in error_type
                or 'time out' in error_text
                or 'timed out' in error_text
            ):
                return {
                    'analyses': [],
                    'error': (
                        'Harvest analysis timed out. You can retry or proceed with manual selection.'
                    ),
                }
            return {'analyses': [], 'error': f'Failed to analyze harvest candidates: {e}'}

    def _select_content_hybrid(
        self,
        job_analysis: Dict,
        customizations: Dict,
        max_skills: Optional[int] = None,
        max_achievements: Optional[int] = None,
        max_publications: Optional[int] = None,
        use_semantic_match: bool = True,
    ) -> Dict:
        """
        Select content using hybrid LLM + rule-based approach.

        Inclusion rules
        ---------------
        Experiences : ALL experiences are included EXCEPT those where the user
            has explicitly approved an "Omit" decision.  The set is sorted by
            relevance score (Emphasize items first) so the most relevant
            content appears first in the generated document.
        Achievements: same blacklist rule.
        Skills      : same blacklist rule; LLM-recommended skills are listed
            first, remaining non-omitted skills follow by score.
        """
        # IDs/names explicitly omitted by the user
        omitted_exp_ids      = set(customizations.get('omitted_experiences', []))
        omitted_skill_names  = set(customizations.get('omitted_skills', []))
        omitted_ach_ids      = set(customizations.get('omitted_achievements', []))

        # IDs carrying an extra relevance boost from user/LLM recommendations
        recommended_exp_ids          = set(customizations.get('recommended_experiences', []))
        recommended_achievement_ids  = set(customizations.get('recommended_achievements', []))
        recommended_skills           = set(customizations.get('recommended_skills', []))

        # Also honour per-item recommendation dicts (LLM structured output)
        for rec in customizations.get('experience_recommendations', []):
            if isinstance(rec, dict):
                if rec.get('recommendation', '').lower() == 'omit':
                    omitted_exp_ids.add(rec.get('id', ''))
                elif rec.get('recommendation', '').lower() in ('emphasize', 'include', 'de-emphasize'):
                    recommended_exp_ids.add(rec.get('id', ''))
        for rec in customizations.get('skill_recommendations', []):
            if isinstance(rec, dict):
                if rec.get('recommendation', '').lower() == 'omit':
                    omitted_skill_names.add(rec.get('name', ''))

        # Get all content
        all_experiences  = self.master_data.get('experience', [])
        session_view = SessionDataView(
            self.master_data,
            customizations,
            customizations,
        )
        all_achievements = session_view.selected_achievements()
        all_skills = []
        for skill in session_view.normalized_skills():
            if isinstance(skill, dict):
                all_skills.append(skill)
            elif isinstance(skill, str):
                all_skills.append({'name': skill})

        # Scoring helpers
        job_keywords     = set(job_analysis.get('ats_keywords', []))
        job_requirements = (
            job_analysis.get('must_have_requirements', []) +
            job_analysis.get('nice_to_have_requirements', [])
        )
        domain = job_analysis.get('domain', '')
        cfg    = get_config()
        max_ach = (
            max_achievements
            if max_achievements is not None
            else cfg.get('generation.max_achievements', 5)
        )
        max_skills = max_skills if max_skills is not None else cfg.get('generation.max_skills', 20)
        max_pubs = (
            max_publications
            if max_publications is not None
            else cfg.get('generation.max_publications', 10)
        )

        # ── Experiences ───────────────────────────────────────────────────────
        # Include ALL experiences; only exclude those explicitly omitted.
        scored_experiences = []
        for exp in all_experiences:
            exp_id = exp.get('id', '')
            if exp_id in omitted_exp_ids:
                continue  # user approved Omit — skip

            # Boost for recommended items
            llm_score     = 10.0 if exp_id in recommended_exp_ids else 0.0
            keyword_score = calculate_relevance_score(exp, job_keywords, job_requirements, domain)
            semantic_score = 0.0
            if self.llm and use_semantic_match:
                semantic_score = self.llm.semantic_match(json.dumps(exp), job_requirements) * 10

            scored_experiences.append((exp, llm_score + keyword_score + semantic_score))

        # Hybrid sort: relevance-primary, recency-secondary within equal scores.
        # "Current", "Present", "", or None are treated as today (float to top among ties).
        _today = _date.today()

        def _parse_end_date(exp: Dict) -> _date:
            raw = str(exp.get('end_date') or exp.get('end') or '').strip()
            if not raw or raw.lower() in ('current', 'present', 'now', 'ongoing'):
                return _today
            for fmt in ('%Y-%m-%d', '%B %Y', '%b %Y', '%Y'):
                try:
                    return datetime.strptime(raw, fmt).date()
                except ValueError:
                    pass
            # Partial match — try extracting a 4-digit year
            m = re.search(r'\b(\d{4})\b', raw)
            if m:
                return _date(int(m.group(1)), 12, 31)
            return _date.min

        scored_experiences.sort(
            key=lambda x: (-x[1], -_parse_end_date(x[0]).toordinal()),
        )
        selected_experiences = [exp for exp, _ in scored_experiences]

        # Sort experiences in reverse chronological order by end date.
        # "Current", "Present", "", or None are treated as today (sorts first).
        _today = _date.today()

        def _parse_end_date(exp: Dict) -> _date:
            raw = str(exp.get('end_date') or exp.get('end') or '').strip()
            if not raw or raw.lower() in ('current', 'present', 'now', 'ongoing'):
                return _today
            for fmt in ('%Y-%m-%d', '%B %Y', '%b %Y', '%Y'):
                try:
                    return datetime.strptime(raw, fmt).date()
                except ValueError:
                    pass
            # Partial match — try extracting a 4-digit year
            m = re.search(r'\b(\d{4})\b', raw)
            if m:
                return _date(int(m.group(1)), 12, 31)
            return _date.min

        # Only apply default chronological sort when the user hasn't manually reordered.
        # The user-override block below will replace this ordering if present.
        selected_experiences = sorted(selected_experiences, key=_parse_end_date, reverse=True)

        # Override: if the user has explicitly reordered experience rows via the UI,
        # apply their ordering stored in customizations['experience_row_order']
        # as a list of experience IDs in the desired display order.
        experience_row_order = customizations.get('experience_row_order', [])
        if experience_row_order:
            order_map = {eid: i for i, eid in enumerate(experience_row_order)}
            selected_experiences = sorted(
                selected_experiences,
                key=lambda e: order_map.get(e.get('id', ''), len(order_map)),
            )

        # ── Per-experience bullet ordering ────────────────────────────────────
        # Default: sort bullets by keyword-overlap relevance.
        # Override: if the user has explicitly reordered bullets via the UI,
        # apply their ordering stored in customizations['achievement_orders']
        # as a list of original indices per experience id.
        achievement_orders = customizations.get('achievement_orders', {})
        ordered_experiences = []
        for exp in selected_experiences:
            exp_id = exp.get('id', '')
            achievements = list(exp.get('achievements') or [])
            if not achievements:
                ordered_experiences.append(exp)
                continue

            if exp_id in achievement_orders:
                user_order = achievement_orders[exp_id]
                reordered = []
                seen_in_order = set()
                for idx in user_order:
                    try:
                        reordered.append(achievements[idx])
                        seen_in_order.add(idx)
                    except IndexError:
                        pass
                for i, a in enumerate(achievements):
                    if i not in seen_in_order:
                        reordered.append(a)
                achievements = reordered
            elif job_keywords:
                def _ach_relevance(ach, _kws=job_keywords):
                    text = (ach.get('text', '') if isinstance(ach, dict) else str(ach)).lower()
                    tokens = set(re.findall(r'\b\w+\b', text))
                    expanded: set = set()
                    for t in tokens:
                        c = self._expansion_index.get(t)
                        if c:
                            expanded.add(c.lower())
                    return len((tokens | expanded) & {kw.lower() for kw in _kws})
                achievements = sorted(achievements, key=_ach_relevance, reverse=True)

            exp = dict(exp)
            exp['ordered_achievements'] = achievements
            ordered_experiences.append(exp)
        selected_experiences = ordered_experiences

        # ── Achievements ──────────────────────────────────────────────────────
        scored_achievements = []
        for ach in all_achievements:
            ach_id = ach.get('id', '')
            if ach_id in omitted_exp_ids or ach_id in omitted_ach_ids:
                continue

            llm_score     = 10.0 if ach_id in recommended_achievement_ids else 0.0
            keyword_score = calculate_relevance_score(ach, job_keywords, job_requirements, domain)
            semantic_score = 0.0
            if self.llm and use_semantic_match:
                semantic_score = self.llm.semantic_match(json.dumps(ach), job_requirements) * 10

            scored_achievements.append((ach, llm_score + keyword_score + semantic_score))

        scored_achievements.sort(key=lambda x: x[1], reverse=True)
        selected_achievements = self._apply_achievement_diversity(
            scored_achievements, max_ach
        )

        # Prepend extra_achievements: LLM-suggested achievements not in master CV that the user approved
        extra_achievements = customizations.get('extra_achievements', [])
        if extra_achievements:
            existing_ach_texts = {(a.get('text', '') if isinstance(a, dict) else str(a)).lower()
                                  for a in selected_achievements}
            prepend_achs = []
            for ach in extra_achievements:
                if isinstance(ach, dict):
                    text = ach.get('description') or ach.get('title', '')
                else:
                    text = str(ach)
                if text and text.lower() not in existing_ach_texts:
                    prepend_achs.append({'text': text, 'id': f'suggested_{len(prepend_achs)}'})
                    existing_ach_texts.add(text.lower())
            selected_achievements = (prepend_achs + selected_achievements)[:max_ach]

        # Prepend extra_achievements: LLM-suggested achievements not in master CV that the user approved
        extra_achievements = customizations.get('extra_achievements', [])
        if extra_achievements:
            existing_ach_texts = {(a.get('text', '') if isinstance(a, dict) else str(a)).lower()
                                  for a in selected_achievements}
            prepend_achs = []
            for ach in extra_achievements:
                if isinstance(ach, dict):
                    text = ach.get('description') or ach.get('title', '')
                else:
                    text = str(ach)
                if text and text.lower() not in existing_ach_texts:
                    prepend_achs.append({'text': text, 'id': f'suggested_{len(prepend_achs)}'})
                    existing_ach_texts.add(text.lower())
            selected_achievements = (prepend_achs + selected_achievements)[:max_ach]

        # ── Skills ────────────────────────────────────────────────────────────
        # Include all non-omitted skills; recommended ones appear first.
        selected_skills: List[Dict] = []
        remaining_skills: List[tuple] = []

        for skill in all_skills:
            skill_name = skill.get('name', '')
            if skill_name in omitted_skill_names:
                continue
            if skill_name in recommended_skills:
                selected_skills.append(skill)
            else:
                skill_score = calculate_skill_score(
                    skill,
                    job_keywords,
                    job_analysis.get('required_skills', [])
                )
                remaining_skills.append((skill, skill_score))

        remaining_skills.sort(key=lambda x: x[1], reverse=True)
        for skill, _ in remaining_skills:
            selected_skills.append(skill)
            if len(selected_skills) >= max_skills:
                break

        # Prepend extra_skills: LLM-suggested skills not in master CV that the user approved
        # and derive years from matched experience entries (user-edited if provided).
        extra_skills = customizations.get('extra_skills', [])
        raw_match_overrides = customizations.get('extra_skill_matches') or {}
        match_overrides: Dict[str, List[str]] = {}
        if isinstance(raw_match_overrides, dict):
            for key, value in raw_match_overrides.items():
                if not isinstance(key, str):
                    continue
                if isinstance(value, str):
                    value = [v.strip() for v in value.split(',') if v.strip()]
                if not isinstance(value, list):
                    continue
                cleaned = [v.strip() for v in value if isinstance(v, str) and v.strip()]
                if cleaned:
                    match_overrides[key] = cleaned

        def _parse_year(value: Any, default_current: bool = False) -> Optional[int]:
            raw = str(value or '').strip()
            if not raw:
                return datetime.now().year if default_current else None
            if raw.lower() in ('current', 'present', 'now', 'ongoing'):
                return datetime.now().year
            m = re.search(r'\b(19|20)\d{2}\b', raw)
            if not m:
                return None
            return int(m.group(0))

        def _experience_years(exp: Dict) -> int:
            start_year = _parse_year(exp.get('start_date') or exp.get('start'))
            end_year = _parse_year(exp.get('end_date') or exp.get('end'), default_current=True)
            if start_year is None and end_year is None:
                return 1
            if start_year is None:
                return 1
            if end_year is None:
                end_year = start_year
            return max(1, end_year - start_year + 1)

        def _derive_years_from_matches(skill_name: str, preferred_ids: Optional[List[str]] = None) -> Optional[int]:
            skill_lower = skill_name.lower().strip()
            matched_ids = set(preferred_ids or [])
            years_total = 0

            for exp in all_experiences:
                if not isinstance(exp, dict):
                    continue
                exp_id = str(exp.get('id') or '').strip()
                if matched_ids:
                    if exp_id not in matched_ids:
                        continue
                else:
                    ach_list = exp.get('achievements') or []
                    ach_text = ' '.join(
                        (a.get('text') if isinstance(a, dict) else str(a))
                        for a in ach_list
                    )
                    haystack = ' '.join([
                        str(exp.get('title') or ''),
                        str(exp.get('company') or ''),
                        ach_text,
                    ]).lower()
                    if skill_lower not in haystack:
                        continue
                years_total += _experience_years(exp)

            return years_total if years_total > 0 else None

        if extra_skills:
            existing_skill_names = {s.get('name', '') for s in all_skills}
            prepend = []
            for raw_skill in extra_skills:
                extra_skill = self._normalize_extra_skill_entry(raw_skill)
                if not extra_skill:
                    continue
                skill_name = extra_skill.get('name', '')
                if skill_name not in omitted_skill_names and skill_name not in existing_skill_names:
                    preferred_ids = match_overrides.get(skill_name, [])
                    years = _derive_years_from_matches(skill_name, preferred_ids)
                    skill_entry: Dict[str, Any] = dict(extra_skill)
                    if years is not None and skill_entry.get('years') is None:
                        skill_entry['years'] = years
                    prepend.append(skill_entry)
            selected_skills = prepend + selected_skills

        # Override: if the user has explicitly reordered skill rows via the UI,
        # apply their ordering stored in customizations['skill_row_order']
        # as a list of skill names in the desired display order.
        skill_row_order = customizations.get('skill_row_order', [])
        if skill_row_order:
            order_map = {name: i for i, name in enumerate(skill_row_order)}
            selected_skills = sorted(
                selected_skills,
                key=lambda s: order_map.get(s.get('name', ''), len(order_map)),
            )

        # Select professional summary from the resolved session/master view.
        summary_view = SessionDataView(
            self.master_data,
            customizations,
            customizations,
        )
        # duckflow:
        #   id: summary_orchestrator_select
        #   kind: orchestrator
        #   timestamp: "2026-03-27T01:23:28Z"
        #   status: shared
        #   reads: ["customizations:summary_focus", "customizations:session_summaries"]
        #   writes: ["cv:selected_content.summary"]
        #   notes: "Resolves the active summary text by overlaying session variants over master variants and selecting the requested key."
        selected_summary = summary_view.selected_summary()

        # Apply an estimated page cap to body content (summary, experience,
        # achievements, skills) before publications are considered.
        max_cv_pages = customizations.get('max_cv_pages')
        if max_cv_pages is None:
            max_cv_pages = cfg.get('generation.max_cv_pages')
        if max_cv_pages is not None:
            chars_per_page = cfg.get('generation.cv_body_chars_per_page', 2500)
            (
                selected_summary,
                selected_experiences,
                selected_achievements,
                selected_skills,
            ) = self._cap_cv_body_to_pages(
                selected_summary,
                selected_experiences,
                selected_achievements,
                selected_skills,
                float(max_cv_pages),
                int(chars_per_page),
            )

        # Select publications — honour user accept/reject decisions if present
        accepted_pubs = customizations.get('accepted_publications')  # list of cite_keys or None
        rejected_pubs = set(customizations.get('rejected_publications') or [])

        # When a page-based publication cap is active, bypass the count limit —
        # _cap_publications_to_pages() (called below) handles trimming instead.
        _pub_page_cap_active = (
            customizations.get('max_publication_pages') is not None
            or cfg.get('generation.max_publication_pages') is not None
        )
        _pub_count_cap = None if _pub_page_cap_active else max_pubs

        if accepted_pubs is not None:
            # User has explicitly selected publications — preserve membership
            # and respect their count exactly (no automatic cap).
            # _pub_count_cap still applies to the fallback path below.
            accepted_set = set(accepted_pubs)
            pub_by_key = {}
            for pub in self._select_publications(job_analysis, max_count=len(self.publications) if self.publications else 50):
                key = pub.get('key', '') or ''
                if key in accepted_set and key not in rejected_pubs:
                    pub_by_key[key] = pub
            selected_publications = [
                pub_by_key[k] for k in accepted_pubs if k in pub_by_key
            ]
        else:
            selected_publications = self._select_publications(
                job_analysis,
                max_count=_pub_count_cap,
            )

        selected_publications = self._sort_selected_publications(
            selected_publications,
            customizations,
        )

        # Apply page-based publication cap when set in session or config.
        # Customizations take precedence; falls back to generation.max_publication_pages.
        max_pub_pages = customizations.get('max_publication_pages')
        if max_pub_pages is None:
            max_pub_pages = cfg.get('generation.max_publication_pages')
        if max_pub_pages is not None:
            chars_per_page = cfg.get('generation.publication_chars_per_page', 1500)
            selected_publications = self._cap_publications_to_pages(
                selected_publications, float(max_pub_pages), int(chars_per_page)
            )

        summary_warnings = self._validate_summary(selected_summary, job_analysis)

        publication_warnings = []
        for pub in selected_publications:
            raw = self.publications.get(pub.get('key') or '')
            if raw and raw.get('venue_warning'):
                publication_warnings.append(
                    f"“{pub.get('title') or pub.get('key') or '?'}”: {raw['venue_warning']}"
                )

        return {
            'personal_info': self.master_data.get('personal_info', {}),
            'summary': selected_summary,
            'summary_warnings': summary_warnings,
            'experiences': selected_experiences,
            'achievements': selected_achievements,
            'skills': selected_skills,
            'skill_category_order': customizations.get('skill_category_order', []),
            'education': self.master_data.get('education', []),
            'certifications': self.master_data.get('certifications', []),
            'publications': selected_publications,
            'publication_warnings': publication_warnings,
            'awards': self.master_data.get('awards', [])
        }

    @staticmethod
    def _estimate_cv_body_chars(
        summary: Any,
        experiences: List[Dict],
        achievements: List[Dict],
        skills: List[Dict],
    ) -> int:
        """Estimate rendered body size using text length plus layout overhead."""
        total = len(str(summary or '').strip())

        for exp in experiences or []:
            if not isinstance(exp, dict):
                continue
            total += 140  # entry-level layout overhead
            total += len(str(exp.get('title') or ''))
            total += len(str(exp.get('company') or ''))
            total += len(str(exp.get('start_date') or exp.get('start') or ''))
            total += len(str(exp.get('end_date') or exp.get('end') or ''))

            bullets = exp.get('ordered_achievements')
            if not isinstance(bullets, list):
                bullets = exp.get('achievements') or []
            for ach in bullets:
                text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                total += max(len(str(text)), 24) + 36

        for ach in achievements or []:
            text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
            total += max(len(str(text)), 20) + 28

        for skill in skills or []:
            name = skill.get('name', '') if isinstance(skill, dict) else str(skill)
            total += max(len(str(name)), 6) + 10

        return total

    @staticmethod
    def _validate_summary(summary: Any, job_analysis: Dict) -> List[str]:
        """Return a list of warning strings for summary quality issues."""
        text = str(summary or '').strip()
        if not text:
            return ['Summary is empty.']
        warnings: List[str] = []

        # Check 1: summary must not open with first-person "I"
        first_word = text.split()[0].rstrip('.,;:')
        if first_word == 'I':
            warnings.append(
                'Summary opens with "I" — avoid first-person pronouns in professional summaries.'
            )

        # Check 2: word count in target range (40–250 words)
        word_count = len(text.split())
        if word_count < 40:
            warnings.append(
                f'Summary is short ({word_count} words) — aim for 40–250 words for a senior candidate.'
            )
        elif word_count > 250:
            warnings.append(
                f'Summary is long ({word_count} words) — aim for 40–250 words to keep recruiter attention.'
            )

        # Check 3: top-3 required skills from job analysis should appear in summary
        required_skills: List[str] = [
            s.lower() for s in (job_analysis or {}).get('required_skills', []) if s
        ]
        if required_skills:
            top3 = required_skills[:3]
            text_lower = text.lower()
            missing = [s for s in top3 if s not in text_lower]
            if missing:
                warnings.append(
                    f'Summary does not mention top required skill(s): {", ".join(missing)}. '
                    'Consider weaving them in naturally.'
                )

        return warnings

    def _cap_cv_body_to_pages(
        self,
        summary: Any,
        experiences: List[Dict],
        achievements: List[Dict],
        skills: List[Dict],
        max_pages: float,
        chars_per_page: int = 2500,
    ) -> Tuple[Any, List[Dict], List[Dict], List[Dict]]:
        """Trim body content until estimated size fits the requested page budget."""
        if max_pages is None:
            return summary, experiences, achievements, skills

        budget = int(max_pages * chars_per_page)
        if budget <= 0:
            return '', [], [], []

        out_summary = summary
        out_experiences = [dict(exp) for exp in (experiences or []) if isinstance(exp, dict)]
        out_achievements = list(achievements or [])
        out_skills = list(skills or [])

        def _current() -> int:
            return self._estimate_cv_body_chars(
                out_summary,
                out_experiences,
                out_achievements,
                out_skills,
            )

        if _current() <= budget:
            return out_summary, out_experiences, out_achievements, out_skills

        # Remove less-critical sections in this order: skills, standalone
        # achievements, then lower-priority experience bullets/entries.
        while out_skills and _current() > budget:
            out_skills.pop()

        while out_achievements and _current() > budget:
            out_achievements.pop()

        while _current() > budget and out_experiences:
            changed = False
            for exp in reversed(out_experiences):
                bullets = exp.get('ordered_achievements')
                key = 'ordered_achievements'
                if not isinstance(bullets, list):
                    bullets = exp.get('achievements')
                    key = 'achievements'
                if isinstance(bullets, list) and len(bullets) > 1:
                    bullets = list(bullets)
                    bullets.pop()
                    exp[key] = bullets
                    changed = True
                    if _current() <= budget:
                        break
            if not changed:
                break

        while len(out_experiences) > 1 and _current() > budget:
            out_experiences.pop()

        if _current() > budget:
            summary_text = str(out_summary or '')
            max_summary_len = max(200, budget // 6)
            if len(summary_text) > max_summary_len:
                out_summary = summary_text[:max_summary_len].rstrip() + '...'

        return out_summary, out_experiences, out_achievements, out_skills
    
    @staticmethod
    def _cap_publications_to_pages(
        pubs: List[Dict],
        max_pages: float,
        chars_per_page: int = 1500,
    ) -> List[Dict]:
        """Trim the publication list to fit within an estimated page budget.

        Uses citation character length as a proxy for rendered height.
        Reads ``formatted_citation`` when available (post-format), otherwise
        falls back to ``formatted`` (raw value from ``_select_publications``).
        The default chars_per_page (1500) is calibrated from observed output:
        ~10 publications spanning ~4 pages at the standard 0.88 em font-size.

        Always keeps at least one publication even if it alone exceeds the
        budget.
        """
        if not pubs or max_pages is None:
            return pubs
        budget = max_pages * chars_per_page
        cumulative = 0
        for i, pub in enumerate(pubs):
            # formatted_citation is available post-format; formatted is the raw
            # value returned by _select_publications before _format_publications runs.
            citation = pub.get('formatted_citation') or pub.get('formatted', '') or ''
            # Minimum 80 chars accounts for formatting overhead even for
            # very short citations (year, authors, title wrapper).
            cumulative += max(len(citation), 80)
            if cumulative > budget:
                return pubs[:max(i, 1)]  # always include at least one entry
        return pubs

    def _select_publications(self, job_analysis: Dict, max_count: Optional[int] = 10) -> List[Dict]:
        """Select most relevant publications, sorted by heuristic relevance score.

        When *max_count* is None, all publications are returned (sorted by score),
        which lets callers partition them into recommended/excluded with full scores.
        """
        if not self.publications:
            return []

        domain = job_analysis.get('domain', '')
        keywords = set(job_analysis.get('ats_keywords', []))

        scored_pubs = []
        for key, pub in self.publications.items():
            score = 0.0
            reasons: list[str] = []

            # Recent publications score higher
            year = self._publication_year_value(pub)
            if year is not None:
                if year >= 2020:
                    score += 30
                    reasons.append(f'recent ({year})')
                elif year >= 2015:
                    score += 20
                    reasons.append(f'recent ({year})')
                elif year >= 2010:
                    score += 10
                    reasons.append(f'{year}')

            # Type bonus
            if pub['type'] == 'article':
                score += 25
                reasons.append('journal article')
            elif pub['type'] in ['inproceedings', 'conference']:
                score += 20
                reasons.append('conference paper')

            # Keyword matches
            title_lower = pub['title'].lower()
            matches = sum(1 for kw in keywords if kw.lower() in title_lower)
            score += matches * 5
            if matches:
                reasons.append(f'{matches} keyword match{"es" if matches > 1 else ""}')

            # Domain-specific
            if domain == 'genomics' and any(
                term in title_lower for term in ['genom', 'gene', 'dna', 'rna']
            ):
                score += 15
                reasons.append('domain match')

            # Normalise to 0–10 scale (raw max ≈ 70)
            normalized = min(10.0, round(score / 7, 1))
            rationale = ('Heuristic: ' + ', '.join(reasons)) if reasons else 'Heuristic: no strong match'

            scored_pubs.append((key, pub, score, normalized, rationale))

        scored_pubs.sort(key=lambda x: x[2], reverse=True)

        limit = max_count  # None → return all
        selected = []
        for key, pub, _raw, normalized, rationale in scored_pubs[:limit]:
            formatted = format_publication(pub, style='apa')
            year_value = pub.get('year', '')
            if not str(year_value or '').strip():
                parsed_year = self._publication_year_value(pub)
                year_value = str(parsed_year) if parsed_year is not None else ''
            selected.append({
                'key': key,
                'formatted': formatted,
                'year': year_value,
                'type': pub['type'],
                'authors': pub.get('authors', ''),
                'title': pub.get('title', ''),
                'journal': pub.get('journal', ''),
                'booktitle': pub.get('booktitle', ''),
                'institution': pub.get('institution', ''),
                'school': pub.get('school', ''),
                'fields': pub.get('fields', {}),
                'relevance_score': normalized,
                'rationale': rationale,
            })

        return selected
    
    def _generate_ats_docx(
        self,
        content: Dict,
        job_analysis: Dict,
        output_dir: Path
    ) -> tuple:
        """Generate ATS-optimized DOCX with enhanced formatting and validation.

        Returns (filepath, ats_score) so callers can persist the score to metadata.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        
        # Set up ATS-optimized styles
        self._setup_ats_styles(doc)
        
        # Header with contact information (ATS-friendly format)
        personal = content['personal_info']
        name = personal.get('name', '')
        
# Candidate name — large bold run (not a Heading style so it does not
        # compete with section Heading 1 paragraphs in the ATS heading hierarchy).
        name_para = doc.add_paragraph()
        name_run  = name_para.add_run(name)
        name_run.bold      = True
        name_run.font.size = Pt(16)
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Contact information — single line, pipe-separated (ATS standard).
        # City/state only (no street address); phone normalized to NNN-NNN-NNNN.
        contact = personal.get('contact', {})
        contact_parts = []

        if contact.get('address_display'):
            contact_parts.append(contact['address_display'])
        elif contact.get('address', {}).get('city'):
            city  = contact['address']['city']
            state = contact['address'].get('state', '')
            contact_parts.append(f"{city}, {state}".strip(', '))
        if contact.get('phone'):
            contact_parts.append(self._normalize_phone(contact['phone']))
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('linkedin'):
            contact_parts.append(contact['linkedin'])

        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add spacing
        doc.add_paragraph()

        # Professional Summary — ATS standard label, Heading 1 style.
        summary_heading = doc.add_paragraph('Professional Summary', style='Heading 1')
        
        summary_text = content.get('summary', '')
        # Enhance summary with job-specific keywords
        enhanced_summary = self._enhance_summary_for_ats(summary_text, job_analysis)
        doc.add_paragraph(enhanced_summary)
        doc.add_paragraph()
        
        # Skills — split hard/soft into "Technical Skills" and "Core Competencies".
        # _optimize_skills_for_ats returns names in priority order; type inferred
        # per skill via _classify_skill_type.
        ats_skill_names = self._optimize_skills_for_ats(content['skills'], job_analysis)
        skill_map = {s.get('name', ''): s for s in content.get('skills', [])}
        hard_skills = [n for n in ats_skill_names
                       if self._classify_skill_type(skill_map.get(n, {})) == 'hard']
        soft_skills = [n for n in ats_skill_names
                       if self._classify_skill_type(skill_map.get(n, {})) == 'soft']

        if hard_skills:
            doc.add_paragraph('Technical Skills', style='Heading 1')
            doc.add_paragraph(' • '.join(hard_skills))
            doc.add_paragraph()

        if soft_skills:
            doc.add_paragraph('Core Competencies', style='Heading 1')
            doc.add_paragraph(' • '.join(soft_skills))
            doc.add_paragraph()

        # Work Experience — ATS standard label, Heading 1 style.
        doc.add_paragraph('Work Experience', style='Heading 1')
        
        for exp in content['experiences']:
            # One-line job entry: Title | Company | Location | Date Range (US-H5).
            loc_parts = []
            if exp.get('location', {}).get('city'):
                loc_parts.append(exp['location']['city'])
            if exp.get('location', {}).get('state'):
                loc_parts.append(exp['location']['state'])
            location_str  = ', '.join(loc_parts) if loc_parts else ''
            date_range     = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            entry_parts    = [exp.get('title', ''), exp.get('company', '')]
            if location_str:
                entry_parts.append(location_str)
            entry_parts.append(date_range)
            entry_line = ' | '.join(p for p in entry_parts if p)

            entry_para = doc.add_paragraph()
            entry_run  = entry_para.add_run(entry_line)
            entry_run.bold      = True
            entry_run.font.size = Pt(11)
            
            # Achievements - Bullet points with quantified results
            if exp.get('achievements'):
                for achievement in exp['achievements']:
                    achievement_text = achievement.get('text', '') if isinstance(achievement, dict) else str(achievement)
                    # Enhance achievement with keywords if needed
                    enhanced_achievement = self._enhance_achievement_for_ats(achievement_text, job_analysis)
                    achievement_para = doc.add_paragraph(enhanced_achievement, style='List Bullet')
                    achievement_para.paragraph_format.left_indent = Pt(18)
            
            doc.add_paragraph()  # Spacing between positions
        
        # Education — ATS standard label, Heading 1 style.
        if content.get('education'):
            doc.add_paragraph('Education', style='Heading 1')
            
            for edu in content['education']:
                degree = edu.get('degree', '')
                field = edu.get('field', '')
                institution = edu.get('institution', '')
                year = edu.get('end_year', '')
                
                degree_line = f"{degree} {field}".strip()
                institution_line = f"{institution}"
                if year:
                    institution_line += f" | {year}"
                
                degree_para = doc.add_paragraph()
                degree_para.add_run(degree_line).bold = True
                doc.add_paragraph(institution_line)
            
            doc.add_paragraph()
        
        # Additional Sections (if present)
        self._add_ats_additional_sections(doc, content, job_analysis)
        
        # Save with ATS-optimized filename
        company = job_analysis.get('company', 'Company').replace(' ', '').replace('/', '-')[:15]
        role = job_analysis.get('title', 'Role').replace(' ', '').replace('/', '-')[:20]
        timestamp = datetime.now().strftime("%Y-%m-%d")
        
        filename = f"CV_{company}_{role}_{timestamp}_ATS.docx"
        filepath = output_dir / filename

        if content.get('ai_attribution'):
            doc.core_properties.keywords = 'AI-assisted'
            doc.core_properties.subject  = 'CV generated with AI assistance using cv-builder'

        doc.save(str(filepath))

        # Validate ATS compatibility
        ats_score = self._validate_ats_compatibility(content, job_analysis)
        logger.info("Generated ATS DOCX: %s (ATS Score: %d/100)", filename, ats_score)

        return filepath, ats_score

    def _setup_ats_styles(self, doc):
        """Set up ATS-optimized document styles."""
        from docx.shared import Pt, RGBColor
        
        # Create custom styles that are ATS-friendly
        styles = doc.styles

        # Heading 1 — used for all section headings in the ATS DOCX.
        # (Candidate name is rendered as a bold run, not a Heading style.)
        try:
            heading1 = styles['Heading 1']
            heading1.font.name = 'Calibri'
            heading1.font.size = Pt(12)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

        # Heading 2 — available for optional sub-sections if needed.
        try:
            heading2 = styles['Heading 2']
            heading2.font.name = 'Calibri'
            heading2.font.size = Pt(11)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

        # Default body paragraph style — sets ATS-safe Calibri across all body runs.
        try:
            normal = styles['Normal']
            normal.font.name = 'Calibri'
            normal.font.size = Pt(11)
        except KeyError:
            pass

        # Clean list style
        try:
            list_bullet = styles['List Bullet']
            list_bullet.font.name = 'Calibri'
            list_bullet.font.size = Pt(10)
        except KeyError:
            pass
    
    def _enhance_summary_for_ats(self, summary: str, job_analysis: Dict) -> str:
        """Return the professional summary unchanged.

        Terminology improvements are handled upstream via
        :meth:`apply_approved_rewrites` before the content reaches this
        stage.  This method is retained as the call site in the ATS DOCX
        generator but no longer mutates the text.

        When no LLM is configured a keyword-gap note is logged so the
        operator is aware of potential ATS misalignment without the output
        being silently altered.
        """
        if not summary:
            return summary

        if not self.llm:
            # Identify missing keywords and log a gap warning.
            summary_lower    = summary.lower()
            key_skills       = job_analysis.get('required_skills', [])
            missing_keywords = [
                s for s in key_skills[:5] if s.lower() not in summary_lower
            ]
            if missing_keywords:
                logger.warning(
                    "_enhance_summary_for_ats: no LLM configured; "
                    "summary may be missing keywords: %s",
                    ', '.join(missing_keywords)
                )
        else:
            logger.info(
                "_enhance_summary_for_ats: summary rewrites are handled "
                "upstream via apply_approved_rewrites — returning unchanged."
            )

        return summary
    
    def _optimize_skills_for_ats(self, skills: List[Dict], job_analysis: Dict) -> List[str]:
        """Return a score-ordered, deduplicated subset of skill names.

        Synonym expansion is applied so that a skill named 'ML' scores a
        match against job keyword 'Machine Learning' and vice versa.
        Only reorders and selects skills — terminology is never renamed here.
        All vocabulary changes must come via :meth:`apply_approved_rewrites`
        before content reaches this method.
        """
        ats_keywords = set(kw.lower() for kw in job_analysis.get('ats_keywords', []))
        required_skills = set(skill.lower() for skill in job_analysis.get('required_skills', []))

        # Expand ATS keywords via synonym map so we can match either direction
        expanded_ats: set = set(ats_keywords)
        for kw in list(ats_keywords):
            canonical = self._expansion_index.get(kw)
            if canonical:
                expanded_ats.add(canonical.lower())
        expanded_required: set = set(required_skills)
        for req in list(required_skills):
            canonical = self._expansion_index.get(req)
            if canonical:
                expanded_required.add(canonical.lower())

        # Priority scoring for skills
        skill_scores = []
        for skill in skills:
            name = skill.get('name', '')
            name_lower = name.lower()
            canonical_lower = self.canonical_skill_name(name).lower()
            years = skill.get('years', 0)

            score = 0
            # High priority for exact keyword matches (direct or via synonym)
            if name_lower in expanded_ats or canonical_lower in expanded_ats:
                score += 50
            if name_lower in expanded_required or canonical_lower in expanded_required:
                score += 40
            # Years of experience bonus
            score += min(years * 2, 20)

            skill_scores.append((name, score))

        # Sort by score and return top skills
        skill_scores.sort(key=lambda x: x[1], reverse=True)

        # Return optimized skill names (top 15 for ATS readability)
        return [skill[0] for skill in skill_scores[:15]]
    
    def _enhance_achievement_for_ats(self, achievement: str, job_analysis: Dict) -> str:
        """Return the achievement text unchanged.

        Checks whether the text opens with a strong action verb and logs a
        warning when it does not, but never modifies the text.  Rewrites are
        handled upstream via :meth:`apply_approved_rewrites`.
        """
        if not achievement:
            return achievement

        text = achievement.strip()
        first_word = self._opening_word_for_verb_check(text)
        if not first_word or first_word.lower() not in self._STRONG_VERBS_LOWER:
            logger.warning(
                "_enhance_achievement_for_ats: bullet does not start with a strong action verb: %r",
                text[:60]
            )

        return text

    # ── Persuasion vocabulary ──────────────────────────────────────────────────

    _STRONG_VERBS: frozenset = frozenset({
        'Accelerated', 'Achieved', 'Architected', 'Automated', 'Built',
        'Championed', 'Coined', 'Conceived', 'Conducted', 'Consolidated',
        'Created', 'Cut', 'Delivered', 'Demonstrated', 'Deployed',
        'Designed', 'Developed', 'Directed', 'Doubled',
        'Drove', 'Enabled', 'Established', 'Expanded', 'Generated',
        'Grew', 'Improved', 'Implemented', 'Increased', 'Integrated',
        'Invented', 'Launched', 'Led', 'Managed', 'Optimized', 'Pioneered',
        'Published',
        'Provided', 'Raised', 'Reduced', 'Refactored', 'Scaled', 'Shipped',
        'Secured', 'Spearheaded', 'Streamlined', 'Taught',
        'Transformed', 'Translated', 'Tripled',
        'Founded',
    })
    _STRONG_VERBS_LOWER: frozenset = frozenset(v.lower() for v in _STRONG_VERBS)

    _WEAK_VERBS: frozenset = frozenset({
        'Assisted', 'Contributed', 'Helped', 'Participated',
        'Supported', 'Supervised', 'Worked', 'Was responsible',
        'Was involved', 'Collaborated', 'Cooperated',
    })
    _WEAK_VERBS_LOWER: frozenset = frozenset(v.lower() for v in _WEAK_VERBS)
    # First-word lookup used in check_persuasion — multi-word entries like
    # 'Was responsible' match on 'was' so passive constructions are caught.
    _WEAK_VERB_FIRST_WORDS_LOWER: frozenset = frozenset(
        v.split()[0].lower() for v in _WEAK_VERBS
    )

    _VAGUE_PHRASES: tuple = (
        'various tasks', 'multiple tasks', 'several tasks',
        'day-to-day', 'various projects', 'multiple projects',
        'various responsibilities', 'general support', 'helped to',
        'assisted with', 'participated in', 'was part of',
        'involved in', 'worked on various', 'worked on multiple',
        'responsible for', 'key player', 'hands-on experience', 'wearing many hats',
    )

    _VAGUE_PHRASES_RE = re.compile(
        r'\b(' + '|'.join(re.escape(phrase) for phrase in _VAGUE_PHRASES) + r')\b',
        re.IGNORECASE,
    )

    _NEGATIVE_FRAMING_RE = re.compile(
        r'\b(despite|although|even though|in spite of|rather than|instead of'
        r'|unfortunately|was not (able|allowed|given)'
        r'|without (adequate|sufficient|proper) (resources?|budget|support|time)'
        r'|limited (budget|resources?|support|headcount)'
        r'|lack of|no (budget|resources?|dedicated|formal))\b',
        re.IGNORECASE,
    )

    @classmethod
    def _opening_text_for_verb_check(cls, text: str) -> str:
        """Return text with an optional leading descriptor label removed.

        Example: "Statistical Genomics: Developed ..." -> "Developed ..."
        """
        stripped = (text or '').strip()
        if not stripped:
            return stripped

        first_word = stripped.split()[0].lower() if stripped.split() else ''
        if first_word in cls._STRONG_VERBS_LOWER or first_word in cls._WEAK_VERB_FIRST_WORDS_LOWER:
            return stripped

        if ':' not in stripped:
            return stripped

        prefix, remainder = stripped.split(':', 1)
        prefix_words = [w for w in re.findall(r"[A-Za-z0-9][A-Za-z0-9'&/.-]*", prefix) if w]
        if not remainder.strip() or not prefix_words:
            return stripped

        if len(prefix_words) > 6:
            return stripped

        # Heuristic: descriptor labels are usually title-case/all-caps noun phrases.
        if all(w.isupper() or w[:1].isupper() or any(ch.isdigit() for ch in w) for w in prefix_words):
            return remainder.strip()

        return stripped

    @classmethod
    def _opening_word_for_verb_check(cls, text: str) -> str:
        """Return the first meaningful opening word for verb checks."""
        normalized = cls._opening_text_for_verb_check(text)
        match = re.search(r"[A-Za-z][A-Za-z'/-]*", normalized)
        return match.group(0) if match else ''

    _METRIC_RE = re.compile(
        r'(?!(?:19|20)\d{2}(?:[–\-]\d{4})?)'  # negative lookahead: exclude year patterns like 2020-2024
        r'((?:\d{1,3}(?:[,\s]\d{3})*|\d+)\s*%?'  # digit-based metric with optional commas/spaces and %
        r'|\$[\d,]+[kmb]?'         # dollar amount
        r'|£[\d,]+[kmb]?'          # pound amount
        r'|€[\d,]+[kmb]?'          # euro amount
        r'|\b\d+\s*x\b'            # multiplier (3x)
        r'|\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fifteen|twenty|hundred|thousand)\b)',  # spelled-out numbers
        re.IGNORECASE,
    )

    # ── ATS skill-type classification ──────────────────────────────────────────

    _SOFT_SKILL_CATEGORIES: frozenset = frozenset({
        'soft', 'soft skills', 'interpersonal', 'leadership', 'communication',
        'core competencies', 'management', 'personal', 'professional skills',
        'collaboration', 'people skills',
    })

    _SOFT_SKILL_NAMES: frozenset = frozenset({
        'communication', 'leadership', 'teamwork', 'collaboration',
        'problem solving', 'critical thinking', 'adaptability', 'creativity',
        'time management', 'organization', 'attention to detail',
        'emotional intelligence', 'empathy', 'conflict resolution',
        'negotiation', 'presentation', 'mentoring', 'coaching',
        'strategic thinking', 'decision making', 'stakeholder management',
        'change management', 'cross-functional collaboration',
    })

    @staticmethod
    def _normalize_phone(phone: str) -> str:
        """Normalize phone to NNN-NNN-NNNN format (no parentheses or spaces)."""
        if not phone:
            return phone
        digits = re.sub(r'\D', '', phone)
        if len(digits) == 11 and digits.startswith('1'):
            digits = digits[1:]
        if len(digits) == 10:
            return f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
        return phone

    @classmethod
    def _classify_skill_type(cls, skill: Dict) -> str:
        """Return 'soft' if skill is a soft/interpersonal skill, else 'hard'.

        Checks ``skill_type`` stored field first (allowing explicit overrides),
        then falls back to category- and name-based heuristics.
        """
        stored = (skill.get('skill_type') or '').lower()
        if stored in ('hard', 'soft'):
            return stored
        category = (skill.get('category') or '').lower().strip()
        name     = (skill.get('name')     or '').lower().strip()
        for soft_cat in cls._SOFT_SKILL_CATEGORIES:
            if soft_cat in category:
                return 'soft'
        if name in cls._SOFT_SKILL_NAMES:
            return 'soft'
        return 'hard'

    def check_persuasion(self, experiences: List[Dict]) -> Dict:
        """Analyse experience bullets for persuasion quality.

        Parameters
        ----------
        experiences:
            List of experience dicts (each with ``id`` and ``achievements``).

        Returns
        -------
        Dict with keys:
          - ``findings``: list of finding dicts (exp_id, bullet_index, text,
            severity, issues)
          - ``summary``: {total_bullets, flagged, strong_count}
        """
        findings = []
        total_bullets = 0
        strong_count  = 0

        for exp in experiences:
            exp_id = exp.get('id', '')
            achievements = exp.get('ordered_achievements') or exp.get('achievements') or []
            exp_first_words: List[Tuple[str, int]] = []   # (first_word_lower, idx)
            finding_by_bullet: Dict[Tuple[str, int], Dict] = {}  # (exp_id, idx) → finding

            for idx, ach in enumerate(achievements):
                text = (ach.get('text', '') if isinstance(ach, dict) else str(ach)).strip()
                if not text:
                    continue
                total_bullets += 1
                issues = []
                first_word = self._opening_word_for_verb_check(text)
                first_word_lower = first_word.lower()

                if first_word_lower:
                    exp_first_words.append((first_word_lower, idx))

                # Weak opening verb — exact first-word match (no prefix collisions)
                if first_word_lower in self._WEAK_VERB_FIRST_WORDS_LOWER:
                    issues.append({
                        'type':       'weak_verb',
                        'severity':   'warning',
                        'suggestion': (
                            f'Replace "{first_word}" with a stronger action verb '
                            '(e.g. Led, Built, Delivered, Reduced, Improved).'
                        ),
                    })
                elif first_word_lower not in self._STRONG_VERBS_LOWER:
                    issues.append({
                        'type':       'no_strong_verb',
                        'severity':   'info',
                        'suggestion': (
                            'Consider opening with a strong action verb '
                            '(e.g. Led, Built, Delivered, Reduced, Improved).'
                        ),
                    })

                # Missing quantification
                if not self._METRIC_RE.search(text):
                    issues.append({
                        'type':       'no_metric',
                        'severity':   'warning',
                        'suggestion': (
                            'Add a quantified result — percentage improvement, '
                            'team size, dollar value, or time saved.'
                        ),
                    })

                # Vague language
                text_lower = text.lower()
                vague_matches = self._VAGUE_PHRASES_RE.findall(text_lower)
                for phrase in vague_matches:
                    issues.append({
                        'type':       'vague_language',
                        'severity':   'warning',
                        'suggestion': (
                            f'Replace vague phrase "{phrase}" with a specific, '
                            'measurable description of impact.'
                        ),
                    })

                # Too short
                if len(text.split()) < 8:
                    issues.append({
                        'type':       'too_short',
                        'severity':   'info',
                        'suggestion': (
                            'Expand this bullet to include context, action, and result '
                            '(aim for 15–25 words).'
                        ),
                    })

                # Negative / defensive framing — positive-sum bullets focus on what was achieved
                neg_match = self._NEGATIVE_FRAMING_RE.search(text)
                if neg_match:
                    issues.append({
                        'type':       'negative_framing',
                        'severity':   'info',
                        'suggestion': (
                            f'Phrase "{neg_match.group(0)}" frames the bullet defensively. '
                            'Rewrite to focus on what you achieved rather than constraints.'
                        ),
                    })

                if not issues:
                    strong_count += 1
                else:
                    finding = {
                        'exp_id':       exp_id,
                        'bullet_index': idx,
                        'text':         text,
                        'severity':     max(
                            (i['severity'] for i in issues),
                            key=lambda s: 0 if s == 'info' else 1,
                        ),
                        'issues': issues,
                    }
                    findings.append(finding)
                    finding_by_bullet[(exp_id, idx)] = finding

            # Repeated opening verb detection: flag 2nd+ occurrences when ≥3 bullets share a verb
            verb_counts = Counter(fw for fw, _ in exp_first_words)
            seen_verb_occurrences: Dict[str, int] = {}
            for fw, idx in exp_first_words:
                total_for_verb = verb_counts[fw]
                if total_for_verb < 3:
                    continue
                occurrence = seen_verb_occurrences.get(fw, 0)
                seen_verb_occurrences[fw] = occurrence + 1
                if occurrence == 0:
                    continue  # First occurrence is fine; only flag repetitions
                suggestion = (
                    f'"{fw.capitalize()}" opens {total_for_verb} bullets in this role. '
                    'Vary your action verbs to show a broader range of contributions.'
                )
                rv_issue = {'type': 'repeated_verb', 'severity': 'warning', 'suggestion': suggestion}
                key = (exp_id, idx)
                if key in finding_by_bullet:
                    entry = finding_by_bullet[key]
                    entry['issues'].append(rv_issue)
                    if entry['severity'] == 'info':
                        entry['severity'] = 'warning'
                else:
                    ach = achievements[idx]
                    ach_text = (ach.get('text', '') if isinstance(ach, dict) else str(ach)).strip()
                    new_finding = {
                        'exp_id':       exp_id,
                        'bullet_index': idx,
                        'text':         ach_text,
                        'severity':     'warning',
                        'issues':       [rv_issue],
                    }
                    findings.append(new_finding)
                    finding_by_bullet[key] = new_finding
                    strong_count -= 1

        # Narrative-thread advisory (GAP-281): warn when ≥3 themes are equally weighted.
        # Count bullets that mention each `relevant_for` tag. If the top 3 tags
        # all fall within 20% of the leading tag's count AND total tagged bullets ≥ 10,
        # emit a narrative_thread advisory in the summary.
        theme_counts: Counter = Counter()
        for exp in experiences:
            achievements = exp.get('ordered_achievements') or exp.get('achievements') or []
            for ach in achievements:
                if not isinstance(ach, dict):
                    continue
                for theme in (ach.get('relevant_for') or []):
                    theme_counts[theme.lower().strip()] += 1
        narrative_thread_advisory = None
        if len(theme_counts) >= 3:
            top_themes = theme_counts.most_common(3)
            top_count = top_themes[0][1]
            tagged_total = sum(theme_counts.values())
            if (
                tagged_total >= 10
                and top_count > 0
                and all(c >= top_count * 0.8 for _, c in top_themes)
            ):
                theme_labels = ', '.join(t for t, _ in top_themes)
                narrative_thread_advisory = {
                    'type':     'narrative_thread',
                    'severity': 'advisory',
                    'detail': (
                        f'Three narrative threads have similar weight: {theme_labels}. '
                        'A focused CV typically emphasises 1–2 primary themes. '
                        'Consider trimming or re-framing bullets to reinforce a clearer story.'
                    ),
                    'theme_counts': {t: c for t, c in top_themes},
                }

        # Narrative-arc advisory: most recent role should show strongest action verbs.
        # Uses start_year/end_year from each experience to determine temporal order.
        narrative_arc_advisory = None
        _current_year = datetime.now().year

        def _verb_strength_score(achs: List) -> Optional[float]:
            total_a, strong_a = 0, 0
            for ach in achs:
                text_a = (ach.get('text', '') if isinstance(ach, dict) else str(ach)).strip()
                if not text_a:
                    continue
                total_a += 1
                fw_a = self._opening_word_for_verb_check(text_a).lower()
                if fw_a in self._STRONG_VERBS_LOWER:
                    strong_a += 1
            return strong_a / total_a if total_a >= 2 else None

        timed_exps: List[Tuple] = []
        for exp in experiences:
            end_y = exp.get('end_year') or (_current_year if exp.get('current') else None)
            start_y = exp.get('start_year')
            sort_key = (end_y or start_y or 0, start_y or 0)
            achs_e = exp.get('ordered_achievements') or exp.get('achievements') or []
            score = _verb_strength_score(achs_e)
            if score is not None and sort_key[0] > 0:
                timed_exps.append((sort_key, score))

        if len(timed_exps) >= 3:
            timed_exps.sort(key=lambda x: x[0], reverse=True)  # most recent first
            recent_score = timed_exps[0][1]
            older_avg = sum(s for _, s in timed_exps[1:]) / (len(timed_exps) - 1)
            if older_avg > 0 and recent_score < older_avg * 0.70:
                narrative_arc_advisory = {
                    'type':     'narrative_arc',
                    'severity': 'advisory',
                    'detail': (
                        f'Your most recent role uses strong action verbs in '
                        f'{round(recent_score * 100)}% of bullets, '
                        f'compared to {round(older_avg * 100)}% for earlier roles. '
                        'A compelling CV shows growing impact in recent roles — '
                        'consider replacing weak verbs in your most recent role with '
                        'stronger action verbs (e.g. Led, Built, Delivered, Drove).'
                    ),
                }

        return {
            'findings': findings,
            'summary':  {
                'total_bullets':             total_bullets,
                'flagged':                   len(findings),
                'strong_count':              strong_count,
                'narrative_thread_advisory': narrative_thread_advisory,
                'narrative_arc_advisory':    narrative_arc_advisory,
            },
        }

    def _add_ats_additional_sections(self, doc, content: Dict, job_analysis: Dict):
        """Add additional sections that improve ATS scoring."""
        
        # Certifications (if present) — Heading 1, title-case ATS label.
        if content.get('certifications'):
            doc.add_paragraph('Certifications', style='Heading 1')
            
            for cert in content['certifications']:
                cert_name = cert.get('name', '')
                cert_issuer = cert.get('issuer', '')
                cert_year = cert.get('year', '')
                
                cert_line = cert_name
                if cert_issuer:
                    cert_line += f" | {cert_issuer}"
                if cert_year:
                    cert_line += f" ({cert_year})"
                
                doc.add_paragraph(cert_line)
            
            doc.add_paragraph()
        
        # Awards (if present and relevant) — Heading 1, title-case ATS label.
        if content.get('awards'):
            doc.add_paragraph('Awards', style='Heading 1')

            for award in content['awards']:
                award_title = award.get('title', '')
                award_year = award.get('year', '')
                award_desc = award.get('description', '')

                award_line = award_title
                if award_year:
                    award_line += f" ({award_year})"

                award_para = doc.add_paragraph()
                award_para.add_run(award_line).bold = True

                if award_desc:
                    doc.add_paragraph(award_desc)

            doc.add_paragraph()

        # Publications (if present) — plain-text citations for ATS keyword coverage.
        publications = content.get('publications', [])
        if publications:
            total_count = len(self.publications) if self.publications else 0
            heading_text = (
                'Selected Publications'
                if total_count and total_count > len(publications)
                else 'Publications'
            )
            doc.add_paragraph(heading_text, style='Heading 1')
            for pub in publications:
                citation = pub.get('formatted_citation', '').strip()
                if citation:
                    doc.add_paragraph(citation)
            doc.add_paragraph()
    
    def _validate_ats_compatibility(self, content: Dict, job_analysis: Dict) -> int:
        """Validate CV for ATS compatibility and return score out of 100."""
        score = 0
        max_score = 100
        
        # Check 1: Contact Information (20 points)
        contact = content.get('personal_info', {}).get('contact', {})
        if contact.get('email'):
            score += 8
        if contact.get('phone'):
            score += 6
        if contact.get('address') or contact.get('address_display'):
            score += 6
        
        # Check 2: Professional Summary (15 points)
        summary = content.get('summary', '')
        if len(summary) > 50:
            score += 10
        if len(summary) > 100:
            score += 5
        
        # Check 3: Skills Match (25 points)
        skills_list = [skill.get('name', '').lower() for skill in content.get('skills', [])]
        required_skills = [skill.lower() for skill in job_analysis.get('required_skills', [])]
        ats_keywords = [kw.lower() for kw in job_analysis.get('ats_keywords', [])]
        
        # Required skills coverage
        matched_required = sum(1 for skill in required_skills if skill in skills_list)
        if required_skills:
            score += int((matched_required / len(required_skills)) * 15)
        
        # ATS keywords coverage  
        matched_keywords = sum(1 for kw in ats_keywords[:10] if kw in ' '.join(skills_list))
        if ats_keywords:
            score += int((matched_keywords / min(len(ats_keywords), 10)) * 10)
        
        # Check 4: Experience Section (25 points)
        experiences = content.get('experiences', [])
        if experiences:
            score += 10
            # Check for quantified achievements
            total_achievements = sum(len(exp.get('achievements', [])) for exp in experiences)
            if total_achievements >= 8:
                score += 10
            elif total_achievements >= 4:
                score += 5
            
            # Check for recent experience
            if any('2023' in exp.get('end_date', '') or '2024' in exp.get('end_date', '') 
                  or exp.get('end_date') == 'Present' for exp in experiences):
                score += 5
        
        # Check 5: Education (10 points)
        if content.get('education'):
            score += 10
        
        # Check 6: Additional Sections (5 points)
        if content.get('certifications'):
            score += 3
        if content.get('awards'):
            score += 2
        
        return min(score, max_score)
    
    def _generate_human_docx(
        self,
        content: Dict,
        job_analysis: Dict,
        output_dir: Path,
        skills_heading: str = 'Technical Skills',
    ) -> Path:
        """Generate human-readable DOCX using python-docx with Calibri, standard margins.

        Sections (all conditional where marked):
        Name / Contact / Summary / Experience / Skills / Education /
        Certifications (if any) / Selected Publications (if any).
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        company   = job_analysis.get('company', 'Company').replace(' ', '')
        role      = job_analysis.get('title', 'Role').replace(' ', '')[:20]
        timestamp = datetime.now().strftime("%Y-%m-%d")
        filename  = f"CV_{company}_{role}_{timestamp}.docx"
        filepath  = output_dir / filename

        doc = Document()

        # ── Page margins (1 inch all sides) ─────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # ── Default paragraph style: Calibri 11 ─────────────────────────────
        style = doc.styles['Normal']
        font  = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # ── Helper functions ─────────────────────────────────────────────────
        def _heading(text: str, level: int = 1):
            style_name = 'Heading 1' if level == 1 else 'Heading 2'
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(13 if level == 1 else 11)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            # Bottom border (thin rule under section heading)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '2c3e50')
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            return p

        def _para(text: str = '', bold: bool = False, italic: bool = False,
                  size: int = 11, indent: float = 0.0, space_after: float = 2):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            if indent:
                p.paragraph_format.left_indent = Inches(indent)
            if text:
                run = p.add_run(text)
                run.bold   = bold
                run.italic = italic
                run.font.size = Pt(size)
            return p

        def _bullet(text: str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_after  = Pt(1)
            p.add_run(text)
            return p

        def _add_hyperlink(paragraph, text: str, url: str, size: int = 10):
            rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), rel_id)

            run = OxmlElement('w:r')
            run_props = OxmlElement('w:rPr')

            color = OxmlElement('w:color')
            color.set(qn('w:val'), '2980B9')
            run_props.append(color)

            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            run_props.append(underline)

            size_el = OxmlElement('w:sz')
            size_el.set(qn('w:val'), str(size * 2))
            run_props.append(size_el)

            run.append(run_props)
            text_el = OxmlElement('w:t')
            text_el.text = text
            run.append(text_el)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)

        # ── Name ─────────────────────────────────────────────────────────────
        personal_info = content.get('personal_info', {})
        name_para = doc.add_paragraph()
        name_run  = name_para.add_run(personal_info.get('name', ''))
        name_run.bold       = True
        name_run.font.size  = Pt(22)
        name_run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)

        # Job title line
        applicant_tagline = str(
            (content.get('template_metadata') or {}).get('applicant_tagline') or ''
        ).strip()
        if applicant_tagline:
            title_para = doc.add_paragraph()
            title_run  = title_para.add_run(applicant_tagline)
            title_run.italic    = True
            title_run.font.size = Pt(12)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(4)

        # ── Contact ──────────────────────────────────────────────────────────
        contact = personal_info.get('contact', {})
        contact_parts = []
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        address = contact.get('address', {})
        if address:
            city  = address.get('city', '')
            state = address.get('state', '')
            if city or state:
                contact_parts.append(f"{city}, {state}".strip(', '))
        if contact.get('linkedin'):
            contact_parts.append(contact['linkedin'].replace('https://', ''))
        if contact_parts:
            cp = doc.add_paragraph(' | '.join(contact_parts))
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(4)
            for run in cp.runs:
                run.font.size = Pt(10)

        # ── Professional Summary ─────────────────────────────────────────────
        summary = content.get('professional_summary', '')
        if summary:
            _heading('Professional Summary')
            _para(summary, space_after=4)

        # ── Experience ───────────────────────────────────────────────────────
        experiences = content.get('experiences', [])
        if experiences:
            _heading('Experience')
            for exp in experiences:
                # Role + Company on same line, dates on right
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(0)
                role_run = p.add_run(exp.get('title', ''))
                role_run.bold = True
                p.add_run('  ')
                co_run = p.add_run(exp.get('company', ''))
                co_run.italic = True
                # Dates as right-aligned run (approximate via tab stop)
                date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
                p.add_run(f"   {date_str}")
                loc = exp.get('location', {})
                if isinstance(loc, dict) and (loc.get('city') or loc.get('state')):
                    _para(f"{loc.get('city', '')}, {loc.get('state', '')}".strip(', '),
                          italic=True, size=10, space_after=1)
                for ach in exp.get('achievements', []):
                    text = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                    if text.strip():
                        _bullet(text)

        # ── Skills ───────────────────────────────────────────────────────────
        skills_by_category = content.get('skills_by_category', [])
        if skills_by_category:
            # duckflow:
            #   id: cv_render.scripts_utils_cv_orchestrator.L3340
            #   kind: artifact
            #   timestamp: "2026-03-27T02:31:32Z"
            #   status: live
            #   reads:
            #     - "artifact:skills_heading"
            #   writes:
            #     - "artifact:human_docx.skills_heading"
            #   notes: "Writes the resolved skills-section heading into the generated human-readable DOCX."
            _heading(skills_heading)
            for cat in skills_by_category:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                cat_run = p.add_run(f"{cat.get('category', '')}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                skills_list = cat.get('skills', [])
                skills_text = ', '.join(
                    s.get('name', s) if isinstance(s, dict) else str(s)
                    for s in skills_list
                )
                skill_run = p.add_run(skills_text)
                skill_run.font.size = Pt(10)

        # ── Education ────────────────────────────────────────────────────────
        education = content.get('education', [])
        if education:
            _heading('Education')
            for edu in education:
                degree = edu.get('degree', '')
                field  = edu.get('field', '')
                inst   = edu.get('institution', '')
                year   = edu.get('end_year') or edu.get('graduation_date', '')
                degree_str = f"{degree}, {field}" if field else degree
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                deg_run = p.add_run(degree_str)
                deg_run.bold = True
                p.add_run(f"  {inst}  ({year})")

        # ── Certifications ───────────────────────────────────────────────────
        certifications = content.get('certifications', [])
        if certifications:
            _heading('Certifications')
            for cert in certifications:
                name   = cert.get('name', '')
                issuer = cert.get('issuer', '')
                year   = cert.get('year', '')
                parts  = [name]
                if issuer:
                    parts.append(issuer)
                if year:
                    parts.append(f"({year})")
                _para(' | '.join(parts), space_after=2)

        # ── Selected Publications ────────────────────────────────────────────
        publications = content.get('publications', [])
        if publications:
            total_count = len(self.publications) if self.publications else 0
            heading_text = 'Selected Publications' if (total_count and total_count > len(publications)) else 'Publications'
            _heading(heading_text)
            for idx, pub in enumerate(publications, 1):
                citation = pub.get('formatted_citation', '')
                if citation:
                    p = doc.add_paragraph(style='List Number')
                    p.paragraph_format.space_after  = Pt(2)
                    p.paragraph_format.left_indent  = Inches(0.25)
                    citation_url = safe_url(pub.get('publication_url', ''))
                    if citation_url:
                        _add_hyperlink(p, citation, citation_url, size=10)
                    else:
                        run = p.add_run(citation)
                        run.font.size = Pt(10)

        # ── Footer: generation timestamp (+ optional AI attribution) ─────────
        _attr_text = '  ·  Generated with AI assistance' if content.get('ai_attribution') else ''
        for sec in doc.sections:
            footer = sec.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = fp.add_run(f"{timestamp}{_attr_text}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.italic = True

        if content.get('ai_attribution'):
            doc.core_properties.keywords = 'AI-assisted'
            doc.core_properties.subject  = 'CV generated with AI assistance using cv-builder'

        doc.save(str(filepath))
        logger.info("Human DOCX: %s", filename)
        return filepath

    # ── Achievement diversity helpers (GAP-243) ───────────────────────────────

    @staticmethod
    def _classify_achievement_impact(text: str) -> str:
        """Classify an achievement bullet into one of six impact-type buckets.

        Buckets: financial, leadership, cost, customer, technical, process.
        Classification is heuristic (keyword-based) and used only for
        diversity capping — it does not affect the rendered output.
        """
        t = (text or '').lower()
        if any(k in t for k in ('revenue', 'sales', ' $ ', '$', 'profit', 'arr',
                                  'mrr', 'usd', 'eur', 'gbp', 'pipeline')):
            return 'financial'
        if any(k in t for k in ('led ', 'managed ', 'team of', 'hired', 'mentored',
                                  'coached', 'headcount', 'direct report', 'people leader')):
            return 'leadership'
        if any(k in t for k in ('cost', 'saved ', 'saving', 'reduced ', 'reduction',
                                  'efficiency gain', 'overhead')):
            return 'cost'
        if any(k in t for k in ('customer', 'user', 'client', 'nps ', 'satisfaction',
                                  'retention', 'churn', 'support ticket')):
            return 'customer'
        if any(k in t for k in ('built ', 'developed ', 'designed ', 'implemented ',
                                  'deployed ', 'architected ', 'launched ')):
            return 'technical'
        return 'process'

    @classmethod
    def _apply_achievement_diversity(
        cls,
        scored: List[tuple],
        max_ach: int,
        max_type_fraction: float = 0.5,
    ) -> List[Dict]:
        """Select up to *max_ach* achievements with a per-impact-type diversity cap.

        When there are at least 3 distinct impact types represented across all
        scored achievements, no single type may account for more than
        *max_type_fraction* (default 50%) of the final selection.  When fewer
        than 3 types are present the cap is not applied (not enough diversity
        to enforce it meaningfully).

        Args:
            scored:            List of (achievement_dict, score) sorted descending.
            max_ach:           Maximum number of achievements to return.
            max_type_fraction: Maximum fraction per impact type (default 0.5).

        Returns:
            List of achievement dicts, length ≤ max_ach.
        """
        if max_ach <= 0:
            return []

        # Classify every candidate
        classified = [
            (ach, score, cls._classify_achievement_impact(
                ach.get('text', '') if isinstance(ach, dict) else str(ach)
            ))
            for ach, score in scored
        ]

        distinct_types = {t for _, _, t in classified}
        if len(distinct_types) < 3:
            # Not enough diversity to apply cap — fall back to straight top-N
            return [ach for ach, _, _ in classified[:max_ach]]

        cap = max(1, int(max_ach * max_type_fraction))
        type_counts: Dict[str, int] = {}
        selected: List[Dict] = []
        overflow: List[Dict] = []

        for ach, _score, itype in classified:
            if len(selected) >= max_ach:
                break
            if type_counts.get(itype, 0) < cap:
                selected.append(ach)
                type_counts[itype] = type_counts.get(itype, 0) + 1
            else:
                overflow.append(ach)

        # Backfill remaining slots with overflow in score order
        for ach in overflow:
            if len(selected) >= max_ach:
                break
            selected.append(ach)

        return selected

    # ── Pre-generation validation ─────────────────────────────────────────────

    @staticmethod
    def _detect_long_bullets(experiences: List[Dict], max_chars: int = 200) -> List[Dict]:
        """Return a list of warnings for experience bullets exceeding max_chars.

        Each warning is {company, title, bullet_text, char_count}.
        Long bullets typically wrap to 3+ lines in the generated DOCX.
        """
        warnings: List[Dict] = []
        for exp in experiences or []:
            if not isinstance(exp, dict):
                continue
            company = exp.get('company', '')
            title   = exp.get('title', '')
            bullets = exp.get('ordered_achievements') or exp.get('achievements') or []
            for bullet in bullets:
                text = bullet.get('text', '') if isinstance(bullet, dict) else str(bullet)
                if len(text) > max_chars:
                    warnings.append({
                        'company':    company,
                        'title':      title,
                        'bullet_text': text[:120] + '…' if len(text) > 120 else text,
                        'char_count': len(text),
                    })
        return warnings

    @staticmethod
    def _detect_sparse_experiences(experiences: List[Dict], min_bullets: int = 2) -> List[Dict]:
        """Return warnings for experience entries that have fewer than min_bullets selected bullets.

        Each warning is {company, title, bullet_count}.
        Entries with 0 or 1 bullets look sparse and may signal incomplete customisation.
        """
        warnings: List[Dict] = []
        for exp in experiences or []:
            if not isinstance(exp, dict):
                continue
            bullets = exp.get('ordered_achievements') or exp.get('achievements') or []
            count = len(bullets)
            if count < min_bullets:
                warnings.append({
                    'company':      exp.get('company', ''),
                    'title':        exp.get('title', ''),
                    'bullet_count': count,
                })
        return warnings

    @staticmethod
    def _detect_year_only_dates(experiences: List[Dict]) -> List[Dict]:
        """Return warnings for experience entries whose start/end dates are year-only (e.g. "2020").

        ATS parsers and recruiters expect month/year format. Year-only dates reduce precision
        and can make chronological ordering ambiguous. Each warning is
        {company, title, field, date_value}.
        """
        import re as _re2
        _year_only = _re2.compile(r'^\d{4}$')
        warnings: List[Dict] = []
        for exp in experiences or []:
            if not isinstance(exp, dict):
                continue
            company = exp.get('company', '')
            title   = exp.get('title', '')
            for field in ('start_date', 'end_date'):
                val = str(exp.get(field) or '').strip()
                if val and _year_only.match(val):
                    warnings.append({
                        'company':    company,
                        'title':      title,
                        'field':      field,
                        'date_value': val,
                    })
        return warnings

    @staticmethod
    def _detect_date_overlaps(experiences: List[Dict]) -> List[Dict]:
        """Return a list of overlap warnings for experience entries with overlapping date ranges.

        Each warning is a dict: {entry_a, entry_b, overlap_description}.
        Only checks entries where both start and end dates can be parsed.
        Overlapping roles at the same company (e.g. promotion) are excluded.
        """
        from datetime import date as _date  # noqa: PLC0415 (already imported at module level)

        _today = _date.today()

        def _parse_date(raw: str, end_of_period: bool = False) -> Optional[_date]:
            raw = (raw or '').strip()
            if not raw or raw.lower() in ('current', 'present', 'now', 'ongoing'):
                return _today if end_of_period else None
            for fmt in ('%Y-%m-%d', '%B %Y', '%b %Y', '%Y-%m', '%Y'):
                try:
                    d = datetime.strptime(raw, fmt).date()
                    # For end dates use last day of the parsed month/year
                    if end_of_period and fmt in ('%B %Y', '%b %Y', '%Y-%m', '%Y'):
                        import calendar  # noqa: PLC0415
                        if fmt == '%Y':
                            d = _date(d.year, 12, 31)
                        else:
                            last_day = calendar.monthrange(d.year, d.month)[1]
                            d = _date(d.year, d.month, last_day)
                    return d
                except ValueError:
                    pass
            m = re.search(r'\b(\d{4})\b', raw)
            if m:
                yr = int(m.group(1))
                return _date(yr, 12, 31) if end_of_period else _date(yr, 1, 1)
            return None

        parsed = []
        for exp in experiences:
            start = _parse_date(str(exp.get('start_date') or exp.get('start') or ''))
            end   = _parse_date(str(exp.get('end_date') or exp.get('end') or ''), end_of_period=True)
            if start and end:
                parsed.append({
                    'exp': exp,
                    'start': start,
                    'end': end,
                    'company': (exp.get('company') or '').strip().lower(),
                })

        warnings: List[Dict] = []
        for i in range(len(parsed)):
            for j in range(i + 1, len(parsed)):
                a, b = parsed[i], parsed[j]
                # Skip same-company overlaps (promotions, parallel roles)
                if a['company'] and b['company'] and a['company'] == b['company']:
                    continue
                # Overlap when one range starts before the other ends
                if a['start'] <= b['end'] and b['start'] <= a['end']:
                    def _fmt(exp_entry: Dict) -> str:
                        title   = exp_entry.get('title') or exp_entry.get('role', 'Unknown role')
                        company = exp_entry.get('company', '')
                        return f"{title} at {company}" if company else title

                    warnings.append({
                        'entry_a': _fmt(a['exp']),
                        'entry_b': _fmt(b['exp']),
                        'overlap_description': (
                            f"{a['start']} – {a['end']} overlaps with {b['start']} – {b['end']}"
                        ),
                    })
        return warnings

    @staticmethod
    def _verify_rewrite_audit_alignment(
        selected_content: Dict,
        rewrite_audit: List[Dict],
    ) -> List[Dict]:
        """Compare accepted/edited rewrite audit entries against the generated content.

        For each audit entry where ``outcome`` is ``'accept'`` or ``'edit'``, the
        expected final text is looked up in *selected_content* at the location
        described by the entry's ``type`` and ``location`` fields.  Any divergence
        between the expected text and the actual rendered text is returned as a
        mismatch warning.

        Returns a list of dicts: {id, type, location, expected, actual}.
        """
        mismatches: List[Dict] = []

        def _norm(text: str) -> str:
            return ' '.join((text or '').split())

        for entry in rewrite_audit or []:
            outcome = entry.get('outcome', '')
            if outcome not in ('accept', 'edit'):
                continue

            if outcome == 'accept':
                expected = entry.get('proposed', '') or ''
            else:
                expected = entry.get('final', '') or ''

            if not expected:
                continue

            kind     = entry.get('type', '')
            loc      = entry.get('location', '')
            entry_id = entry.get('id', '<unknown>')
            actual   = None

            if kind == 'summary' or loc == 'summary':
                actual = selected_content.get('summary', '') or ''

            elif kind == 'bullet':
                m = re.match(r'^([^.]+)\.achievements\[(\d+)\]$', loc)
                if m:
                    exp_id  = m.group(1)
                    ach_idx = int(m.group(2))
                    for exp in selected_content.get('experiences', []):
                        if exp.get('id') == exp_id:
                            achs = exp.get('achievements', [])
                            if 0 <= ach_idx < len(achs):
                                ach = achs[ach_idx]
                                actual = ach.get('text', '') if isinstance(ach, dict) else str(ach)
                            break

            elif kind == 'skill_rename':
                original = entry.get('original', '')
                for skill in selected_content.get('skills', []):
                    if isinstance(skill, dict) and skill.get('name') == expected:
                        actual = expected
                        break
                    if isinstance(skill, dict) and skill.get('name') == original:
                        actual = original
                        break

            if actual is None:
                continue

            if _norm(actual) != _norm(expected):
                mismatches.append({
                    'id':       entry_id,
                    'type':     kind,
                    'location': loc,
                    'expected': expected,
                    'actual':   actual,
                })

        return mismatches


# ── Module-level ATS validation ──────────────────────────────────────────────

def validate_ats_report(output_dir: Path, job_analysis: Dict) -> tuple:
    """Run 16 ATS validation checks on the generated CV files.

    Args:
        output_dir:   Path to the job-specific output directory.
        job_analysis: Job analysis dict (for ATS keyword checks).

    Returns:
        ``(checks, page_count)`` where *checks* is a list of dicts:
        ``{name, label, format, status, detail}`` with status
        ``'pass' | 'warn' | 'fail'`` and *page_count* is an ``int | None``.
        *format* is ``'docx' | 'html' | 'pdf' | 'all'``.
    """
    import re as _re
    import json as _json

    checks: List[Dict] = []

    def _chk(name: str, label: str, fmt: str, status: str, detail: str) -> None:
        checks.append({'name': name, 'label': label, 'format': fmt,
                       'status': status, 'detail': detail})

    # ── locate files ─────────────────────────────────────────────────────────
    ats_docx_files = sorted(output_dir.glob('*_ATS.docx'))
    html_files     = sorted(output_dir.glob('*.html'))
    # Sort PDFs by modification time (newest first) so the most recently finalised
    # PDF is used for the page-count check, not an earlier dated draft.
    pdf_files      = sorted(
        (f for f in output_dir.glob('*.pdf') if '_ATS' not in f.name),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )

    ats_docx  = ats_docx_files[0] if ats_docx_files else None
    html_path = html_files[0]     if html_files     else None
    pdf_path  = pdf_files[0]      if pdf_files      else None

    # ── DOCX checks 1-8, 16 ──────────────────────────────────────────────────
    DOCX_CHECKS = [
        ('docx_text_selectable',       'DOCX text selectable'),
        ('docx_zero_tables',           'No tables in DOCX'),
        ('docx_zero_shapes',           'No text boxes / shapes'),
        ('docx_contact_in_body',       'Contact info in body'),
        ('docx_standard_headings',     'Standard heading text'),
        ('docx_heading1_present',      'Heading 1 style present'),
        ('docx_date_format_consistent','Consistent date formats'),
        ('ats_keyword_presence',       'ATS keyword presence'),
        ('docx_publications_heading',  'Publications heading text'),
    ]

    if ats_docx is None:
        for name, label in DOCX_CHECKS:
            fmt = 'all' if name == 'ats_keyword_presence' else 'docx'
            _chk(name, label, fmt, 'fail', 'ATS DOCX file not found')
    else:
        try:
            from docx import Document as _Document
            doc        = _Document(str(ats_docx))
            paragraphs = doc.paragraphs
            docx_text  = '\n'.join(p.text for p in paragraphs if p.text.strip())

            # 1 — text selectable
            if len(docx_text) > 100:
                _chk('docx_text_selectable', 'DOCX text selectable', 'docx',
                     'pass', f'{len(docx_text):,} characters extracted')
            else:
                _chk('docx_text_selectable', 'DOCX text selectable', 'docx',
                     'fail', 'Little or no text extracted — document may be image-based')

            # 2 — zero tables
            n_tables = len(doc.tables)
            if n_tables == 0:
                _chk('docx_zero_tables', 'No tables in DOCX', 'docx', 'pass', 'No tables found')
            else:
                _chk('docx_zero_tables', 'No tables in DOCX', 'docx', 'fail',
                     f'{n_tables} table(s) — ATS parsers may skip table content')

            # 3 — zero shapes
            # Use Clark notation directly to avoid KeyError on missing nsmap entries
            # (e.g. 'v' namespace absent in some python-docx versions).
            _VML = 'urn:schemas-microsoft-com:vml'
            _MC  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
            shapes = (doc.element.body.findall('.//{%s}textbox' % _VML) +
                      doc.element.body.findall('.//{%s}Fallback' % _MC))
            if not shapes:
                _chk('docx_zero_shapes', 'No text boxes / shapes', 'docx', 'pass', 'No shapes found')
            else:
                _chk('docx_zero_shapes', 'No text boxes / shapes', 'docx', 'warn',
                     f'{len(shapes)} shape element(s) — content may be unreadable by ATS')

            # 4 — contact in body
            email_re = _re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')
            if email_re.search(docx_text):
                _chk('docx_contact_in_body', 'Contact info in body', 'docx', 'pass',
                     'Email address found in document body')
            else:
                _chk('docx_contact_in_body', 'Contact info in body', 'docx', 'fail',
                     'No email address in body text — ATS may miss contact info')

            # 5 & 6 — headings
            STANDARD = frozenset({
                'experience', 'education', 'skills', 'summary', 'publications',
                'certifications', 'achievements', 'awards', 'objective',
                'work experience', 'professional experience', 'technical skills',
                'professional summary', 'selected publications', 'contact',
                'portfolio', 'languages', 'volunteering', 'projects', 'career history',
                'core competencies',
            })
            heading_paras = [p for p in paragraphs if p.style.name.startswith('Heading')]
            heading_texts = [p.text.strip() for p in heading_paras if p.text.strip()]

            # Check if a heading matches a standard heading with word boundaries
            def is_standard_heading(text: str, standards: frozenset) -> bool:
                text_lower = text.lower()
                # Exact match
                if text_lower in standards:
                    return True
                # Word-boundary match: check if any standard heading appears as a complete word
                for standard in standards:
                    if _re.search(r'\b' + _re.escape(standard) + r'\b', text_lower):
                        return True
                return False

            unexpected    = [t for t in heading_texts
                             if not is_standard_heading(t, STANDARD)]
            if not unexpected:
                _chk('docx_standard_headings', 'Standard heading text', 'docx', 'pass',
                     f'{len(heading_texts)} standard section heading(s) found')
            else:
                _chk('docx_standard_headings', 'Standard heading text', 'docx', 'warn',
                     f'Unexpected heading(s): {", ".join(unexpected[:3])}')

            h1_count = sum(1 for p in heading_paras if p.style.name == 'Heading 1')
            if h1_count > 0:
                _chk('docx_heading1_present', 'Heading 1 style present', 'docx', 'pass',
                     f'{h1_count} Heading 1 paragraph(s) found')
            else:
                _chk('docx_heading1_present', 'Heading 1 style present', 'docx', 'warn',
                     'No Heading 1 paragraphs — ATS relies on heading hierarchy')

            # 6b — candidate name casing
            candidate_name = heading_texts[0].strip() if heading_texts else ''
            if candidate_name:
                name_alpha = ''.join(c for c in candidate_name if c.isalpha())
                if name_alpha and name_alpha.isupper():
                    _chk('docx_name_casing', 'Candidate name casing', 'docx', 'warn',
                         f'Name "{candidate_name}" is ALL-CAPS — some ATS systems may mis-parse it')
                elif name_alpha and name_alpha.islower():
                    _chk('docx_name_casing', 'Candidate name casing', 'docx', 'warn',
                         f'Name "{candidate_name}" is all-lowercase — ATS systems expect Title Case')
                else:
                    _chk('docx_name_casing', 'Candidate name casing', 'docx', 'pass',
                         f'Name "{candidate_name}" appears normally cased')

            # 7 — consistent dates
            date_pats = [
                (_re.compile(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b'),
                 'Mon YYYY'),
                (_re.compile(r'\b\d{1,2}/\d{4}\b'), 'MM/YYYY'),
                (_re.compile(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b'),
                 'Full Month YYYY'),
                (_re.compile(r'\b\d{4}-\d{2}(?:-\d{2})?\b'), 'ISO (YYYY-MM or YYYY-MM-DD)'),
                (_re.compile(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December).*?(?:–|-|—).*?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|Present|Current)'),
                 'Date Range'),
                (_re.compile(r'\b(?:Present|Current)\b', _re.IGNORECASE), 'Present/Current'),
            ]
            found_fmts = {name for pat, name in date_pats if pat.search(docx_text)}
            if len(found_fmts) <= 1:
                _chk('docx_date_format_consistent', 'Consistent date formats', 'docx', 'pass',
                     f'Date format: {next(iter(found_fmts), "not detected")}')
            else:
                _chk('docx_date_format_consistent', 'Consistent date formats', 'docx', 'fail',
                     f'Mixed formats — {" and ".join(sorted(found_fmts))} — standardise to one')

            # 8 — ATS keywords (two-tier: required-skill keywords take priority)
            req_skills = [s.lower() for s in job_analysis.get('required_skills', [])[:10]]
            ats_kws    = [k.lower() for k in job_analysis.get('ats_keywords', [])[:15]]
            if not ats_kws and not req_skills:
                _chk('ats_keyword_presence', 'ATS keyword presence', 'all', 'warn',
                     'No ATS keywords defined in job analysis')
            else:
                text_lower = docx_text.lower()

                def _kw_in_text(kw: str, txt: str) -> bool:
                    """Case-insensitive match with hyphen/slash variant normalization."""
                    if kw in txt:
                        return True
                    # Slash form: any part of "ml/mlops" matches "mlops" or "ml"
                    for part in kw.split('/'):
                        part = part.strip()
                        if part and len(part) > 1 and part in txt:
                            return True
                    # Hyphen equivalence: "scikit-learn" matches "scikit learn" or "scitkitlearn"
                    hyph_space = kw.replace('-', ' ')
                    hyph_none  = kw.replace('-', '')
                    if hyph_space != kw and hyph_space in txt:
                        return True
                    if hyph_none != kw and len(hyph_none) > 2 and hyph_none in txt:
                        return True
                    return False

                # Tier 1: required-skill keywords (high-weight; should all be present)
                req_set      = set(req_skills)
                missing_req  = [kw for kw in req_skills if not _kw_in_text(kw, text_lower)]
                # Tier 2: supplemental ATS keywords not already covered by required_skills
                supplemental = [kw for kw in ats_kws if kw not in req_set][:10]
                missing_supp = [kw for kw in supplemental if not _kw_in_text(kw, text_lower)]
                all_checked  = req_skills + supplemental
                missing      = missing_req + missing_supp

                tier_parts = []
                if req_skills:
                    tier_parts.append(f'Required: {len(req_skills) - len(missing_req)}/{len(req_skills)}')
                if supplemental:
                    tier_parts.append(f'Optional: {len(supplemental) - len(missing_supp)}/{len(supplemental)}')
                tier_note = ' | '.join(tier_parts)

                if not missing:
                    _chk('ats_keyword_presence', 'ATS keyword presence', 'all', 'pass',
                         f'All {len(all_checked)} keywords present ({tier_note})')
                elif not missing_req and missing_supp:
                    _chk('ats_keyword_presence', 'ATS keyword presence', 'all', 'warn',
                         f'Required skills all present; {len(missing_supp)} optional keyword(s) missing: '
                         f'{", ".join(missing_supp[:5])}{"…" if len(missing_supp) > 5 else ""}')
                elif len(missing) <= max(1, len(all_checked) // 3):
                    _chk('ats_keyword_presence', 'ATS keyword presence', 'all', 'warn',
                         f'{len(missing)} keyword(s) missing ({tier_note}): {", ".join(missing[:5])}')
                else:
                    _chk('ats_keyword_presence', 'ATS keyword presence', 'all', 'fail',
                         (f'{len(missing)}/{len(all_checked)} keywords missing ({tier_note}): '
                          f'{", ".join(missing[:5])}{"…" if len(missing) > 5 else ""}'))

            # 8b — keyword density (top 5 ATS keywords should appear ≥2 times)
            if ats_kws:
                top_kws = ats_kws[:5]
                thin = []
                for kw in top_kws:
                    # Count occurrences using split to avoid partial matches
                    count = text_lower.count(kw)
                    # Slash/hyphen variants
                    for part in kw.split('/'):
                        part = part.strip()
                        if part and len(part) > 1 and part != kw:
                            count += text_lower.count(part)
                    hyph_space = kw.replace('-', ' ')
                    if hyph_space != kw:
                        count += text_lower.count(hyph_space)
                    if count < 2:
                        thin.append(f'"{kw}" ({count}×)')
                if not thin:
                    _chk('ats_keyword_density', 'ATS keyword density', 'docx', 'pass',
                         f'Top {len(top_kws)} keywords each appear ≥2 times')
                else:
                    _chk('ats_keyword_density', 'ATS keyword density', 'docx', 'warn',
                         f'Low-frequency keywords: {", ".join(thin[:3])} — consider reinforcing')

            # 7b — year-only dates (warn: ATS parsers expect month+year)
            _year_only_re = _re.compile(
                r'(?<!\d)\b(19|20)\d{2}\b(?!\s*[-–—]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d))'
            )
            year_only_matches = _year_only_re.findall(docx_text)
            if year_only_matches:
                _chk('docx_year_only_dates', 'No year-only date entries', 'docx', 'warn',
                     f'{len(year_only_matches)} year-only date(s) detected — add month for better ATS parsing')
            else:
                _chk('docx_year_only_dates', 'No year-only date entries', 'docx', 'pass',
                     'All dates include month and year')

            # 16 — publications heading
            pub_headings = [p for p in heading_paras if 'publication' in p.text.lower()]
            if not pub_headings:
                _chk('docx_publications_heading', 'Publications heading text', 'docx', 'pass',
                     'No publications section (optional)')
            else:
                _allowed = {'Publications', 'Selected Publications'}
                wrong = [p.text.strip() for p in pub_headings
                         if p.text.strip() not in _allowed]
                if not wrong:
                    _chk('docx_publications_heading', 'Publications heading text', 'docx',
                         'pass', 'Heading is "Publications" or "Selected Publications"')
                else:
                    _chk('docx_publications_heading', 'Publications heading text', 'docx',
                         'fail',
                         f'Heading "{wrong[0]}" must be "Publications" or "Selected Publications"')

            # 17 — ATS-safe font compliance (GAP-87)
            _ATS_SAFE_FONTS = frozenset({
                'arial', 'calibri', 'times new roman', 'helvetica',
                'georgia', 'garamond', 'verdana', 'trebuchet ms',
                'courier new', 'palatino', 'book antiqua',
            })
            _font_names: set = set()
            for _p in paragraphs:
                for _run in _p.runs:
                    _fn = (_run.font.name or '').strip().lower()
                    if _fn:
                        _font_names.add(_fn)
            # Also check the document default font
            _def_font = (
                doc.styles['Normal'].font.name if 'Normal' in doc.styles else None
            )
            if _def_font:
                _font_names.add(_def_font.strip().lower())
            if not _font_names:
                _chk('docx_ats_safe_fonts', 'ATS-safe fonts only', 'docx', 'warn',
                     'No explicit font names found in DOCX runs — default theme font assumed')
            else:
                _non_ats = sorted(_fn for _fn in _font_names if _fn not in _ATS_SAFE_FONTS)
                if not _non_ats:
                    _chk('docx_ats_safe_fonts', 'ATS-safe fonts only', 'docx', 'pass',
                         f'All fonts ATS-safe: {", ".join(sorted(_font_names))}')
                else:
                    _chk('docx_ats_safe_fonts', 'ATS-safe fonts only', 'docx', 'warn',
                         (f'Non-standard fonts detected: {", ".join(_non_ats)}'
                          ' — some ATS engines may mis-parse these'))

        except Exception as exc:
            for name, label in DOCX_CHECKS:
                fmt = 'all' if name == 'ats_keyword_presence' else 'docx'
                _chk(name, label, fmt, 'fail', f'DOCX check error: {exc}')

    # ── HTML checks 9-12 ─────────────────────────────────────────────────────
    HTML_CHECKS = [
        ('html_jsonld_present',       'HTML JSON-LD present'),
        ('html_jsonld_valid_person',  'JSON-LD is schema.org/Person'),
        ('html_jsonld_knows_about',   'JSON-LD knowsAbout populated'),
        ('html_required_fields',      'JSON-LD name + email present'),
    ]
    if html_path is None:
        for name, label in HTML_CHECKS:
            _chk(name, label, 'html', 'fail', 'HTML file not found')
    else:
        try:
            from bs4 import BeautifulSoup as _BS
            html_src    = html_path.read_text(encoding='utf-8', errors='replace')
            soup        = _BS(html_src, 'html.parser')
            jsonld_tags = soup.find_all('script', type='application/ld+json')

            if not jsonld_tags:
                for name, label in HTML_CHECKS:
                    _chk(name, label, 'html', 'fail', 'No JSON-LD <script> block found')
            else:
                _chk('html_jsonld_present', 'HTML JSON-LD present', 'html', 'pass',
                     f'{len(jsonld_tags)} JSON-LD block(s) found')
                try:
                    jld = _json.loads(jsonld_tags[0].string or '{}')
                    # 10
                    if (
                        jld.get('@type') == 'Person'
                        and _is_exact_schema_org_context(jld.get('@context'))
                    ):
                        _chk('html_jsonld_valid_person', 'JSON-LD is schema.org/Person',
                             'html', 'pass', '@type: Person with schema.org context')
                    else:
                        _chk('html_jsonld_valid_person', 'JSON-LD is schema.org/Person',
                             'html', 'fail',
                             f'@type="{jld.get("@type","missing")}", expected Person')
                    # 11
                    ka = jld.get('knowsAbout', [])
                    if len(ka) >= 3:
                        _chk('html_jsonld_knows_about', 'JSON-LD knowsAbout populated',
                             'html', 'pass', f'{len(ka)} skills listed')
                    elif ka:
                        _chk('html_jsonld_knows_about', 'JSON-LD knowsAbout populated',
                             'html', 'warn', f'Only {len(ka)} skill(s) in knowsAbout')
                    else:
                        _chk('html_jsonld_knows_about', 'JSON-LD knowsAbout populated',
                             'html', 'fail', 'knowsAbout absent or empty')
                    # 12
                    missing_flds = [f for f in ('name', 'email') if not jld.get(f, '').strip()]
                    if not missing_flds:
                        _chk('html_required_fields', 'JSON-LD name + email present',
                             'html', 'pass',
                             f'name="{jld.get("name","")}", email="{jld.get("email","")}"')
                    else:
                        _chk('html_required_fields', 'JSON-LD name + email present',
                             'html', 'fail',
                             f'Missing required fields: {", ".join(missing_flds)}')
                except _json.JSONDecodeError as exc:
                    for name, label in HTML_CHECKS[1:]:
                        _chk(name, label, 'html', 'fail', f'JSON-LD parse error: {exc}')

        except Exception as exc:
            for name, label in HTML_CHECKS:
                _chk(name, label, 'html', 'fail', f'HTML check error: {exc}')

    # ── PDF render checks 13, 15 — read from already-generated PDF ────────────
    # No re-rendering needed: Chrome already wrote the PDF; pypdf reads it here
    # for page count (reused below) and selectable-text verification.
    page_count: Optional[int] = None
    if pdf_path is None:
        _chk('html_renders_ok', 'PDF generated successfully', 'pdf', 'fail', 'PDF file not found')
        _chk('pdf_has_text',    'PDF has selectable text',    'pdf', 'fail', 'PDF file not found')
    else:
        try:
            import pypdf as _pypdf2
            _reader    = _pypdf2.PdfReader(str(pdf_path))
            page_count = len(_reader.pages)
            _chk('html_renders_ok', 'PDF generated successfully', 'pdf', 'pass',
                 f'PDF has {page_count} page(s)')
            # Check that at least some text is selectable (not a blank/image-only render)
            _pdf_text = ''.join(
                _reader.pages[i].extract_text() or ''
                for i in range(min(page_count, 3))
            ).strip()
            if len(_pdf_text) > 50:
                _chk('pdf_has_text', 'PDF has selectable text', 'pdf', 'pass',
                     f'{len(_pdf_text):,} characters extractable from PDF')
            else:
                _chk('pdf_has_text', 'PDF has selectable text', 'pdf', 'warn',
                     'Little text extractable — PDF may be image-based')
        except Exception as exc:
            _chk('html_renders_ok', 'PDF generated successfully', 'pdf', 'fail',
                 f'PDF read error: {str(exc)[:200]}')
            _chk('pdf_has_text', 'PDF has selectable text', 'pdf', 'fail',
                 'PDF could not be read')

    # ── PDF size check 14 ────────────────────────────────────────────────────
    if pdf_path is None:
        _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'fail', 'PDF file not found')
    else:
        try:
            import pypdf as _pypdf
            reader = _pypdf.PdfReader(str(pdf_path))
            if reader.pages:
                w = float(reader.pages[0].mediabox.width)
                h = float(reader.pages[0].mediabox.height)
                # Normalise to portrait
                w, h = min(w, h), max(w, h)
                if abs(w - 612) < 6 and abs(h - 792) < 6:
                    _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'pass',
                         f'{w:.0f}×{h:.0f} pts — Letter')
                elif abs(w - 595) < 6 and abs(h - 842) < 6:
                    _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'warn',
                         f'{w:.0f}×{h:.0f} pts — appears A4, not US Letter')
                else:
                    _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'warn',
                         f'{w:.0f}×{h:.0f} pts — unexpected page size')
            else:
                _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'fail', 'PDF has no pages')
        except Exception as exc:
            _chk('pdf_us_letter', 'PDF page size is US Letter', 'pdf', 'fail',
                 f'PDF check error: {exc}')

    # ── PDF font embedding check ───────────────────────────────────────────
    if pdf_path is None:
        _chk('pdf_fonts_embedded', 'PDF fonts embedded', 'pdf', 'fail', 'PDF file not found')
    else:
        try:
            import pypdf as _pypdf3
            _reader3 = _pypdf3.PdfReader(str(pdf_path))
            # Walk the /Font resource dictionary for each page
            _unembedded: list = []
            _seen_fonts: set = set()
            for _page in _reader3.pages:
                _res = _page.get('/Resources', {})
                _font_dict = _res.get('/Font', {})
                for _fname, _fobj in _font_dict.items():
                    try:
                        _font_name = str(_fobj.get('/BaseFont', _fname) or _fname)
                        if _font_name in _seen_fonts:
                            continue
                        _seen_fonts.add(_font_name)
                        # A font is embedded if it has a /FontDescriptor with /FontFile*
                        _fd = _fobj.get('/FontDescriptor', {})
                        _has_file = any(
                            k in _fd for k in ('/FontFile', '/FontFile2', '/FontFile3')
                        )
                        if not _has_file:
                            _unembedded.append(_font_name.lstrip('/'))
                    except Exception:
                        pass
            if not _seen_fonts:
                _chk('pdf_fonts_embedded', 'PDF fonts embedded', 'pdf', 'warn',
                     'No font resources found in PDF — font embedding status unknown')
            elif not _unembedded:
                _chk('pdf_fonts_embedded', 'PDF fonts embedded', 'pdf', 'pass',
                     f'All {len(_seen_fonts)} font(s) embedded')
            else:
                _names = ', '.join(_unembedded[:3])
                _chk('pdf_fonts_embedded', 'PDF fonts embedded', 'pdf', 'warn',
                     f'{len(_unembedded)} font(s) not embedded: {_names} — text may not extract correctly')
        except Exception as exc:
            _chk('pdf_fonts_embedded', 'PDF fonts embedded', 'pdf', 'warn',
                 f'Font embedding check failed: {str(exc)[:100]}')

    # ── Page Count Validation ──────────────────────────────────────────────
    # Check CV length against ideal and absolute limits
    cfg = get_config()
    ideal_min = cfg.get('generation.page_count.ideal_min', 2)
    ideal_max = cfg.get('generation.page_count.ideal_max', 3)
    absolute_max = cfg.get('generation.page_count.absolute_max', 4)

    if page_count is None:
        _chk('cv_page_count', 'CV page count', 'pdf', 'fail',
             'Page count could not be determined (HTML render failed)')
    elif page_count == 1:
        _chk('cv_page_count', 'CV page count', 'pdf', 'warn',
             f'{page_count} page — consider {ideal_min}–{ideal_max} pages for senior candidates')
    elif ideal_min <= page_count <= ideal_max:
        _chk('cv_page_count', 'CV page count', 'pdf', 'pass',
             f'{page_count} pages — within ideal {ideal_min}–{ideal_max} page range')
    elif page_count > absolute_max:
        _chk('cv_page_count', 'CV page count', 'pdf', 'fail',
             f'{page_count} pages — exceeds {absolute_max}-page maximum; consider condensing')
    else:  # ideal_max < page_count <= absolute_max
        _chk('cv_page_count', 'CV page count', 'pdf', 'warn',
             f'{page_count} pages — exceeds {ideal_max}-page ideal range')

    return checks, page_count
